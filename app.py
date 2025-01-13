# import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt

import os
import filetype
import docx
import PyPDF2
import re
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.memory import ConversationSummaryMemory
import asyncio
import numpy as np
import json

import google.generativeai as genai
import pathlib
import logging
import sys
import io
import matplotlib.pyplot as plt
import seaborn as sns
# Import things that are needed generically
from langchain.pydantic_v1 import BaseModel, Field
from langchain.tools import BaseTool, StructuredTool, tool
# Define functions to generate investment suggestions :\

import random
from datetime import datetime
import os
import json
from email.mime.text import MIMEText
import smtplib
import jwt

USE_AWS = True  # Set to False to use local storage

# # -------------------------------------Start Aws---------------------
# import paramiko

# # Set up the SSH key file, IP, username, and passphrase
# key_path = "keys/aws_key.pem"  # Path to the converted .pem file
# hostname = "172.31.15.173"  # AWS EC2 public IP address
# username = "pragatidhobe"  # EC2 instance username
# passphrase = "12345678"  # Passphrase, if any

# # Create an SSH client instance
# ssh_client = paramiko.SSHClient()
# ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())

# try:
#     # Load SSH key
#     key = paramiko.RSAKey.from_private_key_file(key_path, password=passphrase)

#     # Connect to the instance
#     ssh_client.connect(hostname=hostname, username=username, pkey=key)

#     # Execute a command (example)
#     stdin, stdout, stderr = ssh_client.exec_command("ls")
#     print(stdout.read().decode())  # Print command output

# except Exception as e:
#     print(f"An error occurred: {e}")

# finally:
#     ssh_client.close()

# # -------------------------------------End Aws---------------------



import boto3
load_dotenv()

# # AWS keys
aws_access_key = os.getenv('aws_access_key')
aws_secret_key = os.getenv('aws_secret_key')
S3_BUCKET_NAME = os.getenv('S3_BUCKET_NAME')
client_summary_folder = os.getenv('client_summary_folder') 
suggestions_folder = os.getenv('suggestions_folder') 
order_list_folder = os.getenv('order_list_folder')
portfolio_list_folder = os.getenv('portfolio_list_folder') 
personality_assessment_folder = os.getenv('personality_assessment_folder') 
login_folder = os.getenv('login_folder')
daily_changes_folder = os.getenv('daily_changes_folder')
signUp_user_folder = os.getenv('signUp_user_folder')
PREDICTIONS_FOLDER = os.getenv('PREDICTIONS_FOLDER')
# Connecting to Amazon S3
s3 = boto3.client(
    's3',
    aws_access_key_id=aws_access_key,
    aws_secret_access_key=aws_secret_key
)

def list_s3_keys(bucket_name, prefix=""):
    try:
        # List objects in the bucket with the given prefix
        response = s3.list_objects_v2(Bucket=bucket_name, Prefix=prefix)
        if 'Contents' in response:
            print("Keys in the S3 folder:")
            for obj in response['Contents']:
                print(obj['Key'])
        else:
            print("No files found in the specified folder.")
    except Exception as e:
        print(f"Error listing objects in S3: {e}")

# Call the function
# list_s3_keys(S3_BUCKET_NAME, signUp_user_folder)
# list_s3_keys(S3_BUCKET_NAME, order_list_folder) 
# list_s3_keys(S3_BUCKET_NAME, portfolio_list_folder) 


####################################################################################

# Download the files :

# def download_json_file(bucket_name, file_key, local_file_path):
#     """
#     Downloads a JSON file from an S3 bucket to a local file.

#     :param bucket_name: Name of the S3 bucket
#     :param file_key: Key (path) of the file in the bucket
#     :param local_file_path: Path to save the file locally
#     """
#     try:
#         # Download the file
#         s3.download_file(bucket_name, file_key, local_file_path)
#         print(f"File '{file_key}' successfully downloaded to '{local_file_path}'.")
        
#         # Load the JSON content to confirm it's valid
#         with open(local_file_path, 'r') as file:
#             data = json.load(file)
#             print("Downloaded JSON content:", data)
#         return data
    
#     except Exception as e:
#         print(f"Error downloading file: {e}")
#         return None

# # Example usage 
# local_file_path = "local_data\orders\orders_SF3648.json"
# downloaded_data = download_json_file(S3_BUCKET_NAME,"order_list_folder/SF3648_orders.json", local_file_path)

# downloaded_data = download_json_file(S3_BUCKET_NAME, "portfolio_list_folder//SF3648.json", local_file_path)

# print(downloaded_data)

######################################################################################

# Upload File :

# def upload_json_file(bucket_name, file_key, local_file_path):
#     """
#     Uploads a JSON file from the local file system to an S3 bucket.

#     :param bucket_name: Name of the S3 bucket
#     :param file_key: Key (path) to upload the file to in the bucket
#     :param local_file_path: Path of the file locally to be uploaded
#     """
#     try:
#         # Upload the file
#         s3.upload_file(local_file_path, bucket_name, file_key, ExtraArgs={'ContentType': 'application/json'})
#         print(f"File '{local_file_path}' successfully uploaded to '{file_key}' in bucket '{bucket_name}'.")
    
#     except Exception as e:
#         print(f"Error uploading file: {e}")

# # Example usage
# local_file_path = "local_data\portfolios\portfolio_SF3648.json"
# upload_json_file(S3_BUCKET_NAME, "portfolio_list_folder//SF3648.json", local_file_path)

# local_file_path = "local_data\orders\orders_SF3648.json"
# upload_json_file(S3_BUCKET_NAME, "order_list_folder/SF3648_orders.json", local_file_path)

# list_s3_keys(S3_BUCKET_NAME, order_list_folder) 


######################################################################################


# delete folders/files from bucket :



# S3 bucket and file details

# FILE_KEY = "order_list_folder/JM4162_orders.json"
# FILE_KEY = "order_list_folder/JR5059_orders.json"
# FILE_KEY = "portfolio_list_folder//JM4162.json"
# FILE_KEY = "portfolio_list_folder//JR5059.json"

# FILE_KEY = "order_list_folder/SF3648_orders.json" 

# FILE_KEY = "order_list_folder/SF3648_orders.json" 


# def delete_file_from_s3(bucket_name, file_key):
#     """
#     Deletes a specified file from an S3 bucket.

#     :param bucket_name: Name of the S3 bucket
#     :param file_key: Key (path) of the file in the bucket
#     """
#     try:
#         # Delete the file
#         response = s3.delete_object(Bucket=bucket_name, Key=file_key)
        
#         # Confirm deletion
#         if response.get('ResponseMetadata', {}).get('HTTPStatusCode') == 204:
#             print(f"File '{file_key}' successfully deleted from bucket '{bucket_name}'.")
#         else:
#             print(f"Failed to delete file '{file_key}' from bucket '{bucket_name}'.")
    
#     except Exception as e:
#         print(f"Error deleting file: {e}")

# # Call the function
# delete_file_from_s3(S3_BUCKET_NAME, FILE_KEY)

# # list_s3_keys(S3_BUCKET_NAME, portfolio_list_folder) 
# list_s3_keys(S3_BUCKET_NAME, order_list_folder) 



# =------------------------------------------------------=






GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

from flask import Flask, request, jsonify

# app = Flask(__name__)
from flask import Flask, request, jsonify, send_file
import asyncio
from flask_cors import CORS
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://localhost:3000"}})

# Configure generativeai with your API key
genai.configure(api_key=GOOGLE_API_KEY)

import markdown

########################### Sign in Sign using aws ###################################################

app.config['JWT_SECRET_KEY'] = os.getenv('JWT_SECRET_KEY')

from flask import Flask, request, jsonify
from flask_bcrypt import Bcrypt
from flask_mail import Mail, Message
import random
import boto3
import json
from datetime import datetime, timedelta,timezone

bcrypt = Bcrypt(app)

# Email configuration
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME') # 'your_email@gmail.com'
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD') #'your_email_password'

mail = Mail(app)

# In-memory storage for email and OTP (for simplicity)
otp_store = {}

# API Endpoints
from flask import Flask, request, jsonify
import random
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# Replace with your email credentials
EMAIL_ADDRESS = os.getenv('MAIL_USERNAME')  #'your-email@gmail.com'
EMAIL_PASSWORD = os.getenv('MAIL_PASSWORD')  #'your-email-password'

 
# Local storage paths
LOCAL_STORAGE_PATH = "local_storage"
os.makedirs(LOCAL_STORAGE_PATH, exist_ok=True)
# otp_store = {}
 
# def load_from_local(filepath):
#     try:
#         if not os.path.exists(filepath):
#             return None
#         with open(filepath, 'r') as file:
#             return json.load(file)
#     except Exception as e:
#         print(f"Error loading file: {e}")
#         return None
   
# def save_to_local(data, filepath):
#     try:
#         os.makedirs(os.path.dirname(filepath), exist_ok=True)
#         with open(filepath, 'w') as file:
#             json.dump(data, file)
#         print(f"Data saved at {filepath}")  # Debug log
#     except Exception as e:
#         print(f"Error saving file: {e}")  # Debug log
#         raise
 
 
# def delete_from_local(filename):
    # file_path = os.path.join(LOCAL_STORAGE_PATH, filename)
    # if os.path.exists(file_path):
    #     os.remove(file_path)
       
# Using AWS and Local Storage :

from botocore.exceptions import NoCredentialsError, PartialCredentialsError


def load_from_local(filepath):
    try:
        if not os.path.exists(filepath):
            return None
        with open(filepath, 'r') as file:
            return json.load(file)
    except Exception as e:
        print(f"Error loading file: {e}")
        return None
 
def save_to_local(data, filepath):
    try:
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        with open(filepath, 'w') as file:
            json.dump(data, file)
        print(f"Data saved at {filepath}")  # Debug log
    except Exception as e:
        print(f"Error saving file: {e}")  # Debug log
        raise
 
def delete_from_local(filename):
    file_path = os.path.join(LOCAL_STORAGE_PATH, filename)
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            print(f"File {filename} deleted from local storage.")  # Debug log
        else:
            print(f"File {filename} not found in local storage.")  # Debug log
    except Exception as e:
        print(f"Error deleting file: {e}")  # Debug log
        raise
 
def save_to_aws(data, filename):
    try:
        s3.put_object(
            Bucket=S3_BUCKET_NAME,
            Key=filename,
            Body=json.dumps(data),
            ContentType='application/json'
        )
        print(f"Data saved to AWS at {filename}")
    except (NoCredentialsError, PartialCredentialsError) as e:
        print(f"AWS credentials error: {e}")
        raise
    except Exception as e:
        print(f"Error saving to AWS: {e}")
        raise
 
def load_from_aws(filename):
    try:
        obj = s3.get_object(Bucket=S3_BUCKET_NAME, Key=filename)
        return json.loads(obj['Body'].read().decode('utf-8'))
    except s3.exceptions.NoSuchKey:
        return None
    except Exception as e:
        print(f"Error loading from AWS: {e}")
        return None
 
def save_user_data(data, email):
    if USE_AWS:
        # Store in AWS under 'signUp_user_folder/<email>.json'
        filename = f"{signUp_user_folder}{email}.json"
        save_to_aws(data, filename)
    else:
        # Store locally under 'users/<email>.json'
        filename = os.path.join(LOCAL_STORAGE_PATH, f"users/{email}.json")
        save_to_local(data, filename)
 
def load_user_data(email):
    if USE_AWS:
        filename = f"{signUp_user_folder}{email}.json"
        return load_from_aws(filename)
    else:
        filename = f"users/{email}.json"
        return load_from_local(os.path.join(LOCAL_STORAGE_PATH, filename))
 
def delete_user_data(email):
    if USE_AWS:
        try:
            filename = f"{signUp_user_folder}{email}.json"
            s3.delete_object(Bucket=S3_BUCKET_NAME, Key=filename)
            print(f"File {filename} deleted from AWS.")  # Debug log
        except Exception as e:
            print(f"Error deleting file from AWS: {e}")  # Debug log
            raise
    else:
        filename = f"users/{email}.json"
        delete_from_local(filename)

 
def send_email(to_email, otp):
    try:
        # Validate email format
        if not re.match(r"[^@]+@[^@]+\.[^@]+", to_email):
            print(f"Invalid email address: {to_email}")
            return False
 
        # Setup email message
        subject = "Your Reset Password Code"
        message = f"Your Reset Password Code is: {otp}"
        msg = MIMEMultipart()
        msg['From'] = EMAIL_ADDRESS
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(message, 'plain'))
 
        # Send email using SMTP
        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.starttls()
            server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            server.sendmail(EMAIL_ADDRESS, to_email, msg.as_string())
        print("Email sent successfully")
        return True
    except Exception as e:
        print(f"Error sending email: {e}")
        return False
 
# Email verification :
 
@app.route('/email-verification', methods=['POST'])
def email_verification():
    try:
        email = request.json.get('email')  # Extract email from the request
        if not email:
            return jsonify({"message": "Email is required"}), 400
 
        print(f"Processing email verification for: {email}")
 
        # Generate the sign-up link
        sign_up_link = f"http://localhost:3000/signUp/{email}"
 
        # Create the email message
        msg = Message(
            "Sign-Up Link - Verify Your Email",
            sender="your_email@gmail.com",
            recipients=[email]
        )
        msg.body = (
            f"Dear User,\n\n"
            f"Congratulations! Your email has been successfully verified. You're just one step away from completing your sign-up process.\n\n"
            f"Click the link below to finish setting up your account:\n"
            f"{sign_up_link}\n\n"
            f"Thank you for choosing us.\n\n"
        )
        # msg.body = (
        #     f"Hello,\n\n"
        #     f"Your email has been successfully verified. Use the following link to complete your sign-up process:\n\n"
        #     f"{sign_up_link}\n\n"
        #     f"If you did not request this verification, please ignore this email.\n\n"
        #     f"Thank you."
        # )
        print(f"Sending email to: {email}\nContent: {msg.body}")
       
        # Send the email
        mail.send(msg)
        print("Email sent successfully.")
 
        return jsonify({"message": "Sign-up link sent successfully"}), 200
 
    except Exception as e:
        print(f"Error sending email: {e}")
        return jsonify({"message": f"Error occurred: {str(e)}"}), 500
 
 
 
@app.route('/verify-otp', methods=['POST'])
def verify_otp():
    data = request.get_json()
    email = data.get('email')
    otp = data.get('otp')
 
    if not email or not otp:
        return jsonify({"error": "Email and OTP are required"}), 400
 
    if otp_store.get(email) == int(otp):
        del otp_store[email]
        return jsonify({"message": "Email verified successfully!"}), 200
    else:
        return jsonify({"error": "Invalid OTP"}), 400
 



# 2. Sign Up

@app.route('/sign-up', methods=['POST'])
def sign_up():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"message": "Invalid JSON"}), 400
 
        email = data.get('email')
        password = data.get('password')
        confirm_password = data.get('confirm_password')
        first_name = data.get('first_name')
        last_name = data.get('last_name')
 
        if not all([email, password, confirm_password, first_name, last_name]):
            return jsonify({"message": "All fields are required"}), 400
 
        if password != confirm_password:
            return jsonify({"message": "Passwords do not match"}), 400
 
        if load_user_data(email):
            return jsonify({"message": "User already exists"}), 400
 
        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
        user_data = {
            "email": email,
            "password": hashed_password,
            "first_name": first_name,
            "last_name": last_name,
            "data": {}
        }
 
        save_user_data(user_data, email)
        return jsonify({"message": "Sign up successful"}), 200
    except Exception as e:
        print(f"Error in sign-up: {e}")
        return jsonify({"message": "Internal server error"}), 500
   
   
# 3. Sign in :

@app.route('/sign-in', methods=['POST'])
def sign_in():
    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
 
        if not all([email, password]):
            return jsonify({"message": "Email and password are required"}), 400
 
        user_data = load_user_data(email)
        if not user_data or not bcrypt.check_password_hash(user_data["password"], password):
            return jsonify({"message": "Invalid email or password"}), 401
 
        token = jwt.encode(
            {"email": email, "exp": datetime.utcnow() + timedelta(hours=5)},
            app.config['JWT_SECRET_KEY'],
            algorithm="HS256"
        )
 
        return jsonify({"message": "Sign in successful", "token": token}), 200
    except Exception as e:
        print(f"Error during sign-in: {e}")
        return jsonify({"message": "Internal server error"}), 500
      

#  # 4. Forgot Password

import traceback

@app.route('/forgot-password', methods=['POST'])
def forgot_password():
    try:
        data = request.get_json()
        print("Data received:", data)
        if not data:
            return jsonify({"message": "Invalid JSON"}), 400
 
        email = data.get('email')
        print("Email extracted:", email)
        if not email:
            return jsonify({"message": "Email is required"}), 400
 
        # Generate a reset code
        reset_code = random.randint(100000, 999999)
        print("Reset code generated:", reset_code)
 
        # Load the user data
        user_data = load_user_data(email)
        if not user_data:
            return jsonify({"message": "User not found"}), 404
 
        # Add the reset code to the user data
        user_data['reset_code'] = reset_code
        user_data['reset_timestamp'] = str(datetime.now())
 
        # Save the updated user data
        save_user_data(user_data, email)
        print("Reset data saved successfully.")
 
        # Send reset code via email
        if send_email(email, reset_code):
            print("Email sent successfully.")
            return jsonify({"message": "Password reset code sent successfully"}), 200
        else:
            print("Failed to send email.")
            return jsonify({"error": "Failed to send reset code"}), 500
 
    except Exception as e:
        traceback.print_exc()  # Logs the full stack trace
        return jsonify({"error": "Internal server error"}), 500
 

# 5. Reset password :

@app.route('/reset-password', methods=['POST'])
def reset_password():
    try:
        data = request.get_json()
        print("Data received:", data)
 
        email = data.get('email')
        reset_code = data.get('reset_code')
        new_password = data.get('new_password')
 
        if not all([email, reset_code, new_password]):
            return jsonify({"message": "All fields are required"}), 400
 
        email = email.lower()
        user_data = load_user_data(email)
 
        if not user_data:
            return jsonify({"message": "User not found"}), 404
 
        user_reset_code = user_data.get('reset_code')
        print(f"Reset code from request: {reset_code}")
        print(f"Reset code from user data: {user_reset_code}")
 
        if str(user_reset_code) != str(reset_code):
            print(f"Reset code from request: {reset_code}")
            print(f"Reset code from user data: {user_reset_code}")
            return jsonify({"message": "Invalid reset code"}), 400
 
        # Update password
        hashed_password = bcrypt.generate_password_hash(new_password).decode('utf-8')
        user_data['password'] = hashed_password
        user_data.pop('reset_code', None)
        user_data.pop('reset_timestamp', None)
 
        save_user_data(user_data, email)
        return jsonify({"message": "Password reset successful"}), 200
 
    except Exception as e:
        traceback.print_exc()
        return jsonify({"message": "Internal server error"}), 500
 
 
 
# ------------------------- With Bearer ------------------------------

# v-2 : update profile photo :

import base64
from werkzeug.utils import secure_filename
from flask import send_from_directory

# Directory to store profile photos
PROFILE_PHOTOS_DIR = "local_data/profile_photos"

from flask import url_for

@app.route('/advisor_profile', methods=['POST', 'GET'])
def advisor_profile():
    if request.method == 'GET':
        token = request.headers.get('Authorization')

        if not token:
            return jsonify({"message": "Token is missing"}), 401

        try:
            token = token.split(" ")[1] if token.startswith("Bearer ") else None
            if not token:
                return jsonify({"message": "Invalid token format"}), 401

            decoded_token = jwt.decode(token, app.config['JWT_SECRET_KEY'], algorithms=["HS256"])
            email = decoded_token.get('email')

            user_data = load_user_data(email)
            if not user_data:
                return jsonify({"message": "User not found"}), 404

            # Fetch client data count
            client_data_url = 'http://localhost:5000/get-all-client-data'
            response = requests.get(client_data_url)

            client_count = 0
            if response.status_code == 200:
                client_list = response.json().get('data', [])
                client_count = len(client_list)

            # Generate URL for profile photo
            filename = f"{email}_profile.jpeg"
            profile_photo_url = url_for('serve_profile_photo', filename=filename, _external=True)

            profile_data = {
                "email": user_data["email"],
                "first_name": user_data["first_name"],
                "last_name": user_data["last_name"],
                "client_count": client_count,
                "profile_photo_url": profile_photo_url
            }

            return jsonify({"message": "Profile retrieved successfully", "profile": profile_data}), 200

        except jwt.ExpiredSignatureError:
            return jsonify({"message": "Token has expired"}), 401
        except jwt.InvalidTokenError:
            return jsonify({"message": "Invalid token"}), 401
        except Exception as e:
            print(f"Error retrieving profile: {e}")
            return jsonify({"message": "Internal server error"}), 500

    if request.method == 'POST':
        try:
            token = request.headers.get('Authorization')

            if not token:
                return jsonify({"message": "Token is missing"}), 401

            token = token.split(" ")[1] if token.startswith("Bearer ") else None
            if not token:
                return jsonify({"message": "Invalid token format"}), 401

            decoded_token = jwt.decode(token, app.config['JWT_SECRET_KEY'], algorithms=["HS256"])
            email = decoded_token.get('email')

            user_data = load_user_data(email)
            if not user_data:
                return jsonify({"message": "User not found"}), 404

            # Check if the request has the 'profile_photo' field
            data = request.get_json()
            if not data or 'profile_photo' not in data:
                return jsonify({"message": "Profile photo is missing"}), 400

            # Decode the base64 image
            profile_photo = data['profile_photo']
            file_extension = "jpg"  # Default file extension
            if profile_photo.startswith("data:image/"):
                file_extension = profile_photo.split(";")[0].split("/")[-1]

            image_data = profile_photo.split(",")[1]
            image_bytes = base64.b64decode(image_data)

            # Save the file locally
            filename = f"{email}_profile.{file_extension}"
            file_path = os.path.join(PROFILE_PHOTOS_DIR, filename)
            with open(file_path, "wb") as f:
                f.write(image_bytes)

            # Generate the profile photo URL
            profile_photo_url = url_for('serve_profile_photo', filename=filename, _external=True)

            # Update user's profile with the new image URL
            user_data["profile_photo_url"] = profile_photo_url
            save_user_data(email, user_data)  # Save updated user data back to storage

            return jsonify({
                "message": "Profile photo uploaded successfully",
                "profile_photo_url": profile_photo_url
            }), 200

        except jwt.ExpiredSignatureError:
            return jsonify({"message": "Token has expired"}), 401
        except jwt.InvalidTokenError:
            return jsonify({"message": "Invalid token"}), 401
        except Exception as e:
            print(f"Error uploading profile photo: {e}")
            return jsonify({"message": "Internal server error"}), 500


@app.route('/profile_photos/<filename>', methods=['GET'])
def serve_profile_photo(filename):
    """
    Serve profile photos from the local directory.
    """
    try:
        return send_from_directory(PROFILE_PHOTOS_DIR, filename)
    except Exception as e:
        print(f"Error serving profile photo: {e}")
        return jsonify({"message": "Error retrieving profile photo"}), 500




# v-1 :
import requests
 
# @app.route('/advisor_profile', methods=['GET'])
# def advisor_profile():
#     token = request.headers.get('Authorization')

#     if not token:
#         return jsonify({"message": "Token is missing"}), 401

#     try:
#         token = token.split(" ")[1] if token.startswith("Bearer ") else None
#         if not token:
#             return jsonify({"message": "Invalid token format"}), 401

#         decoded_token = jwt.decode(token, app.config['JWT_SECRET_KEY'], algorithms=["HS256"])
#         email = decoded_token.get('email')

#         user_data = load_user_data(email)
#         if not user_data:
#             return jsonify({"message": "User not found"}), 404

#         # Fetch client data count
#         client_data_url = 'http://localhost:5000/get-all-client-data'
#         response = requests.get(client_data_url)

#         client_count = 0
#         if response.status_code == 200:
#             client_list = response.json().get('data', [])
#             client_count = len(client_list)

#         profile_data = {
#             "email": user_data["email"],
#             "first_name": user_data["first_name"],
#             "last_name": user_data["last_name"],
#             "client_count": client_count
#         }

#         return jsonify({"message": "Profile retrieved successfully", "profile": profile_data}), 200

#     except jwt.ExpiredSignatureError:
#         return jsonify({"message": "Token has expired"}), 401
#     except jwt.InvalidTokenError:
#         return jsonify({"message": "Invalid token"}), 401
#     except Exception as e:
#         print(f"Error retrieving profile: {e}")
#         return jsonify({"message": "Internal server error"}), 500
    
     
# -------------------------------------------

#  Change Password for Profile

@app.route('/change_password', methods=['POST'])
def change_password():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"message": "Invalid JSON"}), 400

        email = data.get('email')
        old_password = data.get('old_password')
        new_password = data.get('new_password')
        confirm_password = data.get('confirm_password')

        if not all([email, old_password, new_password, confirm_password]):
            return jsonify({"message": "All fields are required"}), 400

        # Check if new password matches the confirm password
        if new_password != confirm_password:
            return jsonify({"message": "Passwords do not match"}), 400

        # Load user data
        if USE_AWS:
            # filename = f"{signUp_user_folder}/{email}.json"
            filename = f"{signUp_user_folder}{email}.json"
            user_data = load_from_aws(filename)
        else:
            filepath = os.path.join(LOCAL_STORAGE_PATH, f"users/{email}.json")
            user_data = load_from_local(filepath)

        if not user_data:
            return jsonify({"message": "User not found"}), 404

        # Check if old password is correct
        if not bcrypt.check_password_hash(user_data["password"], old_password):
            return jsonify({"message": "Old password is incorrect"}), 401

        # Hash the new password
        hashed_new_password = bcrypt.generate_password_hash(new_password).decode('utf-8')

        # Update the user data with the new password
        user_data["password"] = hashed_new_password

        # Save updated user data
        if USE_AWS:
            save_to_aws(user_data, filename)
        else:
            save_to_local(user_data, filepath)

        return jsonify({"message": "Password reset successful"}), 200

    except Exception as e:
        print(f"Error in change_password: {e}")
        return jsonify({"message": "Internal server error"}), 500

 


##########################################################################################################





########################################################################################################################

# def convert_to_markdown(raw_text):
#     # Replace specific text patterns with markdown syntax
#     formatted_text = raw_text.replace('\n', '\n\n')  # Ensure newlines create paragraphs
    
#     # Convert text into markdown format
#     html = markdown.markdown(formatted_text)

#     return html


def markdown_table_to_html(md_table):
    # Split the markdown table by lines
    lines = md_table.strip().split("\n")
    
    # Extract headers and rows
    headers = lines[0].strip('|').split('|')
    rows = [line.strip('|').split('|') for line in lines[2:]]  # Skip the separator line

    # Start creating the HTML table
    html_table = "<table>\n"
    
    # Add headers
    html_table += "  <thead>\n    <tr>\n"
    for header in headers:
        html_table += f"      <th>{header.strip()}</th>\n"
    html_table += "    </tr>\n  </thead>\n"
    
    # Add rows
    html_table += "  <tbody>\n"
    for row in rows:
        html_table += "    <tr>\n"
        for col in row:
            html_table += f"      <td>{col.strip()}</td>\n"
        html_table += "    </tr>\n"
    html_table += "  </tbody>\n</table>"

    return html_table

def process_client_info_and_analysis(content):
    # Identify and extract the client's financial info markdown table part
    client_info_section_start = content.find("| Category | Value |")
    client_info_section_end = content.find("</p>", client_info_section_start) + 4
    
    if client_info_section_start != -1 and client_info_section_end != -1:
        # Extract markdown table
        md_table = content[client_info_section_start:client_info_section_end]
        # Convert only the markdown table portion into an HTML table
        html_table = markdown_table_to_html(md_table)
        # Replace the markdown table part with the generated HTML table
        content = content[:client_info_section_start] + html_table + content[client_info_section_end:]
    
    # Return the rest of the content unchanged
    return content

def generate_final_html(content):
    # Process the content to convert financial info to HTML table while leaving other sections untouched
    html_content = process_client_info_and_analysis(content)
    
    # Any other HTML processing can be done here if needed
    return html_content


# def markdown_table_to_html(md_table):
#     # Split the markdown table by lines
#     lines = md_table.strip().split("\n")
    
#     # Extract headers and rows
#     headers = lines[0].strip('|').split('|')
#     rows = [line.strip('|').split('|') for line in lines[2:]]  # Skip the separator line

#     # Start creating the HTML table
#     html_table = "<table>\n"
    
#     # Add headers
#     html_table += "  <thead>\n    <tr>\n"
#     for header in headers:
#         html_table += f"      <th>{header.strip()}</th>\n"
#     html_table += "    </tr>\n  </thead>\n"
    
#     # Add rows
#     html_table += "  <tbody>\n"
#     for row in rows:
#         html_table += "    <tr>\n"
#         for col in row:
#             html_table += f"      <td>{col.strip()}</td>\n"
#         html_table += "    </tr>\n"
#     html_table += "  </tbody>\n</table>"

#     return html_table



import markdown2
from bs4 import BeautifulSoup

def markdown_to_readable_text(md_text):
    # Convert markdown to HTML
    html = markdown2.markdown(md_text)

    # Parse the HTML
    soup = BeautifulSoup(html, "html.parser")

    # Function to format plain text from tags
    def format_text_from_html(soup):
        formatted_text = ''
        for element in soup:
            if element.name == "h1":
                formatted_text += f"\n\n# {element.text.upper()} #\n\n"
            elif element.name == "h2":
                formatted_text += f"\n\n## {element.text} ##\n\n"
            elif element.name == "h3":
                formatted_text += f"\n\n### {element.text} ###\n\n"
            elif element.name == "strong":
                formatted_text += f"**{element.text}**"
            elif element.name == "em":
                formatted_text += f"_{element.text}_"
            elif element.name == "ul":
                for li in element.find_all("li"):
                    formatted_text += f"\n - {li.text}"
            elif element.name == "ol":
                for idx, li in enumerate(element.find_all("li"), 1):
                    formatted_text += f"\n {idx}. {li.text}"
            elif element.name == "table":
                # Convert markdown table to HTML table
                formatted_text += "<table>\n"
                rows = element.find_all("tr")
                for row in rows:
                    formatted_text += "<tr>\n"
                    cols = row.find_all(["th", "td"])
                    for col in cols:
                        tag = 'th' if col.name == "th" else 'td'
                        formatted_text += f"<{tag}>{col.text.strip()}</{tag}>\n"
                    formatted_text += "</tr>\n"
                formatted_text += "</table>\n"
            else:
                formatted_text += element.text

        return formatted_text.strip()

    return format_text_from_html(soup)



def markdown_to_text(md): # og solution code 
    # Simple conversion for markdown to plain text
    md = md.replace('**', '')
    md = md.replace('*', '')
    md = md.replace('_', '')
    md = md.replace('#', '')
    md = md.replace('`', '')
    return md.strip()



# def extract_responses_from_docx(personality_file):
#     try:
#         doc = docx.Document(personality_file)
#         responses = {}
#         current_question = None

#         # Check paragraphs
#         for para in doc.paragraphs:
#             text = para.text.strip()
#             if text:
#                 # Check if the paragraph contains a question
#                 if "?" in text or text.endswith(":"):
#                     current_question = text
#                 else:
#                     # This is a typed answer
#                     typed_answer = text.strip()
#                     if current_question:
#                         # If the question already has an answer, append to it (handles multiple responses)
#                         if current_question in responses:
#                             responses[current_question] += "; " + typed_answer
#                         else:
#                             responses[current_question] = typed_answer

#         # Check tables for additional responses
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     text = cell.text.strip()
#                     if text:
#                         if "?" in text or text.endswith(":"):
#                             current_question = text
#                         else:
#                             typed_answer = text.strip()
#                             if current_question:
#                                 if current_question in responses:
#                                     responses[current_question] += "; " + typed_answer
#                                 else:
#                                     responses[current_question] = typed_answer

#         return responses

#     except Exception as e:
#         print(f"Error extracting responses: {e}")
#         return None

import docx



# import asyncio
# # from some_generative_ai_library import GenerativeModel  # Replace with actual import

# async def determine_investment_personality(assessment_data):
#     try:
#         # Prepare input text for the chatbot based on assessment data
#         input_text = "User Profile:\n"
#         for question, answer in assessment_data.items():
#             input_text += f"{question}: {answer}\n"

#         # Introduce the chatbot's task and prompt for classification
#         input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
#                       "- Conservative Investor\n" \
#                       "- Moderate Investor\n" \
#                       "- Aggressive Investor\n\n" \
#                       "Please provide the classification below:\n"

#         # Use your generative AI model to generate a response
#         model = GenerativeModel('gemini-1.5-flash')
#         response = await model.generate_content(input_text)

#         # Determine the investment personality from the chatbot's response
#         response_text = response.text.lower()

#         if "conservative investor" in response_text:
#             personality = "Conservative Investor"
#         elif "moderate investor" in response_text:
#             personality = "Moderate Investor"
#         elif "aggressive investor" in response_text:
#             personality = "Aggressive Investor"
#         else:
#             personality = "Unknown"

#         return personality
#     except Exception as e:
#         print(f"Error generating response: {e}")
#         return "Unknown"


# GET Method
async def determine_investment_personality(assessment_data): # proper code 
    try:
        # Prepare input text for the chatbot based on assessment data
        input_text = "User Profile:\n"
        for question, answer in assessment_data.items():
            input_text += f"{question}: {answer}\n"

        # Introduce the chatbot's task and prompt for classification
        input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
                      "- Conservative Investor\n" \
                      "- Moderate Investor\n" \
                      "- Aggressive Investor\n\n" \
                      "Please provide the classification below:\n"

        # Use your generative AI model to generate a response
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(input_text)

        # Determine the investment personality from the chatbot's response
        response_text = response.text.lower()

        if "conservative investor" in response_text:
            personality = "Conservative Investor"
        elif "moderate investor" in response_text:
            personality = "Moderate Investor"
        elif "aggressive investor" in response_text:
            personality = "Aggressive Investor"
        else:
            personality = "Unknown"

        return personality
    except Exception as e:
        print(f"Error generating response: {e}")
        return "Unknown"




#Load the Vector DataBase : # current version :
async def load_vector_db(file_path): # # GET Method 
    try:
        print("Loading vector database...")
        # file_path = os.path.basename(file_path)
        
        # Verify the file path
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        print(f"File path: {file_path}")
        
        # Check file permissions
        if not os.access(file_path, os.R_OK):
            raise PermissionError(f"File is not readable: {file_path}")
        
        # print(file_path)
        
        loader = Docx2txtLoader(file_path)
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        text_chunks = text_splitter.split_documents(documents)
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        # vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        
        vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        # index = faiss.IndexFlatL2(len(embeddings.embed_query("hello world")))

        # vector_store = FAISS(
        #     embedding_function=embeddings,
        #     index=index,
        #     docstore=InMemoryDocstore(),
        #     index_to_docstore_id={},
        # )
        
        print("Vector database loaded successfully.") 
        return vector_store.as_retriever(search_kwargs={"k": 1})
    except Exception as e:
        print(f"Error loading vector database: {e}")
        return None

# import os

# async def load_vector_db(file_storage): 
#     try:
#         # Define the destination folder and ensure it exists
#         destination_folder = 'path/to/your/destination/folder'
#         if not os.path.exists(destination_folder):
#             os.makedirs(destination_folder)
        
#         # Construct the destination file path
#         file_path = os.path.join(destination_folder, file_storage.filename)
        
#         # Save the file to the destination folder
#         file_storage.save(file_path)
        
#         print("Loading vector database...")
#         print(f"File path: {file_path}")
        
#         loader = Docx2txtLoader(file_path)
#         documents = loader.load()
        
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#         text_chunks = text_splitter.split_documents(documents)
        
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
#         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        
#         print("Vector database loaded successfully.") 
#         return vector_store.as_retriever(search_kwargs={"k": 1})
#     except Exception as e:
#         print(f"Error loading vector database: {e}")
#         return None



# investment_personality = "Moderate Investor"
# Retrieval_Chain and Promot Template for Suggestions :
async def make_retrieval_chain(retriever,investmentPersonality,clientName,monthly_investment=10000,investment_period=3): # GET Method
    """
    Create a retrieval chain using the provided retriever.

    Args:
        retriever (RetrievalQA): A retriever object.

    Returns:
        RetrievalQA: A retrieval chain object.
    """
    try:
        # global investment_personality #,summary
        
        print(f"{retriever}\n {investmentPersonality}\n {clientName}\n {monthly_investment}")
        # try:
        #     print(type(investmentPersonality))
        # except Exception as e:
        #     print(f"Error in personality: {e}")
        #     return None
        
        # print(clientName)
        
        llm = ChatGoogleGenerativeAI(
            #model="gemini-pro",
            model = "gemini-1.5-flash",
            temperature = 0.45,
            # temperature=0.7,
            top_p=0.85,
            google_api_key=GOOGLE_API_KEY
        )
        # New Template 
        investment_period = str(investment_period)
        print(investmentPersonality)
        monthly_investment = str(monthly_investment)
        print(monthly_investment)
        print(investment_period)
        
        # New Prompt Template :
        
        prompt_template = """ 
                                You are a Financial Advisor tasked with creating responsible investment suggestions for a client based on their investment personality : """ + investmentPersonality +   "\n" + """ so that the client can reach their Financial Goals, based on their Financial Conditions.
                                Use the following instructions to ensure consistent output:
                                ---

                                ### Required Output Format:
                                
                                #### Client Financial Details:
                                - **Client Name**: """ + clientName + """
                                - **Assets**:
                                - List all asset types, their current values, and annual contributions in a tabular format (columns: "Asset Type", "Current Value", "Annual Contribution").
                                - **Liabilities**:
                                - List all liability types, their balances, interest rates, and monthly payments in a tabular format (columns: "Liability Type", "Balance", "Interest Rate", "Monthly Payment").
                                - **Other Details**:
                                - Retirement plan details, income sources, and goals should be listed in a clear and concise format.
                                - Client's Financial Condition : Analyze the Details and mention the Client's Financial Condition as : Stable/ Currently Stable / Unstable.
                                - **Investment Period** `Z years`
                                
                                #### Investment Allocation:
                                Split investments into **Growth-Oriented Investments** and **Conservative Investments**. Ensure each category includes:
                                - **Investment Type**: Specify the investment type (e.g., "Index Funds", "US Treasury Bonds").
                                - **Allocation Range**: Specify minimum and maximum allocation percentages (e.g., `10% - 20%`).
                                - **Target**: Describe the purpose of the investment.
                                - **How to Invest**: Provide instructions on how to invest in this asset.
                                - **Where to Invest**: Specify platforms or tools for making the investment.

                                **Example**:
                                **Growth-Oriented Investments (Minimum X% - Maximum Y%) **:
                                - **Stocks**: `20% - 30%`
                                - **ETFs**: `10% - 15%`
                                - **Mutual Funds**: `10% - 20%`
                                - **Cryptocurrency**: ` 5% - 10%`
                                - **Real Estates or REITS**: `10% - 20%`
                                - *Target*: Long-term growth potential aligned with the overall market performance tailored to fullfil Client's Financial Goals and manage his Financial Condition.
                                - *How to Invest*: Provide information on how to invest in which market 
                                - *Where to Invest*: Provide Information to buy which assets and how much to invest in terms of amount and percentage(%).Mention 5-6 assets.
                                
                                **Conservative Investments (Minimum X% - Maximum Y%) **:
                                - **High-Yield Savings Account**: `30% - 40%`
                                - **Bonds**: `10% - 20%`
                                - **Commodities**: `5% - 10%`
                                - **Cash**: `5% - 10%`
                                - *Target*: Maintain liquidity for emergencies.
                                - *How to Invest*: Provide information on how to invest.
                                - *Where to Invest*: Mention where to invest and how much to allocate in terms of money and percentage(%). Mention 5-6 assets.

                                #### Returns Overview:
                                - **Minimum Expected Annual Return**: `X% - Y%`
                                - **Maximum Expected Annual Return**: `X% - Y%`
                                - **Minimum Expected Growth in Dollars**: `$X - $Y` (based on the time horizon)
                                - **Maximum Expected Growth in Dollars**: `$X - $Y` (based on the time horizon)
                                - **Time Horizon**: `Z years`

                                ---

                                ### Example Output:
                                
                                #### Client Financial Details: 
                                | Asset Type          | Current Value ($) | Annual Contribution ($) |
                                |----------------------|-------------------|--------------------------|
                                | 401(k), 403(b), 457  | 300               | 15                       |
                                | Traditional IRA      | 200               | 15                       |
                                | Roth IRA             | 500               | 28                       |
                                | Cash/Bank Accounts   | 500,000           | 30,000                   |
                                | Real Estate          | 1,000,000         | -                        |
                                | Total Assets Value   | 1,501,000         | -                        |

                                | Liability Type      | Balance ($) | Interest Rate (%) | Monthly Payment ($) |
                                |---------------------|-------------|--------------------|----------------------|
                                | Mortgage            | 1,000       | 10                | 100                  |
                                | Credit Card         | 400         | 8                 | 400                  |
                                | Other Loans         | 500         | 6                 | 100                  |
                                | Total Liabilities   | 1,900       | -                 | -                    |
                                
                                | Investrment Period | 3 years |
                                
                                **Growth-Oriented Investments (Minimum 40% - Maximum 80%)**:
                                - **Stocks**: `20% - 30%`
                                - **ETFs**: `5% - 10%`
                                - **Mutual Funds**: `5% - 20%`
                                - **Cryptocurrency**: ` 5% - 10%`
                                - **Real Estates or REITS**: `5% - 10%`
                                - *Target*: Long-term growth potential aligned with the market.
                                - *How to Invest*: Purchase low-cost index funds.
                                - *Where to Invest*: Stocks such as NVIDIA,AAPL, Vanguard, LiteCoin.

                                **Conservative Investments (Minimum 40% - Maximum 70%)**:
                                - **High-Yield Savings Account**: `20% - 30%`
                                - **Bonds**: `10% - 20%`
                                - **Commodities**: `5% - 10%`
                                - **Cash**: `5% - 10%`
                                - *Target*: Maintain liquidity for emergencies.
                                - *How to Invest*: Deposit funds into an FDIC-insured account.
                                - *Where to Invest*: Ally Bank, Capital One 360.

                                #### Returns Overview:
                                - **Minimum Expected Annual Return**: `4% - 6%`
                                - **Maximum Expected Annual Return**: `8% - 15%`
                                - **Minimum Expected Growth in Dollars**: `$4,000 - $6,000`
                                - **Maximum Expected Growth in Dollars**: `$8,000 - $15,000`
                                - **Time Horizon**: `3 years`

                                ---

                                Ensure the output strictly follows this structure.


                            ### Rationale for Investment Suggestions:
                            Provide a detailed explanation of why these suggestions align with the client’s financial personality and goals.

                            ---
                            <context>
                            {context}
                            </context>
                            Question: {input}

        """



        # #Wasnt consistent for generating the Bar Graph and Pie Chart :
        # prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
                # Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality and Financial Document provided to you.
                # Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
                # Also give the user detailed information about the investment how to invest,where to invest and how much they
                # should invest in terms of percentage of their investment amount based on the clients Financial Conditions and help them to cover up their Mortgage and Debts if any.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
                # Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
                # investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon and how it can help them accumulate wearlth overtime to achive their Financial  goals.
                # Also give the user minimum and maximum expected growth in dollars for the time horizon .
                # Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
                
                # You are a Financial Advisor for question-answering tasks related to the document. Based on the client's investment personality and financial details provided, generate responsible investment suggestions to achieve their financial goals while managing debts.

                # Step-by-Step Guidance:
                # 1. Assets: Calculate total assets by analyzing the provided financial document in the My Assets section. Ensure you include cash, real estate, retirement accounts, brokerage accounts, and any other relevant asset types from the document.
                # 2. Liabilities: Calculate total liabilities by analyzing the provided financial document in the My Liabilities section. Consider mortgages, credit card debts, student loans, car loans, and other liabilities. 
                # 3. Monthly Investment Feasibility: Use the client's assets and liabilities to assess whether their planned monthly investment is feasible. If not feasible, suggest a more realistic monthly investment amount.
                # 4. Analyze Liabilities: Determine if the client's monthly investment plan is feasible after covering liabilities and expected expenses and also considering some amount for savings. If the client's monthly investment plan is not feasible after covering expenses and savings, generate investment suggestions on a smaller monthly investment plan amount if it can help the client else mention amount is too small for the client's requirementys to be made.
                # 5. Investment Strategy: Suggest a strategy where monthly investments can both generate returns and pay off debts effectively and helps client to achieve their financial goals.
                # 6. Allocation: Provide detailed allocations between growth-oriented investments and conservative investments, ensuring the client can meet their monthly debt obligations and save for their future financial goals.
                # 7. Returns: Include minimum and maximum compounded returns over 5-10 years, along with inflation-adjusted returns for clarity.
                # 8. Suggestions: Offer advice on how to use remaining funds to build wealth after clearing liabilities and achive their financial goal.
                
                
        #         Here's an example for the required Output Format(if there are comments indicated by # in the example output format then thats a side note for your reference dont write it in the response that will be generated ) :
                
        #         Client's Financial Information :(# This is a header line have it in bold) 
                
                
        #         Client Name: """ + clientName + """(# have the client name in underline)

        #         Financial Overview: (#the data presented is just an example for your reference do not consider it as factual refere to the document provided to you and generate data based on the provided data and only when nothing is provided assume some data for analysis, This is a header line have it in bold. The data below it should be displayed in a table format so make sure of that data.There must be 2 columns 1 for Category and second for Value.List down all the assets and liabilities along with its values and then Total of assets,liabilities,etc.)
                
        #         - Total Assets: (# Sum of all client assets and Annual Income . Mention all assets and their respected values.if non consider the example assets)
                
        #         - Total Liabilities: (# Sum of all liabilities. Mention all liabilities and their respected values if non consider the example liabilities)
                
                
        #         - Monthly Liabilities: (# Monthly payments derived from liabilities)
                
        #         - Total Annual Income : (# Sum of all client's anual income)
                
        #         - Monthly Investment Amount : """ + monthly_investment + """ (# if no specific amount is specified to you then only assume  10,000 else consider the amount mention to you and just display the amount)
                
        #         - Investment Period : """ + investment_period + """  (# if no specific period is specified to you then only assume 3 years else consider the period mention to you and just display the period)


        #         Financial Analysis :(#Analyse the assets and liabilities and based on that give a suggestion for analysis generate suggestions for one of the following conditions:)
        #         (#1st condition : Everything is Positive)Based on the given Financial Conditions the client is having a good and stable income with great assets and manageable debt and liabilities.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the clients monthly income.
        #         (# if this condition is true then ignore the other conditions and start with the Investment Suggestions)
                
        #         (#2nd condition : Everything is temporarily Negative) Based on the given Financial Conditions the client is facing a low income for now but have great assets and manageable debt and liabilities.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the client's monthly income but the client might not be able to sustain the monthly investment amount that they are planning.)
        #         Instead I would like to recommend this amount to the client for their monthly investment : (#Mention a feasible amount to the client for monthly investment and start suggesting investments based on this amount and not the previous amount being taken into consideration)
                
        #         (#3rd condition : Everything is Negative) Based on the given Financial Conditions the client is facing a low income and doesnt have good assets to manage the debts and liabilities of the client and in such a condition this monthly investment amount is not feasible.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is not manageable for the client's monthly income and so the client might not be able to sustain the monthly investment amount that they are planning to do.)
        #         I would like to recommend this amount to the client for monthly investment : (# Mention a minimum amount to the client for monthly investment if possible else just say the client should first prioritize on savings and generating more income to manage their debts and liabilities first and so dont give any investment suggestions to the client.)
                
        #         (#If the financial is 1 or 2 only then give investment suggestions to the client)
                
                
                
        #         Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

        #         Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

                
        #         Investment Allocation: (#remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

        #         Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
        #         How to Invest: Diversify across various asset classes like:  (#Give allocations % as well)
                
        #         Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
                
        #         ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
                
        #         Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
                
        #         Consider investing in blue-chip companies or growth sectors like technology. 
                
                
        #         Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


        #         Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
        #         Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

        #         Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
        #         How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
        #         Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
        #         Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
        #         Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
        #         Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
        #         Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
        #         Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.


        #         Time Horizon and Expected Returns:

        #         Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
        #         Minimum Expected Annual Return: 4% - 6% 
                
                
        #         Maximum Expected Annual Return: 8% - 10% 
                
                
        #         Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, (# consider the monthly investment amount and give returns based on that only) $10,000 could grow to approximately 17,908 in 10 years.
        #         Minimum Expected Growth in Dollars: 
                
        #         4,000−6,000 (over 10 years) 
                
                
        #         Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
        #         Inflation Adjusted Returns:(#do not write this part inside the bracket just give answer,assume US inflation rate assume 3% if you dont know, and give the investment returns value that was suggested by you for the considered monthly investment amount after 3,5,10years of growth mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
        #         Rationale for Investment Suggestions:

        #         This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
        #         Important Considerations:

        #         Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

        #         Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

        #         Explain how this suggestions can help the client grow their wealth and improve their financial condition and/or cover up thier loans and in turn achive their Financial goals.
        #         <context>
        #         {context}
        #         </context>
        #         Question: {input}"""

        
        # # Without category and value :
        # prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
        #         Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality and Financial Document provided to you.
        #         Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
        #         Also give the user detailed information about the investment how to invest,where to invest and how much they
        #         should invest in terms of percentage of their investment amount based on the clients Financial Conditions and help them to cover up their Mortgage and Debts if any.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
        #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
        #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon and how it can help them accumulate wearlth overtime to achive their Financial  goals.
        #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
        #         Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
                
        #         You are a Financial Advisor for question-answering tasks related to the document. Based on the client's investment personality and financial details provided, generate responsible investment suggestions to achieve their financial goals while managing debts.

        #         Step-by-Step Guidance:
        #         1. Assets: Calculate total assets by analyzing the provided financial document in the My Assets section. Ensure you include cash, real estate, retirement accounts, brokerage accounts, and any other relevant asset types from the document.
        #         2. Liabilities: Calculate total liabilities by analyzing the provided financial document in the My Liabilities section. Consider mortgages, credit card debts, student loans, car loans, and other liabilities. 
        #         3. Monthly Investment Feasibility: Use the client's assets and liabilities to assess whether their planned monthly investment is feasible. If not feasible, suggest a more realistic monthly investment amount.
        #         4. Analyze Liabilities: Determine if the client's monthly investment plan is feasible after covering liabilities and expected expenses and also considering some amount for savings. If the client's monthly investment plan is not feasible after covering expenses and savings, generate investment suggestions on a smaller monthly investment plan amount if it can help the client else mention amount is too small for the client's requirementys to be made.
        #         5. Investment Strategy: Suggest a strategy where monthly investments can both generate returns and pay off debts effectively and helps client to achieve their financial goals.
        #         6. Allocation: Provide detailed allocations between growth-oriented investments and conservative investments, ensuring the client can meet their monthly debt obligations and save for their future financial goals.
        #         7. Returns: Include minimum and maximum compounded returns over 5-10 years, along with inflation-adjusted returns for clarity.
        #         8. Suggestions: Offer advice on how to use remaining funds to build wealth after clearing liabilities and achive their financial goal.
                
                
        #         Here's an example for the required Output Format(if there are comments indicated by # in the example output format then thats a side note for your reference dont write it in the response that will be generated ) :
                
        #         Client's Financial Information :(# This is a header line have it in bold) 
                
                
        #         Client Name: """ + clientName + """(# have the client name in underline)

        #         (#the data presented is just an example for your reference do not consider it as factual refere to the document provided to you and generate data based on the provided data and only when nothing is provided assume some data for analysis.The data below it should be displayed in a table format so make sure of that data.List down all the assets and liabilities along with its values and then Total of assets,liabilities,etc.)
                
        #         - Total Assets: (# Sum of all client assets and Annual Income . Mention all assets and their respected values.if non consider the example assets)
                
        #         - Total Liabilities: (# Sum of all liabilities. Mention all liabilities and their respected values if non consider the example liabilities)
                
                
        #         - Monthly Liabilities: (# Monthly payments derived from liabilities)
                
        #         - Total Annual Income : (# Sum of all client's anual income)
                
        #         - Monthly Investment Amount : """ + monthly_investment + """ (# if no specific amount is specified to you then only assume  10,000 else consider the amount mention to you and just display the amount)
                
        #         - Investment Period : """ + investment_period + """  (# if no specific period is specified to you then only assume 3 years else consider the period mention to you and just display the period)


        #         Financial Analysis :(#Analyse the assets and liabilities and based on that give a suggestion for analysis generate suggestions for one of the following conditions:)
        #         (#1st condition : Everything is Positive)Based on the given Financial Conditions the client is having a good and stable income with great assets and manageable debt and liabilities.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the clients monthly income.
        #         (# if this condition is true then ignore the other conditions and start with the Investment Suggestions)
                
        #         (#2nd condition : Everything is temporarily Negative) Based on the given Financial Conditions the client is facing a low income for now but have great assets and manageable debt and liabilities.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the client's monthly income but the client might not be able to sustain the monthly investment amount that they are planning.)
        #         Instead I would like to recommend this amount to the client for their monthly investment : (#Mention a feasible amount to the client for monthly investment and start suggesting investments based on this amount and not the previous amount being taken into consideration)
                
        #         (#3rd condition : Everything is Negative) Based on the given Financial Conditions the client is facing a low income and doesnt have good assets to manage the debts and liabilities of the client and in such a condition this monthly investment amount is not feasible.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is not manageable for the client's monthly income and so the client might not be able to sustain the monthly investment amount that they are planning to do.)
        #         I would like to recommend this amount to the client for monthly investment : (# Mention a minimum amount to the client for monthly investment if possible else just say the client should first prioritize on savings and generating more income to manage their debts and liabilities first and so dont give any investment suggestions to the client.)
                
        #         (#If the financial is 1 or 2 only then give investment suggestions to the client)
                
                
        #         Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

        #         Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

        #         Investment Allocation: (#remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

        #         Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
        #         How to Invest: Diversify across various asset classes like:  (#Give allocations % as well)
                
        #         Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
                
        #         ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
                
        #         Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
                
        #         Consider investing in blue-chip companies or growth sectors like technology. 
                
                
        #         Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


        #         Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
        #         Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

        #         Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
        #         How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
        #         Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
        #         Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
        #         Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
        #         Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
        #         Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
        #         Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.


        #         Time Horizon and Expected Returns:

        #         Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
        #         Minimum Expected Annual Return: 4% - 6% 
                
                
        #         Maximum Expected Annual Return: 8% - 10% 
                
                
        #         Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, (# consider the monthly investment amount and give returns based on that only) $10,000 could grow to approximately 17,908 in 10 years.
        #         Minimum Expected Growth in Dollars: 
                
        #         4,000−6,000 (over 10 years) 
                
                
        #         Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
        #         Inflation Adjusted Returns:(#do not write this part inside the bracket just give answer,assume US inflation rate assume 3% if you dont know, and give the investment returns value that was suggested by you for the considered monthly investment amount after 3,5,10years of growth mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
        #         Rationale for Investment Suggestions:

        #         This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
        #         Important Considerations:

        #         Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

        #         Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

        #         Explain how this suggestions can help the client grow their wealth and improve their financial condition and/or cover up thier loans and in turn achive their Financial goals.
        #         <context>
        #         {context}
        #         </context>
        #         Question: {input}"""
                
        print("Retriever Created ")
        print(f"Investment Personality :{investmentPersonality}")
        
                

        llm_prompt = ChatPromptTemplate.from_template(prompt_template)

        document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
        combine_docs_chain = None  

        if retriever is not None :  
            retriever_chain = create_retrieval_chain(retriever,document_chain) 
            # print(retriever_chain)
            return retriever_chain
        else:
            print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
            return None

    except Exception as e:
        print(f"Error in creating chain: {e}")
        return None


import json
import io

# Process_Documents :
async def process_document(file_path): # GET Method
    try:
        print("Processing the document")
        file_type = filetype.guess(file_path)
        if file_type is not None:
            if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # Await the coroutine to extract text and tables
                return await extract_text_and_tables_from_word(file_path)
            elif file_type.mime == "application/pdf":
                return await extract_text_from_pdf(file_path)
        return None
    except Exception as e:
        print(f"Error processing document: {e}")
        return None

# Async function to extract text from a PDF file
async def extract_text_from_pdf(pdf_file_path): # GET Method
    try:
        print("Processing pdf file")
        with open(pdf_file_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file)
            text_content = []
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                text_content.append(page.extract_text())
            return "\n".join(text_content)
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None

# Async function to extract text and tables from a Word document
async def extract_text_and_tables_from_word(docx_file_path): # GET Method
    try:
        print("Extracting text and tables from word file")
        doc = docx.Document(docx_file_path)
        text_content = []
        tables_content = []

        for para in doc.paragraphs:
            text_content.append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables_content.append(table_data)
        print("Extracted text from word file")
        return "\n".join(text_content), tables_content
    except Exception as e:
        print(f"Error extracting text and tables from Word document: {e}")
        return None, None



async def validate_document_content(text, tables):
    """
    Validates the content of the document.

    Args:
        text (str): Extracted text content from the document.
        tables (list): Extracted tables content from the document.

    Returns:
        tuple: Client name and validation errors.
    """
    errors = []
    
    # Extract client name
    client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
    client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

    # Define required sections
    required_sections = [
        "YOUR RETIREMENT GOAL",
        "YOUR OTHER MAJOR GOALS",
        "YOUR ASSETS AND LIABILITIES",
        "MY LIABILITIES",
        "YOUR CURRENT ANNUAL INCOME"
    ]

    # Check for the presence of required sections
    for section in required_sections:
        if section not in text:
            errors.append(f"* {section} section missing.")
    
    # Define table field checks
    table_checks = {
        "YOUR RETIREMENT GOAL": [
            r"When do you plan to retire\? \(age or date\)",
            r"Social Security Benefit \(include expected start date\)",
            r"Pension Benefit \(include expected start date\)",
            r"Other Expected Income \(rental, part-time work, etc.\)",
            r"Estimated Annual Retirement Expense"
        ],
        "YOUR OTHER MAJOR GOALS": [
            r"GOAL", r"COST", r"WHEN"
        ],
        "YOUR ASSETS AND LIABILITIES": [
            r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
            r"Current Value", r"Annual Contributions"
        ],
        "MY LIABILITIES": [
            r"Balance", r"Interest Rate", r"Monthly Payment"
        ]
    }

    # Validate table content
    for section, checks in table_checks.items():
        section_found = False
        for table in tables:
            table_text = "\n".join(["\t".join(row) for row in table])
            if section in table_text:
                section_found = True
                for check in checks:
                    if not re.search(check, table_text, re.IGNORECASE):
                        errors.append(f"* Missing or empty field in {section} section: {check}")
                break
        if not section_found:
            errors.append(f"* {section} section missing.")

    return client_name, errors

####################################################################################################################################

################################################## Extract Numerical Data for Pie Chart, Bar Graph and Line Chart #####################################

import re
from collections import defaultdict
import numpy as np
# Updated for Line Chart :
import re
from collections import defaultdict

# def extract_numerical_data(response):
#     # Define patterns to match different sections and their respective allocations
#     patterns = {
#         'Growth-Oriented Investments': re.compile(r'Growth-Oriented Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
#         'Conservative Investments': re.compile(r'Conservative Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
#         'Time Horizon and Expected Returns': re.compile(r'Time Horizon and Expected Returns:(.*?)$', re.DOTALL)
#     }

#     data = defaultdict(dict)

#     for section, pattern in patterns.items():
#         match = pattern.search(response)
#         if match:
#             investments_text = match.group(1)
#             # Extract individual investment types and their allocations
#             investment_pattern = re.compile(r'(\w[\w\s]+?)\s*\((\d+%)-(\d+%)\)')
#             for investment_match in investment_pattern.findall(investments_text):
#                 investment_type, min_allocation, max_allocation = investment_match
#                 data[section][investment_type.strip()] = {
#                     'min': min_allocation,
#                     'max': max_allocation
#                 }

#     # Extract time horizon and expected returns
#     time_horizon_pattern = re.compile(r'Time Horizon:.*?(\d+)-(\d+) years', re.IGNORECASE)
#     min_return_pattern = re.compile(r'Minimum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
#     max_return_pattern = re.compile(r'Maximum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
#     min_growth_pattern = re.compile(r'Minimum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)
#     max_growth_pattern = re.compile(r'Maximum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)

#     time_horizon_match = time_horizon_pattern.search(response)
#     min_return_match = min_return_pattern.search(response)
#     max_return_match = max_return_pattern.search(response)
#     min_growth_match = min_growth_pattern.search(response)
#     max_growth_match = max_growth_pattern.search(response)

#     if time_horizon_match:
#         data['Time Horizon'] = {
#             'min_years': time_horizon_match.group(1),
#             'max_years': time_horizon_match.group(2)
#         }

#     if min_return_match:
#         data['Expected Annual Return'] = {
#             'min': min_return_match.group(1),
#             'max': min_return_match.group(2)
#         }

#     if max_return_match:
#         data['Expected Annual Return'] = {
#             'min': max_return_match.group(1),
#             'max': max_return_match.group(2)
#         }

#     if min_growth_match:
#         data['Expected Growth in Dollars'] = {
#             'min': min_growth_match.group(1),
#             'max': min_growth_match.group(2)
#         }

#     if max_growth_match:
#         data['Expected Growth in Dollars'] = {
#             'min': max_growth_match.group(1),
#             'max': max_growth_match.group(2)
#         }

#     # Extract inflation-adjusted returns
#     inflation_adjusted_pattern = re.compile(r'Inflation Adjusted Returns:.*?Before Inflation:.*?3 Years: \$(\d+,\d+).*?5 Years: \$(\d+,\d+).*?10 Years: \$(\d+,\d+).*?After Inflation.*?3 Years: \$(\d+,\d+).*?5 Years: \$(\d+,\d+).*?10 Years: \$(\d+,\d+)', re.DOTALL)
#     inflation_adjusted_match = inflation_adjusted_pattern.search(response)

#     if inflation_adjusted_match:
#         data['Inflation Adjusted Returns'] = {
#             'Before Inflation': {
#                 '3 Years': inflation_adjusted_match.group(1),
#                 '5 Years': inflation_adjusted_match.group(2),
#                 '10 Years': inflation_adjusted_match.group(3)
#             },
#             'After Inflation': {
#                 '3 Years': inflation_adjusted_match.group(4),
#                 '5 Years': inflation_adjusted_match.group(5),
#                 '10 Years': inflation_adjusted_match.group(6)
#             }
#         }

#     print(f"DATA extracted from Responses : {data}")
#     return data

# new code:
import re
from collections import defaultdict
import re
from collections import defaultdict

# extract numerical data from responses :

def extract_numerical_data(response):
    data = defaultdict(dict)

    # Match Growth-Oriented Investments and Conservative Investments sections
    growth_pattern = re.compile(r"<strong>Growth-Oriented Investments.*?</strong>:\s*(.*?)(<strong>|<h4>)", re.DOTALL)
    conservative_pattern = re.compile(r"<strong>Conservative Investments.*?</strong>:\s*(.*?)(<strong>|<h4>)", re.DOTALL)
    allocation_pattern = re.compile(r"<strong>(.*?)</strong>:\s*<code>(\d+%)\s*-\s*(\d+%)</code>")

    for category, pattern in [("Growth-Oriented Investments", growth_pattern), 
                               ("Conservative Investments", conservative_pattern)]:
        match = pattern.search(response)
        if match:
            investments_text = match.group(1)
            for investment_match in allocation_pattern.findall(investments_text):
                investment_type, min_allocation, max_allocation = investment_match
                data[category][investment_type.strip()] = {
                    'min': min_allocation.strip('%'),
                    'max': max_allocation.strip('%')
                }

    # Match Returns Overview
    returns_pattern = re.compile(r"<h4>Returns Overview:</h4>\s*(.*?)\s*<h4>", re.DOTALL)
    returns_match = returns_pattern.search(response)
    if returns_match:
        returns_text = returns_match.group(1)

        # Extract returns and growth data
        min_return_match = re.search(r"Minimum Expected Annual Return</strong>:\s*<code>(\d+%)\s*-\s*(\d+%)</code>", returns_text)
        max_return_match = re.search(r"Maximum Expected Annual Return</strong>:\s*<code>(\d+%)\s*-\s*(\d+%)</code>", returns_text)
        min_growth_match = re.search(r"Minimum Expected Growth in Dollars</strong>:\s*<code>\$(\d+,\d+)\s*-\s*\$(\d+,\d+)</code>", returns_text)
        max_growth_match = re.search(r"Maximum Expected Growth in Dollars</strong>:\s*<code>\$(\d+,\d+)\s*-\s*\$(\d+,\d+)</code>", returns_text)
        time_horizon_match = re.search(r"Time Horizon</strong>:\s*<code>(\d+ years)</code>", returns_text)

        if min_return_match:
            data['Expected Annual Return'] = {
                'min': min_return_match.group(1),
                'max': min_return_match.group(2)
            }
        if max_return_match:
            data['Expected Annual Return']['max'] = max_return_match.group(2)

        if min_growth_match:
            data['Expected Growth in Dollars'] = {
                'min': min_growth_match.group(1).replace(',', ''),
                'max': min_growth_match.group(2).replace(',', '')
            }
        if max_growth_match:
            data['Expected Growth in Dollars']['max'] = max_growth_match.group(2).replace(',', '')

        if time_horizon_match:
            data['Time Horizon'] = time_horizon_match.group(1)

    return data



# just prev code :
# def extract_numerical_data(response):
#     # Patterns for different sections
#     patterns = {
#         'Growth-Oriented Investments': re.compile(r'Growth-Oriented Investments.*?Target.*?:(.*?)Where to Invest:', re.DOTALL),
#         'Conservative Investments': re.compile(r'Conservative Investments.*?Target.*?:(.*?)Where to Invest:', re.DOTALL),
#         'Time Horizon and Expected Returns': re.compile(r'Time Horizon and Expected Returns.*?:\s*(.*?)$', re.DOTALL)
#     }

#     data = defaultdict(dict)

#     for section, pattern in patterns.items():
#         match = pattern.search(response)
#         if match:
#             investments_text = match.group(1)
#             # Extract investment details
#             investment_pattern = re.compile(r'([\w\s&/-]+?)\s*\((\d+%)-(\d+%)\)')
#             for investment_match in investment_pattern.findall(investments_text):
#                 investment_type, min_allocation, max_allocation = investment_match
#                 data[section][investment_type.strip()] = {
#                     'min': min_allocation.strip(),
#                     'max': max_allocation.strip()
#                 }

#     # Extract additional details
#     time_horizon_pattern = re.compile(r'Time Horizon.*?(\d+)-(\d+)\s*years', re.IGNORECASE)
#     min_return_pattern = re.compile(r'Minimum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
#     max_return_pattern = re.compile(r'Maximum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
#     min_growth_pattern = re.compile(r'Minimum Expected Growth in Dollars:.*?\$(\d[\d,]*)-\$(\d[\d,]*)', re.IGNORECASE)
#     max_growth_pattern = re.compile(r'Maximum Expected Growth in Dollars:.*?\$(\d[\d,]*)-\$(\d[\d,]*)', re.IGNORECASE)

#     time_horizon_match = time_horizon_pattern.search(response)
#     if time_horizon_match:
#         data['Time Horizon'] = {
#             'min_years': int(time_horizon_match.group(1)),
#             'max_years': int(time_horizon_match.group(2))
#         }

#     min_return_match = min_return_pattern.search(response)
#     if min_return_match:
#         data['Expected Annual Return'] = {
#             'min': min_return_match.group(1),
#             'max': min_return_match.group(2)
#         }

#     max_growth_match = max_growth_pattern.search(response)
#     if max_growth_match:
#         data['Expected Growth in Dollars'] = {
#             'min': int(max_growth_match.group(1).replace(',', '')),
#             'max': int(max_growth_match.group(2).replace(',', ''))
#         }

#     print("Section Data Extracted:", data)
#     print("Growth-Oriented Investments:", data.get('Growth-Oriented Investments', 'Not Found'))
#     print("Conservative Investments:", data.get('Conservative Investments', 'Not Found'))
#     print("Time Horizon Data:", data.get('Time Horizon', 'Not Found'))

#     return data


def normalize_allocations(allocations):
    total = sum(allocations)
    if total == 100:
        return allocations
    return [round((allocation / total) * 100, 2) for allocation in allocations]

# # Updated Line Chart 
import datetime  # Import the datetime module to get the current year

# line chart data code :

def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4):
    try:
        # Get the current year
        # curr_year = datetime.datetime.now().year
        curr_year = datetime.now().year
        

        # Print data_extracted to debug the structure
        print("Data extracted:", data_extracted)

        # Check if 'Expected Annual Return' and 'Time Horizon' exist and have the expected keys
        if 'Expected Annual Return' not in data_extracted:
            print("'Expected Annual Return' missing in data_extracted")
            data_extracted['Expected Annual Return'] = {'min': '8%', 'max': '20%'}
            min_return = 8 #6
            max_return = 20 #8
        else:
            min_return = float(data_extracted['Expected Annual Return'].get('min', '0').strip('%'))
            max_return = float(data_extracted['Expected Annual Return'].get('max', '0').strip('%'))

        min_years = int(data_extracted['Time Horizon'].get('min_years', 1))  # Default to 1 year if missing
        max_years = int(data_extracted['Time Horizon'].get('max_years', 10))  # Default to 10 years if missing

        def calculate_compounded_return(principal, rate, years):
            return principal * (1 + rate / 100) ** years

        def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
            return nominal_return / (1 + inflation_rate / 100) ** years

        # Create labels for the next 10 years starting from the current year
        labels = list(range(curr_year, curr_year + max_years))

        min_compounded = []
        max_compounded = []
        min_inflation_adjusted = []
        max_inflation_adjusted = []

        for year in range(1, max_years + 1):
            # Calculate nominal compounded returns
            min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
            max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

            # Calculate inflation-adjusted compounded returns
            min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
            max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

            # Append results
            min_compounded.append(min_compounded_value)
            max_compounded.append(max_compounded_value)
            min_inflation_adjusted.append(min_inflation_value)
            max_inflation_adjusted.append(max_inflation_value)

        # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
        combined_chart_data = {
            'labels': labels,  # Current year and the next 10 years
            'datasets': [
                {
                    'label': 'Minimum Compounded Return',
                    'data': min_compounded,
                    'borderColor': 'rgb(255, 99, 132)',  # Red color
                    'fill': False
                },
                {
                    'label': 'Maximum Compounded Return',
                    'data': max_compounded,
                    'borderColor': 'rgb(54, 162, 235)',  # Blue color
                    'fill': False
                },
                {
                    'label': 'Min Inflation Adjusted Return',
                    'data': min_inflation_adjusted,
                    'borderColor': 'rgb(75, 192, 192)',  # Light blue
                    'borderDash': [5, 5],  # Dashed line for distinction
                    'fill': False
                },
                {
                    'label': 'Max Inflation Adjusted Return',
                    'data': max_inflation_adjusted,
                    'borderColor': 'rgb(153, 102, 255)',  # Light purple
                    'borderDash': [5, 5],  # Dashed line for distinction
                    'fill': False
                }
            ]
        }
    except KeyError as e:
        print(f"KeyError occurred: {e}")
        return jsonify({'message': f'Key Error: {e}'}), 400
    except Exception as e:
        print(f"Error occurred while preparing data for combined line chart: {e}")
        return jsonify({'message': 'Internal Server Error in creating line chart'}), 500

    print(combined_chart_data)
    return combined_chart_data


# import datetime  # Import the datetime module to get the current year
# # uodated to have current year
# def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4):
#     try:
#         # Get the current year
#         curr_year = datetime.datetime.now().year

#         # Print data_extracted to debug the structure
#         print("Data extracted:", data_extracted)

#         # Check if 'Expected Annual Return' and 'Time Horizon' exist and have the expected keys
#         if 'Expected Annual Return' not in data_extracted:
#             print("'Expected Annual Return' missing in data_extracted")
#             data_extracted['Expected Annual Return']['min'] = 6
#             data_extracted['Expected Annual Return']['max'] = 8
#             min_return = 6
#             max_return = 8
#         else:
#             min_return = float(data_extracted['Expected Annual Return'].get('min', '0').strip('%'))
#             max_return = float(data_extracted['Expected Annual Return'].get('max', '0').strip('%'))

#         min_years = int(data_extracted['Time Horizon'].get('min_years', 1))  # Default to 1 year if missing
#         max_years = int(data_extracted['Time Horizon'].get('max_years', 10))  # Default to 10 years if missing

#         def calculate_compounded_return(principal, rate, years):
#             return principal * (1 + rate / 100) ** years

#         def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
#             return nominal_return / (1 + inflation_rate / 100) ** years

#         # Create labels for the next 10 years starting from the current year
#         labels = list(range(curr_year, curr_year + max_years))

#         min_compounded = []
#         max_compounded = []
#         min_inflation_adjusted = []
#         max_inflation_adjusted = []

#         for year in range(1, max_years + 1):
#             # Calculate nominal compounded returns
#             min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
#             max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

#             # Calculate inflation-adjusted compounded returns
#             min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
#             max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

#             # Append results
#             min_compounded.append(min_compounded_value)
#             max_compounded.append(max_compounded_value)
#             min_inflation_adjusted.append(min_inflation_value)
#             max_inflation_adjusted.append(max_inflation_value)

#         # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
#         combined_chart_data = {
#             'labels': labels,  # Current year and the next 10 years
#             'datasets': [
#                 {
#                     'label': 'Minimum Compounded Return',
#                     'data': min_compounded,
#                     'borderColor': 'rgb(255, 99, 132)',  # Red color
#                     'fill': False
#                 },
#                 {
#                     'label': 'Maximum Compounded Return',
#                     'data': max_compounded,
#                     'borderColor': 'rgb(54, 162, 235)',  # Blue color
#                     'fill': False
#                 },
#                 {
#                     'label': 'Min Inflation Adjusted Return',
#                     'data': min_inflation_adjusted,
#                     'borderColor': 'rgb(75, 192, 192)',  # Light blue
#                     'borderDash': [5, 5],  # Dashed line for distinction
#                     'fill': False
#                 },
#                 {
#                     'label': 'Max Inflation Adjusted Return',
#                     'data': max_inflation_adjusted,
#                     'borderColor': 'rgb(153, 102, 255)',  # Light purple
#                     'borderDash': [5, 5],  # Dashed line for distinction
#                     'fill': False
#                 }
#             ]
#         }
#     except KeyError as e:
#         print(f"KeyError occurred: {e}")
#         return jsonify({'message': f'Key Error: {e}'}), 400
#     except Exception as e:
#         print(f"Error occurred while preparing data for combined line chart: {e}")
#         return jsonify({'message': 'Internal Server Error in creating line chart'}), 500

#     return combined_chart_data





# def plot_investment_allocations(data):
#     # Create subplots with a large figure size
#     fig, axes = plt.subplots(2, 1, figsize= (16,10)) #(28, 15))  # Adjust size as needed

#     # Plot Growth-Oriented Investments
#     growth_data = data['Growth-Oriented Investments']
#     growth_labels = list(growth_data.keys())
#     growth_min = [int(growth_data[label]['min'].strip('%')) for label in growth_labels]
#     growth_max = [int(growth_data[label]['max'].strip('%')) for label in growth_labels]

#     axes[0].bar(growth_labels, growth_min, color='skyblue', label='Min Allocation')
#     axes[0].bar(growth_labels, growth_max, bottom=growth_min, color='lightgreen', label='Max Allocation')
#     axes[0].set_title('Growth-Oriented Investments', fontsize=16)
#     axes[0].set_ylabel('Percentage Allocation', fontsize=14)
#     axes[0].set_xlabel('Investment Types', fontsize=14)
#     axes[0].tick_params(axis='x', rotation=45, labelsize=12)
#     axes[0].tick_params(axis='y', labelsize=12)
#     axes[0].legend()

#     # Plot Conservative Investments
#     conservative_data = data['Conservative Investments']
#     conservative_labels = list(conservative_data.keys())
#     conservative_min = [int(conservative_data[label]['min'].strip('%')) for label in conservative_labels]
#     conservative_max = [int(conservative_data[label]['max'].strip('%')) for label in conservative_labels]

#     axes[1].bar(conservative_labels, conservative_min, color='skyblue', label='Min Allocation')
#     axes[1].bar(conservative_labels, conservative_max, bottom=conservative_min, color='lightgreen', label='Max Allocation')
#     axes[1].set_title('Conservative Investments', fontsize=16)
#     axes[1].set_ylabel('Percentage Allocation', fontsize=14)
#     axes[1].set_xlabel('Investment Types', fontsize=14)
#     axes[1].tick_params(axis='x', rotation=45, labelsize=12)
#     axes[1].tick_params(axis='y', labelsize=12)
#     axes[1].legend()

#     # Tight layout for better spacing
#     plt.tight_layout()
#     plt.show()
#     return fig


# def plot_pie_chart(data):
#     fig, ax = plt.subplots(figsize=(10, 7))  # Increased size

#     # Combine all investment data for pie chart
#     all_data = {**data['Growth-Oriented Investments'], **data['Conservative Investments']}
#     labels = list(all_data.keys())
#     sizes = [int(all_data[label]['max'].strip('%')) for label in labels]
#     colors = plt.cm.Paired(range(len(labels)))

#     wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
#     ax.set_title('Investment Allocation')

#     # Add legend
#     ax.legend(wedges, labels, title="Investment Types", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))

#     return fig



# def bar_chart(data):
#     fig, ax = plt.subplots(figsize=(12, 8))  # Increased size

#     # Data for plotting
#     categories = list(data.keys())
#     values_min = [int(data[cat]['min'].strip('%')) for cat in categories]
#     values_max = [int(data[cat]['max'].strip('%')) for cat in categories]

#     x = range(len(categories))

#     ax.bar(x, values_min, width=0.4, label='Min Allocation', color='skyblue', align='center')
#     ax.bar(x, values_max, width=0.4, label='Max Allocation', color='lightgreen', align='edge')

#     ax.set_xticks(x)
#     ax.set_xticklabels(categories, rotation=45, ha='right')
#     ax.set_xlabel('Investment Categories')
#     ax.set_ylabel('Percentage Allocation')
#     ax.set_title('Investment Allocation')
#     ax.legend()

#     plt.tight_layout()
#     return fig


import random
# generate colors for pie chart :

def generate_colors(n):
    """
    Generate 'n' random RGB colors.

    Args:
        n (int): Number of colors to generate.
    
    Returns:
        list: A list of RGB colors in 'rgb(r, g, b)' format.
    """
    colors = []
    for _ in range(n):
        r = random.randint(0, 255)
        g = random.randint(0, 255)
        b = random.randint(0, 255)
        colors.append(f'rgb({r}, {g}, {b})')
    
    return colors


# import plotly.graph_objects as go
import numpy as np



from datetime import date  # Make sure to import the date class


# Function to parse financial data from the text
import re

def parse_financial_data(text_content):
    assets = []
    liabilities = []

    # Define regex patterns to capture text following headings
    asset_pattern = re.compile(r"MY ASSETS:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)
    liability_pattern = re.compile(r"LIABILITIES:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)

    # Extract assets
    asset_matches = asset_pattern.findall(text_content)
    if asset_matches:
        asset_text = asset_matches[0]
        # Further processing to extract individual asset values if they are detailed
        asset_lines = asset_text.split('\n')
        for line in asset_lines:
            match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
            if match:
                asset_value = float(match.group().replace(",", ""))
                assets.append(asset_value)

    # Extract liabilities
    liability_matches = liability_pattern.findall(text_content)
    if liability_matches:
        liability_text = liability_matches[0]
        # Further processing to extract individual liability values if they are detailed
        liability_lines = liability_text.split('\n')
        for line in liability_lines:
            match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
            if match:
                liability_value = float(match.group().replace(",", ""))
                liabilities.append(liability_value)

    print("Assets Found:", assets)
    print("Liabilities Found:", liabilities)

    return assets, liabilities



# Function to extract numerical values from a text input
def extract_numeric(value):
    try:
        return float(re.sub(r'[^\d.]', '', value))  # Remove non-numeric characters and convert to float
    except ValueError:
        return 0


# plots graph from the details of the form :


def is_float(value):
    try:
        float(value)
        return True
    except ValueError:
        return False



def save_data_to_file(form_data):
    file_path = 'client_data.txt'
    with open(file_path, 'a') as file:
        file.write(str(form_data) + "\n")
    # st.success(f"Form data saved to {file_path}")
    print(f"Form data saved to {file_path}")
    
import math
# calculate compunded amount :
def calculate_compounded_amount(principal, rate, time):
    """
    Calculates the compounded amount using the formula:
    A = P * (1 + r/n)^(nt)
    Assuming n (compounding frequency) is 1 for simplicity (annually).
    """
    if principal == 0 or rate == 0 or time == 0:
        return principal
    else:
        # Using annual compounding
        return principal * (1 + rate / 100) ** time
    
def calculate_totals(assets, liabilities):
    total_assets = sum(extract_numeric(v) for v in assets.values())
    print(f"Total Assets : {total_assets}")
    total_liabilities = 0
    total_liabilities = sum(extract_numeric(v) for v in liabilities.values() )

    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Mortgage']),
    #     liabilities['Annual Mortgage Interest Rate'],
    #     liabilities['Mortagage Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Home Loans']),
    #     liabilities['Home Loans Interest Rate'],
    #     liabilities['Home Loans Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Vehicle Loans']),
    #     liabilities['Vehicle Loans Interest Rate'],
    #     liabilities['Vehicle Loans Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Education Loans']),
    #     liabilities['Education Loans Interest Rate'],
    #     liabilities['Education Loans Time Period']
    # )
    
    # For credit card debt, only calculate compounded amount if interest rate > 0

    # credit_card_balance = extract_numeric(liabilities['Credit Card'])
    # credit_card_interest = liabilities['Credit Card Debt Interest Rate']
    # if credit_card_interest > 0:
    #     # Assuming the time period for credit card debt is 1 year for compounding
    #     total_liabilities += calculate_compounded_amount(credit_card_balance, credit_card_interest, 1)
    # else:
    #     total_liabilities += credit_card_balance
    
    # Miscellaneous debts are taken directly as is
    total_liabilities += extract_numeric(liabilities['Miscellaneous'])
    rounded_liabilities = round(total_liabilities,2)

    print(f"Total liabilities :{total_liabilities}")
    print(f"Rounded of Total liabilities :{rounded_liabilities}")

    return total_assets, rounded_liabilities #total_liabilities



from docx import Document
# Define a helper function to read and extract text from a DOCX file
def read_docx(file_path):
    document = Document(file_path)
    extracted_text = "\n".join([para.text for para in document.paragraphs])
    return extracted_text



class TrieNode:
    def __init__(self):
        self.children = {}
        self.client_ids = []
        self.end_of_name = False  # Marks the end of a client's name

class Trie:
    def __init__(self):
        self.root = TrieNode()

    def insert(self, name, client_id):
        node = self.root
        for char in name:
            if char not in node.children:
                node.children[char] = TrieNode()
            node = node.children[char]
        node.client_ids.append(client_id)
        node.end_of_name = True

    def search(self, prefix):
        node = self.root
        for char in prefix:
            if char in node.children:
                node = node.children[char]
            else:
                return []  # Prefix not found
        return self._get_all_names_from_node(prefix, node)

    def _get_all_names_from_node(self, prefix, node):
        suggestions = []
        if node.end_of_name:
            suggestions.append((prefix, node.client_ids))
        for char, child_node in node.children.items():
            suggestions.extend(self._get_all_names_from_node(prefix + char, child_node))
        return suggestions



def preload_trie():
    trie = Trie()
    clients = {
        "John Doe": "C001",
        "Jane Smith": "C002",
        "James Brown": "C003",
        "Jill Johnson": "C004",
        "Jake White": "C005"
    }
    for name, client_id in clients.items():
        trie.insert(name.lower(), client_id)  # Insert in lowercase for case-insensitive search
    return trie

# generate suggestions :
async def generate_investment_suggestions_for_investor(investment_personality,clientName,financial_data,financial_file,monthly_investment=10000,investment_period=3): # # GET Method for py , for front end its Post API
    
    # retriever = asyncio.run(load_vector_db("uploaded_file"))

    # retriever =  await load_vector_db("uploaded_file")
    try:
        retriever =  await load_vector_db(financial_file)
    except Exception as e :
        print(f"Error : {e}")
        return jsonify("Error : Failed to load vector database and to generate suggestions : {e}"),400
    
    if not retriever:
        # await load_vector_db("data\Financial_Investment_new.docx")
        await load_vector_db("data\EW2400.docx")
        # await load_vector_db("data\Financial_Investment_1_new.docx") # doesnt works
        # await load_vector_db("data\Financial_Investment_1.docx")
        if not retriever:
            raise Exception("Failed to load vector database.")
    
    print("VectorDB is created successfully")
    # retriever = await load_vector_db("data\Financial_Investment_1.docx") 
    
    try:
        chain = await make_retrieval_chain(retriever,investment_personality,clientName,monthly_investment,investment_period)
    except Exception as e :
        print(f"Error : {e}")
        return jsonify("Error : Failed to create retrieval chain and generate suggestions : {e}"),400
    
    if not chain:
        raise Exception("Failed to create retrieval chain.")
    print("Chain is created to generate suggestions ")
    
    # chain = asyncio.run(make_retrieval_chain(retriever))
    
    print(f"Financial Data : {financial_data}")
    try :
        print(type(financial_data))
        query = f"The Investment Personality of {clientName} is : {investment_personality}" + f"Consider the Monthly Investment as {monthly_investment} and Investment period as {investment_period}" + f"Financial Data of client is : {financial_data[0]}"
        print(query)
    except Exception as e :
        print(f"Error : {e}")
        return "Error : Failed to load financial data"
    
    if chain is not None:
        # summary = context
        # query = summary + "\n" + investment_personality
        
        # query = str(investment_personality)
        response = chain.invoke({"input": query})
        
        # format_response = markdown_to_text(response['answer'])
        # return format_response
        
        # html_output = markdown.markdown(response['answer'])
        # return html_output
        
        # readable_text = markdown_to_readable_text(response['answer'])
        # print(readable_text)
        # return readable_text

        # format_text = convert_to_markdown(response['answer'])
        # return format_text
        
        return response['answer']
    
        

        # handle_graph(response['answer'])

    else:
        logging.INFO("response is not generated by llm model")
        return jsonify("response is not generated by llm model"),500
        # st.error("Failed to create the retrieval chain. Please upload a valid document.")

####################################################################################################################
# app begining :
# CORS(app,resources={r"/api/*":{"origins":"*"}})
# CORS(app)

# Initialize the Trie with preloaded clients
trie = preload_trie()

@app.route('/')
def home():
    return "Wealth Advisor Chatbot API"




import random

# Generate unique client ID

# def generate_unique_id(name):
#     name_parts = name.split(" ")
#     first_initial = name_parts[0][0] if len(name_parts) > 0 else ""
#     last_initial = name_parts[1][0] if len(name_parts) > 1 else ""
#     random_number = random.randint(1000, 9999)
#     unique_id = f"{first_initial}{last_initial}{random_number}"
#     return unique_id

# # Save details in a Word file
import docx
import os

# #Curr version :

# Financial Form
def save_to_word_file(data, file_name):
    doc = docx.Document()
    doc.add_heading('Client Details', 0)

    # Adding client details
    client_details = data.get('clientDetail', {})
    doc.add_paragraph(f"Client Name: {client_details.get('clientName', '')}")
    doc.add_paragraph(f"Client Mobile: {client_details.get('clientMoNo', '')}")
    doc.add_paragraph(f"Client Age: {client_details.get('clientAge', '')}")
    doc.add_paragraph(f"Co-Client Name: {client_details.get('coClientName', '')}")
    doc.add_paragraph(f"Co-Client Mobile: {client_details.get('coMobileNo', '')}")
    doc.add_paragraph(f"Co-Client Age: {client_details.get('coClientAge', '')}")

    # Retirement Plan
    retirement_goal = data.get('retirementGoal', {})
    retirement_plan = retirement_goal.get('retirementPlan', {})
    doc.add_paragraph(f"Retirement Plan Client Age: {retirement_plan.get('retirementAgeClient', '')}")
    doc.add_paragraph(f"Retirement Plan Co-Client Age: {retirement_plan.get('retirementAgeCoClient', '')}")
    
    social_benefit = retirement_goal.get('socialBenefit', {})
    doc.add_paragraph(f"Social Benefit Client: {social_benefit.get('socialBenefitClient', '')}")
    doc.add_paragraph(f"Social Benefit Co-Client: {social_benefit.get('socialBenefitCoClient', '')}")
    
    pension_benefit = retirement_goal.get('pensionBenefit', {})
    doc.add_paragraph(f"Pension Benefit Client: {pension_benefit.get('pensionBenefitClient', '')}")
    doc.add_paragraph(f"Pension Benefit Co-Client: {pension_benefit.get('pensionBenefitCoClient', '')}")
    
    otherIncome = retirement_goal.get('otherIncome', {})
    doc.add_paragraph(f"Other IncomeClient Client: {otherIncome.get('otherIncomeClient', '')}")
    doc.add_paragraph(f"Other IncomeClient Co-Client: {otherIncome.get('otherIncomeCoClient', '')}")
   
    # Estimated Annual Retirement Expense ($ or % of current salary)
    annualRetirement = retirement_goal.get('annualRetirement', {})
    doc.add_paragraph(f"Estimated Annual Retirement Expense ($ or % of current salary) Client: {annualRetirement.get('annualRetireClient', '')}")
    doc.add_paragraph(f"Estimated Annual Retirement Expense ($ or % of current salary) Co-Client: {annualRetirement.get('annualRetireCoClient', '')}")
    

    # Assets and Liabilities
    assets_liabilities = data.get('assetsLiabilities', {})
    
    # Assets
    
    for asset_key, asset_info in assets_liabilities.items():
        current_value_key = [key for key in asset_info.keys() if key.startswith("current")][0]
        annual_value_key = [key for key in asset_info.keys() if key.startswith("annual")][0]
        assets_name_key = "assetsName"
        doc.add_paragraph(f"Assets - {asset_info[assets_name_key]} : Current Value - {asset_info[current_value_key]} , Annual Contributions - {asset_info[annual_value_key]}")
        
    # Liabilities
    myLiabilities = data.get('myLiabilities', {})
    for liability_key, liability_info in myLiabilities.items():
        balance_key = [key for key in liability_info.keys() if key.endswith("Balance")][0]
        interest_key = [key for key in liability_info.keys() if key.endswith("Interest")][0]
        monthly_key = [key for key in liability_info.keys() if key.endswith("Monthly")][0]
        liability_name_key = "liabilityName"
        doc.add_paragraph(f"Liabilities - {liability_info[liability_name_key]} : Balance - {liability_info[balance_key]} , Interest - {liability_info[interest_key]} , Monthly - {liability_info[monthly_key]}")
        
    # my_liabilities = data.get('myLiabilities', {})
    # for liability_type, liability_info in my_liabilities.items():
    #     doc.add_paragraph(f"Liabilities - {liability_info.get('liabilityName', '')}: Balance - {liability_info.get('mortgageBalance', '')} Interest - {liability_info.get('mortgageInterest', '')} Monthly - {liability_info.get('mortgageMonthly', '')}")

    # Protection Plan
    protection_plan = data.get('protectionPlan', {})
    doc.add_paragraph(f"Check Will: {protection_plan.get('checkWill', False)}")
    doc.add_paragraph(f"Check Healthcare: {protection_plan.get('checkHealthCare', False)}")
    doc.add_paragraph(f"Check Attorney: {protection_plan.get('checkAttorney', False)}")
    doc.add_paragraph(f"Check Trust: {protection_plan.get('checkTrust', False)}")

    # Insurance Coverage
    insurance_coverage = data.get('insuranceCoverage', {})
    life_insurance_client = insurance_coverage.get('lifeInsuranceClient', {})
    doc.add_paragraph(f"Life Insurance Client: Benefit - {life_insurance_client.get('benefitLIClient', '')} Monthly Pay - {life_insurance_client.get('monthlyPayLIClient', '')}")
    
    life_insurance_co_client = insurance_coverage.get('lifeInsuranceCoClient', {})
    doc.add_paragraph(f"Life Insurance Co-Client: Benefit - {life_insurance_co_client.get('benefitLICoClient', '')} Monthly Pay - {life_insurance_co_client.get('monthlyPayLICoClient', '')}")
 
    disableIncome = insurance_coverage.get('disableIncomeClient', {})
    disableIncomeClient = insurance_coverage.get('disableIncomeClient',{})
    doc.add_paragraph(f"Disable Income Client - {disableIncomeClient.get('benefitDisableClient', '')}")
    
    disableIncomeCoClient = insurance_coverage.get('disableIncomeCoClient', {})
    doc.add_paragraph(f"Disable Income Co-Client - {disableIncomeCoClient.get('benefitDisableCoClient', '')}")
    
    longTermCoClient = insurance_coverage.get('longTermCoClient')
    doc.add_paragraph(f"Long Term Client: Benefit - {longTermCoClient.get('benefitLongTermClient', '')} Monthly Pay - {longTermCoClient.get('monthlyPayLongTermClient', '')}")
    
    investmentAmount = insurance_coverage.get('investmentAmount')
    doc.add_paragraph(f"Investment Amount Available : {investmentAmount}")
                      
    # Goal Fields
    goal_fields = data.get('goalFields', [])
    for goal in goal_fields:
        doc.add_paragraph(f"Goal: {goal.get('goal', '')} Cost: {goal.get('cost', '')} When: {goal.get('when', '')}")

    # Income Fields
    income_fields = data.get('incomeFields', [])
    for income in income_fields:
        doc.add_paragraph(f"Income Source: {income.get('sourceIncome', '')} Amount: {income.get('amountIncome', '')}")

    # funds_investment = data.get('Funds',[]) commented for later use
    # Save file
    file_name = os.path.join("data", file_name)
    doc.save(f"{file_name}.docx")





# Local Folder Path
CLIENT_DATA_DIR = './client_data'
 
# # store client data in aws :
 
@app.route('/submit-client-data', methods=['POST'])
def submit_client_data():
    try:
        # Parse JSON payload
        data = request.get_json()
        if not data:
            return jsonify({'message': 'Invalid or missing request payload'}), 400
 
        # Extract client details
        client_name = data.get('clientDetail', {}).get('clientName')
        unique_id = data.get('uniqueId')
 
        if not client_name or not unique_id:
            return jsonify({'message': 'Client name and unique ID are required'}), 400
 
        print(f"Processing data for client: {client_name}, ID: {unique_id}")
 
        if USE_AWS:
            # AWS Logic
            s3_key = f"{client_summary_folder}client-data/{unique_id}.json"
            try:
                # Check if the client data already exists in S3
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
                existing_data = json.loads(response['Body'].read().decode('utf-8'))
                is_update = True
                print(f"Existing data found for unique ID: {unique_id}")
            except s3.exceptions.NoSuchKey:
                existing_data = {}
                is_update = False
                print(f"No existing data found for unique ID: {unique_id}. Creating new record.")
 
            # Merge or replace the existing data
            if is_update:
                existing_data.update(data)
                data_to_save = existing_data
            else:
                data_to_save = data
 
            # Save to S3
            try:
                s3.put_object(
                    Bucket=S3_BUCKET_NAME,
                    Key=s3_key,
                    Body=json.dumps(data_to_save),
                    ContentType="application/json"
                )
                action = "updated" if is_update else "created"
                print(f"Client data successfully {action} in S3 for unique ID: {unique_id}")
            except Exception as s3_error:
                logging.error(f"Error uploading data to S3: {s3_error}")
                return jsonify({'message': f"Error uploading data to S3: {s3_error}"}), 500
 
        else:
            # Local Storage Logic
            file_path = os.path.join(CLIENT_DATA_DIR, f"client_data/{unique_id}.json")
 
            # Check if the client data already exists
            if os.path.exists(file_path):
                with open(file_path, 'r') as f:
                    existing_data = json.load(f)
                existing_data.update(data)  # Merge the new data
                is_update = True
            else:
                existing_data = data  # New data
                is_update = False
 
            # Save the data to local storage
            with open(file_path, 'w') as f:
                json.dump(existing_data, f, indent=4)
 
            action = "updated" if is_update else "created"
            print(f"Client data successfully {action} for unique ID: {unique_id}")
 
        # Return a success response
        return jsonify({
            'message': f'Client data successfully {action}.',
            'uniqueId': unique_id
        }), 200
 
    except Exception as e:
        logging.error(f"An error occurred: {e}")
        return jsonify({'message': f"An error occurred: {e}"}), 500
 
def process_is_new_client(client_data):
    """
    Adds `isNewClient` flag to the client data based on order existence.
    """
    client_data["isNewClient"] = True  # Default to True
    client_id = client_data.get("uniqueId")
 
    if client_id:
        try:
            order_url = 'http://localhost:5000/show_order_list'  # Replace with actual endpoint
            response = requests.post(order_url, json={'client_id': client_id})
 
            if response.status_code == 200:
                orders = response.json().get("transaction_data", [])
                if len(orders) > 0:
                    client_data["isNewClient"] = False
        except Exception as e:
            print(f"Error checking orders for client {client_id}: {e}")
 
 
# # get client data by id :

# Ensure the directory exists
if not os.path.exists(CLIENT_DATA_DIR):
    os.makedirs(CLIENT_DATA_DIR)
 
# Get client data by client ID
@app.route('/get-client-data-by-id', methods=['GET'])
def get_client_data():
    try:
        # Retrieve client_id from query parameters
        client_id = request.args.get('client_id')
 
        # Validate the client_id
        if not client_id:
            return jsonify({'message': 'client_id is required as a query parameter'}), 400
 
        if USE_AWS:
            # AWS Logic
            s3_key = f"{client_summary_folder}client-data/{client_id}.json"
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
                client_data = json.loads(response['Body'].read().decode('utf-8'))
                return jsonify({
                    'message': 'Client data retrieved successfully.',
                    'data': client_data
                }), 200
            except s3.exceptions.NoSuchKey:
                return jsonify({'message': 'Client data not found for the given client_id.'}), 404
            except Exception as e:
                return jsonify({'message': f"Error retrieving data: {e}"}), 500
        else:
            # Local Storage Logic
            file_path = os.path.join(CLIENT_DATA_DIR, f"client_data/{client_id}.json")
 
            # Check if the file exists and retrieve the data
            if not os.path.exists(file_path):
                return jsonify({'message': 'Client data not found for the given client_id.'}), 404
 
            with open(file_path, 'r') as f:
                client_data = json.load(f)
 
            return jsonify({
                'message': 'Client data retrieved successfully.',
                'data': client_data
            }), 200
 
    except Exception as e:
        return jsonify({'message': f"An error occurred: {e}"}), 500
  
 
# # get all client data :

# Local Folder Path
LOCAL_CLIENT_DATA_FOLDER = './client_data/client_data'
 
# new version :

@app.route('/get-all-client-data', methods=['GET'])
def get_all_client_data():
    try:
        all_data = []
 
        if USE_AWS:
            # AWS Logic
            response = s3.list_objects_v2(Bucket=S3_BUCKET_NAME, Prefix=client_summary_folder)
            if 'Contents' in response:
                for obj in response['Contents']:
                    try:
                        file_key = obj['Key']
                        file_response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=file_key)
                        file_data = file_response['Body'].read().decode('utf-8')
                        data_json = json.loads(file_data)
 
                        # Do not overwrite isNewClient flag if updated elsewhere
                        if 'isNewClient' not in data_json:
                            data_json['isNewClient'] = True  # Default to True if missing
                       
                        all_data.append(data_json)
                    except Exception as e:
                        print(f"Error reading file {obj['Key']}: {e}")
                        continue
            else:
                return jsonify({'message': 'No client data found in S3 bucket.'}), 404
 
        else:
            # Local Storage Logic
            for filename in os.listdir(LOCAL_CLIENT_DATA_FOLDER):
                if filename.endswith(".json"):
                    file_path = os.path.join(LOCAL_CLIENT_DATA_FOLDER, filename)
                    with open(file_path, 'r') as f:
                        client_data = json.load(f)
 
                        # Do not overwrite isNewClient flag if updated elsewhere
                        if 'isNewClient' not in client_data:
                            client_data['isNewClient'] = True  # Default to True if missing
                       
                        all_data.append(client_data)
 
            if not all_data:
                return jsonify({'message': 'No client data found in local storage.'}), 404
 
        # Return combined data
        return jsonify({
            'message': 'All client data retrieved successfully.',
            'data': all_data
        }), 200
 
    except Exception as e:
        print(f"Error occurred while retrieving data: {e}")
        return jsonify({'message': f"Error occurred while retrieving data: {e}"}), 500
  
 
# prev -version: 
# @app.route('/get-all-client-data', methods=['GET'])
# def get_all_client_data():
#     try:
#         all_data = []
 
#         if USE_AWS:
#             # AWS Logic
#             response = s3.list_objects_v2(Bucket=S3_BUCKET_NAME, Prefix=client_summary_folder)
#             if 'Contents' in response:
#                 for obj in response['Contents']:
#                     try:
#                         file_key = obj['Key']
#                         file_response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=file_key)
#                         file_data = file_response['Body'].read().decode('utf-8')
#                         data_json = json.loads(file_data)
 
#                         # Add isNewClient flag
#                         # process_is_new_client(data_json)
#                         all_data.append(data_json)
#                     except Exception as e:
#                         print(f"Error reading file {obj['Key']}: {e}")
#                         continue
#             else:
#                 return jsonify({'message': 'No client data found in S3 bucket.'}), 404
 
#         else:
#             # Local Storage Logic
#             for filename in os.listdir(LOCAL_CLIENT_DATA_FOLDER):
#                 if filename.endswith(".json"):
#                     file_path = os.path.join(LOCAL_CLIENT_DATA_FOLDER, filename)
#                     with open(file_path, 'r') as f:
#                         client_data = json.load(f)
#                         # process_is_new_client(client_data)
#                         all_data.append(client_data)
 
#             if not all_data:
#                 return jsonify({'message': 'No client data found in local storage.'}), 404
 
#         # Return combined data
#         return jsonify({
#             'message': 'All client data retrieved successfully.',
#             'data': all_data
#         }), 200
 
#     except Exception as e:
#         return jsonify({'message': f"Error occurred while retrieving data: {e}"}), 500
    
    # old version:
# # get client data by client id :
# @app.route('/get-client-data-by-id', methods=['GET'])
# def get_client_data():
#     try:
#         # Retrieve client_id from query parameters
#         client_id = request.args.get('client_id')
        
#         # Validate the client_id
#         if not client_id:
#             return jsonify({'message': 'client_id is required as a query parameter'}), 400

#         # Define the S3 key for the object
#         s3_key = f"{client_summary_folder}client-data/{client_id}.json"

#         # Retrieve the object from S3
#         try:
#             response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
#             # Decode and parse the JSON data
#             client_data = json.loads(response['Body'].read().decode('utf-8'))
            
#             return jsonify({
#                 'message': 'Client data retrieved successfully.',
#                 'data': client_data
#             }), 200
#         except s3.exceptions.NoSuchKey:
#             return jsonify({'message': 'Client data not found for the given client_id.'}), 404
#         except Exception as e:
#             return jsonify({'message': f"Error retrieving data: {e}"}), 500

#     except Exception as e:
#         return jsonify({'message': f"An error occurred: {e}"}), 500

#################################################################################################################################
# storing client data using local storage :
# Local storage directories
LOCAL_STORAGE_DIR = "local_storage"
CLIENT_DATA_DIR = os.path.join(LOCAL_STORAGE_DIR, "client_data")
# Ensure directories exist
os.makedirs(CLIENT_DATA_DIR, exist_ok=True)


# @app.route('/submit-client-data', methods=['POST'])
# def submit_client_data():
#     try:
#         # Parse JSON payload
#         data = request.get_json()
#         if not data:
#             return jsonify({'message': 'Invalid or missing request payload'}), 400

#         # Extract client details
#         client_name = data.get('clientDetail', {}).get('clientName')
#         unique_id = data.get('uniqueId')

#         if not client_name or not unique_id:
#             return jsonify({'message': 'Client name and unique ID are required'}), 400

#         print(f"Processing data for client: {client_name}, ID: {unique_id}")

#         # Define the file path for local storage
#         file_path = os.path.join(CLIENT_DATA_DIR, f"client_data/{unique_id}.json")

#         # Check if the client data already exists
#         if os.path.exists(file_path):
#             with open(file_path, 'r') as f:
#                 existing_data = json.load(f)
#             existing_data.update(data)  # Merge the new data
#             is_update = True
#         else:
#             existing_data = data  # New data
#             is_update = False

#         # Save the data to local storage
#         with open(file_path, 'w') as f:
#             json.dump(existing_data, f, indent=4)

#         action = "updated" if is_update else "created"
#         print(f"Client data successfully {action} for unique ID: {unique_id}")

#         return jsonify({
#             'message': f'Client data successfully {action}.',
#             'uniqueId': unique_id
#         }), 200

#     except Exception as e:
#         return jsonify({'message': f"An error occurred: {e}"}), 500


# # Define the directory where client data is stored
# CLIENT_DATA_DIR = './client_data/'
# CLIENT_SUMMARY_DIR = os.path.join(CLIENT_DATA_DIR, "client_data")

# # Ensure the directory exists
# if not os.path.exists(CLIENT_DATA_DIR):
#     os.makedirs(CLIENT_DATA_DIR)

# # Get client data by client ID
# @app.route('/get-client-data-by-id', methods=['GET'])
# def get_client_data():
#     try:
#         # Retrieve client_id from query parameters
#         client_id = request.args.get('client_id')

#         # Validate the client_id
#         if not client_id:
#             return jsonify({'message': 'client_id is required as a query parameter'}), 400

#         # Define the file path for the client data
#         file_path = os.path.join(CLIENT_DATA_DIR, f"client_data/{client_id}.json")

#         # Check if the file exists and retrieve the data
#         if not os.path.exists(file_path):
#             return jsonify({'message': 'Client data not found for the given client_id.'}), 404

#         with open(file_path, 'r') as f:
#             client_data = json.load(f)

#         return jsonify({
#             'message': 'Client data retrieved successfully.',
#             'data': client_data
#         }), 200

#     except Exception as e:
#         return jsonify({'message': f"An error occurred: {e}"}), 500


# import os
# import json
# from flask import Flask, jsonify
# import requests  # This is for calling the show_order_list API
 
 
# @app.route('/get-all-client-data', methods=['GET'])
# def get_all_client_data():
#     try:
#         all_data = []
#         client_data_folder = './client_data/client_data'
 
#         # Iterate over all JSON files in the client data folder
#         for filename in os.listdir(client_data_folder):
#             if filename.endswith(".json"):
#                 file_path = os.path.join(client_data_folder, filename)
               
#                 # Load client data
#                 with open(file_path, 'r') as f:
#                     client_data = json.load(f)
#                     client_id = client_data.get("uniqueId")
 
#                     # Default value for isNewClient
#                     client_data["isNewClient"] = True
 
#                     # Check if the client has orders by calling /show_order_list API
#                     if client_id:
#                         # Make a request to /show_order_list API to check if orders exist
#                         order_url = f'http://localhost:5000/show_order_list'  # Adjust to your actual endpoint
#                         response = requests.post(order_url, json={'client_id': client_id})
 
#                         if response.status_code == 200:
#                             orders = response.json().get("transaction_data", [])
#                             # If orders exist, set isNewClient to False
#                             if len(orders) > 0:
#                                 client_data["isNewClient"] = False
#                         else:
#                             # If no orders, set isNewClient to True
#                             client_data["isNewClient"] = True
 
#                     # Append the client data to the result list
#                     all_data.append(client_data)
 
#         # Handle case when no client data is found
#         if not all_data:
#             return jsonify({'message': 'No client data found in local storage.'}), 404
 
#         # Return all client data
#         return jsonify({
#             'message': 'All client data retrieved successfully.',
#             'data': all_data
#         }), 200
 
#     except Exception as e:
#         return jsonify({'message': f"An error occurred while retrieving data: {e}"}), 500
    
# investment assessment using Local Storage :

import os
import json

# LOCAL_STORAGE_DIR = "local_storage"
CLIENT_DATA_DIR = './client_data/'
CLIENT_SUMMARY_DIR = os.path.join(CLIENT_DATA_DIR, "client_data")
PERSONALITY_ASSESSMENT_DIR = "client_data/personality_assessments"

# Ensure directories exist
os.makedirs(CLIENT_SUMMARY_DIR, exist_ok=True)
os.makedirs(PERSONALITY_ASSESSMENT_DIR, exist_ok=True)



# Personality Assessment using AWS and  Local Storage :


@app.route('/get-personality-assessment', methods=['POST'])
def get_personality_assessment():
    try:
        # Parse incoming request data
        payload = request.json
 
        # Validate the payload
        client_id = payload.get('client_id')
        if not client_id:
            return jsonify({'message': 'client_id is required in the payload.'}), 400
 
        if USE_AWS:
            # Define folder path for S3
            folder_path = f"{personality_assessment_folder}"
            logging.info(f"Looking for files in folder: {folder_path}")
 
            # List objects in the folder
            response = s3.list_objects_v2(Bucket=S3_BUCKET_NAME, Prefix=folder_path)
            logging.debug(f"S3 list_objects_v2 response: {response}")
 
            # Check if the folder contains any objects
            if 'Contents' not in response:
                logging.warning(f"No files found in folder: {folder_path}")
                return jsonify({'message': 'No data found in the specified folder.'}), 404
 
            # Iterate through the files to find the matching client_id
            for obj in response['Contents']:
                file_key = obj['Key']
 
                # Skip the folder itself and non-JSON files
                if file_key == folder_path or not file_key.endswith('.json'):
                    continue
 
                # Fetch file content if the file matches the client_id
                if f"{client_id}.json" in file_key:
                    try:
                        file_response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=file_key)
                        file_content = json.loads(file_response['Body'].read().decode('utf-8'))
                        logging.info(f"Found and retrieved data for client_id {client_id}.")
 
                        return jsonify({
                            'message': 'Data fetched successfully.',
                            'data': file_content  # Ensure the actual client data is nested in 'data'
                        }), 200
                    except Exception as fetch_error:
                        logging.error(f"Error retrieving file {file_key}: {fetch_error}")
                        return jsonify({'message': 'Error retrieving client data from S3.'}), 500
 
            logging.warning(f"No matching data found for client_id {client_id} in folder {folder_path}.")
            return jsonify({'message': 'No data found for the provided client_id.'}), 404
 
        else:
            # Local Storage Logic
            file_path = os.path.join(CLIENT_DATA_DIR, f"personality_assessments/{client_id}.json")
            if not os.path.exists(file_path):
                return jsonify({'message': 'No data found for the provided client_id.'}), 404
 
            with open(file_path, 'r') as f:
                file_content = json.load(f)
 
            return jsonify({
                'message': 'Data fetched successfully.',
                'data': file_content
            }), 200
 
    except Exception as e:
        return jsonify({'message': f"Internal Server Error: {str(e)}"}), 500
 
 
 
# @app.route('/get-personality-assessment', methods=['POST'])
# def get_client_data_by_id():
#     try:
#         # Parse incoming request data
#         payload = request.json

#         # Validate the payload
#         client_id = payload.get('client_id')
#         if not client_id:
#             return jsonify({'message': 'client_id is required in the payload.'}), 400

#         # Locate the client's assessment data
#         file_path = os.path.join(CLIENT_DATA_DIR, f"personality_assessments/{client_id}.json")
#         if not os.path.exists(file_path):
#             return jsonify({'message': 'No data found for the provided client_id.'}), 404

#         with open(file_path, 'r') as f:
#             file_content = json.load(f)

#         return jsonify({
#             'message': 'Data fetched successfully.',
#             'data': file_content
#         }), 200

#     except Exception as e:
#         return jsonify({'message': f'Internal Server Error: {str(e)}'}), 500
    


# # api for generating suggestions with client id using AWS :  

@app.route('/investor-personality-assessment', methods=['POST'])
def investor_personality_assessment():
    try:
        # Parse incoming request data
        data = request.json
        if not data:
            logging.error("No data received in the request.")
            return jsonify({'message': 'Invalid request: No data received.'}), 400
 
        client_id = data.get('client_id')
        assessment_data = data.get('assessment_data')
 
        if not client_id or not assessment_data:
            logging.error("Missing client_id or assessment_data.")
            return jsonify({'message': 'Client ID and assessment data are required.'}), 400
 
        logging.info(f"Received assessment data for client ID: {client_id}")
 
        # Determine the investment personality
        personality = asyncio.run(determine_investment_personality(assessment_data))
        logging.info(f"Determined personality for client ID {client_id}: {personality}")
 
        # Handle data storage based on USE_AWS flag
        if USE_AWS:
            return handle_aws_storage(client_id, assessment_data, personality)
        else:
            return handle_local_storage(client_id, assessment_data, personality)
 
    except Exception as e:
        logging.error(f"Unhandled exception: {e}")
        return jsonify({'message': 'Internal Server Error'}), 500
 
 
def handle_aws_storage(client_id, assessment_data, personality):
    try:
        # Define S3 key for client data
        s3_key = f"{client_summary_folder}client-data/{client_id}.json"
        existing_data = None
 
        # Check if client data already exists in S3
        try:
            response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
            existing_data = json.loads(response['Body'].read().decode('utf-8'))
            logging.info(f"Existing data found for client ID {client_id}: {existing_data}")
        except s3.exceptions.NoSuchKey:
            logging.warning(f"No existing client data found for client ID {client_id}. Creating a new entry.")
 
        # Update existing data or create new data
        if not existing_data:
            existing_data = {}
        existing_data['investment_personality'] = personality
 
        # Save updated data back to S3
        try:
            s3.put_object(
                Bucket=S3_BUCKET_NAME,
                Key=s3_key,
                Body=json.dumps(existing_data),
                ContentType='application/json'
            )
            logging.info(f"Client data successfully updated in S3 for client ID: {client_id}")
        except Exception as e:
            logging.error(f"Error occurred while saving updated client data to S3: {e}")
            return jsonify({'message': f'Error occurred while saving updated data to S3: {e}'}), 500
 
        # Define S3 key for personality assessment file
        file_key = f"{personality_assessment_folder}{client_id}.json"
        existing_file_data = None
 
        # Check if personality assessment file already exists in S3
        try:
            file_response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=file_key)
            file_data = file_response['Body'].read().decode('utf-8')
            existing_file_data = json.loads(file_data)
            logging.info(f"Existing personality assessment file data for client ID {client_id}: {existing_file_data}")
        except s3.exceptions.NoSuchKey:
            logging.info(f"No existing personality assessment file found for client ID {client_id}. Creating a new file.")
 
        # Update or create personality assessment data
        updated_data = {
            'client_id': client_id,
            'assessment_data': assessment_data,
            'investment_personality': personality
        }
 
        if existing_file_data:
            # Merge new data with existing file data
            existing_file_data.update(updated_data)
            updated_data = existing_file_data
            logging.info(f"Updated personality assessment data for client ID {client_id}: {updated_data}")
 
        # Save personality assessment data back to S3
        try:
            s3.put_object(
                Bucket=S3_BUCKET_NAME,
                Key=file_key,
                Body=json.dumps(updated_data),
                ContentType='application/json'
            )
            logging.info(f"Personality assessment data successfully saved to S3 for client ID: {client_id}")
        except Exception as e:
            logging.error(f"Error occurred while saving personality assessment data to S3: {e}")
            return jsonify({'message': f'Error occurred while saving personality assessment data to S3: {e}'}), 500
 
        # Return the updated data
        return jsonify(updated_data), 200
 
    except Exception as e:
        logging.error(f"Unhandled exception: {e}")
        return jsonify({'message': 'Internal Server Error'}), 500
 
def handle_local_storage(client_id, assessment_data, personality):
    try:
        personality_file_path = os.path.join(CLIENT_DATA_DIR, f"personality_assessments/{client_id}.json")
        client_data_dir = os.path.join(CLIENT_DATA_DIR, "client_data")
        client_file_path = os.path.join(client_data_dir, f"{client_id}.json")
        personality_data = {}
 
        if os.path.exists(personality_file_path):
            with open(personality_file_path, 'r') as f:
                personality_data = json.load(f)
 
        personality_data.update({
            "client_id": client_id,
            "assessment_data": assessment_data,
            "investment_personality": personality
        })
 
        with open(personality_file_path, 'w') as f:
            json.dump(personality_data, f, indent=4)
 
        if os.path.exists(client_file_path):
            with open(client_file_path, 'r') as f:
                client_data = json.load(f)
            client_data['investment_personality'] = personality
        else:
            client_data = {
                "client_id": client_id,
                "investment_personality": personality,
            }
 
        with open(client_file_path, 'w') as f:
            json.dump(client_data, f, indent=4)
 
        logging.info(f"Updated client data file for client ID {client_id}")
 
        return jsonify({
            'client_id': client_id,
            'investment_personality': personality
        }), 200
 
    except Exception as e:
        logging.error(f"Error processing investor assessment: {e}")
        return jsonify({'message': 'Internal Server Error'}), 500
 
 

# @app.route('/investor-personality-assessment', methods=['POST'])
# def investor_personality_assessment():
#     try:
#         # Parse incoming data
#         data = request.json
#         client_id = data.get('client_id')
#         assessment_data = data.get('assessment_data')
 
#         if not client_id or not assessment_data:
#             return jsonify({'message': 'Client ID and assessment data are required.'}), 400
 
#         logging.info(f"Received assessment data for client ID: {client_id}")
 
#         # Determine investment personality
#         personality = asyncio.run(determine_investment_personality(assessment_data))
#         logging.info(f"Determined personality for client ID {client_id}: {personality}")
 
#         # Save assessment data and personality in a dedicated file
#         personality_file_path = os.path.join(CLIENT_DATA_DIR, f"personality_assessments/{client_id}.json")
#         client_data_dir = os.path.join(CLIENT_DATA_DIR, "client_data")
#         client_file_path = os.path.join(client_data_dir, f"{client_id}.json")
 
#         # Update or create personality-specific data
#         personality_data = {}
#         if os.path.exists(personality_file_path):
#             with open(personality_file_path, 'r') as f:
#                 personality_data = json.load(f)
 
#         personality_data.update({
#             "client_id": client_id,
#             "assessment_data": assessment_data,
#             "investment_personality": personality
#         })
 
#         with open(personality_file_path, 'w') as f:
#             json.dump(personality_data, f, indent=4)
 
#         # Update the main client data file
#         if os.path.exists(client_file_path):
#             with open(client_file_path, 'r') as f:
#                 client_data = json.load(f)
#             # Update investment personality in the existing client file
#             client_data['investment_personality'] = personality
#         else:
#             # If client file does not exist, create it
#             client_data = {
#                 "client_id": client_id,
#                 "investment_personality": personality,
#             }
 
#         with open(client_file_path, 'w') as f:
#             json.dump(client_data, f, indent=4)
 
#         logging.info(f"Updated client data file for client ID {client_id}")
 
#         return jsonify({
#             'client_id': client_id,
#             'investment_personality': personality
#         }), 200
 
#     except Exception as e:
#         logging.error(f"Error processing investor assessment: {e}")
#         return jsonify({'message': 'Internal Server Error'}), 500
 

##############################################################################################################
 
# @app.route('/personality-assessment', methods=['POST'])
# def personality_selected():
#     try:
#         # Parse incoming data
#         data = request.json
#         if not data:
#             return jsonify({'message': 'Invalid or missing request payload'}), 400

#         investment_personality = data.get('investmentPersonality')
#         client_name = data.get('clientName')
#         client_id = data.get('clientId')

#         print(f"Client Name: {client_name}, Investment Personality: {investment_personality}")

#         # Validate required data
#         if not client_id or not client_name or not investment_personality:
#             return jsonify({'message': 'Missing client_id, clientName, or investmentPersonality.'}), 400

#         # Load client data from local storage
#         file_path = os.path.join(CLIENT_DATA_DIR, f"{client_id}.json")
#         if not os.path.exists(file_path):
#             return jsonify({'message': 'Client data not found for the given client_id.'}), 404

#         with open(file_path, 'r') as f:
#             client_data = json.load(f)

#         print(f"Loaded Client Data: {client_data}")

#         # Generate suggestions
#         try:
#             result, pie_chart_data, bar_chart_data, combined_chart_data = asyncio.run(
#                 make_suggestions_using_clientid(
#                     investment_personality,
#                     client_name,
#                     client_data
#                 )
#             )

#             html_suggestions = markdown.markdown(result)
#             format_suggestions = markdown_to_text(html_suggestions)

#             return jsonify({
#                 "status": 200,
#                 "message": "Success",
#                 "investmentSuggestions": format_suggestions,
#                 "pieChartData": pie_chart_data,
#                 "barChartData": bar_chart_data,
#                 "compoundedChartData": combined_chart_data
#             }), 200

#         except Exception as e:
#             logging.error(f"Error generating suggestions: {e}")
#             return jsonify({'message': f"Error generating suggestions: {e}"}), 500

#     except Exception as e:
#         logging.error(f"Unhandled exception: {e}")
#         return jsonify({'message': 'Internal Server Error'}), 500


###########################################################################################################


# generate pie chart data and bar chart data :
def generate_chart_data(data):
    # Pie Chart Data
    labels = list(data['Growth-Oriented Investments'].keys()) + list(data['Conservative Investments'].keys())
    max_allocations = [
        int(data['Growth-Oriented Investments'][label]['max']) for label in data['Growth-Oriented Investments']
    ] + [
        int(data['Conservative Investments'][label]['max']) for label in data['Conservative Investments']
    ]
    num_labels = len(labels)
    dynamic_colors = generate_colors(num_labels)
    pie_chart_data = {
        'labels': labels,
        'datasets': [{
            'label': 'Investment Allocation',
            'data': max_allocations,
            'backgroundColor': dynamic_colors, #['#FF6384', '#36A2EB', '#FFCE56', '#4BC0C0'],  # Example colors
            'hoverOffset': 4
        }]
    }

    # Bar Chart Data
    min_allocations = [
        int(data['Growth-Oriented Investments'][label]['min']) for label in data['Growth-Oriented Investments']
    ] + [
        int(data['Conservative Investments'][label]['min']) for label in data['Conservative Investments']
    ]
    bar_chart_data = {
        'labels': labels,
        'datasets': [
            {
                'label': 'Allocation for Min returns',
                'data': min_allocations,
                'backgroundColor': 'skyblue'
            },
            {
                'label': 'Allocation for Max returns',
                'data': max_allocations,
                'backgroundColor': 'lightgreen'
            }
        ]
    }

    return pie_chart_data, bar_chart_data

def generate_chart_data_for_assets():
    # Investment Categories and Allocations
    investment_data = {
        "Growth-Oriented Investments": {
            "Stocks": {"min": 30, "max": 40},
            "ETFs": {"min": 15, "max": 20},
            "Growth Stock Mutual Funds": {"min": 10, "max": 15},
            "Real Estate Investment Trusts (REITs)": {"min": 10, "max": 15},
            "Emerging Markets Stocks": {"min": 5, "max": 10},
        },
        "Conservative Investments": {
            "High-Yield Savings Account": {"min": 5, "max": 10},
            "Treasury Bonds": {"min": 5, "max": 10},
        }
    }

    # Generate Pie Chart Data
    labels = list(investment_data['Growth-Oriented Investments'].keys()) + \
             list(investment_data['Conservative Investments'].keys())
    max_allocations = [
        investment_data['Growth-Oriented Investments'][label]['max'] for label in investment_data['Growth-Oriented Investments']
    ] + [
        investment_data['Conservative Investments'][label]['max'] for label in investment_data['Conservative Investments']
    ]
    min_allocations = [
        investment_data['Growth-Oriented Investments'][label]['min'] for label in investment_data['Growth-Oriented Investments']
    ] + [
        investment_data['Conservative Investments'][label]['min'] for label in investment_data['Conservative Investments']
    ]

    num_labels = len(labels)
    dynamic_colors = generate_colors(num_labels)
    dynamic_pie_chart_colors = generate_colors(len(labels))

    # Pie Chart Data
    pie_chart_data = {
        'labels': labels,
        'datasets': [{
            'label': 'Investment Allocation',
            'data': max_allocations,
            'backgroundColor': dynamic_pie_chart_colors,
            'hoverOffset': 4
        }]
    }

    # Bar Chart Data
    bar_chart_data = {
        'labels': labels,
        'datasets': [
            {
                'label': 'Allocation for Min Returns (%)',
                'data': min_allocations,
                'backgroundColor': 'skyblue'
            },
            {
                'label': 'Allocation for Max Returns (%)',
                'data': max_allocations,
                'backgroundColor': 'lightgreen'
            }
        ]
    }

    return pie_chart_data, bar_chart_data


def generate_colors(num_colors):
    # Generate distinct, vibrant colors in hex format
    return ["#" + ''.join(random.choices('0123456789ABCDEF', k=6)) for _ in range(num_colors)]


def generate_pie_chart_data(data):
    labels = list(data['Growth-Oriented Investments'].keys()) + list(data['Conservative Investments'].keys())
    max_allocations = [
        data['Growth-Oriented Investments'][label]['max'] for label in data['Growth-Oriented Investments']
    ] + [
        data['Conservative Investments'][label]['max'] for label in data['Conservative Investments']
    ]

    # Generate dynamic colors for the pie chart
    dynamic_colors = generate_colors(len(labels))
    
    pie_chart_data = {
        'labels': labels,
        'datasets': [{
            'label': 'Investment Allocation',
            'data': max_allocations,
            'backgroundColor': dynamic_colors,
            'hoverOffset': 4
        }]
    }

    return pie_chart_data


#new retrieval_chain code :
# async def generate_prompt_template(retriever,investmentPersonality,clientName,client_data):
#     try:
#         # global investment_personality #,summary
        
#         print(f"{investmentPersonality}\n {clientName}\n {client_data}")
        
        
#         llm = ChatGoogleGenerativeAI(
#             #model="gemini-pro",
#             model = "gemini-1.5-flash",
#             temperature = 0.45,
#             # temperature=0.7,
#             top_p=0.85,
#             google_api_key=GOOGLE_API_KEY
#         )
#         # New Template 
#         investmentPersonality = str(investmentPersonality)
#         print(investmentPersonality)
#         # clientName = str(clientName)
#         print(clientName)
#         context = str(clientName)
        
        
#         # New Prompt Template :
        
#         prompt_template = """ 
#                                 You are a Financial Advisor tasked with creating responsible investment suggestions for a client based on their investment personality : """ + investmentPersonality +   "\n" + """ so that the client can reach their Financial Goals, based on their Financial Conditions.
#                                 Use the following instructions to ensure consistent output:
#                                 ---

#                                 ### Required Output Format:
                                
#                                 #### Client Financial Details:
#                                 - **Client Name**: """ + clientName + f"""
#                                 - **Assets**:
#                                 - List all asset types, their current values, and annual contributions in a tabular format (columns: "Asset Type", "Current Value", "Annual Contribution").
#                                 - **Liabilities**:
#                                 - List all liability types, their balances, interest rates, and monthly payments in a tabular format (columns: "Liability Type", "Balance", "Interest Rate", "Monthly Payment").
#                                 - **Other Details**:
#                                 - Retirement plan details, income sources, and goals should be listed in a clear and concise format.
#                                 - Client's Financial Condition : Analyze the Details and mention the Client's Financial Condition as : Stable/ Currently Stable / Unstable.
#                                 - **Investment Period** `Z years`
                                
#                                 #### Investment Allocation:
#                                 Split investments into **Growth-Oriented Investments** and **Conservative Investments**. Ensure each category includes:
#                                 - **Investment Type**: Specify the investment type (e.g., "Index Funds", "US Treasury Bonds").
#                                 - **Allocation Range**: Specify minimum and maximum allocation percentages (e.g., `10% - 20%`).
#                                 - **Target**: Describe the purpose of the investment.
#                                 - **How to Invest**: Provide instructions on how to invest in this asset.
#                                 - **Where to Invest**: Specify platforms or tools for making the investment.

#                                 **Example**:
#                                 **Growth-Oriented Investments (Minimum X% - Maximum Y%) **:
#                                 - **Stocks**: `20% - 30%`
#                                 - **ETFs**: `10% - 15%`
#                                 - **Mutual Funds**: `10% - 20%`
#                                 - **Cryptocurrency**: ` 5% - 10%`
#                                 - **Real Estates or REITS**: `10% - 20%`
#                                 - *Target*: Long-term growth potential aligned with the overall market performance tailored to fullfil Client's Financial Goals and manage his Financial Condition.
#                                 - *How to Invest*: Provide information on how to invest in which market 
#                                 - *Where to Invest*: Provide Information to buy which assets and how much to invest in terms of amount and percentage(%).Mention 5-6 assets.
                                
#                                 **Conservative Investments (Minimum X% - Maximum Y%) **:
#                                 - **High-Yield Savings Account**: `30% - 40%`
#                                 - **Bonds**: `10% - 20%`
#                                 - **Commodities**: `5% - 10%`
#                                 - **Cash**: `5% - 10%`
#                                 - *Target*: Maintain liquidity for emergencies.
#                                 - *How to Invest*: Provide information on how to invest.
#                                 - *Where to Invest*: Mention where to invest and how much to allocate in terms of money and percentage(%). Mention 5-6 assets.

#                                 #### Returns Overview:
#                                 - **Minimum Expected Annual Return**: `X% - Y%`
#                                 - **Maximum Expected Annual Return**: `X% - Y%`
#                                 - **Minimum Expected Growth in Dollars**: `$X - $Y` (based on the time horizon)
#                                 - **Maximum Expected Growth in Dollars**: `$X - $Y` (based on the time horizon)
#                                 - **Time Horizon**: `Z years`

#                                 ---

#                                 ### Example Output:
                                
#                                 #### Client Financial Details:
#                                 | Asset Type          | Current Value ($) | Annual Contribution ($) |
#                                 |----------------------|-------------------|--------------------------|
#                                 | 401(k), 403(b), 457  | 300               | 15                       |
#                                 | Traditional IRA      | 200               | 15                       |
#                                 | Roth IRA             | 500               | 28                       |
#                                 | Cash/Bank Accounts   | 500,000           | 30,000                   |
#                                 | Real Estate          | 1,000,000         | -                        |
#                                 | Total Assets Value   | 1,501,000         | -                        |

#                                 | Liability Type      | Balance ($) | Interest Rate (%) | Monthly Payment ($) |
#                                 |---------------------|-------------|--------------------|----------------------|
#                                 | Mortgage            | 1,000       | 10                | 100                  |
#                                 | Credit Card         | 400         | 8                 | 400                  |
#                                 | Other Loans         | 500         | 6                 | 100                  |
#                                 | Total Liabilities   | 1,900       | -                 | -                    |
                                
#                                 | Investrment Period | 3 years |
                                
#                                 **Growth-Oriented Investments (Minimum 40% - Maximum 80%)**:
#                                 - **Stocks**: `20% - 30%`
#                                 - **ETFs**: `5% - 10%`
#                                 - **Mutual Funds**: `5% - 20%`
#                                 - **Cryptocurrency**: ` 5% - 10%`
#                                 - **Real Estates or REITS**: `5% - 10%`
#                                 - *Target*: Long-term growth potential aligned with the market.
#                                 - *How to Invest*: Purchase low-cost index funds.
#                                 - *Where to Invest*: Stocks such as NVIDIA,AAPL, Vanguard, LiteCoin.

#                                 **Conservative Investments (Minimum 40% - Maximum 70%)**:
#                                 - **High-Yield Savings Account**: `20% - 30%`
#                                 - **Bonds**: `10% - 20%`
#                                 - **Commodities**: `5% - 10%`
#                                 - **Cash**: `5% - 10%`
#                                 - *Target*: Maintain liquidity for emergencies.
#                                 - *How to Invest*: Deposit funds into an FDIC-insured account.
#                                 - *Where to Invest*: Ally Bank, Capital One 360.

#                                 #### Returns Overview:
#                                 - **Minimum Expected Annual Return**: `4% - 6%`
#                                 - **Maximum Expected Annual Return**: `8% - 15%`
#                                 - **Minimum Expected Growth in Dollars**: `$4,000 - $6,000`
#                                 - **Maximum Expected Growth in Dollars**: `$8,000 - $15,000`
#                                 - **Time Horizon**: `3 years`

#                                 ---

#                                 Ensure the output strictly follows this structure.


#                             ### Rationale for Investment Suggestions:
#                             Provide a detailed explanation of why these suggestions align with the client’s financial personality and goals.

#                             ---
#                             <context>
#                             {context}
#                             </context>
#                             Question: {input}

#         """

#         print(f"Investment Personality :{investmentPersonality}")
        
                

#         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

#         document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
#         combine_docs_chain = None  

#         if retriever is not None :  
#             retriever_chain = create_retrieval_chain(retriever,document_chain) 
#             # print(retriever_chain)
#             return retriever_chain
#         else:
#             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
#             return None

#     except Exception as e:
#         print(f"Error in creating chain: {e}")
#         return None

########################################################################################################



# Create Vector DB for JSON Data from cloud :
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

async def load_vector_db_from_json(json_data):
    try:
        print("Loading vector database from JSON data...")
        
        # Step 1: Convert JSON to a list of Documents
        documents = []
        for key, value in json_data.items():
            if isinstance(value, dict):
                nested_text = "\n".join([f"{nested_key}: {nested_value}" for nested_key, nested_value in value.items()])
                documents.append(Document(page_content=f"{key}:\n{nested_text}"))
            elif isinstance(value, list):
                list_text = "\n".join([str(item) for item in value])
                documents.append(Document(page_content=f"{key}:\n{list_text}"))
            else:
                documents.append(Document(page_content=f"{key}: {value}"))

        print(f"Prepared {len(documents)} documents for FAISS.")

        # Step 2: Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        text_chunks = text_splitter.split_documents(documents)

        # Step 3: Embed and load into FAISS
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        vector_store = FAISS.from_documents(text_chunks, embeddings)

        print("Vector database loaded successfully.")
        return vector_store.as_retriever(search_kwargs={"k": 3})  # Top-3 results
    except Exception as e:
        print(f"Error loading vector database: {e}")
        return None




from langchain.prompts.chat import ChatPromptTemplate
from langchain.chains import StuffDocumentsChain, create_retrieval_chain
from langchain.schema.runnable import RunnableConfig, RunnableSequence
# from langchain.chains import LLMChain
# from langchain.schema.runnable import RunnableMap
# from langchain.schema.runnable import RunnableSequence

async def generate_prompt_with_retriever(retriever, investmentPersonality, clientName):
    try:
        llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            temperature=0.45,
            top_p=0.85,
            google_api_key=GOOGLE_API_KEY
        )
        clientName = clientName
        # Define the prompt
        
        # Previous Version :
        
        prompt_template = """
                            You are a Financial Advisor tasked with creating responsible and detailed investment suggestions for a client based on their investment personality: """ + investmentPersonality + """
                            so that the client can reach their financial goals while considering their financial conditions and market trends. 
                            Use the following instructions to ensure consistent and comprehensive output:
                            ---

                            ### Required Output Format:

                            #### **Investment Allocation:**
                            Split investments into **Growth-Oriented Investments** and **Conservative Investments**. Ensure each category includes:
                            - **Investment Type:** Specify the type of investment (e.g., "Index Funds", "US Treasury Bonds").
                            - **Allocation Range:** Specify minimum and maximum allocation percentages (e.g., `10% - 20%`).
                            - **Target:** Describe the purpose and goal of the investment.
                            - **How to Invest:** Provide practical instructions for investing in this asset.
                            - **Where to Invest:** List platforms or tools for making the investment, and provide **6-7 specific recommendations per asset class** for diversity and balance.

                            #### Example Output:
                            **Growth-Oriented Investments (Minimum X% - Maximum Y%)**:
                            - **Stocks:** `20% - 30%`
                            - *Target:* High growth potential with sector diversification.
                            - *How to Invest:* Direct purchases via a brokerage account; consider market trends.
                            - *Where to Invest:* Examples include large-cap stocks like Apple (AAPL), Tesla (TSLA), emerging market stocks, and small-cap stocks.
                            - **ETFs:** `10% - 15%`
                            - *Target:* Diversified exposure to specific sectors.
                            - *How to Invest:* Invest in sector-specific or international ETFs through platforms like Fidelity or Vanguard.
                            - *Where to Invest:* SPDR S&P 500 ETF (SPY), iShares Core MSCI Emerging Markets ETF (IEMG).

                            #### Returns Overview:
                            - **Minimum Expected Annual Return:** `X% - Y%`
                            - **Maximum Expected Annual Return:** `X% - Y%`
                            - **Minimum Expected Growth in Dollars:** `$X - $Y` (based on the time horizon)
                            - **Maximum Expected Growth in Dollars:** `$X - $Y` (based on the time horizon)
                            - **Time Horizon:** `Z years`

                            ---

                            ### Example Output Format:
                            **Investment Allocation:**

                            **Growth-Oriented Investments (Minimum 70% - Maximum 90%)**:
                            - **Stocks:** `20% - 30%`
                            - *Target:* Capital appreciation from high-growth stocks.
                            - *How to Invest:* Purchase through brokerage accounts like Fidelity.
                            - *Where to Invest:* Tesla (TSLA), NVIDIA (NVDA), Apple (AAPL).

                            **Conservative Investments (Minimum 10% - Maximum 30%)**:
                            - **High-Yield Savings Account:** `5% - 10%`
                            - *Target:* Maintain liquidity for emergencies.
                            - *How to Invest:* Deposit funds into FDIC-insured accounts.
                            - *Where to Invest:* Ally Bank, Marcus by Goldman Sachs.

                            #### Returns Overview:
                            - **Minimum Expected Annual Return:** `6% - 8%`
                            - **Maximum Expected Annual Return:** `12% - 18%`
                            - **Minimum Expected Growth in Dollars:** `$5,000 - $8,000`
                            - **Maximum Expected Growth in Dollars:** `$15,000 - $25,000`
                            - **Time Horizon:** `5 years`

                            ---

                            Ensure the output strictly follows this structure.

                            ### **Rationale for Investment Suggestions:**
                            Provide a detailed explanation of why these suggestions align with the client’s financial personality, goals, and market trends.

                            ---

                            <context>
                            {context}
                            </context>
                            Question: {input}
                            """
        
        print("Prompt Created Successfully")
                

        llm_prompt = ChatPromptTemplate.from_template(prompt_template)

        document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
        combine_docs_chain = None  

        if retriever is not None :  
            retriever_chain = create_retrieval_chain(retriever,document_chain) 
            print("\nRetrieval chain created successfully\n")
            print(retriever_chain)
            return retriever_chain
        else:
            print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
            return None

    except Exception as e:
        print(f"Error in generating prompt or retrieval chain: {e}")
        return None

# V -1 : with client table 

# async def generate_prompt_with_retriever(retriever, investmentPersonality, clientName):
#     try:
#         llm = ChatGoogleGenerativeAI(
#             model="gemini-1.5-flash",
#             temperature=0.45,
#             top_p=0.85,
#             google_api_key=GOOGLE_API_KEY
#         )
#         clientName = clientName
#         # Define the prompt
#         prompt_template = """ 
#                                 You are a Financial Advisor tasked with creating responsible investment suggestions for a client based on their investment personality : """ + investmentPersonality +   "\n" + """ so that the client can reach their Financial Goals, based on their Financial Conditions.
#                                 Use the following instructions to ensure consistent output:
#                                 ---

#                                 ### Required Output Format:
                                
#                                 #### Client Financial Details:
#                                 - **Client Name**: """ + clientName + """
#                                 - **Assets**:
#                                 - List all asset types, their current values, and annual contributions in a tabular format (columns: "Asset Type", "Current Value", "Annual Contribution").
#                                 - **Liabilities**:
#                                 - List all liability types, their balances, interest rates, and monthly payments in a tabular format (columns: "Liability Type", "Balance", "Interest Rate", "Monthly Payment").
#                                 - **Other Details**:
#                                 - Retirement plan details, income sources, and goals should be listed in a clear and concise format.
#                                 - Client's Financial Condition : Analyze the Details and mention the Client's Financial Condition as : Stable/ Currently Stable / Unstable.
#                                 - **Investment Period** `Z years`
                                
#                                 #### Investment Allocation:
#                                 Split investments into **Growth-Oriented Investments** and **Conservative Investments**. Ensure each category includes:
#                                 - **Investment Type**: Specify the investment type (e.g., "Index Funds", "US Treasury Bonds").
#                                 - **Allocation Range**: Specify minimum and maximum allocation percentages (e.g., `10% - 20%`).
#                                 - **Target**: Describe the purpose of the investment.
#                                 - **How to Invest**: Provide instructions on how to invest in this asset.
#                                 - **Where to Invest**: Specify platforms or tools for making the investment.

#                                 **Example**:
#                                 **Growth-Oriented Investments (Minimum X% - Maximum Y%) **:
#                                 - **Stocks**: `20% - 30%`
#                                 - **ETFs**: `10% - 15%`
#                                 - **Mutual Funds**: `10% - 20%`
#                                 - **Cryptocurrency**: ` 5% - 10%`
#                                 - **Real Estates or REITS**: `10% - 20%`
#                                 - *Target*: Long-term growth potential aligned with the overall market performance tailored to fullfil Client's Financial Goals and manage his Financial Condition.
#                                 - *How to Invest*: Provide information on how to invest in which market 
#                                 - *Where to Invest*: Provide Information to buy which assets and how much to invest in terms of amount and percentage(%).Mention 5-6 assets.
                                
#                                 **Conservative Investments (Minimum X% - Maximum Y%) **:
#                                 - **High-Yield Savings Account**: `30% - 40%`
#                                 - **Bonds**: `10% - 20%`
#                                 - **Commodities**: `5% - 10%`
#                                 - **Cash**: `5% - 10%`
#                                 - *Target*: Maintain liquidity for emergencies.
#                                 - *How to Invest*: Provide information on how to invest.
#                                 - *Where to Invest*: Mention where to invest and how much to allocate in terms of money and percentage(%). Mention 5-6 assets.

#                                 #### Returns Overview:
#                                 - **Minimum Expected Annual Return**: `X% - Y%`
#                                 - **Maximum Expected Annual Return**: `X% - Y%`
#                                 - **Minimum Expected Growth in Dollars**: `$X - $Y` (based on the time horizon)
#                                 - **Maximum Expected Growth in Dollars**: `$X - $Y` (based on the time horizon)
#                                 - **Time Horizon**: `Z years`

#                                 ---

#                                 ### Example Output:
                                
#                                 #### Client Financial Details:
#                                 | Asset Type          | Current Value ($) | Annual Contribution ($) |
#                                 |----------------------|-------------------|--------------------------|
#                                 | 401(k), 403(b), 457  | 300               | 15                       |
#                                 | Traditional IRA      | 200               | 15                       |
#                                 | Roth IRA             | 500               | 28                       |
#                                 | Cash/Bank Accounts   | 500,000           | 30,000                   |
#                                 | Real Estate          | 1,000,000         | -                        |
#                                 | Total Assets Value   | 1,501,000         | -                        |

#                                 | Liability Type      | Balance ($) | Interest Rate (%) | Monthly Payment ($) |
#                                 |---------------------|-------------|--------------------|----------------------|
#                                 | Mortgage            | 1,000       | 10                | 100                  |
#                                 | Credit Card         | 400         | 8                 | 400                  |
#                                 | Other Loans         | 500         | 6                 | 100                  |
#                                 | Total Liabilities   | 1,900       | -                 | -                    |
                                
#                                 | Investrment Period | 3 years |
                                
#                                 **Growth-Oriented Investments (Minimum 40% - Maximum 80%)**:
#                                 - **Stocks**: `20% - 30%`
#                                 - **ETFs**: `5% - 10%`
#                                 - **Mutual Funds**: `5% - 20%`
#                                 - **Cryptocurrency**: ` 5% - 10%`
#                                 - **Real Estates or REITS**: `5% - 10%`
#                                 - *Target*: Long-term growth potential aligned with the market.
#                                 - *How to Invest*: Purchase low-cost index funds.
#                                 - *Where to Invest*: Stocks such as NVIDIA,AAPL, Vanguard, LiteCoin.

#                                 **Conservative Investments (Minimum 40% - Maximum 70%)**:
#                                 - **High-Yield Savings Account**: `20% - 30%`
#                                 - **Bonds**: `10% - 20%`
#                                 - **Commodities**: `5% - 10%`
#                                 - **Cash**: `5% - 10%`
#                                 - *Target*: Maintain liquidity for emergencies.
#                                 - *How to Invest*: Deposit funds into an FDIC-insured account.
#                                 - *Where to Invest*: Ally Bank, Capital One 360.

#                                 #### Returns Overview:
#                                 - **Minimum Expected Annual Return**: `4% - 6%`
#                                 - **Maximum Expected Annual Return**: `8% - 15%`
#                                 - **Minimum Expected Growth in Dollars**: `$4,000 - $6,000`
#                                 - **Maximum Expected Growth in Dollars**: `$8,000 - $15,000`
#                                 - **Time Horizon**: `3 years`

#                                 ---

#                                 Ensure the output strictly follows this structure.


#                             ### Rationale for Investment Suggestions:
#                             Provide a detailed explanation of why these suggestions align with the client’s financial personality and goals.

#                             ---
#                             <context>
#                             {context}
#                             </context>
#                             Question: {input}

#         """
                
        
#         print("Prompt Created Successfully")
                

#         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

#         document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
#         combine_docs_chain = None  

#         if retriever is not None :  
#             retriever_chain = create_retrieval_chain(retriever,document_chain) 
#             print("\nRetrieval chain created successfully\n")
#             print(retriever_chain)
#             return retriever_chain
#         else:
#             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
#             return None

#     except Exception as e:
#         print(f"Error in generating prompt or retrieval chain: {e}")
#         return None

# # using aws :

def generate_pie_chart_data(labels, data_values):
    num_labels = len(labels)
    dynamic_colors = generate_colors(num_labels)

    return {
        'labels': labels,
        'datasets': [{
            'label': 'Investment Allocation',
            'data': data_values,
            'backgroundColor': dynamic_colors,
            'hoverOffset': 4
        }]
    }




async def make_suggestions_using_clientid(investmentPersonality, clientName, client_data):
    try:
        print(f"Processing client data for {clientName}.")
        
        # Load vector database
        retriever = await load_vector_db_from_json(client_data)
        if not retriever:
            raise Exception("Failed to load vector database.")

        print(f"Created Retriever : {retriever}")
        # Generate retriever-based prompt
        retrieval_chain = await generate_prompt_with_retriever(retriever, investmentPersonality, clientName)
        if not retrieval_chain:
            raise Exception("Failed to create retrieval chain.")

        # Use the chain to generate a response
        # query = f"""Generate financial suggestions for the client {clientName} based on their investment personality: {investmentPersonality} 
        #         tailored to their Financial Goals and Considering their Financial Situations. Suggest 6-7 assets per category with 6-7 examples per asset."""
        
        query = f"""
                Generate personalized and responsible financial investment suggestions for the client {clientName} based on their investment personality: {investmentPersonality}.
                With the given Client Data : {client_data} Ensure suggestions align with the client's financial goals, current market trends, and financial conditions.
                Provide a **diverse range of 6-7 assets per category** for each investment class (Growth-Oriented and Conservative) and give **6-7 specific examples per asset** with actionable guidance.
                Include:
                1. Allocation ranges, target purpose, and how-to-invest instructions for each asset.
                2. Market insights and rationale for asset selection based on current trends.
                3. A return overview for a realistic time horizon.
                4. Suggestions for rebalancing strategies if market conditions change significantly.
                """

        
        # response = retrieval_chain.invoke(query)
        response = retrieval_chain.invoke({"input": query})
        answer = response['answer']
        print("Suggestions generated successfully.")
        
        # Extract Data from Response

        data_extracted = extract_numerical_data(answer)
        
        min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                        [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
        max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                        [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

        # Normalize allocations
        min_allocations = normalize_allocations(min_allocations)
        max_allocations = normalize_allocations(max_allocations)

        # bar_chart_data,pie_chart_data = generate_chart_data(data_extracted)
        bar_chart_data,_ = generate_chart_data_for_assets()
        
        # Get data for charts
        chart_data = get_chart_data_from_llm(answer)

        # Output the result
        print("CHART DATA FROM LLM :\n")
        print(json.dumps(chart_data, indent=4))
    
        print(f"Bar chart data: {bar_chart_data}")
        
        # pie_chart_data = generate_pie_chart_data(data_extracted)
        
        # Example data usage
        labels = [
            "Stocks", "ETFs", "Growth Stock Mutual Funds",
            "Real Estate Investment Trusts (REITs)", "Emerging Markets Stocks",
            "High-Yield Savings Account", "Treasury Bonds"
        ]
        data_values = [40, 20, 15, 15, 10, 10, 10]

        # Generate Pie Chart Data
        pie_chart_data = generate_pie_chart_data(labels, data_values)
        print("Pie Chart Data:", pie_chart_data)
        print(pie_chart_data)

        print(f"Pie chart data: {pie_chart_data}")

        
        # print(f"Pie Chart Data is : {pie_chart_data}")
        # Prepare the data for the line chart with inflation adjustment
        initial_investment = 10000
        combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
        print(f"\nThe combined chart data is: {combined_chart_data}")
        
        print(f"Suggestions : {answer}")
        
        return answer, pie_chart_data, bar_chart_data, combined_chart_data
            
    except Exception as e:
        print(f"Error generating suggestions: {e}")
        return jsonify({'message': f'Error occurred while generating suggestions: {e}'}), 500


def get_chart_data_from_llm(text):
    """
    Use an LLM to extract data for pie charts and bar graphs from the investment text.

    Args:
        text (str): The investment text input.

    Returns:
        dict: A dictionary containing pie chart data and bar chart data.
    """
    # LLM prompt to extract relevant data
    prompt = f"""
            You are a data analyst. Extract the following data from the provided investment text:

            1. A list for pie chart visualization, where each entry contains:
            - Label: The name of the asset (e.g., "Stocks", "Growth ETFs").
            - Value: The average percentage allocation for the asset.

            2. A list for bar chart visualization, where each entry contains:
            - Label: "Minimum Returns" and "Maximum Returns".
            - Value: The percentage return for the respective category.
            - Dollar Value: The growth in dollars for the respective category.

            Return the output as a JSON object with two keys:
            - "pie_chart_data": [list of pie chart entries].
            - "bar_chart_data": [list of bar chart entries].

            Here is the investment text:
            {text}
            """

    try:
        # Send the prompt to Gemini LLM
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(prompt)

        # Process the response from LLM
        html_suggestions = markdown.markdown(response.text)
        format_suggestions = markdown_to_text(html_suggestions)
        
        # response = gemini.generate_response(
        #     prompt=prompt,
        #     max_tokens=1000,  # Adjust token limit as per requirements
        # )

        # Parse and return the LLM response
        chart_data = json.loads(response.text)
        return chart_data

    except Exception as e:
        print(f"Error fetching data from Gemini LLM: {e}")
        return {"error": str(e)}





# # api for generating suggestions with client id :

@app.route('/personality-assessment', methods=['POST'])
def personality_selected():
    try:
        data = request.json
        try :
            investmentPersonality = data.get('investmentPersonality') # investment_personality
            clientName = data.get('clientName')
            print(f"The clients ClientName is : {clientName} ")
            print(f"InvestmentPersonality received is : {investmentPersonality}")
            logging.info('Recieved Values')
            
        except Exception as e:
            logging.info(f"Error occurred while retrieving client id: {e}")
            return jsonify({'message': f'Error occurred while retrieving client id: {e}'}), 400

        # Retrieve Client Financial Form Information :
        try:
            # Retrieve client_id from query parameters
            clientId = data.get('clientId')
            print(f"Received Client Id : {clientId}")
            # client_id = request.args.get('clientId')
            
            # Validate the client_id
            if not clientId:
                return jsonify({'message': 'client_id is required as a query parameter'}), 400

            # Define the S3 key for the object
            s3_key = f"{client_summary_folder}client-data/{clientId}.json"

            # Retrieve the object from S3
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
                # Decode and parse the JSON data
                client_data = json.loads(response['Body'].read().decode('utf-8'))
                print(f"Received Client Data :\n{client_data}")
                # return jsonify({
                #     'message': 'Client data retrieved successfully.',
                #     'data': client_data
                # }), 200
                
                result,pie_chart_data,bar_chart_data,combined_chart_data = asyncio.run(make_suggestions_using_clientid(investmentPersonality,
                                                                                                                   clientName,client_data))
                
                htmlSuggestions = markdown.markdown(result)
                logging.info(f"Suggestions for investor: \n{result}")
                
                formatSuggestions = markdown_to_text(htmlSuggestions)
                answer = markdown_table_to_html(formatSuggestions)
                print(answer)
                 
                # Return the Results :
                
                # return jsonify({
                #     "status": 200,
                #     "message": "Success",
                #     "investmentSuggestions": answer, #formatSuggestions,
                #     "pieChartData": pie_chart_data,
                #     "barChartData": bar_chart_data,
                #     "compoundedChartData":combined_chart_data
                # }), 200
                
                return jsonify({
                    "status": 200,
                    "message": "Success",
                    "clientdata":client_data,
                    "investmentSuggestions": formatSuggestions,
                    "pieChartData": pie_chart_data,
                    "barChartData": bar_chart_data,
                    "compoundedChartData":combined_chart_data
                }), 200
                
            except s3.exceptions.NoSuchKey:
                return jsonify({'message': 'Client data not found for the given client_id.'}), 404
            except Exception as e:
                return jsonify({'message': f"Error retrieving data: {e}"}), 500

        except Exception as e:
            return jsonify({'message': f"An error occurred: {e}"}), 500
    
    except Exception as e:
        print(f"An error occurred while requesting Data: {e}")
        return jsonify({'message': f"An error occurred while requesting Data :" + str(e)}, 500)
   

#########################################################################################################################

# Route to handle generating investment suggestions
import shutil
import os

def save_file_to_folder(file_storage, destination_folder):
    try:
        # Ensure the destination folder exists
        if not os.path.exists(destination_folder):
            os.makedirs(destination_folder)
        
        # Construct the destination file path
        destination_file_path = os.path.join(destination_folder, file_storage.filename)
        
        # Check if the file already exists
        if not os.path.exists(destination_file_path):
            # Save the file
            file_storage.save(destination_file_path)
            print(f"File saved to {destination_file_path}")
            return destination_file_path
        else:
            print(f"File already exists at {destination_file_path}")
            return destination_file_path
        
    except Exception as e:
        print(f"Error saving file: {e}")


# #Working for both the methods :
# generate_suggestions by taking files as i/p :
@app.route('/generate-investment-suggestions', methods=['POST'])
def generate_investment_suggestions():
    try:
        assessment_file = request.files['assessmentFile']
        financial_file = request.files['financialFile']
        logging.info("Requested files")
        
        responses = extract_responses_from_docx(assessment_file)
        if not responses:
            raise Exception("Failed to extract responses from assessment file.")
        
        destination_folder = 'data'
        file_path = save_file_to_folder(financial_file, destination_folder)
        if not file_path:
            raise Exception("Failed to save financial file.")
        
        financial_data = asyncio.run(process_document(file_path))
        if not financial_data:
            raise Exception("Failed to process financial file.")
        
        logging.info(f"Received Responses from the file {responses}")
        
        personality = asyncio.run(determine_investment_personality(responses))
        if not personality:
            raise Exception("Failed to determine personality.")
        
        logging.info(f"Personality of the user is: {personality}")
        
        clientName = "Rohit Sharma" #"Emilly Watts"
        suggestions = asyncio.run(generate_investment_suggestions_for_investor(personality, clientName, financial_data, file_path))
        if "Error" in suggestions:
            raise Exception(suggestions)
        
        htmlSuggestions = markdown.markdown(suggestions)
        logging.info(f"Suggestions for investor: \n{suggestions}")
        
        formatSuggestions = markdown_to_text(htmlSuggestions)
        answer = markdown_table_to_html(formatSuggestions)
        print(answer)
        
        # need to change the data extraction process : 
        data_extracted = extract_numerical_data(suggestions)
        
        min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                        [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
        max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                        [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

        # Normalize allocations
        min_allocations = normalize_allocations(min_allocations)
        max_allocations = normalize_allocations(max_allocations)

        bar_chart_data,pie_chart_data = generate_chart_data(data_extracted)
        
        # Sometimes Generating Pie Charts and Bar charts : 
        # data_extracted = extract_numerical_data(suggestions)

        # # Fixing pie and bar chart generation
        # growth_investments = data_extracted.get('Growth-Oriented Investments', {})
        # conservative_investments = data_extracted.get('Conservative Investments', {})

        # # Generate normalized allocations
        # min_allocations = [int(growth_investments[label]['min'].strip('%')) for label in growth_investments] + \
        #                 [int(conservative_investments[label]['min'].strip('%')) for label in conservative_investments]
        # max_allocations = [int(growth_investments[label]['max'].strip('%')) for label in growth_investments] + \
        #                 [int(conservative_investments[label]['max'].strip('%')) for label in conservative_investments]

        # # Normalize
        # min_allocations = normalize_allocations(min_allocations)
        # max_allocations = normalize_allocations(max_allocations)

        # # Bar Chart
        # bar_chart_data = {
        #     'labels': list(growth_investments.keys()) + list(conservative_investments.keys()),
        #     'datasets': [
        #         {'label': 'Allocation for Min returns', 'data': min_allocations, 'backgroundColor': 'skyblue'},
        #         {'label': 'Allocation for Max returns', 'data': max_allocations, 'backgroundColor': 'lightgreen'}
        #     ]
        # }

        # # Pie Chart
        # all_labels = list({**growth_investments, **conservative_investments}.keys())
        # num_labels = len(all_labels)
        # max_allocations_for_pie = normalize_allocations(
        #     [int(growth_investments.get(label, {}).get('max', '0').strip('%')) for label in growth_investments] +
        #     [int(conservative_investments.get(label, {}).get('max', '0').strip('%')) for label in conservative_investments]
        # )

        # # Normalize to 100% for pie chart
        # total = sum(max_allocations_for_pie)
        # max_allocations_for_pie = [(value / total) * 100 for value in max_allocations_for_pie]

        # dynamic_colors = generate_colors(num_labels)
        # pie_chart_data = {
        #     'labels': all_labels,
        #     'datasets': [{'label': 'Investment Allocation', 'data': max_allocations_for_pie, 'backgroundColor': dynamic_colors, 'hoverOffset': 4}]
        # }
    #############################################################################################
    
        print(f"Bar chart data: {bar_chart_data}")
        print(f"Pie chart data: {pie_chart_data}")

        # min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
        #                   [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
        # max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
        #                   [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

        # # Normalize allocations
        # min_allocations = normalize_allocations(min_allocations)
        # max_allocations = normalize_allocations(max_allocations)

        # # Update Bar Chart Data
        
        # bar_chart_data = {
        #     'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
        #     'datasets': [{
        #         'label': 'Allocation for Min returns',
        #         'data': min_allocations,
        #         'backgroundColor': 'skyblue'
        #     },
        #     {
        #         'label': 'Allocation for Max returns',
        #         'data': max_allocations,
        #         'backgroundColor': 'lightgreen'
        #     }]
        # }

        # # Similar changes can be made for the Pie Chart Data:
        # all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
        # num_labels = len(all_labels)
        # max_allocations_for_pie = normalize_allocations(
        #     [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
        #     [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']]
        # )
        
        # # Generate colors based on the number of labels
        # dynamic_colors = generate_colors(num_labels)

        # # Update Pie Chart Data
        # pie_chart_data = {
        #     'labels': all_labels,
        #     'datasets': [{
        #         'label': 'Investment Allocation',
        #         'data': max_allocations_for_pie,
        #         'backgroundColor': dynamic_colors,
        #         'hoverOffset': 4
        #     }]
        # }
        
        # print(f"Pie Chart Data is : {pie_chart_data}")
        # Prepare the data for the line chart with inflation adjustment
        initial_investment = 10000
        combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
        print(f"\nThe combined chart data is: {combined_chart_data}")
        
        return jsonify({
            "status": 200,
            "message": "Success",
            "investmentSuggestions":  answer, #formatSuggestions,
            "pieChartData": pie_chart_data,
            "barChartData": bar_chart_data,
            "compoundedChartData": combined_chart_data
        }), 200
        
        # return jsonify({
        #     "status": 200,
        #     "message": "Success",
        #     "investmentSuggestions": htmlSuggestions,
        #     "pieChartData": pie_chart_data,
        #     "barChartData": bar_chart_data,
        #     "compoundedChartData": combined_chart_data
        # }), 200

    except Exception as e:
        logging.info(f"Error in generating investment suggestions: {e}")
        return jsonify({'message': f'Internal Server Error in Generating responses : {e}'}), 500

################################################################-------------------- Stocks Analysis -------------------------------- #################################
# #Stock analysis code :

from flask import Flask, request, jsonify
import yfinance as yf
import pandas as pd
import requests
import os
import logging


NEWS_API_KEY = os.getenv('NEWS_API_KEY')

# Simulate memory using a file
CHAT_HISTORY_FILE = "chat_history.json"
CHAT_ID_TRACKER_FILE = "chat_id_tracker.json"  # File to track chat_id

# Helper to save chat history to a file
def save_chat_history(chat_id, history):
    if os.path.exists(CHAT_HISTORY_FILE):
        with open(CHAT_HISTORY_FILE, 'r') as f:
            chat_data = json.load(f)
    else:
        chat_data = {}

    chat_data[chat_id] = history

    with open(CHAT_HISTORY_FILE, 'w') as f:
        json.dump(chat_data, f, indent=4)

# Helper to load chat history from a file
def load_chat_history(chat_id):
    if os.path.exists(CHAT_HISTORY_FILE):
        with open(CHAT_HISTORY_FILE, 'r') as f:
            chat_data = json.load(f)
        return chat_data.get(str(chat_id), [])
    return []

# Helper to track chat_id and increment it
def get_next_chat_id():
    if os.path.exists(CHAT_ID_TRACKER_FILE):
        with open(CHAT_ID_TRACKER_FILE, 'r') as f:
            chat_id_data = json.load(f)
        chat_id = chat_id_data.get("chat_id", 1)
    else:
        chat_id = 1

    chat_id_data = {"chat_id": chat_id + 1}
    with open(CHAT_ID_TRACKER_FILE, 'w') as f:
        json.dump(chat_id_data, f, indent=4)

    return chat_id


# # Fetch Stock Data :
def get_stock_data(ticker):
    try:
        # Step 1: Fetch Stock Data :
        stock = yf.Ticker(ticker)
        
        data = {}

        company_details = stock.info.get('longBusinessSummary', 'No details available')
        data['Company_Details'] = company_details
        sector = stock.info.get('sector', 'No sector information available')
        data['Sector'] = sector
        prev_close = stock.info.get('previousClose', 'No previous close price available')
        data['Previous_Closing_Price'] = prev_close
        open_price = stock.info.get('open', 'No opening price available')
        data['Today_Opening_Price'] = open_price
         
        hist = stock.history(period="5d")
        if not hist.empty and 'Close' in hist.columns:
            if hist.index[-1].date() == yf.download(ticker, period="1d").index[-1].date():
                close_price = hist['Close'].iloc[-1]
                data['Todays_Closing_Price'] = close_price
            else:
                data['Todays_Closing_Price'] = "Market is open, no closing price available yet."
        else:
            data['Todays_Closing_Price'] = "No historical data available for closing price."

        day_high = stock.info.get('dayHigh', 'No high price available')
        data['Today_High_Price'] = day_high
        day_low = stock.info.get('dayLow', 'No low price available')
        data['Today_Low_Price'] = day_low
        volume = stock.info.get('volume', 'No volume information available')
        data['Today_Volume'] = volume
        dividends = stock.info.get('dividendRate', 'No dividend information available')
        data['Today_Dividends'] = dividends
        splits = stock.info.get('lastSplitFactor', 'No stock split information available')
        data['Today_Stock_Splits'] = splits
        pe_ratio = stock.info.get('trailingPE', 'No P/E ratio available')
        data['PE_Ratio'] = pe_ratio
        market_cap = stock.info.get('marketCap', 'No market cap available')
        data['Market_Cap'] = market_cap

        # Additional KPIs
        data['EPS'] = stock.info.get('trailingEps', 'No EPS information available')
        data['Book_Value'] = stock.info.get('bookValue', 'No book value available')
        data['ROE'] = stock.info.get('returnOnEquity', 'No ROE information available')
        data['ROCE'] = stock.info.get('returnOnAssets', 'No ROCE information available')  # ROCE is not available directly
        
        # Revenue Growth (CAGR) and Earnings Growth would need to be calculated based on historical data
        earnings_growth = stock.info.get('earningsGrowth', 'No earnings growth available')
        revenue_growth = stock.info.get('revenueGrowth', 'No revenue growth available')

        data['Earnings_Growth'] = earnings_growth
        data['Revenue_Growth'] = revenue_growth
        
        
        income_statement = stock.financials
        balance_sheet = stock.balance_sheet
        cashflow = stock.cashflow

        # Step 2: Get News Related to Stock
        try:
            # Fetch Stock News
            news_url = f'https://newsapi.org/v2/everything?q={ticker}&apiKey={NEWS_API_KEY}&pageSize=3'
            news_response = requests.get(news_url, timeout=10)

            if news_response.status_code == 200:
                news_data = news_response.json()
                articles = news_data.get('articles', [])
                if articles:
                    top_news = "\n\n".join([f"{i+1}. {article['title']} - {article['url']}" for i, article in enumerate(articles)])
                    data['Top_News'] = top_news
                else:
                    data['Top_News'] = "No news articles found."
            else:
                error_msg = news_response.json().get("message", "Unknown error occurred.")
                data['Top_News'] = f"Failed to fetch news articles. Error: {error_msg}"
        except requests.exceptions.RequestException as e:
            logging.error(f"Network error fetching news: {e}")
            data['Top_News'] = "Network error occurred while fetching news."

        # news_url = f'https://newsapi.org/v2/everything?q={ticker}&apiKey={NEWS_API_KEY}&pageSize=3'
        # news_response = requests.get(news_url)
        # if news_response.status_code == 200:
        #     news_data = news_response.json()
        #     articles = news_data.get('articles', [])
        #     if articles:
        #         top_news = "\n\n".join([f"{i+1}. {article['title']} - {article['url']}" for i, article in enumerate(articles)])
        #         data['Top_News'] = top_news
        #     else:
        #         data['Top_News'] = "No news articles found."
        # else:
        #     data['Top_News'] = "Failed to fetch news articles."
    except Exception as e:
        logging.info(f"Error occurred while collecting stock data: {e}")
        print(f"Error occurred while collecting stock data: :\n{e}")
        return jsonify({'message': 'Internal Server Error in Stock Data Collection'}), 500
    
    print(data['Top_News'])
    
    try:
            
        # Step 3: Save Financial Data to Excel
        file_path = os.path.join('data', f'{ticker}_financial_data.xlsx')
        with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
            income_statement.to_excel(writer, sheet_name='Income Statement')
            balance_sheet.to_excel(writer, sheet_name='Balance Sheet')
            cashflow.to_excel(writer, sheet_name='Cashflow')

        # Step 4: Perform Analysis
        avg_close = hist['Close'].mean()
        formatted_data = extract_excel_data(file_path)
        return data,formatted_data,avg_close,file_path
    except Exception as e:
        logging.info(f"Error occurred while performing analysis: {e}")
        print(f"Error occurred while performing analysis :\n{e}")
        return jsonify({'message': 'Internal Server Error in Stock Analysis'}), 500



# Helper function to extract a ticker from the query

# # Best Code answers the queries properly :)
def extract_ticker(query):
    # Mapping of popular company names to tickers for demonstration (you can expand this)
    companies_to_tickers = {
        "apple": "AAPL",
        "microsoft": "MSFT",
        "amazon": "AMZN",
        "tesla": "TSLA",
        "google": "GOOGL",
        "nvidia": "NVDA"
    }

    # Split the query into words
    words = query.lower().split()
    
    # Check for known company names or tickers
    for word in words:
        if word in companies_to_tickers:
            # word[0] = word[0].upper()
            # word[1:] = word[1:].lower()
            return companies_to_tickers[word] ,word.capitalize() #word.upper() #word 
    
    # Try to find a valid stock ticker by querying Yahoo Finance
    for word in words:
        if word:  # Ensure the word is not None or empty
            try:
                ticker = yf.Ticker(word.upper())
                if ticker.info.get('regularMarketPrice') is not None:
                    return ticker ,word.upper()  # Return the valid ticker
            except Exception as e:
                continue
    
    # Default fallback if no ticker is found
    print("No valid ticker found in the query.")
    return None,None


def format_chat_history_for_llm(chat_history, new_query):
    # Format chat history as a readable conversation for the model
    conversation = ""
    for entry in chat_history:
        user_query = entry.get('user_query', '')
        message = entry.get('message', '')
        
        # Append user query and model's response to the conversation
        conversation += f"User Query: {user_query}\nResponse: {message}\n\n"
    
    # Append the new query
    conversation += f"User Query: {new_query}\n"
    
    return conversation

from flask import jsonify, send_file, make_response



@app.route('/analyze_stock', methods=['POST'])
def analyze_stock():
    """
    Generate thorough stock analysis using LLM based on company details, market news, and performance.
    """
    try:
        # Fetch input data
        ticker = request.json.get('ticker')
        company = request.json.get('company', None)
        
        if not ticker:
            print("error : Ticker is required")
            ticker = "AMGN"
            # return jsonify({"error": "Ticker is required"}), 400
        
        # Fetch stock data
        data, formatted_data, avg_close, file_path = get_stock_data(ticker)
        
        # Create analysis task prompt for LLM
        task_prompt = f"""
        You are a Stock Market Expert with in-depth knowledge of stock market trends and patterns.
        Analyze the stock performance for {ticker}. The company's details are as follows:{formatted_data}
        Company news : {data.get('Top_News')}
        You have enough data available to analyze the stock and no need to say lack of data or context.

        **Company Name:** 
        **PE Ratio:** {data.get('PE_Ratio')}
        **EPS:** {data.get('EPS')}
        **Book Value:** {data.get('Book_Value')}
        **ROE:** {data.get('ROE')}
        **ROCE:** {data.get('ROCE')}
        **Order Booking:** Not Provided
        **Revenue Growth:** {data.get('Revenue_Growth')}
        **Earnings Growth:** {data.get('Earnings_Growth')}
        **Today's Market Performance:** Closing Price - {data.get('Todays_Closing_Price')}, High Price - {data.get('Today_High_Price')}

        Evaluate the company's income statement, balance sheet, and cash flow. Provide insights into:
        - Whether the stock is overvalued or undervalued.
        - Predictions for its performance in the upcoming quarter.
        - Recommendations for buying, holding, or selling the stock.
        - Give your views on the KPIs in a table format for the Stock:
        PE, EPS, Book Value, ROE, ROCE, Revenue Growth (CAGR), Earnings Growth
        """
        
        # Generate content using LLM model
        model = genai.GenerativeModel('gemini-1.5-flash')
        llm_response = model.generate_content(task_prompt)
        # analysis_response = markdown_to_text(llm_response.text)
        
        # # Extract insights and suggestions from the response
        # formatted_suggestions = markdown.markdown(analysis_response)
        # print(f"\nOutput:\n{formatted_suggestions}")
        
        htmlSuggestions = markdown.markdown(llm_response.text)
        logging.info(f"Suggestions for investor: \n{htmlSuggestions}")
        
        formatSuggestions = markdown_to_text(htmlSuggestions)
        answer = markdown_table_to_html(formatSuggestions)
        print(answer)
        
        stock_price_predictions_data = stock_price_predictions(ticker)
        # Construct response object
        response_data = {
            "ticker": ticker,
            "company": company,
            "average_closing_price": f"${avg_close:.2f}",
            "analysis": answer, # formatted_suggestions,
            "news": data.get("Top_News", "No news available"),
            "graph_url": f"https://finance.yahoo.com/chart/{ticker}",
            "predictions":stock_price_predictions_data
        }

        # Attach the Excel file if available
        # if os.path.exists(file_path):
        #     file_response = send_file(file_path, as_attachment=True, download_name=f'{ticker}_financial_data.xlsx')
        #     file_response.headers['X-Stock-Metadata'] = jsonify(response_data)
        #     return file_response

        return jsonify(response_data)

    except Exception as e:
        logging.error(f"Error generating stock analysis: {e}")
        return jsonify({"error": f"Failed to generate stock analysis: {str(e)}"}), 500

def stock_price_predictions(ticker):
    try:
        # Step 1: Fetch historical stock data
        stock = yf.Ticker(ticker)
        historical_data = stock.history(period="6mo")
        if historical_data.empty:
            return jsonify({"message": f"No historical data found for ticker: {ticker}"}), 404

        # Step 2: Calculate key statistics from historical data
        volatility = compute_volatility(historical_data['Close'])
        sharpe_ratio = compute_sharpe_ratio(historical_data['Close'])
        recent_trend = historical_data['Close'].pct_change().tail(5).mean() * 100  # Last 5-day trend

        # Step 3: Fetch related market and economic news
        news = fetch_news(ticker)
        market_conditions = collect_market_conditions()
        
        if market_conditions == None:
            print("Market Conditions couldnt be determined")
            market_conditions = ""
        
        print(market_conditions)

        # Generate prompt for LLM model
        task = f"""
            You are a top financial analyst tasked with predicting stock price trends for {ticker}.
            Analyze the following:
            - Recent stock price volatility: {volatility:.2f}%
            - Sharpe Ratio: {sharpe_ratio:.2f}
            - Recent price trends (5-day): {recent_trend:.2f}%
            - Market and economic conditions: {market_conditions}
            - Relevant news: {news}

            Predict the expected stock prices for the next month (30 days) under these conditions:
            1. **Best-Case Scenario** (Optimistic market conditions).
            2. **Worst-Case Scenario** (Pessimistic market conditions).
            3. **Confidence Band** (Range of expected prices with 95% confidence).
            
            Introduce **realistic daily ups and downs** caused by market conditions and noise to simulate realistic portfolio performance.

            Example of simulated_response = 
            ### Response Format:
            | Date       | Best-Case Return (%) | Worst-Case Return (%) | Confidence Band (%) | Total Return (%) |
            |------------|-----------------------|-----------------------|---------------------|------------------|
            | 2025-01-01 | 2.5 | -1.0 | 1.0% - 2.0% | 0.75 |
            | 2025-01-15 | 3.0 | -0.5 | 1.5% - 2.5% | 1.25 |
            | 2025-01-31 | 3.5 | 0.0 | 2.0% - 3.0% | 1.75 |
            | 2025-02-01 | 4.0 | 0.5 | 2.5% - 3.5% | 2.25 |
            | 2025-02-15 | 4.5 | 1.0 | 3.0% - 4.0% | 2.75 |
            | 2025-02-28 | 5.0 | 1.5 | 3.5% - 4.5% | 3.25 |
            | 2025-03-01 | 5.5 | 2.0 | 4.0% - 5.0% | 3.75 |
            | 2025-03-15 | 6.0 | 2.5 | 4.5% - 5.5% | 4.25 |
            | 2025-03-31 | 6.5 | 3.0 | 5.0% - 6.0% | 4.75 |

            
            Your Response must be in the above table format no messages is required just table format data.
            """

        # Step 4: Simulate LLM prediction
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(task)

        simulated_response = markdown_to_text(response.text)
        print(simulated_response)

        # Step 5: Extract and refine predictions
        line_chart_data = extract_line_chart_data(simulated_response)
        refined_predictions = add_noise(line_chart_data)

        # Return refined prediction results
        return refined_predictions
    
        # return jsonify({
        #     "ticker": ticker,
        #     "predictions": refined_predictions,
        #     "analysis": simulated_response
        # })

    except Exception as e:
        print(f"Error in predicting stock prices: {e}")
        return jsonify({"message": f"Error predicting stock prices: {e}"}), 500


def collect_market_conditions():
    """
    Fetch and process current market conditions data, including economic indicators,
    news, and trends to assist in stock analysis and prediction.
    
    Returns:
        dict: A dictionary containing market conditions such as interest rates, inflation,
              geopolitical news, and general market sentiment.
    """
    market_conditions = {}

    try:
        # economic_data_url = "https://api.example.com/economic-indicators"
        economic_data_url = f"https://www.alphavantage.co/query?function=REAL_GDP&apikey={ALPHA_VANTAGE_API_KEY}"
        
        market_news_url = f"https://www.alphavantage.co/query?function=SECTOR&apikey={ALPHA_VANTAGE_API_KEY}"

        # market_news_url = "https://api.example.com/market-news"

        # Fetch economic indicators
        # economic_response = requests.get(economic_data_url)
        # if economic_response.status_code == 200:
        #     economic_data = economic_response.json()
        #     market_conditions['interest_rates'] = economic_data.get('interest_rates', 'Data unavailable')
        #     market_conditions['inflation_rate'] = economic_data.get('inflation_rate', 'Data unavailable')
        # else:
        #     market_conditions['interest_rates'] = 'Failed to fetch interest rates'
        #     market_conditions['inflation_rate'] = 'Failed to fetch inflation rate'

        # # Fetch market news
        # news_response = requests.get(market_news_url)
        # if news_response.status_code == 200:
        #     news_data = news_response.json()
        #     market_conditions['market_news'] = [article['title'] for article in news_data.get('articles', [])][:5]
        # else:
        #     market_conditions['market_news'] = 'Failed to fetch market news'

        # # Add other relevant conditions
        # market_conditions['geopolitical_tensions'] = "Moderate tensions observed globally."
        # market_conditions['us_elections'] = "Upcoming elections may influence market trends."
        
        try:
            # Fetch market data from API
            economic_response = requests.get(economic_data_url)
            market_response = requests.get(market_news_url)

            # Check for successful API responses
            if economic_response.status_code == 200 and market_response.status_code == 200:
                market_conditions = {
                    "interest_rates": economic_response.json().get("interest_rates", "Data unavailable"),
                    "inflation_rate": economic_response.json().get("inflation_rate", "Data unavailable"),
                    "market_news": market_response.json().get("news", []),
                    "geopolitical_tensions": "Moderate tensions observed globally.",
                    "us_elections": "Upcoming elections may influence market trends."
                }
            else:
                raise ValueError("API data fetch failed.")

        except Exception as e:
            print(f"Error fetching market conditions: {e}")
            market_conditions = get_default_market_conditions()


    except Exception as e:
        logging.error(f"Error fetching market conditions: {e}")
        market_conditions['error'] = f"Error fetching market conditions: {e}"

    return market_conditions

def get_default_market_conditions():
    default_conditions = {
        "interest_rates": "Stable interest rates at 4.5%.",
        "inflation_rate": "Moderate inflation at 3.1%.",
        "market_news": [
            "Global markets show mixed trends amid economic recovery.",
            "Tech stocks rally as demand for AI-driven solutions increases.",
            "Oil prices stabilize after months of volatility."
        ],
        "geopolitical_tensions": "Moderate tensions observed globally.",
        "us_elections": "After the elections result of Donald Trump winning the electeions may influence market trends in positive way.",
        "global_trade": "Trade agreements show positive progress with new MAGA (Make America Great Again) Policies.",
        "consumer_confidence": "Consumer confidence index steadily increasing."
    }
    return default_conditions


################################################################################
# v-1 :

# @app.route('/analyze_stock', methods=['POST'])
# def analyze_stock():
#     try:
#         ticker = request.json.get('ticker')
#         company = request.json.get('company',None)
#         query = request.json.get('query')
#         chat_id = request.json.get('chat_id', get_next_chat_id())  # Use auto-incrementing chat ID if not provided
#         # chat_id = request.json.get('chat_id', 1)  # Default chat_id to 1 if not provided
        
#         # Load chat history
#         chat_history = load_chat_history(chat_id)

#         # If no ticker provided in the request, try to extract it from the query
#         if not ticker and query:
#             # ticker = extract_ticker(query)
            
#             ticker,company = extract_ticker(query)
        
#         # If a valid ticker is found, fetch stock data
#         if ticker:
#             try:
#                 data, formatted_data, avg_close,file_path = get_stock_data(ticker)
#                 user_query = ticker  # Save the ticker as the user query
#             except Exception as e:
#                 print("Error getting the stock data")
#                 return jsonify({'message': f'Error occurred while fetching stock data: {e}'}), 400
#         else:
#             # No valid ticker found, generate generic suggestions
#             print("No valid ticker found in the query, generating general stock suggestions.")
#             data = {}  # No specific stock data need to check for news
#             formatted_data = ""  # No financial data
#             avg_close = 0
#             user_query = query  # Save the original user query if no ticker is found

#         # If query is empty, set a default query for stock analysis
#         # if not query:
#         #     query = "Generate general stock suggestions based on current market trends and give some stock predictions."
        
        

        
#          # Save the user's query (ticker or original query) to chat history
#         if user_query:
#             chat_history.append({"user_query": user_query, "message": query})
        
#         # Detect if this is a follow-up query based on previous history
#         if chat_history:
#             print("This is a follow-up query. Checking previous chat history.")
#             # The logic here could vary; you might compare the current query with past responses or check patterns
#             query = f"Following up on: {chat_history[-1]['user_query']} \n\n {chat_history[-1]['message']}" + query

#         # Save the user's query (ticker or original query) to chat history
#         chat_history.append({"user_query": user_query, "message": query})
        
      
            
#         # Format the chat history for the LLM
#         try :
#             formatted_history = format_chat_history_for_llm(chat_history, query)
#         except Exception as e:
#             logging.error(f"Error while formatting chat history for LLM: {e}")
#             return jsonify({'message': 'Internal Server Error in Formatting Chat History'}), 500
        
        
#     except Exception as e :
#         logging.error(f"Error while fetching stock data: {e}")
#         return jsonify({'message': 'Internal Server Error in Stock Data Fetch'}), 500
    
#     try:
#         if ticker:
#             # task = f"""You are a Stock Market Expert. You know everything about stock market trends and patterns.Given a stock related query and if the company's details are provided,
#             #             Based on the provided stock data, analyze the stock's performance, including whether it is overvalued or undervalued.
#             #             Give the user details and information of all the KPI's related to the compnay such as PE ratio,EPS,Book Value,ROE,ROCE,Ernings Growth and Revenue Growth and give your views on them.
#             #             Analyse all the stock information and provide the analysis of the company's performance related to Income Statement,Balance Sheet, and Cashflow.
#             #             Predict the stock price range for the next week (if a particular time period is not mentioned) and provide reasons for your prediction.
#             #             Advise whether to buy this stock now or not, with reasons for your advice. If no stock data is provided just answer the user's query.
#             #             If the user asks for some stock suggestions then provide them a list of stock suggestions based on the query.
#             #             If the user has asked a follow up question then provide them a good response by also considering their previous queries
#             #             Do not answer any questions unrelated to the stocks."""
                        
#             task = f"""You are a Stock Market Expert. You know everything about stock market trends and patterns.Given a stock related query and if the company's details are provided,
#                     Based on the provided stock data, analyze the stock's performance, including whether it is overvalued or undervalued.
#                     Give the user details and information of all the KPI's related to the compnay such as PE ratio,EPS,Book Value,ROE,ROCE,Ernings Growth and Revenue Growth and give your views on them.
#                     Analyse all the stock information and provide the analysis of the company's performance related to Income Statement,Balance Sheet, and Cashflow.
#                     Predict the stock price range for the next week (if a particular time period is not mentioned) and provide reasons for your prediction.
#                     Advise whether to buy this stock now or not, with reasons for your advice."""
        

#             query = task + "\nStock Data: " + str(data) + "\nFinancial Data: " + formatted_data + query
        
#         else:
#             task = """You are a Stock Market Expert. You know everything about stock market trends and patterns.Given a stock related query.
#                         You are the best Stock recommendations AI and you give the best recommendations for stocks.Answer to the questions of the users and help them 
#                         with any queries they might have.
#                         If the user asks for some stock suggestions or some good stocks then provide them a list of stock suggestions based on the query give them the well known stocks in that sector or whatever the query asks for .
#                         If the user has asked a follow up question then provide them a good response by also considering their previous queries
#                         Do not answer any questions unrelated to the stocks."""
            
#             query = task + query + "\n\nConversation:\n" + formatted_history #+ chat_history
#             print(f"The formatted chat history passed to llm is : {formatted_history}")
#             print(f"The query passed to llm is : {query}")
#          # task = f"""You are a Stock Market Expert. You know everything about stock market trends and patterns.
#         #             Based on the provided stock data, analyze the stock's performance, including whether it is overvalued or undervalued.
#         #             Predict the stock price range for the next week and provide reasons for your prediction.
#         #             Advise whether to buy this stock now or not, with reasons for your advice."""
        
        
#         # Use your generative AI model for analysis (example with 'gemini-1.5-flash')
#         model = genai.GenerativeModel('gemini-1.5-flash')
#         response = model.generate_content(query)
#         print(response.text)
#         print(data)
    
#     except Exception as e:
#         logging.error(f"Error performing analysis with generative AI: {e}")
#         return jsonify({f"error": "Failed to give analysis of stock data : {e}"}), 500
    
#     # Extract response from the model
#     try:
#         html_suggestions = markdown.markdown(response.text)
        
#         print(f"Html Suggestions : {html_suggestions}")
        
#         logging.info(f"Suggestions for stock: \n{response.text}")
        
#         # format_suggestions = markdown_to_text(response)
#         print(f"Html Suggestions : {html_suggestions}")
#         format_suggestions = markdown_to_text(html_suggestions)
        
#     except Exception as e:
#         logging.error(f"Error extracting text from response: {e}")
#         print(f"Error extracting text from response : {e}")
#         return jsonify({"error": "Failed to analyze stock data"}), 500

#     # Save the assistant's response to chat history
#     chat_history.append({"user_query": user_query, "message": format_suggestions})
#     save_chat_history(chat_id, chat_history)

#     # Increment chat_id for the next follow-up question
#     new_chat_id = get_next_chat_id()
    
#     if data == {}:
#         data['Top_News'] = None
        
#     data['Company'] = company if company else None
      
#     # Return all collected and analyzed data
#       # Create a response dictionary # gave responses in headers :
#     # response_dict = {
#     #     "data": data,
#     #     "average_closing_price": f"${avg_close:.2f}",
#     #     "analysis": format_suggestions,  # Use the response text here
#     #     "news": data.get('Top_News'),
#     #     "graph_url": f"https://finance.yahoo.com/chart/{ticker}"
#     # }
#     # # If the Excel file exists, send it as an attachment along with the response
#     # if os.path.exists(file_path):
#     #         file_response = send_file(file_path, as_attachment=True, download_name=f'{ticker}_financial_data.xlsx')
#     #         file_response.headers['Content-Disposition'] = f'attachment; filename={ticker}_financial_data.xlsx'
#     #         file_response.headers['X-Stock-Metadata'] = json.dumps(response_dict)  # Add metadata as a custom header
#     #         return file_response
#     # else:
#     #     return jsonify(response_dict)
    
#     # if os.path.exists(file_path): # works for either file or response
#     #         # Combine the file response and JSON response
#     #         file_response = send_file(file_path, as_attachment=True, download_name=f'{ticker}_financial_data.xlsx')
#     #         file_response.headers['Content-Disposition'] = f'attachment; filename={ticker}_financial_data.xlsx'
#     #         print("File is passed as attachment")
#     #         return file_response
#     # else:
#     #     print("Data is passed")
#     #     return jsonify(response_dict)
    
#     return jsonify({
#         # "Company": company,
#         "data": data,
#         "average_closing_price": f"${avg_close:.2f}",
#         "analysis": format_suggestions,
#         "news": data['Top_News'],
#         "graph_url": f"https://finance.yahoo.com/chart/{ticker}"
#     }) # "chat_history" : chat_history
#     # # "new_chat_id" : new_chat_id

def extract_excel_data(file_path):
    financial_data = ""
    xls = pd.ExcelFile(file_path)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        financial_data += f"\n\nSheet: {sheet_name}\n"
        financial_data += df.to_string()
    
    print(f"Financial data of excel file : {financial_data}")
    return financial_data

#########------------------- Portfolio Analysis --------------------------------################

# Fetch the Asset Data from the Selected Market :

# Local and AWS Directory :

MARKET_ASSETS_FOLDER = "market_assets/"
LOCAL_STORAGE_PATH = "local_data"

# Helper function for Timestamp serialization

def serialize_timestamp(obj):
    """
    Convert Timestamp or datetime objects to ISO format strings.
    """
    if isinstance(obj, (pd.Timestamp, datetime)):
        return obj.isoformat()
    raise TypeError(f"Type {type(obj)} not serializable")


def save_to_aws_with_timestamp(data, filename):
    try:
        serialized_data = json.dumps(data, default=serialize_timestamp)  # Use the custom serializer
        s3.put_object(
            Bucket=S3_BUCKET_NAME,
            Key=filename,
            Body=serialized_data,
            ContentType='application/json'
        )
        print(f"Data saved to AWS at {filename}")
    except (NoCredentialsError, PartialCredentialsError) as e:
        print(f"AWS credentials error: {e}")
        raise
    except Exception as e:
        print(f"Error saving to AWS: {e}")
        raise


# import yfinance as yf

ALPHA_VANTAGE_API_KEY = os.getenv('ALPHA_VANTAGE_API_KEY')

import requests

# Fetch Stocks for NASDAQ,NYSE,S&P500 and DOW JONES :

# # V-3 :

# prev : working :
# Updated :

def fetch_all_assets_by_preference(market_name, preference=None):
    """
    Fetch assets for a given market and filter by type if preference is provided.
    Handles NASDAQ, NYSE, S&P500, and Dow Jones dynamically.
    Preferences: "stocks", "etfs", "bonds", "commodities", "mutual funds".
    """
    try:
        market_name = market_name.lower()
        preference = preference.lower() if preference else None
        assets = []

        # Fetch for NASDAQ and NYSE using Alpha Vantage
        if market_name in ["nasdaq", "nyse"]:
            exchange_code = "NASDAQ" if market_name == "nasdaq" else "NYSE"
            url = f"https://www.alphavantage.co/query?function=LISTING_STATUS&apikey={ALPHA_VANTAGE_API_KEY}"
            response = requests.get(url)
            print(response)
            if response.status_code == 200:
                stocks = response.text.splitlines()  # Alpha Vantage returns CSV data
                print(stocks)
                for row in stocks[1:]:  # Skip header row
                    data = row.split(",")
                    if len(data) > 2 and data[2].strip() == exchange_code:
                        symbol = data[0]
                        name = data[1]

                        # Filter based on preference
                        # Determine asset type (heuristics for ETF or stock)
                        asset_type = "stock"
                        if "ETF" in name.upper() or "TRUST" in name.upper() or symbol.endswith("O"):
                            asset_type = "etf"
                            assets.append({"name": name, "symbol": symbol, "type": "ETF"})
                            
                        # Filter based on preference
                        # if not preference or preference == asset_type:
                        elif preference:
                            assets.append({"name": name, "symbol": symbol, "type": preference})
                            
                        # if preference:  # Assuming preference is handled externally
                        #     assets.append({"name": name, "symbol": symbol, "type": preference})
                        
                print(f"Assets in {market_name} :\n{assets}")
                return assets
            else:
                print(f"Alpha Vantage API error: {response.status_code}")
                return []

        # Fetch for S&P500 using Wikipedia
        elif market_name == "s&p500":
            url = "https://en.wikipedia.org/wiki/List_of_S%26P_500_companies"
            response = requests.get(url)
            if response.status_code == 200:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(response.content, "html.parser")
                table = soup.find("table", {"id": "constituents"})
                rows = table.find_all("tr")[1:]  # Skip header row
                for row in rows:
                    cols = row.find_all("td")
                    symbol = cols[0].text.strip()
                    name = cols[1].text.strip()
                    if not preference or preference == "stocks":
                        assets.append({"name": name, "symbol": symbol, "type": "stock"})
                return assets
            else:
                print(f"Failed to fetch S&P500 data from Wikipedia: {response.status_code}")
                return []
            
        # If no matching market is found
        return []

    except Exception as e:
        print(f"Error fetching assets for {market_name}: {e}")
        return []





@app.route('/market-assets', methods=['POST'])
def market_assets():
    try:
        data = request.get_json()
        market_name = data.get("market_name")
        preference = data.get("preference")
        print(market_name)
        print(preference)
        # preference = "stocks"
        
        if not market_name:
            return jsonify({"message": "Market name is required"}), 400

        # Define filename for storage
        filename = f"{MARKET_ASSETS_FOLDER}{market_name.lower()}_assets.json"
        if USE_AWS:
            assets = load_from_aws(filename)
        else:
            assets = load_from_local(os.path.join(LOCAL_STORAGE_PATH, filename))

        # Fetch updated assets for the market
        
        # updated_assets = fetch_all_stocks_for_market_dynamic(market_name)
        updated_assets = fetch_all_assets_by_preference(market_name,preference)
        
        if not updated_assets:
            return jsonify({"message": f"No data found for the market: {market_name}"}), 404

        # Check if there are new assets
        if not assets or updated_assets != assets:
            # Update the assets list
            if USE_AWS:
                save_to_aws_with_timestamp(updated_assets, filename)
            else:
                save_to_local(updated_assets, os.path.join(LOCAL_STORAGE_PATH, filename))
            message = "Assets list updated successfully"
        else:
            message = "Assets list is up-to-date"

        print(f"\nUpdated Assets :\n{updated_assets}")
        
        return jsonify({
            "message": message,
            "market": market_name,
            "assets": updated_assets
        }), 200

    except Exception as e:
        print(f"Error in market-assets API: {e}")
        return jsonify({"message": f"Internal server error: {e}"}), 500


################################################ Fetch Bonds from Categories #################################################

# def fetch_treasure_bonds():
# def fetch_treasury_bonds():
#     """
#     Fetch bond data (Treasury Yields) from Alpha Vantage.
#     """
#     try:
#         url = f"https://www.alphavantage.co/query?function=TREASURY_YIELD&interval=monthly&maturity=10year&apikey={ALPHA_VANTAGE_API_KEY}"
#         response = requests.get(url)

#         if response.status_code == 200:
#             data = response.json()
#             if "data" in data:
#                 bonds = []
#                 for item in data["data"]:
#                     maturity_date = item.get("maturityDate", "N/A")
#                     yield_rate = item.get("value", "N/A")
#                     bonds.append({"name": f"10 Year Treasury", "symbol": "10Y", "yield": yield_rate, "maturity": maturity_date})

#                 return bonds
#             else:
#                 print("No bond data available.")
#                 return []
#         else:
#             print(f"Alpha Vantage API error: {response.status_code}")
#             return []
#     except Exception as e:
#         print(f"Error fetching bonds: {e}")
#         return []
    

# # Fetch Treasury Bonds

def fetch_treasury_bonds():
    try:
        url = f"https://www.alphavantage.co/query?function=TREASURY_YIELD&interval=monthly&maturity=10year&apikey={ALPHA_VANTAGE_API_KEY}"
        response = requests.get(url)

        if response.status_code == 200:
            data = response.json()
            if "data" in data:
                bonds = []
                for item in data["data"]:
                    maturity_date = item.get("maturityDate", "N/A")
                    yield_rate = item.get("value", "N/A")
                    bonds.append({
                        "name": "10 Year Treasury",
                        "symbol": "10Y",
                        "yield": yield_rate,
                        # "maturity": maturity_date
                    })
                return bonds
            else:
                print("No bond data available.")
                return []
        else:
            print(f"Alpha Vantage API error: {response.status_code}")
            return []
    except Exception as e:
        print(f"Error fetching Treasury bonds: {e}")
        return []

# Fetch Corporate Bonds
def fetch_corporate_bonds():
    try:
        url = f"https://www.alphavantage.co/query?function=CORPORATE_BOND&apikey={ALPHA_VANTAGE_API_KEY}"
        response = requests.get(url)

        if response.status_code == 200:
            data = response.json()
            if "data" in data:
                bonds = [
                    {
                        "name": item.get("name", "N/A"),
                        "symbol": item.get("symbol", "N/A"),
                        "yield": item.get("yield", "N/A"),
                        "maturity": item.get("maturityDate", "N/A")
                    }
                    for item in data["data"]
                ]
                return bonds
            else:
                print("No corporate bond data available.")
                return []
        else:
            print(f"Alpha Vantage API error: {response.status_code}")
            return []
    except Exception as e:
        print(f"Error fetching Corporate bonds: {e}")
        return []

# Fetch Mortgage-Related Bonds
def fetch_mortgage_related_bonds():
    try:
        url = f"https://www.alphavantage.co/query?function=MORTGAGE_RELATED_BONDS&apikey={ALPHA_VANTAGE_API_KEY}"
        response = requests.get(url)

        if response.status_code == 200:
            data = response.json()
            if "data" in data:
                bonds = [
                    {
                        "name": item.get("name", "N/A"),
                        "symbol": item.get("symbol", "N/A"),
                        "yield": item.get("yield", "N/A"),
                        "maturity": item.get("maturityDate", "N/A")
                    }
                    for item in data["data"]
                ]
                return bonds
            else:
                print("No mortgage-related bond data available.")
                return []
        else:
            print(f"Alpha Vantage API error: {response.status_code}")
            return []
    except Exception as e:
        print(f"Error fetching Mortgage-Related bonds: {e}")
        return []

# Fetch Municipal Bonds
def fetch_municipal_bonds():
    try:
        url = f"https://www.alphavantage.co/query?function=MUNICIPAL_BONDS&apikey={ALPHA_VANTAGE_API_KEY}"
        response = requests.get(url)

        if response.status_code == 200:
            data = response.json()
            if "data" in data:
                bonds = [
                    {
                        "name": item.get("name", "N/A"),
                        "symbol": item.get("symbol", "N/A"),
                        "yield": item.get("yield", "N/A"),
                        "maturity": item.get("maturityDate", "N/A")
                    }
                    for item in data["data"]
                ]
                return bonds
            else:
                print("No municipal bond data available.")
                return []
        else:
            print(f"Alpha Vantage API error: {response.status_code}")
            return []
    except Exception as e:
        print(f"Error fetching Municipal bonds: {e}")
        return []


# Fetch Money Market Bonds
def fetch_money_market_bonds():
    """
    Fetch Money Market bond data from Alpha Vantage.
    """
    try:
        url = f"https://www.alphavantage.co/query?function=MONEY_MARKET_BONDS&apikey={ALPHA_VANTAGE_API_KEY}"
        response = requests.get(url)

        if response.status_code == 200:
            data = response.json()
            if "data" in data:
                bonds = [
                    {
                        "name": item.get("name", "N/A"),
                        "symbol": item.get("symbol", "N/A"),
                        "yield": item.get("yield", "N/A"),
                        "maturity": item.get("maturityDate", "N/A")
                    }
                    for item in data["data"]
                ]
                return bonds
            else:
                print("No money market bond data available.")
                return []
        else:
            print(f"Alpha Vantage API error: {response.status_code}")
            return []
    except Exception as e:
        print(f"Error fetching Money Market bonds: {e}")
        return []

##############################################################################

# api to fetch currencies : currently on hold

# # Set API Key and Secret
# XE_Currency_API_key = os.getenv('XE_Currency_API_key')
# XE_API_SECRET = ""  
# XE_BASE_URL = "https://xecdapi.xe.com/v1"

# def fetch_all_currencies():
#     """
#     Fetch a list of all currencies using XE API.
#     """
#     try:
#         url = f"{XE_BASE_URL}/currencies"
#         response = requests.get(url, auth=HTTPBasicAuth(XE_Currency_API_key, XE_API_SECRET))

#         if response.status_code != 200:
#             print(f"Failed to fetch currencies. Status code: {response.status_code}")
#             print(f"Response: {response.text}")
#             return []

#         currency_data = response.json()
#         currencies = [{"symbol": currency["code"], "name": currency["name"]} for currency in currency_data["currencies"]]
#         return currencies

#     except Exception as e:
#         print(f"Error fetching currencies: {e}")
#         return []

# def fetch_currency_prices(base_currency="USD"):
#     """
#     Fetch exchange rates of all currencies relative to a base currency using XE API.
#     """
#     try:
#         url = f"{XE_BASE_URL}/convert_from.json/?from={base_currency}&amount=1"
#         response = requests.get(url, auth=HTTPBasicAuth(XE_Currency_API_key, XE_API_SECRET))

#         if response.status_code != 200:
#             print(f"Failed to fetch currency prices. Status code: {response.status_code}")
#             print(f"Response: {response.text}")
#             return []

#         price_data = response.json()
#         prices = [{"symbol": rate["code"], "price": rate["mid"]} for rate in price_data["rates"]]
#         return prices

#     except Exception as e:
#         print(f"Error fetching currency prices: {e}")
#         return []

# # Flask API Route
# # @app.route('/fetch-global-currencies', methods=['GET'])
# def fetch_global_currencies():
#     """
#     API to fetch all global currencies and their exchange rates.
#     """
#     try:
#         base_currency = request.args.get("base_currency", "USD")
#         currencies = fetch_all_currencies()
#         if not currencies:
#             print("No currencies fetched.")
#             return jsonify({"message": "Failed to fetch currencies"}), 500

#         prices = fetch_currency_prices(base_currency)
#         if not prices:
#             print("No prices fetched.")
#             return jsonify({"message": "Failed to fetch prices"}), 500

#         # Merge currency names and prices
#         merged_data = []
#         for currency in currencies:
#             price_entry = next((p for p in prices if p["symbol"] == currency["symbol"]), None)
#             merged_data.append({
#                 "symbol": currency["symbol"],
#                 "name": currency["name"],
#                 "price": price_entry["price"] if price_entry else "N/A"
#             })

#         print(merged_data)
#         return jsonify({"currencies": merged_data}), 200

#     except Exception as e:
#         print(f"Error in fetch-global-currencies API: {e}")
#         return jsonify({"message": f"Internal server error: {e}"}), 500


##########################################################################################

# # Api to fetch bonds :

# #V2 : working properly 

@app.route('/get-bonds', methods=['POST'])
def get_bonds():
    """
    Fetches the List of Bonds for various categories
    """
    try:
        data = request.get_json()
        category = data.get("category", "").lower()
        treasury_bonds = [
            {"name":"13 WEEK TREASURY BILL" ,"symbol":"^IRX"},
            {"name":"Treasury Yield 5 Years" ,"symbol":"^FVX"},
            {"name":"CBOE Interest Rate 10 Year T No" ,"symbol":"^TNX"},
            {"name":"Treasury Yield 30 Years" ,"symbol":"^TYX"}
        ]
        corporate_bonds = [
            {"name":"BlackRock High Yield Port Svc","symbol":"BHYSX"},
            {"name":"American Funds American High-Inc F2","symbol":"AHIFX"},
            {"name":"PGIM High Yield R6","symbol":"PHYQX"},
            {"name":"Federated Hermes Instl High Yield Bd IS","symbol":"FIHBX"}
        ]
        if category == "treasury":
            return jsonify({"bonds":treasury_bonds}), 200
        elif category == "corporate":
            return jsonify({"bonds":corporate_bonds}), 200
        else:
            return jsonify({"message": "Invalid category. Choose between 'treasury' or 'corporate'."}), 400
            
    except Exception as e:
        print(f"Error in get-bonds API: {e}")
        return jsonify({"message": f"Error in get-bonds API: {e}"}), 500


@app.route('/fetch-bonds', methods=['POST'])
def fetch_bonds():
    """
    Fetches the price of a specific bond using the ticker provided in the request payload.
    """
    try:
        data = request.get_json()
        category = data.get("category", "").lower()

        if not category:
            return jsonify({
                "message": "Ticker not provided in the request.",
                "price": "N/A"
            }), 400
            
        selected_ticker = data.get("ticker")

        # testing mutual funds :
        fetch_MutualFunds()
        
        # Function to fetch the latest closing price
        def fetch_price(ticker):
            try:
                # Fetch historical data for the bond
                data = yf.download(ticker, period="1mo", interval="1d")
                if not data.empty:
                    # Get the latest closing price
                    return round(data['Close'].iloc[-1], 2)
                else:
                    return "Price not available"
            except Exception as e:
                print(f"Error fetching data for ticker {ticker}: {e}")
                return "Price not available"

        # Fetch price for the selected Bond
        price = fetch_price(selected_ticker)
        print(f"Bond price for {selected_ticker}:\n{price}")

        return jsonify({
            "message": f"Price for {selected_ticker} fetched successfully",
            "ticker": selected_ticker,
            "price": price
        }), 200

    except Exception as e:
        print(f"Error fetching bond prices: {e}")
        return jsonify({"message": f"Internal server error: {e}"}), 500

# only treasury is working but it gives 10 Y US Tereasury bonds

# def fetch_bonds_from_yahoo(category):
#     """
#     Fetch bonds data from Yahoo Finance for a specific category.
#     Returns a list of bonds for the given category.
#     """
#     try:
#         if category == "money_market":
#             print("Fetching currencies for money market...")
#             # currencies = fetch_global_currencies()
#             # return currencies

#         url = "https://finance.yahoo.com/markets/bonds/"
#         response = requests.get(url)
#         if response.status_code != 200:
#             print(f"Failed to fetch bonds page. Status code: {response.status_code}")
#             return []

#         soup = BeautifulSoup(response.content, "html.parser")

#         bond_categories = {
#             "treasury": "Treasury Bonds",
#             "corporate": "Corporate Bonds",
#             "municipal": "Municipal Bonds",
#             "money_market": "Money Market",
#         }

#         if category not in bond_categories:
#             print(f"Invalid category: {category}.")
#             return []

#         section_name = bond_categories[category]
#         section = soup.find("section", {"aria-label": section_name})
#         if not section:
#             print(f"No data found for {section_name}.")
#             return []

#         table = section.find("table")
#         if not table:
#             print(f"No table found for {section_name}.")
#             return []

#         rows = table.find_all("tr")[1:]  # Skip header row
#         bonds = []
#         for row in rows:
#             cols = row.find_all("td")
#             if len(cols) < 4:  # Ensure required columns are present
#                 continue
#             symbol = cols[0].text.strip()
#             name = cols[1].text.strip()
#             price = cols[2].text.strip()
#             yield_rate = cols[3].text.strip()

#             bonds.append({
#                 "symbol": symbol,
#                 "name": name,
#                 "price": price,
#                 "yield": yield_rate,
#             })

#         return bonds

#     except Exception as e:
#         print(f"Error fetching bonds: {e}")
#         return []


# @app.route('/fetch-bonds', methods=['POST'])
# def fetch_bonds():
#     """
#     API endpoint to fetch bonds data by category.
#     """
#     try:
        # data = request.get_json()
        # category = data.get("category", "").lower()

#         if not category:
#             return jsonify({"message": "Category is required."}), 400

#         print(f"Fetching bonds for category: {category}")
#         bonds = fetch_bonds_from_yahoo(category)

#         if not bonds:
#             return jsonify({"message": f"No bonds found for category '{category}'."}), 404

#         return jsonify({"category": category, "bonds": bonds}), 200

#     except Exception as e:
#         print(f"Error in fetch-bonds API: {e}")
#         return jsonify({"message": f"Internal server error: {e}"}), 500





###################################################### Fetch Commodities #######################################################


# API to Fetch commodity data :

# #V2 : working properly 
@app.route('/fetch-commodities', methods=['POST'])
def fetch_commodities():
    """
    Fetches the price of a specific commodity using the ticker provided in the request payload.
    """
    try:
        data = request.get_json()
        selected_ticker = data.get("commodities")  

        if not selected_ticker:
            return jsonify({
                "message": "Ticker not provided in the request.",
                "price": "N/A"
            }), 400

        # Function to fetch the latest closing price
        def fetch_price(ticker):
            try:
                # Fetch historical data for the commodity
                data = yf.download(ticker, period="1d", interval="1d")
                if not data.empty:
                    # Get the latest closing price
                    return round(data['Close'].iloc[-1], 2)
                else:
                    return "Price not available"
            except Exception as e:
                print(f"Error fetching data for ticker {ticker}: {e}")
                return "Price not available"

        # Fetch price for the selected commodity
        price = fetch_price(selected_ticker)
        print(f"Commodity price for {selected_ticker}:\n{price}")

        return jsonify({
            "message": f"Price for {selected_ticker} fetched successfully",
            "symbol": selected_ticker,
            "price": price
        }), 200

    except Exception as e:
        print(f"Error fetching commodity prices: {e}")
        return jsonify({"message": f"Internal server error: {e}"}), 500


# V-1 : working properly for all

# @app.route('/fetch-commodities', methods=['POST'])
# def fetch_commodities():
#     """
#     Fetches the price of specific commodities: WTI Crude, Brent Crude, Gold, Silver, Natural Gas.
#     """
#     try:
#         data = request.get_json()
#         selected_commodity = data.get("commodity")  # Single commodity or None

#         # Commodity tickers for Yahoo Finance
#         commodity_tickers = {
#             "WTI Crude": "CL=F",
#             "Brent Crude": "BZ=F",
#             "Gold": "GC=F",
#             "Silver": "SI=F",
#             "Natural Gas": "NG=F"
#         }

#         # Function to fetch the latest closing price
#         def fetch_price(ticker):
#             try:
#                 # Fetch historical data for the commodity
#                 data = yf.download(ticker, period="1d", interval="1d")
#                 if not data.empty:
#                     # Get the latest closing price
#                     return round(data['Close'].iloc[-1], 2)
#                 else:
#                     return "Price not available"
#             except Exception as e:
#                 print(f"Error fetching data for ticker {ticker}: {e}")
#                 return "Price not available"

#         # If a specific commodity is requested
#         if selected_commodity:
#             ticker = commodity_tickers[selected_commodity] #commodity_tickers.get(selected_commodity)
#             print(ticker)
#             if not ticker:
#                 return jsonify({
#                     "message": f"Commodity '{selected_commodity}' not found",
#                     "prices": {}
#                 }), 404

#             # Fetch price for the selected commodity
#             price = fetch_price(ticker)
#             print(f"Commodity price for {selected_commodity} :\n{price}")
            
#             return jsonify({
#                 "message": f"Price for {selected_commodity} fetched successfully",
#                 "prices": {price}
#             }), 200

#         # #Fetch prices for all commodities if no specific one is provided
#         commodity_prices = {}
#         for name, ticker in commodity_tickers.items():
#             commodity_prices[name] = fetch_price(ticker)

#         print("Commodity prices :\n", commodity_prices)

#         return jsonify({
#             "message": "Commodity prices fetched successfully",
#             "prices": commodity_prices
#         }), 200

#     except Exception as e:
#         print(f"Error fetching commodity prices: {e}")
#         return jsonify({"message": f"Internal server error: {e}"}), 500

##################################################### Fetch Cryptocurrencies from Exchanges ####################################

# v-2 :
@app.route('/crypto-assets', methods=['POST'])
def fetch_cryptos_from_exchange():
    """
    Fetch the list of cryptocurrencies available on a given exchange.
    Supported exchanges: CoinGecko, Binance, Binance.US, Coincheck.
    """
    try:
        data = request.get_json()
        exchange_name = data.get("exchange_name", "").lower()
        cryptos = []
        #test reits :
        fetch_reits()
        if exchange_name == "coingecko":
            # Fetch data from CoinGecko
            url = "https://api.coingecko.com/api/v3/coins/markets"
            params = {
                "vs_currency": "usd",
                "order": "market_cap_desc",
                "per_page": 250,
                "page": 1,
            }
            response = requests.get(url, params=params)
            if response.status_code == 200:
                data = response.json()
                for coin in data:
                    symbol = coin["symbol"].upper()
                    cryptos.append({"name": coin["name"], "symbol": f"{symbol}-USD" })
                    
                    # if coin["name"] == "Bitcoin" or coin["name"] == "Ethereum":
                    #     symbol = coin["symbol"].upper()
                    #     cryptos.append({"name": coin["name"], "symbol": f"{symbol}-USD" })
                    # else:
                    #     cryptos.append({"name": coin["name"], "symbol": coin["symbol"].upper()})
            else:
                return jsonify({"message": f"Failed to fetch data from CoinGecko: {response.status_code}"}), 500

        elif exchange_name == "binance":
            # Fetch data from Binance
            url = "https://api.binance.com/api/v3/exchangeInfo"
            response = requests.get(url)
            if response.status_code == 200:
                data = response.json()
                for symbol_info in data["symbol"]:
                    base_asset = symbol_info["baseAsset"]
                    quote_asset = symbol_info["quoteAsset"]
                    
                    # if base_asset == 'ETH' or base_asset == 'BTC':
                    if base_asset == 'ETH' or base_asset == 'BTC' or base_asset == 'XRP' or base_asset == 'USDT' or 'BNB' :
                        cryptos.append({
                            "symbol": f"{base_asset}-USD",
                            "name": f"{base_asset}"
                        })
                    elif base_asset == 'ALGO':
                        cryptos.append({
                            "symbol": f"{base_asset}-INR",
                            "name": f"{base_asset}"
                        })
                    else:
                        cryptos.append({
                            "symbol": base_asset,
                            "name": f"{base_asset}"
                        })
            else:
                return jsonify({"message": f"Failed to fetch data from Binance: {response.status_code}"}), 500

        elif exchange_name == "binance.us":
            # Fetch data from Binance.US
            url = "https://api.binance.us/api/v3/exchangeInfo"
            response = requests.get(url)
            if response.status_code == 200:
                data = response.json()
                for symbol_info in data["symbol"]:
                    base_asset = symbol_info["baseAsset"]
                    quote_asset = symbol_info["quoteAsset"]
                    
                    if base_asset == 'ETH' or base_asset == 'BTC' or base_asset == 'XRP' or base_asset == 'USDT' or 'BNB' :
                        cryptos.append({
                            "symbol": f"{base_asset}-USD",
                            "name": f"{base_asset}"
                        })
                    elif base_asset == 'ALGO':
                        cryptos.append({
                            "symbol": f"{base_asset}-INR",
                            "name": f"{base_asset}"
                        })
                    else:
                        cryptos.append({
                            "symbol": base_asset,
                            "name": f"{base_asset}"
                        })
            else:
                return jsonify({"message": f"Failed to fetch data from Binance.US: {response.status_code}"}), 500

        else:
            return jsonify({"message": "Exchange not supported."}), 404

        # Return the list of cryptos
        return jsonify({
            "message": "Cryptos list fetched successfully.",
            "exchange_name": exchange_name,
            "cryptos": cryptos
        }), 200

    except Exception as e:
        return jsonify({"message": f"Internal server error: {e}"}), 500


####################################################################################

# Fetch REITS :


FINNHUB_API_KEY = os.getenv('FINNHUB_API_KEY')

# v-2 :best version

@app.route("/fetch-reits", methods=['POST'])
def fetch_reits():
    try:
        # AWS key for the REIT list
        reit_list_key = "reits/reit_list.json"

        # Check if the REIT list already exists in AWS
        try:
            response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=reit_list_key)
            reit_list = json.loads(response['Body'].read().decode('utf-8'))
            logging.info("REIT list loaded from AWS.")
            return jsonify({"message": "REITs loaded from AWS successfully", "data": reit_list}), 200
        except s3.exceptions.NoSuchKey:
            logging.info("REIT list not found in AWS. Fetching from Finnhub.")

        # Fetch the list of US stocks from Finnhub
        url = f"https://finnhub.io/api/v1/stock/symbol?exchange=US&token={FINNHUB_API_KEY}"
        response = requests.get(url)

        if response.status_code != 200:
            return jsonify({"message": "Failed to fetch REITs from Finnhub", "status_code": response.status_code}), 500

        data = response.json()
        valid_reits = []

        # Filter REITs and fetch prices in a single loop
        for item in data:
            if "REIT" in item.get("description", "") or "Real Estate" in item.get("description", ""):
                symbol = item["symbol"]
                name = item["description"]

                # Fetch the price for the current REIT
                price_url = f"https://finnhub.io/api/v1/quote?symbol={symbol}&token={FINNHUB_API_KEY}"
                price_response = requests.get(price_url)

                if price_response.status_code == 200:
                    price_data = price_response.json()
                    price = price_data.get("c", 0)  # "c" is the current price key

                    # Only add to the list if the price is greater than 0
                    if price > 0:
                        valid_reits.append({"symbol": symbol, "name": name, "price": price})

        # Save the valid REITs list to AWS
        s3.put_object(
            Bucket=S3_BUCKET_NAME,
            Key=reit_list_key,
            Body=json.dumps(valid_reits),
            ContentType='application/json'
        )
        logging.info("REIT list saved to AWS.")

        return jsonify({"message": "REITs fetched and saved successfully", "data": valid_reits}), 200

    except Exception as e:
        logging.error(f"Error fetching REITs from Finnhub: {e}")
        return jsonify({"message": "An error occurred while fetching REITs", "error": str(e)}), 500

@app.route("/get-reit-price", methods=['POST'])
def get_reit_price():
    try:
        # Ensure Content-Type is application/json
        if not request.is_json:
            return jsonify({"message": "Invalid Content-Type. Please set 'Content-Type: application/json'."}), 415

        # Parse the REIT symbol from the request
        symbol = request.json.get("symbol")
        if not symbol or not isinstance(symbol, str) or not symbol.strip():
            return jsonify({"message": "Invalid symbol provided. Must be a non-empty string."}), 400

        symbol = symbol.strip()  # Remove any leading/trailing spaces

        # Load the REIT list from AWS
        response = s3.get_object(Bucket=S3_BUCKET_NAME, Key="reits/reit_list.json")
        reit_list = json.loads(response['Body'].read())

        # Validate if the symbol exists in the REIT list
        reit = next((item for item in reit_list if item["symbol"] == symbol), None)
        if not reit:
            return jsonify({"message": f"Symbol {symbol} is not a valid REIT."}), 400

        # Fetch yield for the REIT
        yield_url = f"https://finnhub.io/api/v1/stock/metric?symbol={symbol}&metric=dividends&token={FINNHUB_API_KEY}"
        yield_response = requests.get(yield_url)

        if yield_response.status_code == 200:
            yield_data = yield_response.json()
            dividend_yield = yield_data.get("metric", {}).get("dividendYieldIndicatedAnnual", 5)
        else:
            dividend_yield = 5 # default value

        # Return the REIT information
        return jsonify({
            "message": "Price and yield fetched successfully",
            "reit_info": {"symbol": reit["symbol"], "name": reit["name"], "price": reit["price"], "yield": dividend_yield}
        }), 200

    except Exception as e:
        logging.error(f"Error fetching REIT prices: {e}")
        return jsonify({"message": "An error occurred while fetching REIT prices", "error": str(e)}), 500

####################################################################################


# Fetch Mutual Funds :
# V-2 : best version fetches all the list and prices

# Function to fetch the latest price
def fetch_price(ticker):
    """
    Fetch the latest price of a mutual fund using Yahoo Finance.
    :param ticker: The symbol of the mutual fund.
    :return: The latest closing price or "Price not available".
    """
    try:
        # Fetch data for the mutual fund with valid period and interval
        data = yf.download(ticker, period="1mo", interval="1d")
        if not data.empty:
            # Get the latest closing price
            return round(data['Close'].iloc[-1], 2)
        else:
            return "Price not available"
    except Exception as e:
        print(f"Error fetching data for ticker {ticker}: {e}")
        return "Price not available"

def fetch_mutual_funds_from_yahoo():
    """
    Fetch a list of mutual funds and their prices from Yahoo Finance's Mutual Funds Gainers page.
    :return: List of mutual funds with symbol, name, and price.
    """
    try:
        url = "https://finance.yahoo.com/markets/mutualfunds/gainers/"
        response = requests.get(url)
        if response.status_code != 200:
            print(f"Failed to fetch Yahoo Finance page. Status code: {response.status_code}")
            return []

        soup = BeautifulSoup(response.content, "html.parser")
        table = soup.find("table")  # Locate the main table with mutual fund data

        if not table:
            print("No table found on the Yahoo Finance page.")
            return []

        rows = table.find_all("tr")[1:]  # Skip the header row
        mutual_funds = []

        for row in rows:
            cols = row.find_all("td")
            if len(cols) < 3:  # Ensure required columns are present
                continue

            symbol = cols[0].text.strip()
            name = cols[1].text.strip()
            price = fetch_price(symbol)  # Fetch the price dynamically

            # Skip mutual funds where the price is not available
            if price == "Price not available":
                continue

            mutual_funds.append({
                "symbol": symbol,
                "name": name,
                "price": price
            })

        return mutual_funds

    except Exception as e:
        print(f"Error fetching mutual funds from Yahoo Finance: {e}")
        return []

# Endpoint to fetch mutual funds
@app.route('/fetch-MutualFunds', methods=['POST'])
def fetch_MutualFunds():
    """
    Fetch and return mutual funds and their details.
    """
    try:
        mutual_funds = fetch_mutual_funds_from_yahoo()

        print(f"Mutual funds: {mutual_funds}")
        return jsonify({
            "message": "Mutual funds fetched successfully.",
            "mutual_funds": mutual_funds
        }), 200

    except Exception as e:
        print(f"Failed to fetch mutual funds: {e}")
        return jsonify({"error": "Failed to fetch mutual funds"}), 500


# v-1 : fetched list but not price

# def fetch_mutual_funds_from_yahoo():
#     """
#     Fetch a list of mutual funds and their prices from Yahoo Finance's Mutual Funds Gainers page.
#     """
#     try:
#         url = "https://finance.yahoo.com/markets/mutualfunds/gainers/"
#         response = requests.get(url)
#         if response.status_code != 200:
#             print(f"Failed to fetch Yahoo Finance page. Status code: {response.status_code}")
#             return []

#         soup = BeautifulSoup(response.content, "html.parser")
#         table = soup.find("table")  # Locate the main table with mutual fund data

#         if not table:
#             print("No table found on the Yahoo Finance page.")
#             return []

#         rows = table.find_all("tr")[1:]  # Skip the header row
#         mutual_funds = []

#         for row in rows:
#             cols = row.find_all("td")
#             if len(cols) < 3:  # Ensure required columns are present
#                 continue

#             symbol = cols[0].text.strip()
#             name = cols[1].text.strip()
#             price = cols[2].text.strip()

#             mutual_funds.append({
#                 "symbol": symbol,
#                 "name": name,
#                 "price": price
#             })

#         return mutual_funds

#     except Exception as e:
#         print(f"Error fetching mutual funds from Yahoo Finance: {e}")
#         return []

# def fetch_mutual_fund_details(symbol):
#     """
#     Fetch detailed information about a mutual fund using yfinance.
#     :param symbol: Mutual fund symbol (e.g., "VASGX").
#     :return: Detailed mutual fund information.
#     """
#     try:
#         mutual_fund = yf.Ticker(symbol)
#         info = mutual_fund.info
#         return {
#             "symbol": symbol,
#             "name": info.get("shortName", "N/A"),
#             "price": info.get("regularMarketPrice", "Price not available"),
#             "category": info.get("category", "N/A"),
#             "fundFamily": info.get("fundFamily", "N/A")
#         }
#     except Exception as e:
#         print(f"Error fetching data for ticker {symbol}: {e}")
#         return {}

# # Endpoint to fetch mutual funds
# # @app.route('/fetch-MutualFunds', methods=['GET'])
# def fetch_MutualFunds():
#     """
#     Fetch and return mutual funds and their details.
#     """
#     try:
#         mutual_funds = fetch_mutual_funds_from_yahoo()
#         detailed_mutual_funds = []

#         for fund in mutual_funds:
#             print(f"Fetching details for {fund['symbol']}...")
#             details = fetch_mutual_fund_details(fund["symbol"])
#             if details:
#                 detailed_mutual_funds.append(details)
                
#         print(f"Mutual funds {detailed_mutual_funds}")
#         # return jsonify({
#         #     "message": "Mutual funds fetched successfully.",
#         #     "mutual_funds": detailed_mutual_funds
#         # }), 200

#     except Exception as e:
#         print(f"Failed to fetch mutual funds: {e}")
#         return jsonify({"error": "Failed to fetch mutual funds"}), 500


####################################################################################
# Fetch Current Stock Price :

@app.route('/current_stock_price', methods=['POST'])
def current_stock_price():
    try:
        ticker = request.json.get('ticker')
        stock = yf.Ticker(ticker)
        # Fetch the current stock price using the 'regularMarketPrice' field
        current_price = stock.info.get('regularMarketPrice')
        
        if not current_price:
            print(f"Failed to retrieve the current price for {ticker}.\nExtracting closing Price of the Stock")
            current_price = stock.history(period='1d')['Close'].iloc[-1]
            return jsonify({"current_price":current_price})
        
        if current_price is None:
            # If still None, check for mutual fund-specific fields
            print(f"Attempting to retrieve price for Mutual Fund {ticker}...")
            fund_close_price = stock.history(period="1d")['Close']
            if len(fund_close_price) > 0:
                current_price = fund_close_price.iloc[-1]  # Last available closing price
            return jsonify({"current_price":current_price})

        # If everything fails, raise an error
        if current_price is None:
            raise ValueError(f"Unable to retrieve price for {ticker}.")

        return jsonify({"current_price":current_price})
    
    except Exception as e:
        print(f"Failed to retrieve the current price for {ticker} : {e}")
        return jsonify({"error": f"Failed to retrieve the current price for {ticker}"}), 500


@app.route('/dividend_yield', methods=['POST'])
def dividend_yield():
    
    ticker_name = request.json.get('ticker')
    # Create a Ticker object using yfinance
    stock = yf.Ticker(ticker_name)
    
    # Fetch the stock information, including dividend yield
    try:
        dividend_yield = stock.info.get('dividendYield')
        sector = stock.info.get('sector')
        industry = stock.info.get('industry')

        if dividend_yield is not None:
            dividend_yield_percent = dividend_yield * 100  # Convert to percentage
            print(f"The dividend yield for {ticker_name} is: {dividend_yield_percent:.2f}%")
        else:
            print(f"No dividend yield information available for {ticker_name}.")
        
        # Additional information check to verify it's a REIT or commercial real estate company
        if industry and ('reit' in industry.lower() or 'real estate' in industry.lower()):
            print(f"{ticker_name} belongs to the {industry} industry.")
        else:
            print(f"{ticker_name} may not be a REIT or a commercial real estate company.")
        
        return jsonify({'dividend_yield_percent': float(dividend_yield_percent) , "status": 200})
    except Exception as e:
        print(f"Error occurred while fetching data for {ticker_name}: {e}")

# # Works well for real estate as well : 
## Direct Ownership :
def calculate_direct_property_ownership(vacancy_rate, capex, cap_rate, market_value, 
                                        property_management_fees, maintenance_repairs, 
                                        property_taxes, insurance, utilities, hoa_fees):
    # 1. Calculate the Gross Rental Income (assuming 100% occupancy)
    gross_rental_income = market_value * cap_rate
    
    # 2. Adjust for vacancy
    effective_rental_income = gross_rental_income * (1 - vacancy_rate)
    
    # 3. Total Operating Expenses
    operating_expenses = (property_management_fees + maintenance_repairs + property_taxes + 
                          insurance + utilities + hoa_fees)
    
    # 4. Net Operating Income (NOI)
    noi = effective_rental_income - operating_expenses
    
    # 5. Capital Expenditures (CapEx)
    # CapEx are large expenses that increase property value but are not part of NOI
    cash_flow_before_financing = noi - capex
    
    # 6. Return on Investment (ROI) assuming market value as initial investment
    roi = (cash_flow_before_financing / market_value) * 100
    
    # Return a dictionary with all key metrics
    return gross_rental_income,effective_rental_income,operating_expenses,noi,cash_flow_before_financing,roi
    # return {
    #     'Gross Rental Income': gross_rental_income,
    #     'Effective Rental Income': effective_rental_income,
    #     'Operating Expenses': operating_expenses,
    #     'Net Operating Income (NOI)': noi,
    #     'Cash Flow Before Financing': cash_flow_before_financing,
    #     'Return on Investment (ROI)': roi
    # }


# # # Updated Local Storage Code :

LOCAL_STORAGE_PATH = "data/orders/"

# @app.route('/order_placed', methods=['POST'])
# def order_placed():
#     try:
#         # Extract data from the request
#         order_data = request.json.get('order_data')
#         client_name = request.json.get('client_name', 'Rohit Sharma')  # Default client name
#         client_id = request.json.get('client_id', 'RS4603')  # Default client ID if not provided
#         funds = request.json.get('funds')  # Example extra data if needed
#         print(f"Received order for client: {client_name} ({client_id}), Available Funds: {funds}")

#         # Local file path for storing orders
#         order_file_path = os.path.join(LOCAL_STORAGE_PATH, f"{client_id}_orders.json")

#         # Load existing data from local storage if available
#         if os.path.exists(order_file_path):
#             with open(order_file_path, 'r') as file:
#                 client_transactions = json.load(file)
#             print(f"Loaded existing transactions for client {client_id}")
#         else:
#             # Initialize a new transaction list if the file doesn't exist
#             client_transactions = []
#             print(f"No existing transactions for client {client_id}. Initializing new list.")

#         # Process Real Estate or other assets based on asset class
#         assetClass = order_data.get('assetClass')
#         print(f"Processing Asset Class: {assetClass}")
        
#         if assetClass == 'Real Estate':
#             ownership = order_data.get('ownership')
#             if ownership in ['REIT/Fund', 'Commercial Real Estate (Triple Net Lease)']:
#                 # Real estate REIT/fund or commercial real estate transaction
#                 new_transaction = {
#                     "AssetClass": assetClass,
#                     "Ownership": ownership,
#                     "Date": order_data.get('date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
#                     "Name": order_data.get('name'),
#                     "TransactionAmount": order_data.get('investmentAmount'),
#                     "DividendYield": order_data.get('dividendYield')
#                 }
#             else:
#                 # Direct real estate transaction
#                 new_transaction = {
#                     "AssetClass": assetClass,
#                     "Ownership": ownership,
#                     "Date": order_data.get('date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
#                     "Name": order_data.get('name'),
#                     "EstimatedAnnualIncome": order_data.get('estimated_annual_income'),
#                     "EstimatedYield": order_data.get('estimated_yield')
#                 }
#         else:
#             # Standard transaction for Stocks, Bonds, etc.
#             new_transaction = {
#                 "Market": order_data.get('market'),
#                 "AssetClass": assetClass,
#                 "Date": order_data.get('date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
#                 "Action": order_data.get('buy_or_sell'),
#                 "Name": order_data.get('name'),
#                 "Symbol": order_data.get('symbol'),
#                 "Units": order_data.get('units'),
#                 "UnitPrice": order_data.get('unit_price'),
#                 "TransactionAmount": order_data.get('transactionAmount')
#             }

#         # Append the new transaction to the client's transaction list
#         client_transactions.append(new_transaction)
#         print(f"Appended transaction for client {client_id}: {new_transaction}")

#         # Save the updated data back to local storage
#         with open(order_file_path, 'w') as file:
#             json.dump(client_transactions, file, indent=4)
#         print(f"Saved updated transactions for client {client_id} in local storage.")

#         return jsonify({"message": "Order placed successfully", "status": 200})

#     except Exception as e:
#         print(f"Error occurred while placing order: {e}")
#         return jsonify({"message": f"Error occurred while placing order: {str(e)}"}), 500


# Using AWS to Place Order :

def create_transaction(order_data, asset_class):
    """Helper function to create a transaction entry based on asset class"""
    if asset_class == 'Real Estate':
        ownership = order_data.get('ownership')
        if ownership in ['REIT/Fund', 'Commercial Real Estate (Triple Net Lease)']:
            # Real estate REIT/fund or commercial real estate transaction
            return {
                "AssetClass": asset_class,
                "Ownership": ownership,
                "Date": order_data.get('date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                "Name": order_data.get('name'),
                "UnitPrice": order_data.get('price'),
                "TransactionAmount": order_data.get('TransactionAmount'),  # "TransactionAmount": order_data.get('investmentAmount'),
                "Action": order_data.get('buy_or_sell'),
                "Units": order_data.get('units'),
                "DividendYield": order_data.get('dividendYield')
            }
        else:
            # Direct real estate transaction
            return {
                "AssetClass": asset_class,
                "Ownership": ownership,
                "Date": order_data.get('date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                "Name": order_data.get('name'),
                "EstimatedAnnualIncome": order_data.get('estimated_annual_income'),
                "EstimatedYield": order_data.get('estimated_yield')
            }
    else:
        # Standard transaction for Stocks, Bonds, etc.
        return {
            "Market": order_data.get('market'),
            "AssetClass": asset_class,
            "Date": order_data.get('date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            "Action": order_data.get('buy_or_sell'),
            "Name": order_data.get('name'),
            "Symbol": order_data.get('symbol'),
            "Units": order_data.get('units'),
            "UnitPrice": order_data.get('unit_price'),
            "TransactionAmount": order_data.get('transactionAmount')
        }
 

def save_data_to_aws(client_id, client_transactions, order_list_key, order_data):
    """Helper function to save data to AWS S3"""
    asset_class = order_data.get('assetClass')
    new_transaction = create_transaction(order_data, asset_class)
 
    # Append the new transaction to the client's transaction list
    client_transactions.append(new_transaction)
    print(f"Appended transaction for client {client_id}: {new_transaction}")
 
    # Save the updated data back to AWS S3
    updated_data = json.dumps(client_transactions, indent=4)
    s3.put_object(Bucket=S3_BUCKET_NAME, Key=order_list_key, Body=updated_data)
    print(f"Saved updated transactions for client {client_id} in S3 bucket.")
 
 
def save_data_to_local(client_id, client_transactions, order_file_path, order_data):
    """Helper function to save data to local storage"""
    asset_class = order_data.get('assetClass')
    new_transaction = create_transaction(order_data, asset_class)
 
    # Append the new transaction to the client's transaction list
    client_transactions.append(new_transaction)
    print(f"Appended transaction for client {client_id}: {new_transaction}")
 
    # Save the updated data back to local storage
    os.makedirs(os.path.dirname(order_file_path), exist_ok=True)
    with open(order_file_path, 'w') as file:
        json.dump(client_transactions, file, indent=4)
    print(f"Saved updated transactions for client {client_id} in local storage.")
 


# Local storage path (if USE_AWS is False)
LOCAL_STORAGE_PATH = "data/orders/"

# new version :

@app.route('/order_placed', methods=['POST'])
def order_placed():
    try:
        # Extract data from the request
        order_data = request.json.get('order_data')
        client_name = request.json.get('client_name', 'Rohit Sharma')  # Default client name
        client_id = request.json.get('client_id', 'RS4603')  # Default client ID if not provided
        funds = request.json.get('funds')  # Example extra data if needed
        print(f"Received order for client: {client_name} ({client_id}), Available Funds: {funds}")
 
        # Check whether to use AWS or local storage
        if USE_AWS:
            # AWS S3 file key for the order list
            order_list_key = f"{order_list_folder}{client_id}_orders.json"
            # client_summary_key = f"{client_summary_folder}{client_id}_summary.json"
            client_summary_key = f"{client_summary_folder}client-data/{client_id}.json"
           
           
            # Load existing transactions
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=order_list_key)
                client_transactions = json.loads(response['Body'].read().decode('utf-8'))
                print(f"Loaded existing transactions for client {client_id} from S3")
            except s3.exceptions.NoSuchKey:
                client_transactions = []
                print(f"No existing transactions for client {client_id}. Initializing new list.")
           
            # Save new order
            save_data_to_aws(client_id, client_transactions, order_list_key, order_data)
 
            # Update client summary file to set isNewClient to False
            print(f"summary_file_path keys {client_summary_key}")
            try:
                summary_response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=client_summary_key)
                client_summary = json.loads(summary_response['Body'].read().decode('utf-8'))
                print(f"summary_file_path {client_summary}")
                client_summary['isNewClient'] = False  # Set isNewClient to False
                s3.put_object(Bucket=S3_BUCKET_NAME, Key=client_summary_key, Body=json.dumps(client_summary))
                print(f"Updated client summary for {client_id} to set isNewClient as False in S3")
            except s3.exceptions.NoSuchKey:
                print(f"No summary file found for client {client_id} in S3.")
 
        else:
            # Local storage paths
            order_file_path = os.path.join(LOCAL_STORAGE_PATH, f"{client_id}_orders.json")
            summary_file_path = os.path.join(LOCAL_CLIENT_DATA_FOLDER, f"{client_id}_summary.json")
           
            # Load existing transactions
            if os.path.exists(order_file_path):
                with open(order_file_path, 'r') as file:
                    client_transactions = json.load(file)
                print(f"Loaded existing transactions for client {client_id} from local storage")
            else:
                client_transactions = []
                print(f"No existing transactions for client {client_id}. Initializing new list.")
 
            # Save new order
            save_data_to_local(client_id, client_transactions, order_file_path, order_data)
 
            # Update client summary file to set isNewClient to False
            if os.path.exists(summary_file_path):
                with open(summary_file_path, 'r+') as summary_file:
                    client_summary = json.load(summary_file)
                    client_summary['isNewClient'] = False
                    summary_file.seek(0)
                    summary_file.write(json.dumps(client_summary))
                    summary_file.truncate()
                print(f"Updated client summary for {client_id} to set isNewClient as False in local storage")
            else:
                print(f"No summary file found for client {client_id} in local storage.")
 
        return jsonify({"message": "Order placed successfully", "status": 200})
 
    except Exception as e:
        print(f"Error occurred while placing order: {e}")
        return jsonify({"message": f"Error occurred while placing order: {str(e)}"}), 500
 

# ## Using AWS to Show Order :
@app.route('/show_order_list', methods=['POST'])
def show_order_list():
    try:
        # Get client_id from the request
        client_id = request.json.get('client_id')
 
        if not client_id:
            return jsonify({"message": "Client ID is required", "status": 400})
 
        if USE_AWS:
            # Define the S3 file key for the given client ID
            order_list_key = f"{order_list_folder}{client_id}_orders.json"
            print(f"Fetching transactions for client {client_id} from AWS S3")
 
            try:
                # Fetch the file from the S3 bucket
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=order_list_key)
                file_content = response['Body'].read().decode('utf-8')
 
                # Parse the file content as JSON
                client_transactions = json.loads(file_content)
                print(f"Retrieved transactions for client {client_id}: {client_transactions}")
 
                return jsonify({"transaction_data": client_transactions, "status": 200})
 
            except s3.exceptions.NoSuchKey:
                # Handle case where the file does not exist in S3
                print(f"No transactions found for client ID: {client_id}")
                return jsonify({"message": "No transactions found for the provided client ID", "status": 404})
 
            except Exception as e:
                print(f"Error occurred while fetching data from S3: {e}")
                return jsonify({"message": f"Error occurred while fetching data from S3: {str(e)}"}), 500
 
        else:
            # Local file path for storing orders
            order_file_path = os.path.join(LOCAL_STORAGE_PATH, f"{client_id}_orders.json")
 
            # Check if the order file exists
            if os.path.exists(order_file_path):
                # Load transactions from the local file
                with open(order_file_path, 'r') as file:
                    client_transactions = json.load(file)
                print(f"Retrieved transactions for client {client_id}: {client_transactions}")
                return jsonify({"transaction_data": client_transactions, "status": 200})
            else:
                print(f"No transactions found for client ID: {client_id}")
                return jsonify({"message": "No transactions found for the provided client ID", "status": 404})
 
    except Exception as e:
        print(f"Error occurred while retrieving the order list: {e}")
        return jsonify({"message": f"Error occurred while retrieving order list: {str(e)}"}), 500
    
    
    
# Updated Show Order List for Local Storage :

# @app.route('/show_order_list', methods=['POST'])
# def show_order_list():
#     try:
#         # Get client_id from the request
#         client_id = request.json.get('client_id')

#         if not client_id:
#             return jsonify({"message": "Client ID is required", "status": 400})

#         # Local file path for storing orders
#         order_file_path = os.path.join(LOCAL_STORAGE_PATH, f"{client_id}_orders.json")

#         # Check if the order file exists
#         if os.path.exists(order_file_path):
#             # Load transactions from the local file
#             with open(order_file_path, 'r') as file:
#                 client_transactions = json.load(file)
#             print(f"Retrieved transactions for client {client_id}: {client_transactions}")
#             return jsonify({"transaction_data": client_transactions, "status": 200})
#         else:
#             print(f"No transactions found for client ID: {client_id}")
#             return jsonify({"message": "No transactions found for the provided client ID", "status": 404})

#     except Exception as e:
#         print(f"Error occurred while retrieving the order list: {e}")
#         return jsonify({"message": f"Error occurred while retrieving order list: {str(e)}"}), 500


### Using AWS to show Portfolio of the user :

# Constants for file paths
LOCAL_STORAGE_PATH = "local_data"
ORDER_LIST_PATH = os.path.join(LOCAL_STORAGE_PATH, "orders")
DAILY_CHANGES_PATH = os.path.join(LOCAL_STORAGE_PATH, "daily_changes")
PORTFOLIO_PATH = os.path.join(LOCAL_STORAGE_PATH, "portfolios")

# Ensure directories exist for local storage
os.makedirs(ORDER_LIST_PATH, exist_ok=True)
os.makedirs(DAILY_CHANGES_PATH, exist_ok=True)
os.makedirs(PORTFOLIO_PATH, exist_ok=True)

@app.route('/portfolio', methods=['POST'])
def portfolio():
    try:
        # Extract client ID and current date
        client_id = request.json.get('client_id')
        curr_date = request.json.get('curr_date', datetime.now().strftime('%Y-%m-%d'))

        if not client_id:
            return jsonify({"message": "Client ID is required"}), 400

        # Load client orders
        if USE_AWS:
            # Load orders from AWS S3
            order_list_key = f"{order_list_folder}{client_id}_orders.json"
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=order_list_key)
                client_orders = json.loads(response['Body'].read().decode('utf-8'))
                logging.info(f"Fetched orders for client_id: {client_id}")
            except s3.exceptions.NoSuchKey:
                return jsonify({"message": f"No orders found for client_id: {client_id}"}), 404
            except Exception as e:
                logging.error(f"Error fetching orders from AWS: {e}")
                return jsonify({"message": f"Error fetching orders: {e}"}), 500
        else:
            # Load orders from local storage
            order_file_path = os.path.join(ORDER_LIST_PATH, f"{client_id}_orders.json")
            if not os.path.exists(order_file_path):
                return jsonify({"message": f"No orders found for client_id: {client_id}"}), 404
            with open(order_file_path, 'r') as file:
                client_orders = json.load(file)

        # Initialize portfolio data and metrics
        portfolio_data = []
        portfolio_current_value = 0
        porfolio_daily_change = 0
        portfolio_investment_gain_loss = 0

        # Load or initialize daily changes
        if USE_AWS:
            daily_changes_key = f"{daily_changes_folder}/{client_id}_daily_changes.json"
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=daily_changes_key)
                daily_changes = json.loads(response['Body'].read().decode('utf-8'))
            except s3.exceptions.NoSuchKey:
                daily_changes = {}
            except Exception as e:
                logging.error(f"Error fetching daily changes from AWS: {e}")
                return jsonify({"message": f"Error fetching daily changes: {e}"}), 500
        else:
            daily_changes_file = os.path.join(DAILY_CHANGES_PATH, f"{client_id}_daily_changes.json")
            if os.path.exists(daily_changes_file):
                with open(daily_changes_file, 'r') as file:
                    daily_changes = json.load(file)
            else:
                daily_changes = {}

        # Process client orders
        for order in client_orders:
            asset_class = order.get('AssetClass', 'N/A')
            name = order.get('Name', 'N/A')
            symbol = order.get('Symbol', 'N/A')
            units = order.get('Units', 0)
            bought_price = order.get('UnitPrice', 0)
            transaction_amount = order.get('TransactionAmount', 0)

            # Fetch current stock price
            def fetch_current_stock_price(ticker):
                stock = yf.Ticker(ticker)
                try:
                    current_price = stock.history(period='1d')['Close'].iloc[-1]
                    return current_price
                except Exception as e:
                    logging.error(f"Error fetching stock price for {ticker}: {e}")
                    return 0

            current_price = fetch_current_stock_price(symbol)
            diff_price = current_price - bought_price
            daily_price_change = diff_price
            daily_value_change = daily_price_change * units
            current_value = current_price * units

            # Calculate investment gain/loss and other metrics
            investment_gain_loss = diff_price * units
            investment_gain_loss_per = round((investment_gain_loss / transaction_amount) * 100, 2) if transaction_amount > 0 else 0

            # Append data to portfolio
            portfolio_data.append({
                "assetClass": asset_class,
                "name": name,
                "symbol": symbol,
                "Quantity": units,
                "Delayed_Price": current_price,
                "current_value": current_value,
                "Daily_Price_Change": daily_price_change,
                "Daily_Value_Change": daily_value_change,
                "Amount_Invested_per_Unit": bought_price,
                "Amount_Invested": transaction_amount,
                "Investment_Gain_or_Loss_percentage": investment_gain_loss_per,
                "Investment_Gain_or_Loss": investment_gain_loss,
                "Time_Held": order.get('Date', 'N/A'),
            })

            # Update portfolio metrics
            portfolio_current_value += current_value
            porfolio_daily_change += daily_price_change
            portfolio_investment_gain_loss += investment_gain_loss

        # Calculate daily change percentages
        portfolio_daily_change_perc = round((porfolio_daily_change / portfolio_current_value) * 100, 2) if portfolio_current_value > 0 else 0
        portfolio_investment_gain_loss_perc = round((portfolio_investment_gain_loss / portfolio_current_value) * 100, 4) if portfolio_current_value > 0 else 0

        # Update daily changes for the current date
        daily_changes[curr_date] = {
            "portfolio_current_value": portfolio_current_value,
            "porfolio_daily_change": porfolio_daily_change,
            "portfolio_daily_change_perc": portfolio_daily_change_perc,
            "portfolio_investment_gain_loss": portfolio_investment_gain_loss,
            "portfolio_investment_gain_loss_perc": portfolio_investment_gain_loss_perc,
        }

        # Save daily changes and portfolio data
        if USE_AWS:
            # Save daily changes to AWS
            try:
                s3.put_object(
                    Bucket=S3_BUCKET_NAME,
                    Key=daily_changes_key,
                    Body=json.dumps(daily_changes),
                    ContentType='application/json'
                )
                logging.info(f"Updated daily changes for client_id: {client_id} in AWS.")
            except Exception as e:
                logging.error(f"Error saving daily changes to AWS: {e}")
                return jsonify({"message": f"Error saving daily changes: {e}"}), 500

            # Save portfolio data to AWS
            portfolio_key = f"{portfolio_list_folder}/{client_id}.json"
            try:
                s3.put_object(
                    Bucket=S3_BUCKET_NAME,
                    Key=portfolio_key,
                    Body=json.dumps(portfolio_data),
                    ContentType='application/json'
                )
                logging.info(f"Saved portfolio data for client_id: {client_id} in AWS.")
            except Exception as e:
                logging.error(f"Error saving portfolio data to AWS: {e}")
                return jsonify({"message": f"Error saving portfolio data: {e}"}), 500
        else:
            # Save daily changes locally
            with open(daily_changes_file, 'w') as file:
                json.dump(daily_changes, file, indent=4)

            # Save portfolio data locally
            portfolio_file_path = os.path.join(PORTFOLIO_PATH, f"portfolio_{client_id}.json")
            with open(portfolio_file_path, 'w') as file:
                json.dump(portfolio_data, file, indent=4)

        # Response data
        portfolio_response = {
            "portfolio_current_value": portfolio_current_value,
            "porfolio_daily_change": porfolio_daily_change,
            "portfolio_daily_change_perc": portfolio_daily_change_perc,
            "portfolio_investment_gain_loss": portfolio_investment_gain_loss,
            "portfolio_investment_gain_loss_perc": portfolio_investment_gain_loss_perc,
            "daily_changes": daily_changes,
            "portfolio_data": portfolio_data,
        }

        return jsonify(portfolio_response), 200

    except Exception as e:
        logging.error(f"Error in portfolio: {e}")
        return jsonify({"message": f"Error occurred: {str(e)}"}), 500


# Updated Portfolio List using Local Storage :
@app.route('/download_excel', methods=['GET'])
def download_excel():
    file_path = request.args.get('file_path')
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return jsonify({"message": "File not found"}), 404


## Collect live news for stocks in portfolio :
# Define a function to fetch news for a given query 
def fetch_news(query):
    news_url = f'https://newsapi.org/v2/everything?q={query}&apiKey={NEWS_API_KEY}&pageSize=3'
    news_response = requests.get(news_url)
    
    if news_response.status_code == 200:
        news_data = news_response.json()
        articles = news_data.get('articles', [])
        if articles:
            top_news = "\n\n".join([f"{i+1}. {article['title']} - {article['url']}" for i, article in enumerate(articles)])
        else:
            top_news = "No news articles found."
    else:
        top_news = "Failed to fetch news articles."
    
    return top_news

# Function to collect news for each asset in the portfolio
def collect_portfolio_news(portfolio_data):
    portfolio_news = {}
    
    for asset in portfolio_data:
        asset_class = asset.get("AssetClass", "Unknown")
        name = asset.get("Name", "")
        symbol = asset.get("Symbol", None)
        
        # Generate a news query based on the asset class and name/symbol
        if asset_class == "Stocks" or asset_class == "Bonds":
            query = symbol if symbol else name
        elif asset_class == "cryptocurrency":
            query = asset.get("Name", "")
        elif asset_class == "Real Estate":
            query = asset.get("Name", "")
        else:
            query = asset.get("Name", "")
        
        # Fetch news for the query
        news = fetch_news(query)
        portfolio_news[name] = news
    
    return portfolio_news

# aws method :

# Paths for local storage
LOCAL_STORAGE_PATH = "local_data"
ORDER_LIST_PATH = os.path.join(LOCAL_STORAGE_PATH, "orders")
DAILY_CHANGES_PATH = os.path.join(LOCAL_STORAGE_PATH, "daily_changes")
PORTFOLIO_PATH = os.path.join(LOCAL_STORAGE_PATH, "portfolios")

# Ensure directories exist for local storage
os.makedirs(ORDER_LIST_PATH, exist_ok=True)
os.makedirs(DAILY_CHANGES_PATH, exist_ok=True)
os.makedirs(PORTFOLIO_PATH, exist_ok=True)

@app.route('/analyze_portfolio', methods=['POST'])
def analyze_portfolio():
    try:
        # Retrieve the requested asset type and other input data
        assetName = request.json.get('assetName', 'all')
        client_name = request.json.get('client_name')
        funds = request.json.get('funds')
        client_id = request.json.get('client_id')
        investor_personality = request.json.get('investor_personality', 'Aggressive Investor Personality')

        # Initialize economic news to pass to LLM
        topics = ["rising interest rates", "U.S. inflation", "geopolitical tensions", "US Elections", "Global Wars"]
        economic_news = {topic: fetch_news(topic) for topic in topics}

        # Validate the client_id
        if not client_id:
            return jsonify({'message': 'client_id is required as a query parameter'}), 400

        # Load portfolio data (using local or AWS storage based on USE_AWS)
        if USE_AWS:
            portfolio_key = f"{portfolio_list_folder}/{client_id}.json"
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_key)
                portfolio_data = json.loads(response['Body'].read().decode('utf-8'))
            except s3.exceptions.NoSuchKey:
                return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
        else:
            portfolio_file_path = os.path.join(PORTFOLIO_PATH, f"portfolio_{client_id}.json")
            if not os.path.exists(portfolio_file_path):
                return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
            with open(portfolio_file_path, 'r') as file:
                portfolio_data = json.load(file)

        # Verify portfolio data is a list
        if not isinstance(portfolio_data, list):
            return jsonify({"message": "Portfolio data is not in the expected format"}), 500

        # Initialize variables to calculate portfolio-level metrics
        portfolio_current_value = sum(asset["current_value"] for asset in portfolio_data)
        portfolio_daily_change = sum(asset["Daily_Value_Change"] for asset in portfolio_data)
        portfolio_investment_gain_loss = sum(asset["Investment_Gain_or_Loss"] for asset in portfolio_data)

        if portfolio_current_value != 0:
            portfolio_daily_change_perc = (portfolio_daily_change / portfolio_current_value) * 100
            portfolio_investment_gain_loss_perc = (portfolio_investment_gain_loss / portfolio_current_value) * 100
        else:
            portfolio_daily_change_perc = 0
            portfolio_investment_gain_loss_perc = 0

        # Filter portfolio data if a specific asset type is requested
        if assetName != 'all':
            filtered_portfolio_data = [
                asset for asset in portfolio_data if asset["assetClass"].lower() == assetName.lower()
            ]
        else:
            filtered_portfolio_data = portfolio_data

        # Load client financial data (from AWS or local based on USE_AWS)
        
        if USE_AWS:
            client_data_key = f"{client_summary_folder}client-data/{client_id}.json"
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=client_data_key)
                client_data = json.loads(response['Body'].read().decode('utf-8'))
            except Exception as e:
                logging.error(f"Error occurred while retrieving client data from AWS: {e}")
                return jsonify({'message': f'Error occurred while retrieving client data from S3: {e}'}), 500
        else:
            client_data_file_path = os.path.join("client_data", "client_data", f"{client_id}.json")
            if not os.path.exists(client_data_file_path):
                return jsonify({"message": f"No client data found for client ID: {client_id}"}), 404
            with open(client_data_file_path, 'r') as f:
                client_data = json.load(f)

        portfolio_news = collect_portfolio_news(filtered_portfolio_data)

        task = f"""
                You are the best Stock Market Expert and Portfolio Analyst working for a Wealth Manager on the client: {client_name}.
                The portfolio contains several stocks and investments.
                Based on the portfolio data provided:

                - The available funds for the client are {funds}.
                - The current value of the portfolio is {portfolio_current_value}.
                - The portfolio's daily change is {portfolio_daily_change}.
                - The daily percentage change is {portfolio_daily_change_perc:.2f}%.
                - The total gain/loss in the portfolio is {portfolio_investment_gain_loss}.
                - The percentage gain/loss in the portfolio is {portfolio_investment_gain_loss_perc:.2f}%.
                - The risk tolerance of the client based on their investment personality is {investor_personality}.

                Given the Clients Financial Data: {client_data} determine the Financial Situation based on the Assets,Liabilities and Debts of of the Client as : Stable,Currently Stable or Unstable.
                Based on the Client's Financial Situation and the Client's Financial Goals,
                Provide an in-depth analysis of the portfolio, including an evaluation of performance, suggestions for improvement, 
                and detailed stock recommendations to the Wealth Manager for the client based on the Client's Financial Situation and in order to achive their Financial Goal's and the Client's risk tolerance for the given portfolio : {portfolio_data}
                and top news of each holdings in the portfolio : {portfolio_news} and the economic news of the US Market : {economic_news}

                - If the client has a conservative investment personality, give stocks and low risk assets recommendations that could provide returns with minimal risk.
                - If the client has a moderate investment personality, give stocks and medium risk assets recommendations that could provide returns with a moderate level of risk.
                - If the client has an aggressive investment personality, give stocks,Real Estate,cryptocurrency,or any High Risk High Reward Assets recommendations that could provide higher returns with higher risk. 
                Also, help the Wealth Manager rearrange the funds, including which stocks to sell and when to buy them.

                Provide detailed reasons for each stock recommendation based on the funds available to the client and their investor personality in order for the Client to achive their Financial Goals. Include specific suggestions on handling the portfolio, such as when to buy, when to sell, and in what quantities, to maximize the client's profits. Highlight the strengths and weaknesses of the portfolio, and give an overall performance analysis.
                Given the Clients Financial Data: {client_data} determine the Financial Situation based on the Assets, Liabilities, and Debts of the Client as: Stable, Currently Stable or Unstable irrespective of their current investmnets or remaining funds or current returns.If the user has a good portfolio mention that as well.
                Additionally, provide:

                1. A risk assessment of the current portfolio composition.
                2. Give a proper Analysis and Performance of the current portfolio holdings by considering its current news.
                3. Funds Rearrangement of the portfolio if required and give stocks that would give better returns to the client.
                4. Recommendations for sector allocation to balance risk and return as per the investor personality and suggest stocks accordingly.
                5. Strategies for tax efficiency in the portfolio management.
                6. Insights on market trends and current economic news that could impact the portfolio.
                7. Explain in brief the Contingency plans for different market scenarios (bullish, bearish, and volatile markets) and suggest some stocks/assets and sectors from which the client can benefit .
                8. Explain How the client can achieve their Financial Goals of the client that they have mentioned and whether they can  achieve it/them till the time(if mentioned) they are planning of achieving it/them.

                Ensure the analysis is comprehensive and actionable, helping the Wealth Manager make informed decisions to optimize the client's portfolio.
                Dont give any Disclaimer as you are providing all the information to a Wealth Manager who is a Financial Advisor and has good amount of knowledge and experience in managing Portfolios.
                """
                
        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(task)

            # Process the response from LLM
            html_suggestions = markdown.markdown(response.text)
            format_suggestions = markdown_to_text(html_suggestions)

            # Return the analysis response
            return jsonify({
                "portfolio_current_value": portfolio_current_value,
                "portfolio_daily_change": portfolio_daily_change,
                "portfolio_daily_change_perc": f"{portfolio_daily_change_perc:.2f}%",
                "portfolio_investment_gain_loss": portfolio_investment_gain_loss,
                "portfolio_investment_gain_loss_perc": f"{portfolio_investment_gain_loss_perc:.2f}%",
                "suggestion": format_suggestions,
                "assetClass": assetName
            }), 200

        except Exception as e:
            logging.error(f"Error generating suggestions from LLM: {e}")
            return jsonify({"message": f"Error occurred while analyzing the portfolio: {e}"}), 500

    except Exception as e:
        logging.error(f"Error in analyzing portfolio: {e}")
        return jsonify({"message": f"Error analyzing portfolio: {e}"}), 500



#########################################################################################################################
# Analyzing the Portfolio using Local Storage :
# New Version :
# New Version :

# File paths for local storage
LOCAL_STORAGE_PATH = "local_data"
ORDER_LIST_PATH = os.path.join(LOCAL_STORAGE_PATH, "orders")
DAILY_CHANGES_PATH = os.path.join(LOCAL_STORAGE_PATH, "daily_changes")
PORTFOLIO_PATH = os.path.join(LOCAL_STORAGE_PATH, "portfolios")

# Ensure directories exist
os.makedirs(ORDER_LIST_PATH, exist_ok=True)
os.makedirs(DAILY_CHANGES_PATH, exist_ok=True)
os.makedirs(PORTFOLIO_PATH, exist_ok=True)



#####################################################################################################################
# Actual vs predicted comparison using local storage :


# Local storage directories
BASE_DIR = "local_data"
DAILY_CHANGES_DIR = os.path.join(BASE_DIR, "daily_changes")
PREDICTIONS_DIR = os.path.join(BASE_DIR, "predictions")
COMPARISONS_DIR = os.path.join(BASE_DIR, "comparisons")
PORTFOLIO_DIR = "local_data/portfolios"
# CLIENT_SUMMARY_DIR = os.path.join(BASE_DIR, "client_summary") 

# Ensure all directories exist
os.makedirs(DAILY_CHANGES_DIR, exist_ok=True)
os.makedirs(PREDICTIONS_DIR, exist_ok=True)
os.makedirs(COMPARISONS_DIR, exist_ok=True)
# os.makedirs(CLIENT_SUMMARY_DIR, exist_ok=True)

# Helper: Fetch current date and determine the start of the quarter
def get_start_of_quarter():
    current_date = datetime.now()
    quarter_start_months = [1, 4, 7, 10]  # January, April, July, October
    start_month = quarter_start_months[(current_date.month - 1) // 3]
    return datetime(current_date.year, start_month, 1)

# Helper: Save data to a file
def save_to_file(filepath, data):
    with open(filepath, 'w') as f:
        json.dump(data, f, indent=4)

# Helper: Load data from a file
def load_from_file(filepath):
    if os.path.exists(filepath):
        with open(filepath, 'r') as f:
            return json.load(f)
    return None

# Function to update daily return data
def update_daily_returns(client_id, portfolio_daily_change, current_date):
    daily_changes_file = os.path.join(DAILY_CHANGES_DIR, f"{client_id}_daily_changes.json")
    
    # Load existing data or initialize if not found
    daily_changes = load_from_file(daily_changes_file) or {
        "start_of_quarter": str(get_start_of_quarter()), 
        "daily_returns": []
    }
    
    # Check if today's return exists and update it if necessary
    last_recorded_date = daily_changes["daily_returns"][-1]["date"] if daily_changes["daily_returns"] else None
    if last_recorded_date == current_date:
        if daily_changes["daily_returns"][-1]["price"] != portfolio_daily_change:
            daily_changes["daily_returns"][-1]["price"] = portfolio_daily_change
    else:
        # Add new entry for today's return
        daily_changes["daily_returns"].append({"date": current_date, "price": portfolio_daily_change})
    
    # Save updated daily changes back to file
    save_to_file(daily_changes_file, daily_changes)

# Function to calculate actual returns
def calculate_actual_returns(client_id):
    start_of_quarter = get_start_of_quarter()
    current_date = datetime.now().strftime("%Y-%m-%d")
    daily_changes_file = os.path.join(DAILY_CHANGES_DIR, f"{client_id}_daily_changes.json")
    
    # Load daily changes
    daily_changes = load_from_file(daily_changes_file)
    if not daily_changes:
        return {"message": "No daily return data found for the client."}
    
    # Filter data from the start of the quarter
    quarter_data = [entry for entry in daily_changes["daily_returns"] if datetime.strptime(entry["date"], "%Y-%m-%d") >= start_of_quarter]
    
    # Calculate total and percentage returns
    if quarter_data:
        initial_price = quarter_data[0]["price"]
        final_price = quarter_data[-1]["price"]
        total_return = sum(entry["price"] for entry in quarter_data)
        percentage_return = (final_price - initial_price) / initial_price * 100
        return {"total_return": total_return, "percentage_return": percentage_return, "daily_returns": quarter_data}
    
    return {"message": "No valid data for the current quarter."}


# Actual vs Predicted Using AWS and Local Storage :


# Define directories for local storage
PORTFOLIO_DIR = "local_data/portfolios"
PREDICTIONS_DIR = "local_data/predictions"
COMPARISONS_DIR = "local_data/comparisons"

os.makedirs(PORTFOLIO_DIR, exist_ok=True)
os.makedirs(PREDICTIONS_DIR, exist_ok=True)
os.makedirs(COMPARISONS_DIR, exist_ok=True)


def save_predictions(client_id, current_quarter, refined_line_chart_data):
    """
    Save predictions line chart data to AWS S3 or locally.
    
    Args:
        client_id (str): Unique client identifier.
        current_quarter (str): Current quarter for prediction data.
        refined_line_chart_data (dict): Refined line chart data to save.
    """
    try:
        if USE_AWS:
            # Save predictions to AWS S3
            predictions_key = f"{PREDICTIONS_FOLDER}/{client_id}_{current_quarter}_line_chart.json"
            try:
                s3.put_object(
                    Bucket=S3_BUCKET_NAME,
                    Key=predictions_key,
                    Body=json.dumps(refined_line_chart_data),
                    ContentType='application/json'
                )
                print(f"Saved predictions for client_id: {client_id}, quarter: {current_quarter} in AWS.")
                logging.info(f"Saved predictions for client_id: {client_id}, quarter: {current_quarter} in AWS.")
            except Exception as e:
                logging.error(f"Error saving predictions to AWS: {e}")
                print(f"Error saving predictions to AWS: {e}")
                return {"message": f"Error saving predictions to AWS: {e}"}, 500
        else:
            # Save predictions locally
            prediction_file_path = os.path.join(PREDICTIONS_DIR, f"{client_id}_{current_quarter}_line_chart.json")
            try:
                with open(prediction_file_path, 'w') as file:
                    json.dump(refined_line_chart_data, file, indent=4)
                logging.info(f"Saved predictions for client_id: {client_id}, quarter: {current_quarter} locally.")
            except Exception as e:
                logging.error(f"Error saving predictions locally: {e}")
                return {"message": f"Error saving predictions locally: {e}"}, 500
    except Exception as e:
        logging.error(f"Unexpected error in saving predictions: {e}")
        return {"message": f"Internal server error: {e}"}, 500



# Actual vs Predicted Endpoint

# test version :

# @app.route('/actual_vs_predicted', methods=['POST'])
# def actual_vs_predicted():
#     try:
#         # Retrieve client ID and current portfolio daily change
#         client_id = request.json.get('client_id')
#         portfolio_daily_change = request.json.get('portfolio_daily_change')
#         current_date = datetime.now().strftime("%Y-%m-%d")
        
#         client_name = request.json.get("client_name")
#         funds = request.json.get("funds")
#         investor_personality = request.json.get("investor_personality", "Aggressive Investor Personality")
        
#         # Get current quarter
#         current_quarter = get_current_quarter()

#         # Define file paths and S3 keys
#         predicted_file_path = os.path.join(PREDICTIONS_DIR, f"{client_id}_{current_quarter}_line_chart.json")
#         predicted_s3_key = f"{PREDICTIONS_FOLDER}/{client_id}_{current_quarter}_line_chart.json"
#         portfolio_predictions_key = f"{PREDICTIONS_FOLDER}/{client_id}_{current_quarter}_portfolio.json"

#         # Load previously predicted line chart data
#         # if USE_AWS:
#         #     try:
#         #         response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=predicted_s3_key)
#         #         predicted_line_chart_data = json.loads(response['Body'].read().decode('utf-8'))
#         #         print("\nFound Prediction Line Chart Data \n")
#         #     except s3.exceptions.NoSuchKey:
#         #         # Create Prediction Line Chart as it wasn't created before
#                 # predicted_line_chart_data = create_current_prediction_line_chart(client_id, client_name, funds, investor_personality)
#                 # print("\nSaving the Predictions Line Chart\n")
#                 # save_predictions(client_id, current_quarter, predicted_line_chart_data)
#         # else:
#         #     predicted_line_chart_data = load_from_file(predicted_file_path, predicted_s3_key)
#         #     if not predicted_line_chart_data:
#         #         return jsonify({"message": f"No previous predictions found for this client."}), 404
        
#         predicted_line_chart_data = create_current_prediction_line_chart(client_id, client_name, funds, investor_personality)
#         print("\nSaving the Predictions Line Chart\n")
#         save_predictions(client_id, current_quarter, predicted_line_chart_data)

#         # Fetch and process portfolio data
#         if USE_AWS:
#             portfolio_key = f"{portfolio_list_folder}/{client_id}.json"
#             try:
#                 response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_key)
#                 current_portfolio_data = json.loads(response['Body'].read().decode('utf-8'))
#             except s3.exceptions.NoSuchKey:
#                 return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
            
#             # Load previous portfolio predictions data
#             try:
#                 response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_predictions_key)
#                 previous_portfolio_data = json.loads(response['Body'].read().decode('utf-8'))
#             except s3.exceptions.NoSuchKey:
#                 previous_portfolio_data = None
#         else:
#             portfolio_file = os.path.join(PORTFOLIO_DIR, f"portfolio_{client_id}.json")
#             current_portfolio_data = load_from_file(portfolio_file)
#             if not current_portfolio_data:
#                 return jsonify({"message": f"No portfolio data found for client ID: {client_id}"}), 404
            
#             portfolio_predictions_file = os.path.join(PREDICTIONS_DIR, f"{client_id}_{current_quarter}_portfolio.json")
#             previous_portfolio_data = load_from_file(portfolio_predictions_file)

#         # Check for changes in the portfolio
#         # if current_portfolio_data != previous_portfolio_data:
#         #     print("Portfolio data has changed. Updating predictions.")
#         #     predicted_line_chart_data = create_current_prediction_line_chart(client_id, client_name, funds, investor_personality)

#         #     # Save updated portfolio and predictions
#         #     if USE_AWS:
#         #         s3.put_object(
#         #             Bucket=S3_BUCKET_NAME,
#         #             Key=portfolio_predictions_key,
#         #             Body=json.dumps(current_portfolio_data),
#         #             ContentType='application/json'
#         #         )
#         #         save_predictions(client_id, current_quarter, predicted_line_chart_data)
#         #     else:
#         #         save_to_file(portfolio_predictions_file, current_portfolio_data)
#         #         save_to_file(predicted_file_path, predicted_line_chart_data)
#         # else:
#         #     print("No changes in portfolio. Using existing predictions.")

#         # Process daily changes data
#         daily_changes_key = f"{daily_changes_folder}/{client_id}_daily_changes.json"
#         if USE_AWS:
#             try:
#                 response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=daily_changes_key)
#                 raw_daily_changes_data = json.loads(response['Body'].read().decode('utf-8'))
#                 print(f"Raw Daily Changes Data :\n{raw_daily_changes_data}")
#             except s3.exceptions.NoSuchKey:
#                 logging.warning(f"No daily changes data found for client ID: {client_id} in AWS.")
#                 return jsonify({"message": "No daily changes data found."}), 404
#         else:
#             daily_changes_file = os.path.join(PORTFOLIO_DIR, f"{client_id}_daily_changes.json")
#             if os.path.exists(daily_changes_file):
#                 with open(daily_changes_file, 'r') as file:
#                     raw_daily_changes_data = json.load(file)
#             else:
#                 logging.warning(f"No daily changes data found locally for client ID: {client_id}.")
#                 return jsonify({"message": "No daily changes data found."}), 404

#         # Process daily changes data for only current quarter :
#         daily_changes_data = []
#         current_quarter_date = get_current_quarter_dates()
#         start_date = current_quarter_date[0] #"2025-01-01"  # Define the starting date for actual data
#         print(f"Current Quarter Start Date :{start_date}")
        
#         for timestamp, details in raw_daily_changes_data.items():
#             try:
#                 # Normalize the date
#                 date = datetime.strptime(timestamp.split(',')[0], "%m/%d/%Y").strftime("%Y-%m-%d")

#                 # Safely get the correct daily change value
#                 value = details.get("portfolio_daily_change") or details.get("porfolio_daily_change", 0)

#                 # Append if the value is not zero and is after or on the start date
#                 if value != 0 and date >= start_date:
#                     daily_changes_data.append({"date": date, "value": value})
#             except Exception as e:
#                 logging.warning(f"Skipping malformed entry {timestamp}: {e}")

#         # Remove duplicates, retaining the latest value for each date
#         unique_daily_changes = {}
#         for entry in daily_changes_data:
#             unique_daily_changes[entry["date"]] = entry["value"]

#         # Convert back to a sorted list of values and dates starting from the current date
#         sorted_actual_dates = sorted(unique_daily_changes.keys())
#         actual_line_chart_data = [unique_daily_changes[date] for date in sorted_actual_dates]

#         # Debugging output
#         print("Processed Daily Changes Data:", daily_changes_data)
#         print("Unique Daily Changes:", unique_daily_changes)
#         print("Actual Line Chart Data:", actual_line_chart_data)
#         print("Actual Data Dates:", sorted_actual_dates)

        
#         return jsonify({
#             "client_id": client_id,
#             "comparison_chart_data": {
#                 "actual_dates": sorted_actual_dates,
#                 "actual_values": actual_line_chart_data,
#                 "predicted": predicted_line_chart_data,
#             }
#         }), 200

#     except Exception as e:
#         print(f"Error generating comparison: {e}")
#         return jsonify({"message": f"Error generating comparison: {e}"}), 500

# v-2 : Very Fast and Also checks changes in Portfolio :

@app.route('/actual_vs_predicted', methods=['POST'])
def actual_vs_predicted():
    try:
        # Retrieve client ID and current portfolio daily change
        client_id = request.json.get('client_id')
        portfolio_daily_change = request.json.get('portfolio_daily_change')
        current_date = datetime.now().strftime("%Y-%m-%d")
        
        client_name = request.json.get("client_name")
        funds = request.json.get("funds")
        investor_personality = request.json.get("investor_personality", "Aggressive Investor Personality")
        
        # Get current quarter
        current_quarter = get_current_quarter()

        # Define file paths and S3 keys
        predicted_file_path = os.path.join(PREDICTIONS_DIR, f"{client_id}_{current_quarter}_line_chart.json")
        predicted_s3_key = f"{PREDICTIONS_FOLDER}/{client_id}_{current_quarter}_line_chart.json"
        portfolio_predictions_key = f"{PREDICTIONS_FOLDER}/{client_id}_{current_quarter}_portfolio.json"

        # Load previously predicted line chart data
        if USE_AWS:
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=predicted_s3_key)
                predicted_line_chart_data = json.loads(response['Body'].read().decode('utf-8'))
                print("\nFound Prediction Line Chart Data \n")
            except s3.exceptions.NoSuchKey:
                # Create Prediction Line Chart as it wasn't created before
                predicted_line_chart_data = create_current_prediction_line_chart(client_id, client_name, funds, investor_personality)
                print("\nSaving the Predictions Line Chart\n")
                save_predictions(client_id, current_quarter, predicted_line_chart_data)
        else:
            predicted_line_chart_data = load_from_file(predicted_file_path, predicted_s3_key)
            if not predicted_line_chart_data:
                return jsonify({"message": f"No previous predictions found for this client."}), 404

        # Fetch and process portfolio data
        if USE_AWS:
            portfolio_key = f"{portfolio_list_folder}/{client_id}.json"
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_key)
                current_portfolio_data = json.loads(response['Body'].read().decode('utf-8'))
            except s3.exceptions.NoSuchKey:
                return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
            
            # Load previous portfolio predictions data
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_predictions_key)
                previous_portfolio_data = json.loads(response['Body'].read().decode('utf-8'))
            except s3.exceptions.NoSuchKey:
                previous_portfolio_data = None
        else:
            portfolio_file = os.path.join(PORTFOLIO_DIR, f"portfolio_{client_id}.json")
            current_portfolio_data = load_from_file(portfolio_file)
            if not current_portfolio_data:
                return jsonify({"message": f"No portfolio data found for client ID: {client_id}"}), 404
            
            portfolio_predictions_file = os.path.join(PREDICTIONS_DIR, f"{client_id}_{current_quarter}_portfolio.json")
            previous_portfolio_data = load_from_file(portfolio_predictions_file)

        # Check for changes in the portfolio
        if current_portfolio_data != previous_portfolio_data:
            print("Portfolio data has changed. Updating predictions.")
            predicted_line_chart_data = create_current_prediction_line_chart(client_id, client_name, funds, investor_personality)

            # Save updated portfolio and predictions
            if USE_AWS:
                s3.put_object(
                    Bucket=S3_BUCKET_NAME,
                    Key=portfolio_predictions_key,
                    Body=json.dumps(current_portfolio_data),
                    ContentType='application/json'
                )
                save_predictions(client_id, current_quarter, predicted_line_chart_data)
            else:
                save_to_file(portfolio_predictions_file, current_portfolio_data)
                save_to_file(predicted_file_path, predicted_line_chart_data)
        else:
            print("No changes in portfolio. Using existing predictions.")

        # Process daily changes data
        daily_changes_key = f"{daily_changes_folder}/{client_id}_daily_changes.json"
        if USE_AWS:
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=daily_changes_key)
                raw_daily_changes_data = json.loads(response['Body'].read().decode('utf-8'))
                print(f"Raw Daily Changes Data :\n{raw_daily_changes_data}")
            except s3.exceptions.NoSuchKey:
                logging.warning(f"No daily changes data found for client ID: {client_id} in AWS.")
                return jsonify({"message": "No daily changes data found."}), 404
        else:
            daily_changes_file = os.path.join(PORTFOLIO_DIR, f"{client_id}_daily_changes.json")
            if os.path.exists(daily_changes_file):
                with open(daily_changes_file, 'r') as file:
                    raw_daily_changes_data = json.load(file)
            else:
                logging.warning(f"No daily changes data found locally for client ID: {client_id}.")
                return jsonify({"message": "No daily changes data found."}), 404

        # Process daily changes data for only current quarter :
        daily_changes_data = []
        current_quarter_date = get_current_quarter_dates()
        start_date = current_quarter_date[0] #"2025-01-01"  # Define the starting date for actual data
        print(f"Current Quarter Start Date :{start_date}")
        
        for timestamp, details in raw_daily_changes_data.items():
            try:
                # Normalize the date
                date = datetime.strptime(timestamp.split(',')[0], "%m/%d/%Y").strftime("%Y-%m-%d")

                # Safely get the correct daily change value
                value = details.get("portfolio_daily_change") or details.get("porfolio_daily_change", 0)

                # Append if the value is not zero and is after or on the start date
                if value != 0 and date >= start_date:
                    daily_changes_data.append({"date": date, "value": value})
            except Exception as e:
                logging.warning(f"Skipping malformed entry {timestamp}: {e}")

        # Remove duplicates, retaining the latest value for each date
        unique_daily_changes = {}
        for entry in daily_changes_data:
            unique_daily_changes[entry["date"]] = entry["value"]

        # Convert back to a sorted list of values and dates starting from the current date
        sorted_actual_dates = sorted(unique_daily_changes.keys())
        actual_line_chart_data = [unique_daily_changes[date] for date in sorted_actual_dates]

        # Debugging output
        print("Processed Daily Changes Data:", daily_changes_data)
        print("Unique Daily Changes:", unique_daily_changes)
        print("Actual Line Chart Data:", actual_line_chart_data)
        print("Actual Data Dates:", sorted_actual_dates)

        
        return jsonify({
            "client_id": client_id,
            "comparison_chart_data": {
                "actual_dates": sorted_actual_dates,
                "actual_values": actual_line_chart_data,
                "predicted": predicted_line_chart_data,
            }
        }), 200

    except Exception as e:
        print(f"Error generating comparison: {e}")
        return jsonify({"message": f"Error generating comparison: {e}"}), 500

############################################################################################

# v-1 : correctly working for actual data for the current quarter

# @app.route('/actual_vs_predicted', methods=['POST'])
# def actual_vs_predicted():
#     try:
#         # Retrieve client ID and current portfolio daily change
#         client_id = request.json.get('client_id')
#         portfolio_daily_change = request.json.get('portfolio_daily_change')
#         current_date = datetime.now().strftime("%Y-%m-%d")
        
#         client_name = request.json.get("client_name")
#         funds = request.json.get("funds")
#         investor_personality = request.json.get("investor_personality", "Aggressive Investor Personality")
        
#         # current_quarter = "2025_Q1"
#         current_quarter = get_current_quarter()

#         # Define file paths and S3 keys
#         predicted_file_path = os.path.join(PREDICTIONS_DIR, f"{client_id}_{current_quarter}_line_chart.json")
        
#         predicted_s3_key = f"{PREDICTIONS_FOLDER}/{client_id}_{current_quarter}_line_chart.json"
        
#         # predicted_s3_key = f"predictions/{client_id}_{current_quarter}_line_chart.json"

#         # Load previously predicted line chart data
#         if USE_AWS:
#             try:
#                 response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=predicted_s3_key)
#                 predicted_line_chart_data = json.loads(response['Body'].read().decode('utf-8'))
#                 print("\nFound Prediction Line Chart Data \n")
#             except s3.exceptions.NoSuchKey:
#                 # Create Prediction Line Chart as it wasnt created before:
#                 predicted_line_chart_data = create_current_prediction_line_chart(client_id,client_name,funds,investor_personality)
                
#                 # Save Prediction Line Chart Data :
#                 print("\nSaving the Predictions Line Chart\n")
#                 save_predictions(client_id,current_quarter,predicted_line_chart_data)
#                 # return jsonify({"message": f"Predicted line chart file not found for client ID: {client_id}"}), 404
#         else:
#             predicted_line_chart_data = load_from_file(predicted_file_path, predicted_s3_key)
#             if not predicted_line_chart_data:
#                 return jsonify({"message": f"No previous predictions found for this client."}), 404

#         # Fetch and process portfolio data
#         # Load portfolio data
#         if USE_AWS:
#             # portfolio_key = f"{portfolio_list_folder}{client_id}.json"
#             portfolio_key = f"{portfolio_list_folder}/{client_id}.json"
#             try:
#                 response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_key)
#                 portfolio_data = json.loads(response['Body'].read().decode('utf-8'))
#             except s3.exceptions.NoSuchKey:
#                 return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
#         else:
#             portfolio_file = os.path.join(PORTFOLIO_DIR, f"portfolio_{client_id}.json")
#             portfolio_data = load_from_file(portfolio_file)
#             if not portfolio_data:
#                 return jsonify({"message": f"No portfolio data found for client ID: {client_id}"}), 404

#         # Update daily returns 
#         # daily_changes_key = f"{daily_changes_folder}/{client_id}_daily_changes.json"
        
#         # Load daily changes data
#         daily_changes_key = f"{daily_changes_folder}/{client_id}_daily_changes.json"
#         if USE_AWS:
#             try:
#                 response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=daily_changes_key)
#                 raw_daily_changes_data = json.loads(response['Body'].read().decode('utf-8'))
#                 print(f"Raw Daily Changes Data :\n{raw_daily_changes_data}")
                
#             except s3.exceptions.NoSuchKey:
#                 logging.warning(f"No daily changes data found for client ID: {client_id} in AWS.")
#                 return jsonify({"message": "No daily changes data found."}), 404
#         else:
#             daily_changes_file = os.path.join(PORTFOLIO_DIR, f"{client_id}_daily_changes.json")
#             if os.path.exists(daily_changes_file):
#                 with open(daily_changes_file, 'r') as file:
#                     raw_daily_changes_data = json.load(file)
#             else:
#                 logging.warning(f"No daily changes data found locally for client ID: {client_id}.")
#                 return jsonify({"message": "No daily changes data found."}), 404

#         # Process daily changes data for only current quarter :
#         daily_changes_data = []
#         current_quarter_date = get_current_quarter_dates()
#         start_date = current_quarter_date[0] #"2025-01-01"  # Define the starting date for actual data
#         print(f"Current Quarter Start Date :{start_date}")
        
#         for timestamp, details in raw_daily_changes_data.items():
#             try:
#                 # Normalize the date
#                 date = datetime.strptime(timestamp.split(',')[0], "%m/%d/%Y").strftime("%Y-%m-%d")

#                 # Safely get the correct daily change value
#                 value = details.get("portfolio_daily_change") or details.get("porfolio_daily_change", 0)

#                 # Append if the value is not zero and is after or on the start date
#                 if value != 0 and date >= start_date:
#                     daily_changes_data.append({"date": date, "value": value})
#             except Exception as e:
#                 logging.warning(f"Skipping malformed entry {timestamp}: {e}")

#         # Remove duplicates, retaining the latest value for each date
#         unique_daily_changes = {}
#         for entry in daily_changes_data:
#             unique_daily_changes[entry["date"]] = entry["value"]

#         # Convert back to a sorted list of values and dates starting from the current date
#         sorted_actual_dates = sorted(unique_daily_changes.keys())
#         actual_line_chart_data = [unique_daily_changes[date] for date in sorted_actual_dates]

#         # Debugging output
#         print("Processed Daily Changes Data:", daily_changes_data)
#         print("Unique Daily Changes:", unique_daily_changes)
#         print("Actual Line Chart Data:", actual_line_chart_data)
#         print("Actual Data Dates:", sorted_actual_dates)

        
#         comparison_data = {
#             "actual_dates": sorted_actual_dates,
#             "actual_values": actual_line_chart_data,
#             "predicted": predicted_line_chart_data,
#         }

#         # Process daily changes data : working correctly
#         # daily_changes_data = []

#         # for timestamp, details in raw_daily_changes_data.items():
#         #     try:
#         #         # Normalize the date
#         #         date = datetime.strptime(timestamp.split(',')[0], "%m/%d/%Y").strftime("%Y-%m-%d")

#         #         # Safely get the correct daily change value
#         #         value = details.get("portfolio_daily_change") or details.get("porfolio_daily_change", 0)

#         #         # Append if the value is not zero
#         #         if value != 0:
#         #             daily_changes_data.append({"date": date, "value": value})
#         #     except Exception as e:
#         #         logging.warning(f"Skipping malformed entry {timestamp}: {e}")

#         # # Remove duplicates, retaining the latest value for each date
#         # unique_daily_changes = {}
#         # for entry in daily_changes_data:
#         #     unique_daily_changes[entry["date"]] = entry["value"]

#         # # Convert back to a sorted list of values
#         # actual_line_chart_data = [unique_daily_changes[date] for date in sorted(unique_daily_changes)]

#         # # Debugging output
#         # print("Processed Daily Changes Data:", daily_changes_data)
#         # print("Unique Daily Changes:", unique_daily_changes)
#         # print("Actual Line Chart Data:", actual_line_chart_data)

#         # # Combine actual and predicted data
#         # comparison_data = {
#         #     "actual": actual_line_chart_data,
#         #     "predicted": predicted_line_chart_data
#         # }

#         # Save comparison data
#         # Save line chart data locally or to AWS based on storage flag
#         if USE_AWS:
#             comparison_s3_key = f"comparisons/{client_id}_{current_quarter}_comparison_chart.json"
#             try:
#                 s3.put_object(
#                     Bucket=S3_BUCKET_NAME,
#                     Key=comparison_s3_key,
#                     Body=json.dumps(comparison_data),
#                     ContentType='application/json'
#                 )
#                 print("Saved Comparison Prediction Data to AWS")
#                 logging.info(f"Saved Comparison prediction data to AWS")
                
#             except Exception as e:
#                 logging.error(f"Error saving prediction data to AWS: {e}")
#                 return jsonify({"message": "Error saving prediction data to AWS."}), 500
#         else:
#             comparison_file_path = os.path.join(COMPARISONS_DIR, f"{client_id}_{current_quarter}_comparison_chart.json")
#             save_to_file(comparison_file_path, comparison_data)
#             # comparison_file = os.path.join(COMPARISONS_DIR, f"{client_id}_{current_quarter}_line_chart.json")
#             # save_to_file(comparison_file, comparison_data)

#         # Return the comparison data
#         return jsonify({
#             "client_id": client_id,
#             "comparison_chart_data": comparison_data
#         }), 200

#     except Exception as e:
#         print(f"Error generating comparison: {e}")
#         return jsonify({"message": f"Error generating comparison: {e}"}), 500

#####################################################################################################################

def create_current_prediction_line_chart(client_id,client_name,funds,investor_personality) :
    try:
        # Retrieve client and portfolio details
        client_id = request.json.get("client_id")
        client_name = request.json.get("client_name")
        funds = request.json.get("funds")
        investor_personality = request.json.get("investor_personality", "Aggressive Investor Personality")

        # Load portfolio data (from AWS or local storage)
        if USE_AWS:
            portfolio_key = f"{portfolio_list_folder}/{client_id}.json"
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_key)
                portfolio_data = json.loads(response["Body"].read().decode("utf-8"))
            except s3.exceptions.NoSuchKey:
                return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
        else:
            portfolio_file_path = os.path.join(PORTFOLIO_PATH, f"portfolio_{client_id}.json")
            if not os.path.exists(portfolio_file_path):
                return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
            with open(portfolio_file_path, "r") as file:
                portfolio_data = json.load(file)

        # Load market data for beta calculation
        market_returns = fetch_historical_returns(MARKET_INDEX)

        # Prepare date intervals for the current quarter
        current_quarter = get_current_quarter()
        date_intervals = get_current_quarter_dates()
        print(f"Current Quarter: {current_quarter}")

        confidence_data = []

        # Process each asset in the portfolio
        for asset in portfolio_data:
            ticker = asset.get("symbol")
            if not ticker or ticker == "N/A":
                continue

            # Fetch historical returns
            historical_returns = fetch_historical_returns(ticker)
            if historical_returns.empty:
                asset["volatility"] = 0.8
                asset["sharpe_ratio"] = 0.7
                asset["beta"] = 0.5
                asset["forecasted_returns"] = [0] * FORECAST_DAYS
                asset["simulated_returns"] = [0] * FORECAST_DAYS
                continue

            # Metrics Calculation
            asset["volatility"] = compute_volatility(historical_returns)
            asset["sharpe_ratio"] = compute_sharpe_ratio(historical_returns)
            asset["beta"] = compute_beta(historical_returns, market_returns)
            asset["forecasted_returns"] = arima_forecast(historical_returns).tolist()
            asset["simulated_returns"] = simulate_fluctuations(asset["forecasted_returns"][0], asset["volatility"])
            
        # Load client financial data
        if USE_AWS:
            # client_summary_key = f"{client_summary_folder}{client_id}.json"
            client_summary_key = f"{client_summary_folder}client-data/{client_id}.json"
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=client_summary_key)
                client_financial_data = json.loads(response['Body'].read().decode('utf-8'))
            except Exception as e:
                client_summary_file = os.path.join(CLIENT_SUMMARY_DIR, f"{client_id}.json")
                client_financial_data = load_from_file(client_summary_file)
                if not client_financial_data:
                    return jsonify({"message": f"No client financial data found for client ID in local: {client_id}"}), 404
                logging.error(f"Error retrieving client financial data from AWS,will extract File from Local if Present: {e}")
                # return jsonify({"message": f"No client financial data found for client ID: {client_id}"}), 404
        else:
            client_summary_file = os.path.join(CLIENT_SUMMARY_DIR, f"{client_id}.json")
            client_financial_data = load_from_file(client_summary_file)
            if not client_financial_data:
                return jsonify({"message": f"No client financial data found for client ID: {client_id}"}), 404

        # Initialize economic news to pass to LLM
        topics = ["rising interest rates", "U.S. inflation", "geopolitical tensions", "US Elections", "Global Wars"]
        economic_news = {topic: fetch_news(topic) for topic in topics}
        portfolio_news = collect_portfolio_news(portfolio_data)
        
        # Process daily changes data
        daily_changes_key = f"{daily_changes_folder}/{client_id}_daily_changes.json"
        if USE_AWS:
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=daily_changes_key)
                raw_daily_changes_data = json.loads(response['Body'].read().decode('utf-8'))
                print(f"Raw Daily Changes Data :\n{raw_daily_changes_data}")
            except s3.exceptions.NoSuchKey:
                logging.warning(f"No daily changes data found for client ID: {client_id} in AWS.")
                return jsonify({"message": "No daily changes data found."}), 404
        else:
            daily_changes_file = os.path.join(PORTFOLIO_DIR, f"{client_id}_daily_changes.json")
            if os.path.exists(daily_changes_file):
                with open(daily_changes_file, 'r') as file:
                    raw_daily_changes_data = json.load(file)
            else:
                logging.warning(f"No daily changes data found locally for client ID: {client_id}.")
                return jsonify({"message": "No daily changes data found."}), 404
            
        # Process daily changes data for only current quarter :
        daily_changes_data = []
        current_quarter_date = get_current_quarter_dates()
        start_date = current_quarter_date[0] #"2025-01-01"  # Define the starting date for actual data
        print(f"Current Quarter Start Date :{start_date}")
        
        for timestamp, details in raw_daily_changes_data.items():
            try:
                # Normalize the date
                date = datetime.strptime(timestamp.split(',')[0], "%m/%d/%Y").strftime("%Y-%m-%d")

                # Safely get the correct daily change value
                value = details.get("portfolio_daily_change") or details.get("porfolio_daily_change", 0)

                # Append if the value is not zero and is after or on the start date
                if value != 0 and date >= start_date:
                    daily_changes_data.append({"date": date, "value": value})
            except Exception as e:
                logging.warning(f"Skipping malformed entry {timestamp}: {e}")

        # Remove duplicates, retaining the latest value for each date
        unique_daily_changes = {}
        for entry in daily_changes_data:
            unique_daily_changes[entry["date"]] = entry["value"]

        # Convert back to a sorted list of values and dates starting from the current date
        sorted_actual_dates = sorted(unique_daily_changes.keys())
        actual_line_chart_data = [unique_daily_changes[date] for date in sorted_actual_dates]
        
        # Generate task for LLM with refined prediction ranges
        min_actual_return = min(actual_line_chart_data)
        max_actual_return = max(actual_line_chart_data)
        buffer = 0.05 * (max_actual_return - min_actual_return)


        # Generate task for LLM
        task = f"""
            You are the world's best financial advisor tasked with predicting the current quarter's (3-month) returns for a client's portfolio containing various assets.
            You are the best at Predicting client's portfolio returns and nothing is impossible for you to predict, and you consider all realistic market factors such as:
            - Random market fluctuations (up and down movements).
            - Economic volatility.
            - Natural noise in predictions.
            The client, {client_name}, has the following portfolio:

            Portfolio Details: {portfolio_data}
            Portfolio Analyis: {asset}
            Financial Situation: {client_financial_data}
            Available Funds: ${funds}
            Investor Personality: {investor_personality}
            Portfolio News: {portfolio_news}
            Economic News: {economic_news}
            Portfolio Daily Dates : {sorted_actual_dates}
            Portfolio Daily Returns: {actual_line_chart_data}
                     
            Analyze the portfolio and each assets in the portfolio properly and also refer to the Portfolio news and Economic News for your reference and Performance of the assets.
            Alongside this you are passed with it you may or may not be provided with the actual daily returns of that portfolio.
            If Provided try to align the returns with what the current daily returns are dont give unrealistic returns.If returns are in negative showcase that and give the predictions in that negative range if their returns dont seem to go positive only after you analyze all the information.
            If the Provided Daily Returns are very high then see if they can sustain these returns and try to predict as per the current daily returns after you analyze all the information.
            Based on the given provided information :
            Predict the expected returns (in percentages and dollar amounts) for the overall portfolio at the following dates:
            {date_intervals}

            Predict the portfolio's **daily returns** in this quarter(3 months). Include:
            1. **Best-Case Scenario** (High returns under favorable conditions).
            2. **Worst-Case Scenario** (Low returns under unfavorable conditions).
            3. **Confidence Band** (Range of returns at 95% confidence level).
            
            1. Ensure predicted returns reflect realistic market conditions by keeping.

            2. Avoid predicting sudden, unrealistic spikes or crashes unless explicitly indicated by the actual returns.

            3. Dynamically align predictions based on the latest actual market trends and fluctuations provided in the data set.

            4. Introduce natural noise, but maintain predicted returns within a reasonable range close to actual returns for gradual, smooth portfolio changes.You amy refer to the daily returns so far {raw_daily_changes_data}.

            
            Introduce **realistic daily ups and downs** caused by market conditions and noise to simulate realistic portfolio performance.

            The client, {client_name}, has a portfolio characterized by the following constraints:

            - The actual portfolio daily returns range between {min_actual_return}% and {max_actual_return}%.
            - Best-case scenario returns must not exceed {max_actual_return + 5}% under normal conditions within {buffer}
            - Worst-case scenario returns should not fall below {min_actual_return - 5}% within {buffer}
            - Introduce realistic fluctuations in predictions, but align the trends smoothly with recent market conditions

            Example of simulated_response = 
            ### Response Format:
            | Date       | Best-Case Return (%) | Worst-Case Return (%) | Confidence Band (%) | Total Return (%) |
            |------------|-----------------------|-----------------------|---------------------|------------------|
            | 2025-01-01 | 2.5 | -1.0 | 1.0% - 2.0% | 0.75 |
            | 2025-01-15 | 3.0 | -0.5 | 1.5% - 2.5% | 1.25 |
            | 2025-01-31 | 3.5 | 0.0 | 2.0% - 3.0% | 1.75 |
            | 2025-02-01 | 4.0 | 0.5 | 2.5% - 3.5% | 2.25 |
            | 2025-02-15 | 4.5 | 1.0 | 3.0% - 4.0% | 2.75 |
            | 2025-02-28 | 5.0 | 1.5 | 3.5% - 4.5% | 3.25 |
            | 2025-03-01 | 5.5 | 2.0 | 4.0% - 5.0% | 3.75 |
            | 2025-03-15 | 6.0 | 2.5 | 4.5% - 5.5% | 4.25 |
            | 2025-03-31 | 6.5 | 3.0 | 5.0% - 6.0% | 4.75 |

            
            Your Response must be in the above table format no messages is required just table format data.
        """

        # Simulate LLM response
        model = genai.GenerativeModel("gemini-1.5-flash")
        response = model.generate_content(task)
        simulated_response = markdown_to_text(response.text)
        print(simulated_response)

        # Extract line chart data
        line_chart_data = extract_line_chart_data(simulated_response)
        print(f"Line Chart Data: {line_chart_data}")

        refined_line_chart_data = add_noise(line_chart_data)
        print(f"Refined Line Chart Data: {refined_line_chart_data}")
        
        # Adjust refined predictions to align with actual returns
        refined_line_chart_data = adjust_predictions_to_actual_range(actual_line_chart_data, refined_line_chart_data)
        
        refined_line_chart_data = add_dynamic_fluctuations(refined_line_chart_data, actual_line_chart_data)
        print(f"Adjusted Refined Line Chart Data: {refined_line_chart_data}")

        # Save predictions
        
        # prediction_file = os.path.join(PREDICTIONS_DIR, f"{client_id}_{current_quarter}_line_chart.json")
        # save_to_file(prediction_file, refined_line_chart_data)

        # Return the response
        return refined_line_chart_data
        # return jsonify({
        #     "client_id": client_id,
        #     "client_name": client_name,
        #     "predicted_returns": simulated_response,
        #     "line_chart_data": refined_line_chart_data,
        # }), 200

    except Exception as e:
        print(f"Error in predicting returns: {e}")
        return jsonify({"message": f"Error predicting returns: {e}"}), 500
    


def adjust_predictions_to_actual_range(actual_data, predicted_data):
    """
    Adjust predicted returns to match the range and variability of actual returns,
    while managing extreme losses and unrealistic predictions.
    
    :param actual_data: List of actual portfolio daily returns
    :param predicted_data: Dict containing predicted returns and confidence band data
    :return: Updated predicted_data with adjusted values
    """
    if not actual_data:
        print("No actual data available to adjust predictions.")
        return predicted_data

    # Calculate key statistics from actual data
    actual_mean = np.mean(actual_data)
    actual_std_dev = np.std(actual_data)
    min_actual = min(actual_data)
    max_actual = max(actual_data)

    # Set thresholds to clamp predictions within realistic ranges
    lower_bound = min_actual * 1.2  # 20% buffer for losses
    upper_bound = max_actual * 1.2  # 20% buffer for gains

    def clamp(value):
        return max(lower_bound, min(value, upper_bound))

    # Adjust predicted values using scaling, bias correction, and clamping
    scale_factor = actual_std_dev / np.std(predicted_data["best_case"]) if np.std(predicted_data["best_case"]) > 0 else 1
    bias_adjustment = actual_mean - np.mean(predicted_data["best_case"])

    predicted_data["best_case"] = [clamp((x * scale_factor) + bias_adjustment) for x in predicted_data["best_case"]]
    predicted_data["worst_case"] = [clamp((x * scale_factor) + bias_adjustment) for x in predicted_data["worst_case"]]
    predicted_data["total_returns"]["percentages"] = [clamp((x * scale_factor) + bias_adjustment) for x in predicted_data["total_returns"]["percentages"]]

    # Adjust confidence bands similarly
    predicted_data["confidence_band"] = [
        (clamp((lower * scale_factor) + bias_adjustment), clamp((upper * scale_factor) + bias_adjustment))
        for lower, upper in predicted_data["confidence_band"]
    ]

    print("Predictions adjusted and clamped to manage large losses.")
    return predicted_data

import numpy as np

def add_dynamic_fluctuations(predicted_data, actual_data, fluctuation_factor=0.05):
    """
    Adjusts predicted values by introducing dynamic fluctuations to prevent flat-line behavior.
    
    :param predicted_data: Dict containing predicted best_case, worst_case, and total_returns
    :param actual_data: List of actual daily returns
    :param fluctuation_factor: Factor to control the degree of random fluctuations (default 5%)
    :return: Updated predicted_data with dynamic fluctuations
    """
    def introduce_fluctuation(value):
        noise = np.random.uniform(-fluctuation_factor, fluctuation_factor) * abs(value)
        return value + noise

    # Apply fluctuation to avoid flat-line predictions
    predicted_data["best_case"] = [introduce_fluctuation(x) for x in predicted_data["best_case"]]
    predicted_data["worst_case"] = [introduce_fluctuation(x) for x in predicted_data["worst_case"]]
    predicted_data["total_returns"]["percentages"] = [introduce_fluctuation(x) for x in predicted_data["total_returns"]["percentages"]]

    # Apply fluctuation to confidence bands while maintaining order
    predicted_data["confidence_band"] = [
        (min(introduce_fluctuation(lower), introduce_fluctuation(upper)), 
         max(introduce_fluctuation(lower), introduce_fluctuation(upper)))
        for lower, upper in predicted_data["confidence_band"]
    ]

    return predicted_data

import calendar
from datetime import datetime, timedelta

def get_current_quarter():
    now = datetime.now()
    quarter = (now.month - 1) // 3 + 1
    return f"Q{quarter}-{now.year}"


def get_current_quarter_dates():
    now = datetime.now()
    quarter = (now.month - 1) // 3 + 1
    start_month = 3 * (quarter - 1) + 1
    start_date = datetime(now.year, start_month, 1)
    _, days_in_month = calendar.monthrange(now.year, start_month + 2)  # End of quarter
    end_date = datetime(now.year, start_month + 2, days_in_month)

    dates = [start_date + timedelta(days=i) for i in range(0, (end_date - start_date).days + 1, 7)]
    return [date.strftime("%Y-%m-%d") for date in dates]


# Actual vs Predicted Endpoint : Original Code :

# @app.route('/actual_vs_predicted', methods=['POST'])
# def actual_vs_predicted():
#     try:
#         # Retrieve client ID and current portfolio daily change
#         client_id = request.json.get('client_id')
#         portfolio_daily_change = request.json.get('porfolio_daily_change')
#         current_date = datetime.now().strftime("%Y-%m-%d")

#         current_quarter = "2024_Q4"
        
#         # Load previously predicted line chart data
#         predicted_file = os.path.join(PREDICTIONS_DIR, f"{client_id}_{current_quarter}_line_chart.json")
#         predicted_line_chart_data = load_from_file(predicted_file)
#         if not predicted_line_chart_data:
#             return jsonify({'message': 'No previous predictions found for this client.'}), 404

#         # Fetch and process portfolio data
#         PORTFOLIO_DIR_file = os.path.join(PORTFOLIO_DIR, f"portfolio_{client_id}.json")
#         portfolio_data = load_from_file(PORTFOLIO_DIR_file)
#         if not portfolio_data:
#             return jsonify({'message': 'Portfolio data not found for this client.'}), 404

#         # Update daily returns if there's a change
#         # update_daily_returns(client_id, portfolio_daily_change, current_date)

#         # Calculate actual returns
#         # actual_line_chart_data = calculate_actual_returns(client_id)
        
#         # actual_line_chart_data = [2582.1 - 2209.48 + 2469.66*2 - 4709.36,
#         #                           2199.4 - 2209.48 + 2613.03*2 - 4709.36,
#         #                           2501.9 - 2209.48 + 2517.45*2 - 4709.36,
#         #                           2490.6 - 2209.48 + 2517.45*2 - 4709.36,
#         #                           3225.6 - 2209.48 + 3131.91*2 - 4709.36,
#         #                           3463.1 - 2209.48 + 3705.41*2 - 4709.36,
#         #                           3463.1 - 2209.48 + 3719.06*2 - 4709.36,
#         #                           4191.1 - 2209.48 + 3898.17*2 - 4709.36,
#         #                           ]
        
#         actual_line_chart_data = [602.58,506.62,618.96,606.66,1570.58,3955.08,3982.38,4068.60]
                
                                  
                                
#         # Combine actual and predicted data
#         comparison_data = {
#             "actual": actual_line_chart_data,
#             "predicted": predicted_line_chart_data
#         }

#         # Save comparison data locally
#         comparison_file = os.path.join(COMPARISONS_DIR, f"{client_id}_{current_quarter}_comparison_chart.json")
#         save_to_file(comparison_file, comparison_data)

#         # Return the comparison data
#         return jsonify({
#             "client_id": client_id,
#             "comparison_chart_data": comparison_data
#         }), 200

#     except Exception as e:
#         print(f"Error generating comparison: {e}")
#         return jsonify({"message": f"Error generating comparison: {e}"}), 500


######################################################################################################################
# Portfolio Return on Investment Prediction for next quarter Local Storage:
import os
import json
from flask import Flask, request, jsonify
from datetime import datetime, timedelta
import calendar
import markdown


# Generate next quarter's dates
def get_next_quarter_dates():
    current_date = datetime.now()
    current_month = current_date.month

    # Determine the starting month of the next quarter
    if current_month in [1, 2, 3]:  # Q1
        start_month = 4  # Q2
    elif current_month in [4, 5, 6]:  # Q2
        start_month = 7  # Q3
    elif current_month in [7, 8, 9]:  # Q3
        start_month = 10  # Q4
    else:  # Q4
        start_month = 1  # Q1 of the next year

    # Determine the year of the next quarter
    next_quarter_year = current_date.year if start_month != 1 else current_date.year + 1

    # Generate dates for the next quarter
    next_quarter_dates = []
    for month in range(start_month, start_month + 3):
        # Get the first, 15th, and last day of the month
        first_day = datetime(next_quarter_year, month, 1)
        fifteenth_day = datetime(next_quarter_year, month, 15)
        last_day = datetime(next_quarter_year, month, calendar.monthrange(next_quarter_year, month)[1])

        next_quarter_dates.extend([first_day.strftime("%Y-%m-%d"),
                                   fifteenth_day.strftime("%Y-%m-%d"),
                                   last_day.strftime("%Y-%m-%d")])

    return next_quarter_dates

def get_next_quarter():
    current_date = datetime.now()
    current_month = current_date.month

    # Determine the starting month of the next quarter
    if current_month in [1, 2, 3]:  # Q1
        start_month = 4  # Q2
        next_quarter = str(current_date.year) + "_Q2"
    elif current_month in [4, 5, 6]:  # Q2
        start_month = 7  # Q3
        next_quarter = str(current_date.year) + "_Q3"
    elif current_month in [7, 8, 9]:  # Q3
        start_month = 10  # Q4
        next_quarter = str(current_date.year) + "_Q4"
    else:  # Q4
        start_month = 1  # Q1 of the next year
        next_quarter = str(current_date.year + 1) + "_Q1"
        
    # # Determine the year of the next quarter
    # next_quarter_year = current_date.year if start_month != 1 else current_date.year + 1
    
    return next_quarter
    
    

# Function to save data to a file
def save_to_file(filepath, data):
    with open(filepath, 'w') as f:
        json.dump(data, f, indent=4)

# Function to load data from a file
def load_from_file(filepath):
    if os.path.exists(filepath):
        with open(filepath, 'r') as f:
            return json.load(f)
    return None





#########################################################################################
# # Updated New Implementation Version :

#  getting error

from statsmodels.tsa.stattools import adfuller
from pmdarima import auto_arima


# Constants
FORECAST_DAYS = 63  # 3 months (approx business days)
MARKET_INDEX = '^GSPC'  # S&P 500
# PREDICTIONS_DIR = "predictions/"

# Local storage directories
BASE_DIR = "local_data"
PREDICTIONS_DIR = os.path.join(BASE_DIR, "predictions")
os.makedirs(PREDICTIONS_DIR, exist_ok=True)

# PORTFOLIO_DIR = "portfolios/"
CLIENT_SUMMARY_DIR =  "client_data/client_data" #"client_summary/"

# --- Utility Functions ---


##########################################################################################
# # Prediction Improvements :

# #1. Fetch historical returns for the past 3 months
import yfinance as yf

def fetch_historical_returns(ticker, period='3mo'):
    """Fetch historical returns using yfinance."""
    stock = yf.Ticker(ticker)
    data = stock.history(period=period)
    data['returns'] = data['Close'].pct_change()

    # Clean the data: Drop NaN, inf, or -inf values
    data = data.dropna(subset=['returns'])
    data = data.replace([np.inf, -np.inf], np.nan).dropna()

    # Ensure index frequency for ARIMA
    if data.index.inferred_freq is None:
        data = data.asfreq('B', method='pad')  # Business day frequency

    return data['returns']


# 2. Add Quantitative Metrics :
import numpy as np

def compute_volatility(returns):
    """Calculate standard deviation of returns (volatility)."""
    return np.std(returns)

def compute_sharpe_ratio(returns, risk_free_rate=0.0):
    """Calculate Sharpe Ratio (risk-adjusted return)."""
    mean_return = np.mean(returns)
    std_dev = np.std(returns)
    return (mean_return - risk_free_rate) / std_dev if std_dev != 0 else 0


def compute_beta(asset_returns, market_returns):
    """Calculate Beta (sensitivity to market)."""
    # Align lengths of asset_returns and market_returns
    min_length = min(len(asset_returns), len(market_returns))
    asset_returns = asset_returns.iloc[-min_length:]
    market_returns = market_returns.iloc[-min_length:]

    # Calculate covariance and beta
    covariance = np.cov(asset_returns, market_returns)[0][1]
    market_variance = np.var(market_returns)
    return covariance / market_variance if market_variance != 0 else 0


# # Fetch market index returns (S&P 500)
# market_returns = fetch_historical_returns('^GSPC')


# 3. Add ARIMA Forecasting for Returns
from statsmodels.tsa.arima.model import ARIMA
from statsmodels.tsa.stattools import adfuller
from statsmodels.tools.sm_exceptions import ConvergenceWarning
import warnings

# def check_stationarity(data):
#     """Placeholder function to check if the series is stationary."""
#     from statsmodels.tsa.stattools import adfuller
#     result = adfuller(data)
#     return result[1] <= 0.05  # p-value <= 0.05 implies stationarity

# def arima_forecast(returns, forecast_days=92):
#     """
#     Use ARIMA to forecast future returns with robust error handling.
#     """
#     if not isinstance(returns, pd.Series):
#         returns = pd.Series(returns)

#     # Clean data
#     returns = returns.replace([np.inf, -np.inf], np.nan).dropna()

#     # Ensure enough data points
#     if len(returns) < 10:
#         print("Insufficient data for ARIMA, using mean-based forecast.")
#         return pd.Series([np.mean(returns)] * forecast_days)

#     # Ensure the index has a frequency for time series modeling
#     if returns.index.inferred_freq is None:
#         returns.index = pd.date_range(start=returns.index[0], periods=len(returns), freq='B')

#     # Check stationarity and apply differencing if needed
#     if not check_stationarity(returns):
#         returns = returns.diff().dropna()

#     # Suppress convergence warnings
#     warnings.filterwarnings("ignore", category=ConvergenceWarning)

#     # ARIMA model with error handling
#     try:
#         model = ARIMA(returns, order=(1, 1, 1), enforce_stationarity=False)
#         model_fit = model.fit()
#         forecast = model_fit.forecast(steps=forecast_days)
#         return forecast
#     except Exception as e:
#         print(f"ARIMA failed: {e}")
#         # Fallback: Return constant mean forecast
#         return pd.Series([np.mean(returns)] * forecast_days)


def check_stationarity(series):
    """Perform ADF test to check stationarity."""
    result = adfuller(series)
    return result[1] < 0.05  # Stationary if p-value < 0.05



def arima_forecast(returns, forecast_days=92):
    """Use ARIMA to forecast future returns with error handling."""
    if not isinstance(returns, pd.Series):
        returns = pd.Series(returns)

    # Clean data
    returns = returns.replace([np.inf, -np.inf], np.nan).dropna()

    # Ensure enough data points
    if len(returns) < 10:
        return pd.Series([np.mean(returns)] * forecast_days)

    # Ensure the index has a frequency
    if returns.index.inferred_freq is None:
        returns = returns.asfreq('B')  # Business day frequency

    # Check stationarity and apply differencing if needed
    if not check_stationarity(returns):
        returns = returns.diff().dropna()

    # ARIMA model with error handling
    try:
        model = ARIMA(returns, order=(1, 1, 1), enforce_stationarity=False)
        model_fit = model.fit()
        forecast = model_fit.forecast(steps=forecast_days)
        return forecast
    except Exception as e:
        print(f"ARIMA failed: {e}")
        # Fallback: Return constant mean forecast
        return pd.Series([np.mean(returns)] * forecast_days)


# 4. Simulate Realistic Fluctuations
import random

def adf_test(returns):
    """Perform stationarity test."""
    result = adfuller(returns)
    return "Stationary" if result[1] < 0.05 else "Non-Stationary"

def simulate_fluctuations(base_value, volatility, days=30):
    """Simulate fluctuations based on volatility."""
    simulated = [base_value]
    for _ in range(1, days):
        noise = random.uniform(-volatility, volatility)
        simulated.append(simulated[-1] * (1 + noise))
    return simulated


# 5. Validate Predictions
def validate_predictions(predictions, historical_returns, threshold=0.1):
    """Smooth out predictions exceeding a threshold deviation."""
    validated = []
    # last_known = historical_returns[-1]
    last_known = historical_returns.iloc[-1]
    for pred in predictions:
        if abs(pred - last_known) > threshold:
            pred = (pred + last_known) / 2  # Smooth deviations
        validated.append(pred)
        last_known = pred
    return validated


# 6. Visualize Predictions with Confidence Bands
import matplotlib.pyplot as plt

# #Not that Important to be plotted :
def plot_with_confidence(asset_name, best_case, worst_case):
    """Plot predictions with confidence bands."""
    days = len(best_case)
    dates = [datetime.datetime.today() + datetime.timedelta(days=i) for i in range(days)]
    
    lower_bound = [min(b, w) for b, w in zip(best_case, worst_case)]
    upper_bound = [max(b, w) for b, w in zip(best_case, worst_case)]
    
    plt.figure(figsize=(10, 6))
    plt.fill_between(dates, lower_bound, upper_bound, color='pink', alpha=0.2, label="Confidence Band")
    plt.plot(dates, best_case, color='red', label='Best Case')
    plt.plot(dates, worst_case, color='blue', linestyle='dashed', label='Worst Case')
    plt.title(f"{asset_name} Return Predictions")
    plt.xlabel("Date")
    plt.ylabel("Returns")
    plt.legend()
    plt.show()

import matplotlib.pyplot as plt

# # 7. Extract Line Chart and Plot Return Predictions :

# Extract line chart data from LLM responseimport re
from datetime import datetime
from bs4 import BeautifulSoup
import re

# Line Chart V-3 :
import re


def extract_line_chart_data(response_text):
    """Parses table text into structured line chart data."""
    dates = []
    best_case = []
    worst_case = []
    confidence_band = []
    total_returns = []

    lines = response_text.strip().split("\n")
    
    # Iterate through the table rows starting after the header
    for line in lines[2:]:  # Skip the first two header lines
        # Remove excess whitespace and split the line by "|"
        columns = [col.strip() for col in line.split("|")[1:-1]]  # Ignore leading/trailing "|"
        
        # Check if the line has valid table data
        if len(columns) == 5:
            dates.append(columns[0])
            best_case.append(float(columns[1]))
            worst_case.append(float(columns[2]))
            
            # Extract confidence band (e.g., "0.2% - 1.2%")
            confidence_match = re.match(r"([\d\.\-]+)% - ([\d\.\-]+)%", columns[3])
            if confidence_match:
                confidence_low = float(confidence_match.group(1))
                confidence_high = float(confidence_match.group(2))
                confidence_band.append((confidence_low, confidence_high))
            
            total_returns.append(float(columns[4]))

    return {
        "dates": dates,
        "best_case": best_case,
        "worst_case": worst_case,
        "confidence_band": confidence_band,
        "total_returns": {"percentages": total_returns}
    }
 
def extract_next_quarter_line_chart_data(response_text):
    """
    Parses table text into structured data for monetary-based predictions.
    Ensures the correct handling of best-case, worst-case, and confidence band values.
    """
    dates, best_case, worst_case, confidence_band, total_returns_amounts = [], [], [], [], []

    # Split the response text into lines and process rows after headers
    lines = response_text.strip().split("\n")

    for line in lines[2:]:  # Skip headers
        columns = [col.strip() for col in line.split("|")[1:-1]]  # Remove outer '|'

        if len(columns) == 6:  # Ensure row has required columns
            try:
                dates.append(columns[1])
                best_case.append(float(columns[2].replace("$", "").strip()))
                worst_case.append(float(columns[3].replace("$", "").strip()))
                
                # Extract confidence band values
                confidence_low = float(columns[4].replace("$", "").strip())
                confidence_high = float(columns[5].replace("$", "").strip())
                confidence_band.append((confidence_low, confidence_high))
                
                total_returns_amounts.append(float(columns[5].replace("$", "").strip()))
            except ValueError as e:
                print(f"Skipping row due to value error: {e}")

    # Ensure compatibility with expected output format
    return {
        "dates": dates,
        "best_case": best_case,
        "worst_case": worst_case,
        "confidence_band": confidence_band,
        "total_returns": {"percentages": [], "amounts": total_returns_amounts}
    }



def plot_return_predictions(line_chart_data):
    dates = line_chart_data["dates"]
    best_case = line_chart_data["best_case"]
    worst_case = line_chart_data["worst_case"]
    confidence_band = line_chart_data["confidence_band"]
    total_returns = line_chart_data["total_returns"]["percentages"]

    # Extract confidence intervals
    confidence_low = [low for low, _ in confidence_band]
    confidence_high = [high for _, high in confidence_band]

    plt.figure(figsize=(12, 6))

    # Plot total returns
    plt.plot(dates, total_returns, color="black", label="Overall Returns", linewidth=2)

    # Plot best-case and worst-case lines
    plt.plot(dates, best_case, "r-", label="Best Case")
    plt.plot(dates, worst_case, "b--", label="Worst Case")

    # Plot confidence band
    plt.fill_between(dates, confidence_low, confidence_high, color="pink", alpha=0.5, label="Confidence Band")

    plt.title("Portfolio Return Predictions")
    plt.xlabel("Date")
    plt.ylabel("Returns (%)")
    plt.legend()
    plt.grid()
    plt.xticks(rotation=45)
    plt.tight_layout()
    plt.show()
    
# 8. Add Noise

def add_noise(data, noise_level=0.3):
    """Add random noise to simulate market fluctuations."""
    data_with_noise = {
        "dates": data["dates"],
        "best_case": [x + np.random.normal(0, noise_level) for x in data["best_case"]],
        "worst_case": [x + np.random.normal(0, noise_level) for x in data["worst_case"]],
        "confidence_band": [
            (low + np.random.normal(0, noise_level), high + np.random.normal(0, noise_level))
            for low, high in data["confidence_band"]
        ],
        "total_returns": {
            "percentages": [x + np.random.normal(0, noise_level) for x in data["total_returns"]["percentages"]],
            "amounts": data["total_returns"].get("amounts", [])
        }
    }
    return data_with_noise

# 9. Refine Line Chart Data

def plot_refined_data(refined_data):
    """Plots refined line chart data with best-case, worst-case, and total returns."""
    dates = refined_data['dates']
    best_case = refined_data['best_case']
    worst_case = refined_data['worst_case']
    confidence_band = refined_data['confidence_band']
    total_returns = refined_data['total_returns']['percentages']
    
    # Unpack confidence band into lower and upper bounds
    confidence_lower = [band[0] for band in confidence_band]
    confidence_upper = [band[1] for band in confidence_band]
    
    plt.figure(figsize=(12, 6))
    
    # Plot Best-Case, Worst-Case, and Total Returns
    plt.plot(dates, best_case, label='Best Case (%)', color='green', linestyle='--', marker='o')
    plt.plot(dates, worst_case, label='Worst Case (%)', color='red', linestyle='--', marker='o')
    plt.plot(dates, total_returns, label='Total Returns (%)', color='blue', marker='o')
    
    # Fill Confidence Band
    plt.fill_between(dates, confidence_lower, confidence_upper, color='gray', alpha=0.3, label='Confidence Band')
    
    # Add labels and title
    plt.xlabel('Date')
    plt.ylabel('Returns (%)')
    plt.title('Portfolio Return Predictions with Fluctuations')
    plt.xticks(rotation=45)
    plt.legend()
    plt.tight_layout()
    plt.show()



#################################################################################################




# # --- Flask Endpoint --- #######################################################################################

############################################################################################
# Endpoint to predict returns

# Final Predict Returns for Next Quarter :

# v-2 :
@app.route('/predict_returns', methods=['POST'])
def predict_returns():
    try:
        # Retrieve client and portfolio details
        client_id = request.json.get('client_id')
        client_name = request.json.get('client_name')
        funds = request.json.get('funds')
        investor_personality = request.json.get('investor_personality', 'Aggressive Investor Personality')
        
        # current_quarter = get_current_quarter()
        next_quarter = get_next_quarter()
        
        predicted_file_path = os.path.join(PREDICTIONS_DIR, f"{client_id}_{next_quarter}_line_chart.json")
        next_predicted_s3_key = f"{PREDICTIONS_FOLDER}/{client_id}_{next_quarter}_line_chart.json"
        portfolio_predictions_key = f"{PREDICTIONS_FOLDER}/{client_id}_{next_quarter}_portfolio.json"
        
        simulated_response = None  # Initialize simulated_response in case no change in Portfolio
        
        # Load previously predicted line chart data
        if USE_AWS:
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=next_predicted_s3_key)
                refined_line_chart_data = json.loads(response['Body'].read().decode('utf-8'))
                print("\nFound Prediction Line Chart Data \n")
            except s3.exceptions.NoSuchKey:
                # Create Next Quarter Prediction Line Chart as it wasn't created before
                
                simulated_response,refined_line_chart_data = create_next_quarter_prediction_line_chart(client_id, client_name, funds, investor_personality)
                
                print("\nSaving the Next Quarter Returns Predictions Line Chart\n")
                save_predictions(client_id, next_quarter, refined_line_chart_data)
        else:
            refined_line_chart_data = load_from_file(predicted_file_path, next_predicted_s3_key)
            if not refined_line_chart_data:
                return jsonify({"message": f"No previous predictions found for this client."}), 404

        # Fetch and process portfolio data
        if USE_AWS:
            portfolio_key = f"{portfolio_list_folder}/{client_id}.json"
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_key)
                current_portfolio_data = json.loads(response['Body'].read().decode('utf-8'))
            except s3.exceptions.NoSuchKey:
                return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
            
            # Load previous portfolio predictions data
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_predictions_key)
                previous_portfolio_data = json.loads(response['Body'].read().decode('utf-8'))
            except s3.exceptions.NoSuchKey:
                previous_portfolio_data = None
        else:
            portfolio_file = os.path.join(PORTFOLIO_DIR, f"portfolio_{client_id}.json")
            current_portfolio_data = load_from_file(portfolio_file)
            if not current_portfolio_data:
                return jsonify({"message": f"No portfolio data found for client ID: {client_id}"}), 404
            
            portfolio_predictions_file = os.path.join(PREDICTIONS_DIR, f"{client_id}_{next_quarter}_portfolio.json")
            previous_portfolio_data = load_from_file(portfolio_predictions_file)

        # Check for changes in the portfolio
        # if current_portfolio_data == previous_portfolio_data:
        if current_portfolio_data != previous_portfolio_data:
        
            print("Portfolio data has changed. Updating predictions for next quarter returns.")
            
            simulated_response,refined_line_chart_data = create_next_quarter_prediction_line_chart(client_id, client_name, funds, investor_personality)

            # Save updated portfolio and predictions
            if USE_AWS:
                s3.put_object(
                    Bucket=S3_BUCKET_NAME,
                    Key=portfolio_predictions_key,
                    Body=json.dumps(current_portfolio_data),
                    ContentType='application/json'
                )
                save_predictions(client_id, next_quarter, refined_line_chart_data)
            else:
                save_to_file(portfolio_predictions_file, current_portfolio_data)
                save_to_file(predicted_file_path, refined_line_chart_data)
        else:
            print("No changes in portfolio. Using existing predictions.")


        # Return the response
        return jsonify({
            "client_id": client_id,
            "client_name": client_name,
            "predicted_returns": simulated_response,
            "line_chart_data": refined_line_chart_data
        }), 200

    except Exception as e:
        print(f"Error in predicting returns: {e}")
        return jsonify({"message": f"Error predicting returns: {e}"}), 500
    


# V-2 :

def create_next_quarter_prediction_line_chart(client_id,client_name,funds,investor_personality):
    try:

        # Load portfolio data (using local or AWS storage based on USE_AWS)
        if USE_AWS:
            portfolio_key = f"{portfolio_list_folder}/{client_id}.json"
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_key)
                portfolio_data = json.loads(response['Body'].read().decode('utf-8'))
            except s3.exceptions.NoSuchKey:
                return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
        else:
            portfolio_file_path = os.path.join(PORTFOLIO_PATH, f"portfolio_{client_id}.json")
            if not os.path.exists(portfolio_file_path):
                return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
            with open(portfolio_file_path, 'r') as file:
                portfolio_data = json.load(file)
        
        # portfolio_file = os.path.join(PORTFOLIO_DIR, f"portfolio_{client_id}.json")
        # portfolio_data = load_from_file(portfolio_file)
        # if portfolio_data is None:
        #     return jsonify({"message": f"No portfolio data found for client ID: {client_id}"}), 404

        # Load market data for beta calculation
        market_returns = fetch_historical_returns(MARKET_INDEX)

        # Prepare date intervals
        next_quarter = get_next_quarter()
        print(f"Next Quarter: {next_quarter}")

        confidence_data = []
        
        # Iterate over each asset in the portfolio
        
        for asset in portfolio_data:  # Iterate directly over the list of dictionaries
            ticker = asset.get('symbol')  # Use .get() to safely retrieve the 'symbol' key
            if not ticker:
                continue
            if ticker == 'N/A':
                continue

            # Fetch historical returns
            historical_returns = fetch_historical_returns(ticker)
            if historical_returns.empty:
                print(f"No valid returns for {ticker}. Assigning defaults.")
                asset['volatility'] = 0.8
                asset['sharpe_ratio'] = 0.7
                asset['beta'] = 0.5
                asset['forecasted_returns'] = [0] * FORECAST_DAYS
                asset['simulated_returns'] = [0] * FORECAST_DAYS
                continue

            # Metrics Calculation
            volatility = compute_volatility(historical_returns)
            print(volatility)
            sharpe_ratio = compute_sharpe_ratio(historical_returns)
            print(sharpe_ratio)
            beta = compute_beta(historical_returns, market_returns)
            print(beta)
            stationarity = adf_test(historical_returns)
            print(stationarity)

            # Forecasting
            forecasted_returns = arima_forecast(historical_returns)
            print(forecasted_returns)
            simulated_returns = simulate_fluctuations(forecasted_returns.iloc[0], volatility)
            print(simulated_returns)

            # Save metrics back to the portfolio
            asset['volatility'] = volatility
            asset['sharpe_ratio'] = sharpe_ratio
            asset['beta'] = beta
            asset['stationarity'] = stationarity
            asset['forecasted_returns'] = forecasted_returns.tolist()
            asset['simulated_returns'] = simulated_returns
            
         # Process daily changes data
        daily_changes_key = f"{daily_changes_folder}/{client_id}_daily_changes.json"
        if USE_AWS:
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=daily_changes_key)
                raw_daily_changes_data = json.loads(response['Body'].read().decode('utf-8'))
                print(f"Raw Daily Changes Data :\n{raw_daily_changes_data}")
            except s3.exceptions.NoSuchKey:
                logging.warning(f"No daily changes data found for client ID: {client_id} in AWS.")
                return jsonify({"message": "No daily changes data found."}), 404
        else:
            daily_changes_file = os.path.join(PORTFOLIO_DIR, f"{client_id}_daily_changes.json")
            if os.path.exists(daily_changes_file):
                with open(daily_changes_file, 'r') as file:
                    raw_daily_changes_data = json.load(file)
            else:
                logging.warning(f"No daily changes data found locally for client ID: {client_id}.")
                return jsonify({"message": "No daily changes data found."}), 404
            
        # Process daily changes data for only current quarter :
        daily_changes_data = []
        current_quarter_date = get_current_quarter_dates()
        start_date = current_quarter_date[0] #"2025-01-01"  # Define the starting date for actual data
        print(f"Current Quarter Start Date :{start_date}")
        
        for timestamp, details in raw_daily_changes_data.items():
            try:
                # Normalize the date
                date = datetime.strptime(timestamp.split(',')[0], "%m/%d/%Y").strftime("%Y-%m-%d")

                # Safely get the correct daily change value
                value = details.get("portfolio_daily_change") or details.get("porfolio_daily_change", 0)

                # Append if the value is not zero and is after or on the start date
                if value != 0 and date >= start_date:
                    daily_changes_data.append({"date": date, "value": value})
            except Exception as e:
                logging.warning(f"Skipping malformed entry {timestamp}: {e}")

        # Remove duplicates, retaining the latest value for each date
        unique_daily_changes = {}
        for entry in daily_changes_data:
            unique_daily_changes[entry["date"]] = entry["value"]

        # Convert back to a sorted list of values and dates starting from the current date
        sorted_actual_dates = sorted(unique_daily_changes.keys())
        actual_line_chart_data = [unique_daily_changes[date] for date in sorted_actual_dates]
        
        # Generate task for LLM with refined prediction ranges
        min_actual_return = min(actual_line_chart_data)
        max_actual_return = max(actual_line_chart_data)
        buffer = 0.05 * (max_actual_return - min_actual_return)


        # Load client financial data
        if USE_AWS:
            # client_summary_key = f"{client_summary_folder}{client_id}.json"
            client_summary_key = f"{client_summary_folder}client-data/{client_id}.json"
            try:
                response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=client_summary_key)
                client_financial_data = json.loads(response['Body'].read().decode('utf-8'))
            except Exception as e:
                client_summary_file = os.path.join(CLIENT_SUMMARY_DIR, f"{client_id}.json")
                client_financial_data = load_from_file(client_summary_file)
                if not client_financial_data:
                    return jsonify({"message": f"No client financial data found for client ID in local: {client_id}"}), 404
                logging.error(f"Error retrieving client financial data from AWS,will extract File from Local if Present: {e}")
                # return jsonify({"message": f"No client financial data found for client ID: {client_id}"}), 404
        else:
            client_summary_file = os.path.join(CLIENT_SUMMARY_DIR, f"{client_id}.json")
            client_financial_data = load_from_file(client_summary_file)
            if not client_financial_data:
                return jsonify({"message": f"No client financial data found for client ID: {client_id}"}), 404

        # Initialize economic news to pass to LLM
        topics = ["rising interest rates", "U.S. inflation", "geopolitical tensions", "US Elections", "Global Wars"]
        economic_news = {topic: fetch_news(topic) for topic in topics}
        portfolio_news = collect_portfolio_news(portfolio_data)

        # Generate date intervals for next quarter
        date_intervals = get_next_quarter_dates()
        next_quarter = get_next_quarter()

        task = f"""
            You are the world's best financial advisor tasked with predicting the next quarter's (3-month) returns for a client's portfolio containing various assets.
            You are the best at Predicting client's portfolio returns and nothing is impossible for you to predict, and you consider all realistic market factors such as:
            - Random market fluctuations (up and down movements).
            - Economic volatility.
            - Natural noise in predictions.
            The client, {client_name}, has the following portfolio:

            Portfolio Details: {portfolio_data}
            Portfolio Analyis: {asset}
            Financial Situation: {client_financial_data}
            Available Funds: ${funds}
            Investor Personality: {investor_personality}
            Portfolio News: {portfolio_news}
            Economic News: {economic_news}
            Overall Daily Changes Data : {raw_daily_changes_data}
                     
            Analyze the portfolio and each assets in the portfolio properly and also refer to the Portfolio news and Economic News for your reference and Performance of the assets.
            Predict the expected returns (in percentages and dollar amounts) for the overall portfolio at the following dates:
            {date_intervals}

           Predict the portfolio's **daily returns** in the next quarter(3 months). Include:
            1. **Best-Case Scenario** (High returns under favorable conditions).
            2. **Worst-Case Scenario** (Low returns under unfavorable conditions).
            3. **Confidence Band** (Range of returns at 95% confidence level).
            
            1. Ensure predicted returns reflect realistic market conditions by keeping.

            2. Avoid predicting sudden, unrealistic spikes or crashes unless explicitly indicated by the actual returns.

            3. Dynamically align predictions based on the latest actual market trends and fluctuations provided in the data set.

            4. Introduce natural noise, but maintain predicted returns within a reasonable range close to actual returns for gradual, smooth portfolio changes.You amy refer to the daily returns so far {raw_daily_changes_data}.

            
            Introduce **realistic daily ups and downs** caused by market conditions and noise to simulate realistic portfolio performance.
            Refer to the daily portfolio changes {raw_daily_changes_data} to get estimations of the Projected Portfolios Returns in the Next Quarter.Try to match the Current Daily Returns Range so that the Projections become more accurate and realistic.

            The client, {client_name}, has a portfolio characterized by the following constraints:

            - The actual portfolio daily returns range between {min_actual_return}% and {max_actual_return}%.
            - Best-case scenario returns must not exceed {max_actual_return + 5}% under normal conditions within {buffer}
            - Worst-case scenario returns should not fall below {min_actual_return - 5}% within {buffer}
            - Introduce realistic fluctuations in predictions, but align the trends smoothly with recent market conditions

            Example of simulated_response = 
            ### Response Format:
            | Date       | Best-Case Return (%) | Worst-Case Return (%) | Confidence Band (%) | Total Return (%) |
            |------------|-----------------------|-----------------------|---------------------|------------------|
            | 2025-01-01 | 22.5                  | -1.0                  | 6.0% - 20.0%        | 14.0             |
            | 2025-01-15 | 30.0                  | 5.0                   | 7.5% - 24.0%        | 18.5             |
            | 2025-01-31 | 35.0                  | 11.0                  | 14.0% - 28.0%       | 20.0             |
            | 2025-02-01 | 25.0                  | 3.0                   | -1.8% - 20.0%       | 16.5             |
            | 2025-02-15 | 33.5                  | 6.0                   | 10.0% - 26.0%       | 27.5             |
            | 2025-02-28 | 45.0                  | 11.0                  | 14.0% - 28.0%       | 20.0             |
            | 2025-03-01 | 50.5                  | 12.0                  | 20.0% - 34.0%       | 33.75            |
            | 2025-03-15 | 46.0                  | 8.5                   | 26.5% - 39.0%       | 34.25            |
            | 2025-03-31 | 50.5                  | 11.0                  | 30.0% - 44.0%       | 36.75            |

            
            Your Response must be in the above table format no messages is required just table format data.

            Make Sure all the table contents have values and no null/none/blank value is passed.
        """
            # | Date       | Best-Case Return (%) | Worst-Case Return (%) | Confidence Band (%) | Total Return (%) |
            # |------------|-----------------------|-----------------------|---------------------|------------------|
            # | 2025-01-01 | 22.5                  | -1.0                  | 20.0% - 6.0%        | 14.0             |
            # | 2025-01-15 | 30.0                  | 5.0                   | 24.0% - 7.5%        | 18.5             |
            # | 2025-01-31 | 35.0                  | 11.0                  | 28.0% - 14.0%       | 20.0             |
            # | 2025-02-01 | 25.0                  | 3.0                   | 20.0% - -1.8%       | 16.5             |
            # | 2025-02-15 | 33.5                  | 6.0                   | 26.0% - 10.0%       | 27.5             |
            # | 2025-02-28 | 45.0                  | 11.0                  | 28.0% - 14.0%       | 20.0             |
            # | 2025-03-01 | 50.5                  | 12.0                  | 34.0% - 20.0%       | 33.75            |
            # | 2025-03-15 | 46.0                  | 8.5                   | 39.0% - 26.5%       | 34.25            |
            # | 2025-03-31 | 50.5                  | 11.0                  | 44.0% - 30.0%       | 36.75            |

            # Example of simulated_response = 
            # ### Response Format:
            

            
        # Simulate LLM prediction
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(task)
        simulated_response = markdown_to_text(response.text)
        print(simulated_response)
        
        line_chart_data = extract_line_chart_data(simulated_response)
        # line_chart_data = extract_next_quarter_line_chart_data(simulated_response)
        print(f"\nLine Chart Data :{line_chart_data}")
        
        refined_line_chart_data = add_noise(line_chart_data)
        print(f"\nRefined Line Chart Data :{refined_line_chart_data}")

        # Save predictions
        # save_predictions(client_id,next_quarter,refined_line_chart_data)

        return simulated_response,refined_line_chart_data

    except Exception as e:
        print(f"Error in predicting returns: {e}")
        return jsonify({"message": f"Error predicting returns: {e}"}), 500
        
    
# Gives Percentages : V-1 :
# def create_next_quarter_prediction_line_chart(client_id,client_name,funds,investor_personality):
#     try:

#         # Load portfolio data (using local or AWS storage based on USE_AWS)
#         if USE_AWS:
#             portfolio_key = f"{portfolio_list_folder}/{client_id}.json"
#             try:
#                 response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_key)
#                 portfolio_data = json.loads(response['Body'].read().decode('utf-8'))
#             except s3.exceptions.NoSuchKey:
#                 return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
#         else:
#             portfolio_file_path = os.path.join(PORTFOLIO_PATH, f"portfolio_{client_id}.json")
#             if not os.path.exists(portfolio_file_path):
#                 return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
#             with open(portfolio_file_path, 'r') as file:
#                 portfolio_data = json.load(file)
        
#         # portfolio_file = os.path.join(PORTFOLIO_DIR, f"portfolio_{client_id}.json")
#         # portfolio_data = load_from_file(portfolio_file)
#         # if portfolio_data is None:
#         #     return jsonify({"message": f"No portfolio data found for client ID: {client_id}"}), 404

#         # Load market data for beta calculation
#         market_returns = fetch_historical_returns(MARKET_INDEX)

#         # Prepare date intervals
#         next_quarter = get_next_quarter()
#         print(f"Next Quarter: {next_quarter}")

#         confidence_data = []
        
#         # Iterate over each asset in the portfolio
        
#         for asset in portfolio_data:  # Iterate directly over the list of dictionaries
#             ticker = asset.get('symbol')  # Use .get() to safely retrieve the 'symbol' key
#             if not ticker:
#                 continue
#             if ticker == 'N/A':
#                 continue

#             # Fetch historical returns
#             historical_returns = fetch_historical_returns(ticker)
#             if historical_returns.empty:
#                 print(f"No valid returns for {ticker}. Assigning defaults.")
#                 asset['volatility'] = 0.8
#                 asset['sharpe_ratio'] = 0.7
#                 asset['beta'] = 0.5
#                 asset['forecasted_returns'] = [0] * FORECAST_DAYS
#                 asset['simulated_returns'] = [0] * FORECAST_DAYS
#                 continue

#             # Metrics Calculation
#             volatility = compute_volatility(historical_returns)
#             print(volatility)
#             sharpe_ratio = compute_sharpe_ratio(historical_returns)
#             print(sharpe_ratio)
#             beta = compute_beta(historical_returns, market_returns)
#             print(beta)
#             stationarity = adf_test(historical_returns)
#             print(stationarity)

#             # Forecasting
#             forecasted_returns = arima_forecast(historical_returns)
#             print(forecasted_returns)
#             simulated_returns = simulate_fluctuations(forecasted_returns.iloc[0], volatility)
#             print(simulated_returns)

#             # Save metrics back to the portfolio
#             asset['volatility'] = volatility
#             asset['sharpe_ratio'] = sharpe_ratio
#             asset['beta'] = beta
#             asset['stationarity'] = stationarity
#             asset['forecasted_returns'] = forecasted_returns.tolist()
#             asset['simulated_returns'] = simulated_returns

#         # Load client financial data
#         if USE_AWS:
#             # client_summary_key = f"{client_summary_folder}{client_id}.json"
#             client_summary_key = f"{client_summary_folder}client-data/{client_id}.json"
#             try:
#                 response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=client_summary_key)
#                 client_financial_data = json.loads(response['Body'].read().decode('utf-8'))
#             except Exception as e:
#                 client_summary_file = os.path.join(CLIENT_SUMMARY_DIR, f"{client_id}.json")
#                 client_financial_data = load_from_file(client_summary_file)
#                 if not client_financial_data:
#                     return jsonify({"message": f"No client financial data found for client ID in local: {client_id}"}), 404
#                 logging.error(f"Error retrieving client financial data from AWS,will extract File from Local if Present: {e}")
#                 # return jsonify({"message": f"No client financial data found for client ID: {client_id}"}), 404
#         else:
#             client_summary_file = os.path.join(CLIENT_SUMMARY_DIR, f"{client_id}.json")
#             client_financial_data = load_from_file(client_summary_file)
#             if not client_financial_data:
#                 return jsonify({"message": f"No client financial data found for client ID: {client_id}"}), 404

#         # Initialize economic news to pass to LLM
#         topics = ["rising interest rates", "U.S. inflation", "geopolitical tensions", "US Elections", "Global Wars"]
#         economic_news = {topic: fetch_news(topic) for topic in topics}
#         portfolio_news = collect_portfolio_news(portfolio_data)

#         # Generate date intervals for next quarter
#         date_intervals = get_next_quarter_dates()
#         next_quarter = get_next_quarter()

#         task = f"""
#             You are the world's best financial advisor tasked with predicting the next quarter's (3-month) returns for a client's portfolio containing various assets.
#             You are the best at Predicting client's portfolio returns and nothing is impossible for you to predict, and you consider all realistic market factors such as:
#             - Random market fluctuations (up and down movements).
#             - Economic volatility.
#             - Natural noise in predictions.
#             The client, {client_name}, has the following portfolio:

#             Portfolio Details: {portfolio_data}
#             Portfolio Analyis: {asset}
#             Financial Situation: {client_financial_data}
#             Available Funds: ${funds}
#             Investor Personality: {investor_personality}
#             Portfolio News: {portfolio_news}
#             Economic News: {economic_news}
                     
#             Analyze the portfolio and each assets in the portfolio properly and also refer to the Portfolio news and Economic News for your reference and Performance of the assets.
#             Predict the expected returns (in percentages and dollar amounts) for the overall portfolio at the following dates:
#             {date_intervals}

#             Predict the portfolio's **daily returns** over the next 3 months. Include:
#             1. **Best-Case Scenario** (High returns under favorable conditions).
#             2. **Worst-Case Scenario** (Low returns under unfavorable conditions).
#             3. **Confidence Band** (Range of returns at 95% confidence level).
            
#             Introduce **realistic daily ups and downs** caused by market conditions and noise to simulate realistic portfolio performance.

#             Example of simulated_response = 
#             ### Response Format:
#             | Date       | Best-Case Return (%) | Worst-Case Return (%) | Confidence Band (%) | Total Return (%) |
#             |------------|-----------------------|-----------------------|---------------------|------------------|
#             | 2025-01-01 | 2.5 | -1.0 | 1.0% - 2.0% | 0.75 |
#             | 2025-01-15 | 3.0 | -0.5 | 1.5% - 2.5% | 1.25 |
#             | 2025-01-31 | 3.5 | 0.0 | 2.0% - 3.0% | 1.75 |
#             | 2025-02-01 | 4.0 | 0.5 | 2.5% - 3.5% | 2.25 |
#             | 2025-02-15 | 4.5 | 1.0 | 3.0% - 4.0% | 2.75 |
#             | 2025-02-28 | 5.0 | 1.5 | 3.5% - 4.5% | 3.25 |
#             | 2025-03-01 | 5.5 | 2.0 | 4.0% - 5.0% | 3.75 |
#             | 2025-03-15 | 6.0 | 2.5 | 4.5% - 5.5% | 4.25 |
#             | 2025-03-31 | 6.5 | 3.0 | 5.0% - 6.0% | 4.75 |

            
#             Your Response must be in the above table format no messages is required just table format data.
#         """
        
#         # Simulate LLM prediction
#         model = genai.GenerativeModel('gemini-1.5-flash')
#         response = model.generate_content(task)
#         simulated_response = markdown_to_text(response.text)
#         print(simulated_response)
#         line_chart_data = extract_line_chart_data(simulated_response)
#         print(f"\nLine Chart Data :{line_chart_data}")
        
#         refined_line_chart_data = add_noise(line_chart_data)
#         print(f"\nRefined Line Chart Data :{refined_line_chart_data}")

#         # Save predictions
#         # save_predictions(client_id,next_quarter,refined_line_chart_data)

#         return simulated_response,refined_line_chart_data

#     except Exception as e:
#         print(f"Error in predicting returns: {e}")
#         return jsonify({"message": f"Error predicting returns: {e}"}), 500

# V-1 :
    
# @app.route('/predict_returns', methods=['POST'])
# def predict_returns():
#     try:
#         # Retrieve client and portfolio details
#         client_id = request.json.get('client_id')
#         client_name = request.json.get('client_name')
#         funds = request.json.get('funds')
#         investor_personality = request.json.get('investor_personality', 'Aggressive Investor Personality')

#         # Load portfolio data (using local or AWS storage based on USE_AWS)
#         if USE_AWS:
#             portfolio_key = f"{portfolio_list_folder}/{client_id}.json"
#             try:
#                 response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_key)
#                 portfolio_data = json.loads(response['Body'].read().decode('utf-8'))
#             except s3.exceptions.NoSuchKey:
#                 return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
#         else:
#             portfolio_file_path = os.path.join(PORTFOLIO_PATH, f"portfolio_{client_id}.json")
#             if not os.path.exists(portfolio_file_path):
#                 return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404
#             with open(portfolio_file_path, 'r') as file:
#                 portfolio_data = json.load(file)
        
#         # portfolio_file = os.path.join(PORTFOLIO_DIR, f"portfolio_{client_id}.json")
#         # portfolio_data = load_from_file(portfolio_file)
#         # if portfolio_data is None:
#         #     return jsonify({"message": f"No portfolio data found for client ID: {client_id}"}), 404

#         # Load market data for beta calculation
#         market_returns = fetch_historical_returns(MARKET_INDEX)

#         # Prepare date intervals
#         next_quarter = get_next_quarter()
#         print(f"Next Quarter: {next_quarter}")

#         confidence_data = []
        
#         # Iterate over each asset in the portfolio
        
#         for asset in portfolio_data:  # Iterate directly over the list of dictionaries
#             ticker = asset.get('symbol')  # Use .get() to safely retrieve the 'symbol' key
#             if not ticker:
#                 continue
#             if ticker == 'N/A':
#                 continue

#             # Fetch historical returns
#             historical_returns = fetch_historical_returns(ticker)
#             if historical_returns.empty:
#                 print(f"No valid returns for {ticker}. Assigning defaults.")
#                 asset['volatility'] = 0.8
#                 asset['sharpe_ratio'] = 0.7
#                 asset['beta'] = 0.5
#                 asset['forecasted_returns'] = [0] * FORECAST_DAYS
#                 asset['simulated_returns'] = [0] * FORECAST_DAYS
#                 continue

#             # Metrics Calculation
#             volatility = compute_volatility(historical_returns)
#             print(volatility)
#             sharpe_ratio = compute_sharpe_ratio(historical_returns)
#             print(sharpe_ratio)
#             beta = compute_beta(historical_returns, market_returns)
#             print(beta)
#             stationarity = adf_test(historical_returns)
#             print(stationarity)

#             # Forecasting
#             forecasted_returns = arima_forecast(historical_returns)
#             print(forecasted_returns)
#             simulated_returns = simulate_fluctuations(forecasted_returns.iloc[0], volatility)
#             print(simulated_returns)

#             # Save metrics back to the portfolio
#             asset['volatility'] = volatility
#             asset['sharpe_ratio'] = sharpe_ratio
#             asset['beta'] = beta
#             asset['stationarity'] = stationarity
#             asset['forecasted_returns'] = forecasted_returns.tolist()
#             asset['simulated_returns'] = simulated_returns

#         # Load client financial data
#         if USE_AWS:
#             # client_summary_key = f"{client_summary_folder}{client_id}.json"
#             client_summary_key = f"{client_summary_folder}client-data/{client_id}.json"
#             try:
#                 response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=client_summary_key)
#                 client_financial_data = json.loads(response['Body'].read().decode('utf-8'))
#             except Exception as e:
#                 client_summary_file = os.path.join(CLIENT_SUMMARY_DIR, f"{client_id}.json")
#                 client_financial_data = load_from_file(client_summary_file)
#                 if not client_financial_data:
#                     return jsonify({"message": f"No client financial data found for client ID in local: {client_id}"}), 404
#                 logging.error(f"Error retrieving client financial data from AWS,will extract File from Local if Present: {e}")
#                 # return jsonify({"message": f"No client financial data found for client ID: {client_id}"}), 404
#         else:
#             client_summary_file = os.path.join(CLIENT_SUMMARY_DIR, f"{client_id}.json")
#             client_financial_data = load_from_file(client_summary_file)
#             if not client_financial_data:
#                 return jsonify({"message": f"No client financial data found for client ID: {client_id}"}), 404

#         # Initialize economic news to pass to LLM
#         topics = ["rising interest rates", "U.S. inflation", "geopolitical tensions", "US Elections", "Global Wars"]
#         economic_news = {topic: fetch_news(topic) for topic in topics}
#         portfolio_news = collect_portfolio_news(portfolio_data)

#         # Generate date intervals for next quarter
#         date_intervals = get_next_quarter_dates()
#         next_quarter = get_next_quarter()

#         task = f"""
#             You are the world's best financial advisor tasked with predicting the next quarter's (3-month) returns for a client's portfolio containing various assets.
#             You are the best at Predicting client's portfolio returns and nothing is impossible for you to predict, and you consider all realistic market factors such as:
#             - Random market fluctuations (up and down movements).
#             - Economic volatility.
#             - Natural noise in predictions.
#             The client, {client_name}, has the following portfolio:

#             Portfolio Details: {portfolio_data}
#             Portfolio Analyis: {asset}
#             Financial Situation: {client_financial_data}
#             Available Funds: ${funds}
#             Investor Personality: {investor_personality}
#             Portfolio News: {portfolio_news}
#             Economic News: {economic_news}
                     
#             Analyze the portfolio and each assets in the portfolio properly and also refer to the Portfolio news and Economic News for your reference and Performance of the assets.
#             Predict the expected returns (in percentages and dollar amounts) for the overall portfolio at the following dates:
#             {date_intervals}

#             Predict the portfolio's **daily returns** over the next 3 months. Include:
#             1. **Best-Case Scenario** (High returns under favorable conditions).
#             2. **Worst-Case Scenario** (Low returns under unfavorable conditions).
#             3. **Confidence Band** (Range of returns at 95% confidence level).
            
#             Introduce **realistic daily ups and downs** caused by market conditions and noise to simulate realistic portfolio performance.

#             Example of simulated_response = 
#             ### Response Format:
#             | Date       | Best-Case Return (%) | Worst-Case Return (%) | Confidence Band (%) | Total Return (%) |
#             |------------|-----------------------|-----------------------|---------------------|------------------|
#             | 2025-01-01 | 2.5 | -1.0 | 1.0% - 2.0% | 0.75 |
#             | 2025-01-15 | 3.0 | -0.5 | 1.5% - 2.5% | 1.25 |
#             | 2025-01-31 | 3.5 | 0.0 | 2.0% - 3.0% | 1.75 |
#             | 2025-02-01 | 4.0 | 0.5 | 2.5% - 3.5% | 2.25 |
#             | 2025-02-15 | 4.5 | 1.0 | 3.0% - 4.0% | 2.75 |
#             | 2025-02-28 | 5.0 | 1.5 | 3.5% - 4.5% | 3.25 |
#             | 2025-03-01 | 5.5 | 2.0 | 4.0% - 5.0% | 3.75 |
#             | 2025-03-15 | 6.0 | 2.5 | 4.5% - 5.5% | 4.25 |
#             | 2025-03-31 | 6.5 | 3.0 | 5.0% - 6.0% | 4.75 |

            
#             Your Response must be in the above table format no messages is required just table format data.
#         """
        
#         # Simulate LLM prediction
#         model = genai.GenerativeModel('gemini-1.5-flash')
#         response = model.generate_content(task)
#         simulated_response = markdown_to_text(response.text)
#         print(simulated_response)
#         line_chart_data = extract_line_chart_data(simulated_response)
#         print(f"\nLine Chart Data :{line_chart_data}")
        
#         refined_line_chart_data = add_noise(line_chart_data)
#         print(f"\nRefined Line Chart Data :{refined_line_chart_data}")

#         # Save predictions
#         save_predictions(client_id,next_quarter,refined_line_chart_data)

#         # Return the response
#         return jsonify({
#             "client_id": client_id,
#             "client_name": client_name,
#             "predicted_returns": simulated_response,
#             "line_chart_data": refined_line_chart_data
#         }), 200

#     except Exception as e:
#         print(f"Error in predicting returns: {e}")
#         return jsonify({"message": f"Error predicting returns: {e}"}), 500
    
    
######################################################################################################################
###################################             Dashboard Analysis ###################################################

# # Helper Functions :

# Get top and low performing Portfolio : 

def get_top_low_portfolios(insights):
    """
    Generate tables for the top 10 high-performing and bottom 10 low-performing portfolios.

    Args:
        insights (list): List of client insights containing net worth.

    Returns:
        dict: Tables for top and low performers.
    """
    # Sort clients by net worth
    sorted_insights = sorted(insights, key=lambda x: x["net_worth"], reverse=True)

    # Top 10 performers
    top_performers = sorted_insights[:10]
    top_table = [
        {
            "Rank": i + 1,
            "Client Name": client["client_name"],
            "Net Worth": client["net_worth"]
        }
        for i, client in enumerate(top_performers)
    ]

    # Bottom 10 performers
    low_performers = sorted_insights[-10:]
    low_table = [
        {
            "Rank": len(sorted_insights) - i,
            "Client Name": client["client_name"],
            "Net Worth": client["net_worth"]
        }
        for i, client in enumerate(reversed(low_performers))
    ]

    return {
        "top_performers": top_table,
        "low_performers": low_table
    }

# # Get Top Low Performance API :

#  # Best version :

@app.route('/get_top_low_performers', methods=['POST'])
def get_top_low_performers_api():
    try:
        # Fetch all clients' financial data
        clients = request.json.get("data")
        if not clients:
            return jsonify({"message": "No client data provided"}), 400

        client_performance = []
        for client in clients:
            client_id = client.get("uniqueId")
            if not client_id:
                continue

            # Load client financial data (from AWS or local based on USE_AWS)
            if USE_AWS:
                client_data_key = f"{client_summary_folder}client-data/{client_id}.json"
                try:
                    response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=client_data_key)
                    client_data = json.loads(response['Body'].read().decode('utf-8'))
                except Exception as e:
                    logging.error(f"Error occurred while retrieving client data from AWS: {e}")
                    continue
            else:
                client_data_file_path = os.path.join("client_data", "client_data", f"{client_id}.json")
                if not os.path.exists(client_data_file_path):
                    continue
                with open(client_data_file_path, 'r') as f:
                    client_data = json.load(f)

            investment_personality = client_data.get("investment_personality", "Unknown")
            
            # Load portfolio data
            if USE_AWS:
                portfolio_key = f"{portfolio_list_folder}/{client_id}.json"
                try:
                    response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=portfolio_key)
                    portfolio_data = json.loads(response['Body'].read().decode('utf-8'))
                except s3.exceptions.NoSuchKey:
                    continue
            else:
                portfolio_file_path = os.path.join(PORTFOLIO_PATH, f"portfolio_{client_id}.json")
                if not os.path.exists(portfolio_file_path):
                    continue
                with open(portfolio_file_path, 'r') as file:
                    portfolio_data = json.load(file)

            # Parse numeric values safely
            funds = float(client_data.get("investmentAmount", 0))  # Convert to float
            if funds == 0:
                continue

            invested_amount = sum(float(asset.get("Amount_Invested", 0)) for asset in portfolio_data)
            available_funds = funds - invested_amount

            # Calculate portfolio performance
            portfolio_current_value = sum(float(asset.get("current_value", 0)) for asset in portfolio_data)
            portfolio_investment_gain_loss = sum(float(asset.get("Investment_Gain_or_Loss", 0)) for asset in portfolio_data)

            if portfolio_current_value != 0:
                portfolio_investment_gain_loss_perc = (portfolio_investment_gain_loss / portfolio_current_value) * 100
            else:
                portfolio_investment_gain_loss_perc = 0

            # Calculate high asset class
            asset_class_returns = {}
            for asset in portfolio_data:
                asset_class = asset.get("assetClass", "Unknown")
                asset_return = float(asset.get("Investment_Gain_or_Loss", 0))
                if asset_class != "Unknown":
                    asset_class_returns[asset_class] = asset_class_returns.get(asset_class, 0) + asset_return
                else:
                    print(f"Missing asset_class for asset: {asset}")  # Debug log

            # Log asset_class_returns for debugging
            print(f"Asset class returns for client {client_id}: {asset_class_returns}")

            # Find the asset class with the highest return
            if asset_class_returns:
                high_asset_class = max(asset_class_returns, key=asset_class_returns.get)
            else:
                high_asset_class = "Unknown"

            # Append client performance
            client_performance.append({
                "client_id": client_id,
                "client_name": client_data["clientDetail"]["clientName"],
                "funds": funds,
                "investment_personality": investment_personality,
                "available_funds": available_funds,
                "invested_amount": invested_amount,
                "portfolio_current_value": portfolio_current_value,
                "portfolio_investment_gain_loss_perc": portfolio_investment_gain_loss_perc,
                "portfolio_investment_gain_loss": portfolio_investment_gain_loss,
                "high_asset_class": high_asset_class
            })

        # Sort clients by portfolio investment gain/loss percentage
        client_performance.sort(key=lambda x: x["portfolio_investment_gain_loss_perc"], reverse=True)

        # Get top 10 and low 10 performers
        top_performers = client_performance[:10]
        low_performers = client_performance[::-1][:10]  # Reverse to get bottom 10

        return jsonify({
            "message": "Top and Low Performers Retrieved Successfully",
            "client_performance": client_performance,
            "top_performers": top_performers,
            "low_performers": low_performers
        }), 200

    except Exception as e:
        print(f"Error in get_top_low_performers_api: {e}")
        return jsonify({"message": f"Error in get_top_low_performers_api: {e}"}), 500


#########################################################################################################
# # Generate Inforgraphics for Dashboard :

# fetch portfolio data quickly :

# # V-1 :

from concurrent.futures import ThreadPoolExecutor

def fetch_consolidated_portfolio(client_ids):
    """
    Fetch consolidated portfolio data for multiple clients from AWS S3 bucket.
    Returns a dictionary where keys are client IDs and values are portfolio data.
    """
    portfolios = {}
    try:
        # List all objects in the portfolio folder
        response = s3.list_objects_v2(Bucket=S3_BUCKET_NAME, Prefix=portfolio_list_folder)
        if "Contents" not in response:
            print("No portfolios found in the specified folder.")
            return portfolios

        # Filter objects for the requested client IDs
        portfolio_keys = [
            obj["Key"] for obj in response["Contents"]
            if obj["Key"].split("/")[-1].replace(".json", "").split("//")[-1] in client_ids
        ]

        if not portfolio_keys:
            print("No matching portfolios found for the given client IDs.")
            return portfolios

        print(f"Portfolio keys to fetch: {portfolio_keys}")

        # Function to fetch a single portfolio
        def fetch_portfolio(key):
            try:
                obj = s3.get_object(Bucket=S3_BUCKET_NAME, Key=key)
                client_id = key.split("/")[-1].replace(".json", "")
                data = json.loads(obj["Body"].read().decode("utf-8"))
                print(f"Successfully fetched portfolio for client_id: {client_id}")
                return client_id, data
            except Exception as e:
                print(f"Error fetching portfolio {key}: {e}")
                return None, None

        # Use ThreadPoolExecutor for parallel fetching
        with ThreadPoolExecutor(max_workers=10) as executor:
            results = list(executor.map(fetch_portfolio, portfolio_keys))

        # Consolidate results into a dictionary
        portfolios = {client_id: data for client_id, data in results if client_id and data}

        print(f"Consolidated Portfolios:\n{portfolios}")

        return portfolios

    except Exception as e:
        print(f"Error fetching portfolios from S3: {e}")
        return portfolios



# API to get the asset-class infographics :

# Best Version :

@app.route('/asset_class_infographics', methods=['POST'])
def asset_class_infographics():
    try:
        # Fetch all clients' financial data
        clients = request.json.get("data")
        if not clients:
            return jsonify({"message": "No client data provided"}), 400

        # Create a dictionary for quick lookup of client data by uniqueId
        client_map = {client.get("uniqueId"): client for client in clients if client.get("uniqueId")}
        client_ids = list(client_map.keys())
        if not client_ids:
            return jsonify({"message": "No valid client IDs found"}), 400

        # Fetch consolidated portfolio data
        portfolios = fetch_consolidated_portfolio(client_ids)
        if not portfolios:
            return jsonify({"message": "No portfolio data found for provided clients"}), 404

        # Initialize asset class data
        asset_class_data = {}

        # Process portfolios for asset class aggregation
        for client_id, portfolio in portfolios.items():
            # Handle portfolio as a list of assets (fallback for invalid structure)
            if isinstance(portfolio, list):
                print(f"Portfolio for client_id {client_id} is a list. Wrapping it in a dictionary.")
                portfolio = {"assets": portfolio}

            # Ensure portfolio is a dictionary
            if not isinstance(portfolio, dict):
                print(f"Invalid portfolio format for client_id: {client_id}")
                continue

            # Safely get assets from the portfolio
            assets = portfolio.get("assets", [])
            if not isinstance(assets, list):
                print(f"'assets' is not a list for client_id: {client_id}")
                continue

            print(f"Processing assets for client_id {client_id}: {assets}")

            for asset in assets:
                # Debugging: Print each asset being processed
                print(f"Processing asset for client {client_id}: {asset}")

                asset_class = asset.get("assetClass", "Other").capitalize()
                asset_class = "ETF" if asset_class == "Etf" else asset_class

                if asset_class not in asset_class_data:
                    asset_class_data[asset_class] = {
                        "clients": set(),  # Use a set to prevent duplicate client counts
                        "total_invested": 0,
                        "total_returns": 0
                    }

                # Update asset class data
                invested_amount = float(asset.get("Amount_Invested", 0))
                returns = float(asset.get("Investment_Gain_or_Loss", 0))

                asset_class_data[asset_class]["clients"].add(client_id)  # Add client ID to the set
                asset_class_data[asset_class]["total_invested"] += invested_amount
                asset_class_data[asset_class]["total_returns"] += returns

        # Convert sets to counts and prepare final asset class data
        for asset_class, data in asset_class_data.items():
            data["num_clients"] = len(data["clients"])  # Count the number of unique clients
            data["clients"] = list(data["clients"])  # Convert back to a list if needed

        # Debugging: Print the aggregated asset_class_data
        print(f"Aggregated Asset Class Data: {asset_class_data}")

        # Prepare pie chart data
        pie_chart_data = {
            "labels": list(asset_class_data.keys()),
            "datasets": [{
                "data": [data["total_invested"] for data in asset_class_data.values()],
                "label": "Total Invested"
            }]
        }
        print(f"Pie Chart Data: {pie_chart_data}")

        # Format the response for the frontend
        response = {
            "message": "Asset Class Infographics generated successfully",
            "pie_chart_data": pie_chart_data,
            "details": asset_class_data
        }

        return jsonify(response), 200

    except Exception as e:
        print(f"Error in asset_class_infographics: {e}")
        return jsonify({"message": f"Error in asset_class_infographics: {e}"}), 500


#########################################################################################################

# Asset Class Table :

# Best Version :

@app.route('/get_best_performing_assets', methods=['POST'])
def get_best_performing_assets_api():
    try:
        # Fetch all clients' financial data
        request_data = request.json
        clients = request_data.get("data")
        asset_class = request_data.get("asset_class")
        asset_class = "etf" if asset_class == "ETF" else asset_class
        
        if not clients:
            return jsonify({"message": "No client data provided"}), 400
        if not asset_class:
            return jsonify({"message": "No asset class provided"}), 400

        # Create a map for client data by uniqueId
        client_map = {client.get("uniqueId"): client for client in clients if client.get("uniqueId")}
        client_ids = list(client_map.keys())
        if not client_ids:
            return jsonify({"message": "No valid client IDs found"}), 400

        # Fetch consolidated portfolio data
        portfolios = fetch_consolidated_portfolio(client_ids)
        if not portfolios:
            return jsonify({"message": "No portfolio data found for provided clients"}), 404

        # Initialize best-performing assets list
        best_performing_assets = []

        for client_id, portfolio in portfolios.items():
            # Handle portfolio as a list of assets (fallback for invalid structure)
            if isinstance(portfolio, list):
                print(f"Portfolio for client_id {client_id} is a list. Wrapping it in a dictionary.")
                portfolio = {"assets": portfolio}

            # Ensure portfolio is a dictionary
            if not isinstance(portfolio, dict):
                print(f"Invalid portfolio format for client_id: {client_id}")
                continue

            # Safely get assets from the portfolio
            assets = portfolio.get("assets", [])
            if not isinstance(assets, list):
                print(f"'assets' is not a list for client_id: {client_id}")
                continue

            for asset in assets:
                # Check if the asset belongs to the selected asset class
                if asset.get("assetClass") != asset_class:
                    continue

                # Extract required fields and calculate returns percentage
                invested_amount = float(asset.get("Amount_Invested", 0))
                returns = float(asset.get("Investment_Gain_or_Loss", 0))
                returns_perc = (returns / invested_amount * 100) if invested_amount > 0 else 0

                client_data = client_map.get(client_id, {})
                best_performing_assets.append({
                    "Client Id": client_id,
                    "Client Name": client_data.get("clientDetail", {}).get("clientName", "Unknown"),
                    "Funds": float(client_data.get("investmentAmount", 0)),
                    "Invested Amount": invested_amount,
                    "Asset": asset.get("name", "Unknown"),
                    "Returns ($)": returns,
                    "Returns (%)": returns_perc
                })

        # Sort assets by returns percentage in descending order
        best_performing_assets.sort(key=lambda x: x["Returns (%)"], reverse=True)

        # Format the response
        response = {
            "message": f"Top assets for asset class '{asset_class}' retrieved successfully",
            "asset_class": asset_class,
            "performance_table": best_performing_assets
        }
        
        print(f"Top assets for asset class :{asset_class}\n{best_performing_assets}")

        return jsonify(response), 200

    except Exception as e:
        print(f"Error in get_best_performing_assets_api: {e}")
        return jsonify({"message": f"Error in get_best_performing_assets_api: {e}"}), 500


##########################################################################################################
#########################################################################################################

# Analyze Dashboard :

def fetch_portfolios(client_ids):
    """
    Fetch consolidated portfolio data for multiple clients from AWS S3 bucket.
    Returns a dictionary where keys are client IDs and values are portfolio data.
    """
    portfolios = {}
    try:
        # List all objects in the portfolio folder
        response = s3.list_objects_v2(Bucket=S3_BUCKET_NAME, Prefix=portfolio_list_folder)
        if "Contents" not in response:
            print("No portfolios found in the specified folder.")
            return portfolios

        # Filter objects for the requested client IDs
        portfolio_keys = [
            obj["Key"] for obj in response["Contents"]
            if obj["Key"].split("/")[-1].replace(".json", "").split("//")[-1] in client_ids
        ]

        if not portfolio_keys:
            print("No matching portfolios found for the given client IDs.")
            return portfolios

        print(f"Portfolio keys to fetch: {portfolio_keys}")

        # Function to fetch a single portfolio
        def fetch_portfolio(key):
            try:
                obj = s3.get_object(Bucket=S3_BUCKET_NAME, Key=key)
                client_id = key.split("/")[-1].replace(".json", "")
                data = json.loads(obj["Body"].read().decode("utf-8"))
                # print(f"Successfully fetched portfolio for client_id: {client_id}")
                return client_id, data
            except Exception as e:
                print(f"Error fetching portfolio {key}: {e}")
                return None, None

        # Use ThreadPoolExecutor for parallel fetching
        with ThreadPoolExecutor(max_workers=10) as executor:
            results = list(executor.map(fetch_portfolio, portfolio_keys))

        # Consolidate results into a dictionary
        portfolios = {client_id: data for client_id, data in results if client_id and data}

        # print(f"Portfolios Collected :\n{portfolios}")
        
        return portfolios

    except Exception as e:
        print(f"Error fetching portfolios from S3: {e}")
        return portfolios
    
###########################################################################################

#  # v-4 : More insights and metrics :

@app.route('/analyze_dashboard', methods=['POST'])
def analyze_dashboard():
    try:
        # Fetch all clients' financial data
        clients = request.json.get("data")
        if not clients:
            return jsonify({"message": "No client data provided"}), 400

        client_ids = [client.get("uniqueId") for client in clients if client.get("uniqueId")]
        if not client_ids:
            return jsonify({"message": "No valid client IDs found in the request"}), 400

        # call infographics for testing purposes :
        # dashboard_infographics(clients)
        
        # Fetch portfolios for all clients in bulk
        portfolios = fetch_portfolios(client_ids)

        insights = []
        performance_list = []
        stress_test_results = []
        client_goal_progress = {}
        client_rebalancing_recommendations = {}
        
        for client in clients:
            client_id = client.get("uniqueId")
            if not client_id or client_id not in portfolios:
                continue

            # Load client financial data (from AWS or local based on USE_AWS)
            if USE_AWS:
                client_data_key = f"{client_summary_folder}client-data/{client_id}.json"
                try:
                    response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=client_data_key)
                    client_data = json.loads(response['Body'].read().decode('utf-8'))
                except Exception as e:
                    logging.error(f"Error occurred while retrieving client data from AWS: {e}")
                    continue
            else:
                client_data_file_path = os.path.join("client_data", "client_data", f"{client_id}.json")
                if not os.path.exists(client_data_file_path):
                    continue
                with open(client_data_file_path, 'r') as f:
                    client_data = json.load(f)

            # Ensure funds are numeric
            funds = float(client_data.get("investmentAmount", 0) or 0)
            if funds <= 0:
                continue

            # Get portfolio data from the fetched portfolios
            portfolio_data = portfolios.get(client_id, [])

            # Safely parse assets and liabilities
            assets = sum(float(value or 0) for value in client_data["assetsDatasets"]["datasets"][0]["data"])
            liabilities = sum(float(value or 0) for value in client_data["liabilityDatasets"]["datasets"][0]["data"])
            net_worth = assets - liabilities
            
            # Savings Rate
            annual_income = sum(float(inc.get("amountIncome", 0) or 0) for inc in client_data.get("incomeFields", []))
            savings_rate = (funds / annual_income * 100) if annual_income > 0 else 0
            print(f"Savings Rate :{savings_rate}")
            
            # Calculate Monthly Expenses
            liability_monthly_payments = sum(
                float(client_data["myLiabilities"][key].get("mortgageMonthly", 0) or 0)
                for key in client_data["myLiabilities"]
                if "Monthly" in client_data["myLiabilities"][key]
            )

            insurance_monthly_premiums = sum(
                float(client_data["insuranceCoverage"][key].get("monthlyPayLIClient", 0) or 0)
                for key in client_data["insuranceCoverage"]
                if "monthlyPayLIClient" in client_data["insuranceCoverage"][key]
            )

            # Monthly Expenses = Liabilities + Insurance - Savings (Investments)
            monthly_expenses = liability_monthly_payments + insurance_monthly_premiums

            # Use monthly expenses for emergency fund coverage calculation
            emergency_fund_coverage = (liquid_assets / monthly_expenses) if monthly_expenses > 0 else 0

            # Additional Metrics
            debt_to_asset_ratio = (liabilities / assets * 100) if assets > 0 else 0

            investment_personality = client_data.get("investment_personality", "Unknown")
            retirement_age = client_data["retirementGoal"]["retirementPlan"]["retirementAgeClient"]
            retirement_goals = client_data.get("retirementGoal", "Unknown")
            # print(f"Retirement Goal :\n{retirement_goals}")
            financial_goals = client_data.get("goalFields","Unknown")
            # print(f"Financial goals :\n{financial_goals}")
            client_income = client_data.get("incomeFields","Unknown")
            # print(f"Client income :\n{client_income}")

            invested_amount = sum(float(asset["Amount_Invested"] or 0) for asset in portfolio_data)
            available_funds = funds - invested_amount

            # Initialize variables to calculate portfolio-level metrics
            portfolio_current_value = sum(float(asset["current_value"] or 0) for asset in portfolio_data)
            portfolio_daily_change = sum(float(asset["Daily_Value_Change"] or 0) for asset in portfolio_data)
            portfolio_investment_gain_loss = sum(float(asset["Investment_Gain_or_Loss"] or 0) for asset in portfolio_data)
            
            if portfolio_current_value != 0:
                portfolio_daily_change_perc = (portfolio_daily_change / portfolio_current_value) * 100
                portfolio_investment_gain_loss_perc = (portfolio_investment_gain_loss / portfolio_current_value) * 100
            else:
                portfolio_daily_change_perc = 0
                portfolio_investment_gain_loss_perc = 0
            
            # Calculate portfolio time held
            time_held_values = []
            for asset in portfolio_data:
                time_held_str = asset.get("Time_Held", "")
                if time_held_str:
                    try:
                        # Adjusted format string for American date format (MM/DD/YYYY)
                        time_held_dt = datetime.strptime(time_held_str, "%m/%d/%Y, %H:%M:%S")
                        time_held_values.append(time_held_dt)
                    except ValueError as ve:
                        print(f"Error parsing Time_Held for asset {asset.get('name', 'Unknown')}: {time_held_str} - {ve}")

            if time_held_values:
                oldest_time_held = min(time_held_values)
                portfolio_time_held = (datetime.now() - oldest_time_held).days
                # print(f"Oldest Time Held: {oldest_time_held}, Portfolio Time Held: {portfolio_time_held} days")
            else:
                portfolio_time_held = 0
                print("No valid Time_Held data found in portfolio.")

            # print(f"Portfolio Time Held: {portfolio_time_held} days")
            
            # Calculate high asset class
            asset_class_returns = {}
            for asset in portfolio_data:
                asset_class = asset.get("assetClass", "Unknown")
                asset_return = float(asset.get("Investment_Gain_or_Loss", 0))
                if asset_class != "Unknown":
                    asset_class_returns[asset_class] = asset_class_returns.get(asset_class, 0) + asset_return
                else:
                    print(f"Missing asset_class for asset: {asset}")  # Debug log

            # Log asset_class_returns for debugging
            # print(f"Asset class returns for client {client_id}: {asset_class_returns}")

            # Find the asset class with the highest return
            if asset_class_returns:
                high_asset_class = max(asset_class_returns, key=asset_class_returns.get)
            else:
                high_asset_class = "Unknown"
                
            # Asset Allocation
            asset_allocation = {}
            for asset in portfolio_data:
                asset_class = asset.get("assetClass", "Other")
                allocation = float(asset.get("Amount_Invested", 0))
                asset_allocation[asset_class] = asset_allocation.get(asset_class, 0) + allocation

            total_allocation = sum(asset_allocation.values())
            asset_allocation = {k: v / total_allocation * 100 for k, v in asset_allocation.items()}  # Convert to percentage
            
            # print(f"Asset Allocation :{asset_allocation} \nTotal Allocation :{total_allocation}")
            
            # Risk-reward metrics
            sharpe_ratio = (portfolio_investment_gain_loss / invested_amount) if invested_amount > 0 else 0
            volatility = portfolio_daily_change_perc  # Assuming daily change percentage as volatility for simplicity
            volatility = volatility if volatility > 0 else volatility * (-1)
            # print(f"sharpe ratio :{sharpe_ratio}")
            
            # Diversification analysis
            diversification_score = len(asset_allocation)  # More asset classes indicate higher diversification
            # print(f"Diversification Score :{diversification_score}")
            
            # Stress Testing
            stress_scenarios = ["recession", "high inflation", "market crash"]
            stress_test_result = {
                "client_id": client_id,
                "scenarios": {}
            }
            for scenario in stress_scenarios:
                factor = {
                    "recession": -0.2,
                    "high inflation": -0.15,
                    "market crash": -0.3
                }.get(scenario, -0.1)
                stressed_portfolio_value = portfolio_data[0].get("current_value", 0) * (1 + factor)
                stress_test_result["scenarios"][scenario] = round(stressed_portfolio_value, 2)
            stress_test_results.append(stress_test_result)
            
            print(f"Stress Test Results : {stress_test_results}")

            
            # Liquidity analysis :
            
            # Extract assets and liabilities data
            assets_labels = client_data["assetsDatasets"]["labels"]
            assets_data = client_data["assetsDatasets"]["datasets"][0]["data"]
            liabilities_data = client_data["liabilityDatasets"]["datasets"][0]["data"]

            # Map liquid asset categories
            liquid_asset_categories = [
                "Cash/bank accounts",
                "Brokerage/non-qualified accounts",
                "529 Plans",
                "Roth IRA, Roth 401(k)"
            ]

            # Calculate liquid assets
            liquid_assets = sum(
                float(assets_data[i] or 0)
                for i, label in enumerate(assets_labels)
                if label in liquid_asset_categories
            )

            # Calculate total liabilities
            total_liabilities = sum(float(liability or 0) for liability in liabilities_data)

            # Calculate liquidity ratio
            liquidity_ratio = (liquid_assets / total_liabilities * 100) if total_liabilities > 0 else 0
            
            print(f"Liquid assets :{liquid_assets}")
            print(f"Liquidity Ratio :{liquidity_ratio}")

            # Progress Towards Financial Goals
            progress_to_goals_score = {}
            if isinstance(financial_goals, list):
                for goal in financial_goals:
                    try:
                        goal_cost = float(goal.get("cost", 0) or 0)  # Convert cost to float
                    except ValueError:
                        goal_cost = 0  # Fallback if cost is invalid

                    progress = (net_worth / goal_cost) * 100 if goal_cost > 0 else 0
                    progress_to_goals_score[goal.get("goal", "Unknown Goal")] = round(progress, 4)
                    print(f"Progress to goals score :\n{progress_to_goals_score}")    
            else:
                print("Financial goals are not in the expected format.")
            
            client_goal_progress[client_id] = progress_to_goals_score
            
            # Recommendations
            rebalancing_recommendations = []
            for asset_class, percentage in asset_allocation.items():
                if percentage > 50:
                    rebalancing_recommendations.append(f"Reduce allocation in {asset_class} to improve diversification.")

            client_rebalancing_recommendations[client_id] = rebalancing_recommendations
            
            # Append to performance list
            performance_list.append({
                "client_id": client_id,
                "client_name": client_data["clientDetail"]["clientName"],
                "retirement_goals" : retirement_goals,
                "retirement_age" : retirement_age,
                "financial_goals" : financial_goals,
                "client_income" : client_income,
                "assets": assets,
                "liabilities": liabilities,
                "net_worth": net_worth,
                "funds": funds,
                "available_funds": available_funds,
                "invested_amount": invested_amount,
                "portfolio_current_value": portfolio_current_value,
                "portfolio_investment_gain_loss_perc": portfolio_investment_gain_loss_perc,
                "portfolio_investment_gain_loss": portfolio_investment_gain_loss,
                "high_asset_class":high_asset_class,
                "portfolio_time_held" : portfolio_time_held ,
                "asset_allocation":asset_allocation,
                "total_allocation":total_allocation,
                "sharpe_ratio": sharpe_ratio,
                "savings_ratio":savings_rate,
                "Liability Monthly Payments":liability_monthly_payments,
                "Insurance Monthly Premiums": insurance_monthly_premiums,
                "Calculated Monthly Expenses": monthly_expenses,
                "Emergency Funds Coverage":emergency_fund_coverage,
                "Debt To Asset Ratio": debt_to_asset_ratio,
                "volatility": volatility,
                "diversification_score": diversification_score,
                "liquidity_ratio": liquidity_ratio,
                "progress_to_goals_score": progress_to_goals_score,
                "stress_test_results": stress_test_results,
                "rebalancing_recommendations": rebalancing_recommendations
            })

            # Generate insights using Gemini LLM
            task = f"""Your task is to highlight the Client Data, their Portfolio Data, and its Performance. 
                    Analyze the financial data of the client {client_data['clientDetail']['clientName']}:
                    - Total Assets: ${assets}
                    - Total Liabilities: ${liabilities}
                    - Net Worth: ${net_worth}
                    - Investment Personality: {investment_personality}
                    - Retirement Age Goal: {retirement_age}
                    - Retirement Goal: {retirement_goals}
                    - Client Income: {client_income}
                    - Financial Goals: {financial_goals}
                    - Funds: {funds}
                    - Invested Amount: ${invested_amount}
                    - Available Funds: ${available_funds}
                    - Current Portfolio Value: ${portfolio_current_value}
                    - Portfolio Daily Change: ${portfolio_daily_change}
                    - Portfolio Daily Change Percentage: {portfolio_daily_change_perc}%
                    - Portfolio Returns: ${portfolio_investment_gain_loss}
                    - Portfolio Returns Percentage: {portfolio_investment_gain_loss_perc}%
                    - Longest Portfolio Time Held: {portfolio_time_held} days/weeks/months/years
                    - Best Performing Asset Class: {high_asset_class}
                    - Best Performing Asset: {portfolio_data}
                    - Asset Allocation (Breakdown): {asset_allocation}
                    - Total Allocation: ${total_allocation}
                    - Sharpe Ratio: {sharpe_ratio}
                    - Volatility: {volatility}%
                    - Diversification Score: {diversification_score} (Higher scores indicate more diversification.)
                    - Liquid Assets: ${liquid_assets}
                    - Liquidity Ratio: {liquidity_ratio}%
                    - Savings Ratio: {savings_rate}%
                    - Liability Monthly Payments: ${liability_monthly_payments}
                    - Insurance Monthly Premiums: ${insurance_monthly_premiums}
                    - Monthly Expenses: ${monthly_expenses}
                    - Progress to Goals Score: {progress_to_goals_score} 
                    - Stress Test Results: {stress_test_results}
                    - Rebalancing Recommendations: {rebalancing_recommendations}

                    Provide detailed insights into:
                    1. **Asset Class Breakdown**: Highlight how each asset class has performed daily and its contribution to the portfolio.
                    2. **Performance Analysis**: Explain why specific asset classes or investments are performing well or poorly.
                    3. **Risk Exposure**: Assess the client's exposure to high-risk or volatile assets like cryptocurrency.
                    4. **Goal Alignment**: Evaluate the client's progress toward financial and retirement goals.
                    5. **Emergency Fund Adequacy**: Discuss whether the client has sufficient liquidity for emergencies.
                    6. **Debt Management**: Highlight potential risks arising from high liabilities or monthly payments.
                    7. **Stress Testing**: Assess how the portfolio would fare under adverse market conditions like inflation or recession.
                    8. **Rebalancing Needs**: Provide actionable recommendations for rebalancing the portfolio to reduce risk and align with goals.

                    Generate specific insights for the client on:
                    - How their portfolio aligns with their stated investment personality.
                    - Strategies to improve diversification, liquidity, and overall portfolio health.
                    - Long-term viability of their financial plan given current metrics.
                    """

            model = genai.GenerativeModel("gemini-1.5-flash")
            response = model.generate_content(task)
            insights_text = markdown.markdown(response.text)

            # Append insights
            insights.append({
                "client_id": client_id,
                "client_name": client_data["clientDetail"]["clientName"],
                "insights": insights_text,
                "net_worth": net_worth,
            })

        # Analyze the Dashboard:
        assets = sum(map(lambda x: x["net_worth"], insights))
        liabilities = sum(map(lambda x: x["net_worth"], insights))
        net_worth = assets - liabilities
        insights.sort(key=lambda x: x["net_worth"], reverse=True)
        top_client_id = insights[0]["client_id"]
        top_client_name = insights[0]["client_name"]

        # Identify top and low performers
        performance_list.sort(key=lambda x: x["portfolio_investment_gain_loss_perc"], reverse=True)
        top_performers = performance_list[:10]
        low_performers = performance_list[-10:]

        curr_date = datetime.now()
        
        task = f"""
            You are working for a skilled Wealth Manager who is an expert in Portfolio and Asset Management. 
            Your task is to provide comprehensive insights, actionable recommendations, and analysis to help the Wealth Manager better manage the portfolios of a large client base.
            The Wealth Manager needs detailed, actionable, and reliable information to ease their workload, improve portfolio performance, and ensure client satisfaction.
            
            ### Formatting Guidelines:
            1. Use **bolded section headings** to clearly delineate sections of the report.
            2. Organize information into structured categories such as **Insights**, **Comparative Analysis**, **Recommendations**, and **Key Metrics**.
            3. Avoid raw table syntax like '|---|---|', and instead use visually separated sections.
            4. Ensure actionable recommendations are **concise** and organized in bullet points under each section.
            5. Use readable percentages, dollar values, and ratios where applicable.
            6. The report should have a **summary section** at the end with actionable next steps.
            7. Avoid redundancy and ensure actionable recommendations are clearly highlighted.
            8. Use headings (e.g., **"1. Top Performing Clients"**) to separate main sections.
            9. Include all the Sections mentioned below in the report.
            
            ### Dashboard Insights and Performance Analysis

            Date : {curr_date}
            
            #### **Key Deliverables**:
            Refer to the insights here: {insights}.
            Provide a **clear, structured, and actionable report** covering the following areas:

            ---

            ### 1. **Top Performing Clients**
            - Analyze and Highlight patterns among top-performing clients ({top_performers}), focusing on:
            - Net Worth, Liquidity Ratios, Asset Allocation,Monthly Expenses,Savings Rates and Risk Metrics  (e.g., Sharpe Ratio, Volatility).
            - Best-performing assets and asset classes.
            - Strategies or characteristics that contributed to their success.
            - Provide actionable recommendations to replicate successful strategies for other clients.
            
            ---

            ### 2. **Low Performing Clients**
            - Identify and Examine weaknesses among low-performing clients ({low_performers}):
            - Common gaps in their financial profiles, such as:
                - High Debt-to-Asset Ratios, Negative Net Worth, Low Diversification Scores.
                - Over-reliance on volatile or underperforming assets.
                - Insufficient Emergency Funds
                - Misaligned Risk Tolerance and Portfolio Composition
                - High Monthly Expenses and Low Savings Rates
            - Poor Goal Alignment or Misaligned Risk Tolerance.
            - Suggest actionable improvements:
            - Debt Management Plans, Diversification Strategies, Goal-Based Investment Recommendations.
            - Suggest ways and Methods to reduce Monthly Expenses and Increase Savings by Identifying some expenses that can be avoided based on whatever data is available. 
            - Rebalancing Recommendations for Portfolios to address underperforming areas.

            ---

            ### 3. **Comparative Analysis of Clients**
            - Highlight differences and similarities between top-performing and low-performing clients ({top_performers}, {low_performers}):
            - Highlight Portfolio Characteristics: Asset Allocation, Investor Personality,Financial Situation,Monthly Expenses,Savings Patterns, risk metrics,goal alignment. etc.
            - Identify common Vulnerabilities and Trends: Common weaknesses and opportunities for improvement.
            - Common Trends and Development for both.
            - Provide insights into successful strategies and common pitfalls to avoid.  
            - Provide actionable steps to address commonalities and strengthen the client base and to strengthen portfolios against future risks..
            - Suggest strategies to uplift low-performing clients using trends observed in top-performing portfolios.

            ---

            ### 4. **Stress Testing Results**
            - Simulate adverse market scenarios (e.g., recession, inflation spike) using {stress_test_results}.
            - Analyze:
            - Portfolio Resilience Rankings for all clients.
            - Potential Financial Impacts on Net Worth and Liquidity.
            - Provide actionable steps to strengthen and fortify portfolios against future risks.

            ---

            ### 5. **Rebalancing Recommendations**
            - Provide Tailored portfolio rebalancing actions and suggestionsfor underperforming clients :({client_rebalancing_recommendations}):
            - Address over-concentration in certain asset classes or sectors.
            - Highlight opportunities in underweighted sectors (e.g., Bonds, Index Funds, Real Estate).
            - Align portfolios with client financial goals and risk profiles.
            - If Rebalancing Recommendations are valid mention it again for the Wealth Manager in brief for the clients who require Rebalancing Recommendations.

            ---

            ### 6. **Asset Allocation Trends**
            - Analyze trends in asset allocation across all clients:
            - Highlight Dominant Asset Classes: Cryptocurrency, Stocks, Bonds, etc., and their risks and returns.
            - Suggest Opportunities to align and optimize allocation based on market conditions and client objectives.

            ---

            ### 7. **Goal Achievement Progress**
            - Assess the progress of client goals ({client_goal_progress}):
            - Evaluate how many :
                - Goals Achieved(we need to determine that it wont be mentioned direclty),
                - Goals on Track(including an estimate of when the goal will be achieved for most clients),
                - Goals at Risk of Delay.
                - Unreachable Goals,help clients to achieve reasonable Goals or give new Target Goal Year
            - Provide Recommendations steps to help clients stay on track or accelerate progress or achieve reasonable goals or adjust target timelines.

            ---

            ### 8. **Risk and Diversification Metrics**
            - Calculate and Evaluate:
            - Diversification Scores (Average and Individual Client Scores).
            - Risk Metrics (including Sharpe Ratios, Portfolio Volatility).
            - Liquidity Ratios and Emergency Fund Coverage.
            - Provide actionable strategies and steps to enhance and improve overall portfolio diversification and risk management.

            ---

            ### 9. **Common Patterns and Trends**
            - Identify patterns among top-performing clients:
            - Traits like High Liquidity, Consistent Savings, Diversified Portfolios.
            - Highlight patterns among low-performing clients:
            - High Liabilities, Poor Goal Alignment, Over-reliance on volatile assets.
            - Provide actionable insights to address vulnerabilities and capitalize on strengths.

            ---

            ### 10. **Goal-Based Insights**
            - How many goals have been achieved, are on track, or at risk?
            - Suggest steps to:
            - Accelerate progress for at-risk goals.
            - Adjust timelines for unreachable goals.

            ---

            ### 11. **Most Gaining and Losing Assets**
            - Highlight the **most gaining** and **most losing assets** for:
            - Individual Clients.
            - Across all portfolios.
            - Recommend strategies to leverage high-performing assets or mitigate losses.

            ---
            
            ### 12. **Longest and Shortest Investment Periods**
            - Identify clients with the:
            - Longest Investment Period and its impact on performance.
            - Shortest Investment Period and how it reflects their strategies.

            ---

            ### 13. **Overall Portfolio Management Score**
            - Assign a score out of 100 for the Wealth Manager based on:
            - Diversification, Liquidity, Goal Progress, Returns Generated, Risk Management.
            - Provide reasons for the score and actionable steps for improvement.

            ---

            ### 14. **Custom Recommendations**
            - Identify and Suggest specific investment opportunities including:
            - Low-risk alternatives like Bonds, ETFs, or Index Funds.
            - High-potential sectors based on current market trends.
            - Provide personalized strategies to optimize portfolios for both high-net-worth and low-performing clients.

            ---
            
            ### 15. **Actionable Summary**
            - Provide key insights, recommendations, and patterns to help the Wealth Manager:
            - Improve portfolio performance.
            - Align client portfolios with financial goals and market trends.
            - Strengthen risk management and diversification.

            #### **Note**:
            Ensure your response is structured, clear, and comprehensive.Make sure all the above 15 points are covered in your report
            Give in a Report Format and Avoid stating "insufficient data" or vague assumptions. 
            Use all the provided data to generate meaningful and actionable insights.
            Make sure there is no repeatations in your answer making it unnecessarily lenghty and wordy.
            """

        response = model.generate_content(task)
        dashboard_analysis = markdown_to_text_new(response.text)
        # Process the response from LLM
        html_suggestions = markdown.markdown(dashboard_analysis)
        print(html_suggestions)
        dashboard_suggestions = markdown_to_text_new(html_suggestions)
        print(dashboard_suggestions)
        # Return the response
        return jsonify({
            "message": "Dashboard Analysis generated successfully",
            "dashboard_analysis": dashboard_suggestions,
            "performance": insights,
            "performance_list": performance_list
        }), 200

    except Exception as e:
        print(f"Error in analyzing dashboard: {e}")
        return jsonify({"message": f"Error in analyzing dashboard: {e}"}), 500


def markdown_to_text_new(md):
    """
    Converts markdown to readable plain text with improved formatting.
    """
    # Replace markdown headers with uppercase text
    md = re.sub(r'#{1,6}\s*(.*)', r'\n\1\n' + ('-' * 50), md)

    # Convert bullet points to simple dashes
    md = md.replace('* ', '- ')
    md = md.replace('**', '')  # Remove bold markers

    # Ensure proper spacing for readability
    md = md.replace('\n', '\n\n')

    return md.strip()

################################################################################################
# Dashboard Infographics :

def rebalancing_strategy(before):
    """
    Example rebalancing strategy: Reduce categories above 50% allocation to 50% and 
    distribute the remainder evenly across other categories.
    """
    max_allocation = 50  # Max allocation threshold in %
    excess_allocation = 0
    after = {}

    # Cap allocations above the threshold and calculate excess
    for category, percentage in before.items():
        if percentage > max_allocation:
            excess_allocation += percentage - max_allocation
            after[category] = max_allocation
        else:
            after[category] = percentage

    # Redistribute excess allocation proportionally
    total_under_allocated = sum(1 for v in after.values() if v < max_allocation)
    redistribution = excess_allocation / total_under_allocated if total_under_allocated > 0 else 0

    for category in after:
        if after[category] < max_allocation:
            after[category] += redistribution
            after[category] = min(after[category], max_allocation)

    return after


# Calculate metrics and prepare data for dashboard infographics:

# # v-2 :
# def calculate_dashboard_metrics(clients, portfolios, client_summary_folder, use_aws=True, s3=s3, s3_bucket_name=S3_BUCKET_NAME):
#     performance_list = []
#     stress_test_results = []
#     client_goal_progress = {}
#     client_rebalancing_recommendations = {}
#     comparative_analysis_data = []
#     portfolio_management_scores = []
#     risk_and_liquidity_data = []

#     def get_safe_field(dictionary, keys, default=None):
#         """Helper function to safely get a nested field from a dictionary."""
#         try:
#             for key in keys:
#                 dictionary = dictionary.get(key, {})
#             return dictionary or default
#         except AttributeError:
#             return default

#     def rebalancing_strategy(before):
#         """
#         Example rebalancing strategy: Cap allocations above 50% and redistribute excess proportionally.
#         """
#         max_allocation = 50  # Max allocation threshold in %
#         excess_allocation = 0
#         after = {}

#         # Cap allocations above the threshold and calculate excess
#         for category, percentage in before.items():
#             if percentage > max_allocation:
#                 excess_allocation += percentage - max_allocation
#                 after[category] = max_allocation
#             else:
#                 after[category] = percentage

#         # Redistribute excess allocation proportionally
#         total_under_allocated = sum(1 for v in after.values() if v < max_allocation)
#         redistribution = excess_allocation / total_under_allocated if total_under_allocated > 0 else 0

#         for category in after:
#             if after[category] < max_allocation:
#                 after[category] += redistribution
#                 after[category] = min(after[category], max_allocation)

#         return after

#     for client in clients:
#         client_id = client.get("uniqueId")
#         if not client_id or client_id not in portfolios:
#             logging.warning(f"Client ID {client_id} is invalid or not in portfolios. Skipping.")
#             continue

#         # Safely fetch client name
#         client_name = get_safe_field(client, ["clientDetail", "clientName"], "Unknown")
#         if client_name == "Unknown":
#             logging.warning(f"Missing client name for client ID {client_id}. Defaulting to 'Unknown'.")

#         # Fetch client data
#         client_data = None
#         if use_aws:
#             client_data_key = f"{client_summary_folder}client-data/{client_id}.json"
#             try:
#                 response = s3.get_object(Bucket=s3_bucket_name, Key=client_data_key)
#                 client_data = json.loads(response['Body'].read().decode('utf-8'))
#             except Exception as e:
#                 logging.error(f"Error retrieving client data from AWS for {client_id}: {e}")
#                 continue
#         else:
#             client_data_path = os.path.join("client_data", f"{client_id}.json")
#             if os.path.exists(client_data_path):
#                 with open(client_data_path, 'r') as file:
#                     client_data = json.load(file)
#         if not client_data:
#             logging.warning(f"No client data found for client ID {client_id}. Skipping.")
#             continue

#         # Process and calculate metrics
#         funds = float(client_data.get("investmentAmount", 0) or 0)
#         assets = sum(float(value or 0) for value in client_data["assetsDatasets"]["datasets"][0]["data"])
#         liabilities = sum(float(value or 0) for value in client_data["liabilityDatasets"]["datasets"][0]["data"])
#         net_worth = assets - liabilities

#         # Savings Rate
#         annual_income = sum(float(inc.get("amountIncome", 0) or 0) for inc in client_data.get("incomeFields", []))
#         savings_rate = (funds / annual_income * 100) if annual_income > 0 else 0

#         # Emergency Fund Coverage
#         monthly_expenses = sum(
#             float(liab.get("mortgageMonthly", 0) or 0) for liab in client_data["myLiabilities"].values()
#         ) + sum(
#             float(ins.get("monthlyPayLIClient", 0) or 0) for ins in client_data["insuranceCoverage"].values()
#         )
#         assets_labels = client_data["assetsDatasets"]["labels"]
#         assets_data = client_data["assetsDatasets"]["datasets"][0]["data"]

#         liquid_asset_categories = [
#             "Cash/bank accounts",
#             "Brokerage/non-qualified accounts",
#             "529 Plans",
#             "Roth IRA, Roth 401(k)"
#         ]
#         liquid_assets = sum(
#             float(assets_data[i] or 0)
#             for i, label in enumerate(assets_labels)
#             if label in liquid_asset_categories
#         )
#         total_liabilities = liabilities
#         liquidity_ratio = (liquid_assets / total_liabilities * 100) if total_liabilities > 0 else 0
#         emergency_fund_coverage = (liquid_assets / monthly_expenses) if monthly_expenses > 0 else 0

#         # Debt-to-Asset Ratio
#         debt_to_asset_ratio = (liabilities / assets * 100) if assets > 0 else 0

#         # Asset Allocation
#         portfolio_data = portfolios.get(client_id, [])
#         asset_allocation = {}
#         for asset in portfolio_data:
#             asset_class = asset.get("assetClass", "Other")
#             allocation = float(asset.get("Amount_Invested", 0))
#             asset_allocation[asset_class] = asset_allocation.get(asset_class, 0) + allocation
#         total_allocation = sum(asset_allocation.values())
#         asset_allocation = {k: v / total_allocation * 100 for k, v in asset_allocation.items()}

#         # Portfolio Metrics
#         diversification_score = len(asset_allocation)
#         sharpe_ratio = client_data.get("sharpeRatio", 0)
#         volatility = client_data.get("volatility", 0)
#         portfolio_current_value = sum(float(asset["current_value"] or 0) for asset in portfolio_data)
#         portfolio_daily_change = sum(float(asset["Daily_Value_Change"] or 0) for asset in portfolio_data)
#         portfolio_investment_gain_loss = sum(float(asset["Investment_Gain_or_Loss"] or 0) for asset in portfolio_data)
#         portfolio_investment_gain_loss_perc = (
#             (portfolio_investment_gain_loss / portfolio_current_value) * 100 if portfolio_current_value != 0 else 0
#         )

#         # Progress Towards Financial Goals
#         financial_goals = client_data.get("goalFields", [])
#         progress_to_goals_score = {}
#         for goal in financial_goals:
#             goal_cost = float(goal.get("cost", 0) or 0)
#             progress = (net_worth / goal_cost * 100) if goal_cost > 0 else 0
#             progress_to_goals_score[goal.get("goal", "Unknown Goal")] = round(progress, 2)
#         client_goal_progress[client_id] = progress_to_goals_score

#         # Append Data
#         comparative_analysis_data.append({
#             "client_name": client_name,
#             "assets": assets,
#             "liabilities": liabilities,
#             "net_worth": net_worth,
#             "profit_loss": portfolio_investment_gain_loss,
#             "profit_loss_percentage": portfolio_investment_gain_loss_perc,
#         })
#         risk_and_liquidity_data.append({
#             "client_name": client_name,
#             "liquidity_ratio": liquidity_ratio,
#             "debt_to_asset_ratio": debt_to_asset_ratio,
#             "net_worth": net_worth,
#             "emergency_fund_coverage": emergency_fund_coverage,
#         })
#         portfolio_management_scores.append({
#             "client_id": client_id,
#             "client_name": client_name,
#             "sharpe_ratio": sharpe_ratio,
#             "savings_rate": savings_rate,
#             "volatility": volatility,
#             "diversification_score": diversification_score,
#             "progress_to_goals_score": progress_to_goals_score,
#         })

#     print(f"Risk and Liquidity Data: {risk_and_liquidity_data}")

#     dashboard_data= {"recommendations_impact": {}}
#     for client_id, portfolio_data in portfolios.items():
#         # Aggregate current asset allocation (calculate 'before')
#         asset_allocation = {}
#         for asset in portfolio_data:
#             asset_class = asset.get("assetClass", "Other")
#             allocation = float(asset.get("Amount_Invested", 0))
#             asset_allocation[asset_class] = asset_allocation.get(asset_class, 0) + allocation

#         # Normalize asset allocation to percentages
#         total_allocation = sum(asset_allocation.values())
#         before_allocation = {k: v / total_allocation * 100 for k, v in asset_allocation.items()}

#         # # Calculate 'after' allocation using the rebalancing strategy
#         after_allocation = rebalancing_strategy(before_allocation)

#         # # Generate recommendations
#         recommendations = [
#             f"Reduce allocation in {cat} to below 50%." for cat, perc in before_allocation.items() if perc > 50
#         ]
#         client_rebalancing_recommendations[client_id] = recommendations

#         # # Prepare chart data for recommendations impact
#         chart_data = {
#             "categories": list(before_allocation.keys()),
#             "before_values": list(before_allocation.values()),
#             "after_values": [after_allocation.get(cat, 0) for cat in before_allocation.keys()],
#         }
#         dashboard_data["recommendations_impact"][client_id] = chart_data

#     return {
#         "performance_list": performance_list,
#         "stress_test_results": stress_test_results,
#         "client_goal_progress": client_goal_progress,
#         "client_rebalancing_recommendations": client_rebalancing_recommendations,
#         "comparative_analysis_data": comparative_analysis_data,
#         "portfolio_management_scores": portfolio_management_scores,
#         "risk_and_liquidity_data": risk_and_liquidity_data,
        # "recommendations_impact": dashboard_data["recommendations_impact"][client_id],
#     }



# prev :

# Updated :

def calculate_dashboard_metrics(clients, portfolios, client_summary_folder, use_aws=True, s3=s3, s3_bucket_name=S3_BUCKET_NAME):
    performance_list = []
    stress_test_results = []
    client_goal_progress = {}
    client_rebalancing_recommendations = {}
    comparative_analysis_data = []
    portfolio_management_scores = []
    risk_and_liquidity_data = []

    def get_safe_field(dictionary, keys, default=None):
        """Helper function to safely get a nested field from a dictionary."""
        try:
            for key in keys:
                dictionary = dictionary.get(key, {})
            return dictionary or default
        except AttributeError:
            return default

    def rebalancing_strategy(before):
        """
        Example rebalancing strategy: Cap allocations above 50% and redistribute excess proportionally.
        """
        max_allocation = 50  # Max allocation threshold in %
        excess_allocation = 0
        after = {}

        # Cap allocations above the threshold and calculate excess
        for category, percentage in before.items():
            if percentage > max_allocation:
                excess_allocation += percentage - max_allocation
                after[category] = max_allocation
            else:
                after[category] = percentage

        # Redistribute excess allocation proportionally
        total_under_allocated = sum(1 for v in after.values() if v < max_allocation)
        redistribution = excess_allocation / total_under_allocated if total_under_allocated > 0 else 0

        for category in after:
            if after[category] < max_allocation:
                after[category] += redistribution
                after[category] = min(after[category], max_allocation)

        return after

    for client in clients:
        client_id = client.get("uniqueId")
        if not client_id or client_id not in portfolios:
            logging.warning(f"Client ID {client_id} is invalid or not in portfolios. Skipping.")
            continue

        # Safely fetch client name
        client_name = get_safe_field(client, ["clientDetail", "clientName"], "Unknown")
        if client_name == "Unknown":
            logging.warning(f"Missing client name for client ID {client_id}. Defaulting to 'Unknown'.")

        # Fetch client data
        client_data = None
        if use_aws:
            client_data_key = f"{client_summary_folder}client-data/{client_id}.json"
            try:
                response = s3.get_object(Bucket=s3_bucket_name, Key=client_data_key)
                client_data = json.loads(response['Body'].read().decode('utf-8'))
            except Exception as e:
                logging.error(f"Error retrieving client data from AWS for {client_id}: {e}")
                continue
        else:
            client_data_path = os.path.join("client_data", f"{client_id}.json")
            if os.path.exists(client_data_path):
                with open(client_data_path, 'r') as file:
                    client_data = json.load(file)
        if not client_data:
            logging.warning(f"No client data found for client ID {client_id}. Skipping.")
            continue

        # Process and calculate metrics
        funds = float(client_data.get("investmentAmount", 0) or 0)
        assets = sum(float(value or 0) for value in client_data["assetsDatasets"]["datasets"][0]["data"])
        liabilities = sum(float(value or 0) for value in client_data["liabilityDatasets"]["datasets"][0]["data"])
        net_worth = assets - liabilities

        # Savings Rate
        annual_income = sum(float(inc.get("amountIncome", 0) or 0) for inc in client_data.get("incomeFields", []))
        savings_rate = (funds / annual_income * 100) if annual_income > 0 else 0

        # Emergency Fund Coverage
        monthly_expenses = sum(
            float(liab.get("mortgageMonthly", 0) or 0) for liab in client_data["myLiabilities"].values()
        ) + sum(
            float(ins.get("monthlyPayLIClient", 0) or 0) for ins in client_data["insuranceCoverage"].values()
        )
        assets_labels = client_data["assetsDatasets"]["labels"]
        assets_data = client_data["assetsDatasets"]["datasets"][0]["data"]

        liquid_asset_categories = [
            "Cash/bank accounts",
            "Brokerage/non-qualified accounts",
            "529 Plans",
            "Roth IRA, Roth 401(k)"
        ]
        liquid_assets = sum(
            float(assets_data[i] or 0)
            for i, label in enumerate(assets_labels)
            if label in liquid_asset_categories
        )
        total_liabilities = liabilities
        liquidity_ratio = (liquid_assets / total_liabilities * 100) if total_liabilities > 0 else 0
        emergency_fund_coverage = (liquid_assets / monthly_expenses) if monthly_expenses > 0 else 0

        # Debt-to-Asset Ratio
        debt_to_asset_ratio = (liabilities / assets * 100) if assets > 0 else 0

        # Asset Allocation
        portfolio_data = portfolios.get(client_id, [])
        asset_allocation = {}
        for asset in portfolio_data:
            asset_class = asset.get("assetClass", "Other")
            allocation = float(asset.get("Amount_Invested", 0))
            asset_allocation[asset_class] = asset_allocation.get(asset_class, 0) + allocation
        total_allocation = sum(asset_allocation.values())
        asset_allocation = {k: v / total_allocation * 100 for k, v in asset_allocation.items()}

        # Portfolio Metrics
        diversification_score = len(asset_allocation)
        sharpe_ratio = client_data.get("sharpeRatio", 0)
        volatility = client_data.get("volatility", 0)
        portfolio_current_value = sum(float(asset["current_value"] or 0) for asset in portfolio_data)
        portfolio_daily_change = sum(float(asset["Daily_Value_Change"] or 0) for asset in portfolio_data)
        portfolio_investment_gain_loss = sum(float(asset["Investment_Gain_or_Loss"] or 0) for asset in portfolio_data)
        portfolio_investment_gain_loss_perc = (
            (portfolio_investment_gain_loss / portfolio_current_value) * 100 if portfolio_current_value != 0 else 0
        )

        # Progress Towards Financial Goals
        financial_goals = client_data.get("goalFields", [])
        progress_to_goals_score = {}
        for goal in financial_goals:
            goal_cost = float(goal.get("cost", 0) or 0)
            progress = (net_worth / goal_cost * 100) if goal_cost > 0 else 0
            progress_to_goals_score[goal.get("goal", "Unknown Goal")] = round(progress, 2)
            
        # client_goal_progress[client_id] = progress_to_goals_score
        client_goal_progress[client_name] = progress_to_goals_score
        

        # Append Data
        comparative_analysis_data.append({
            "client_name": client_name,
            "assets": assets,
            "liabilities": liabilities,
            "net_worth": net_worth,
            "profit_loss": portfolio_investment_gain_loss,
            "profit_loss_percentage": portfolio_investment_gain_loss_perc,
        })
        risk_and_liquidity_data.append({
            "client_name": client_name,
            "liquidity_ratio": liquidity_ratio,
            "debt_to_asset_ratio": debt_to_asset_ratio,
            "net_worth": net_worth,
            "emergency_fund_coverage": emergency_fund_coverage,
        })
        portfolio_management_scores.append({
            "client_id": client_id,
            "client_name": client_name,
            "sharpe_ratio": sharpe_ratio,
            "savings_rate": savings_rate,
            "volatility": volatility,
            "diversification_score": diversification_score,
            "progress_to_goals_score": progress_to_goals_score,
        })

    # Prepare rebalancing and recommendations
    dashboard_data = {"recommendations_impact": {}}
    for client_id, portfolio_data in portfolios.items():
        # Aggregate current asset allocation (calculate 'before')
        asset_allocation = {}
        for asset in portfolio_data:
            asset_class = asset.get("assetClass", "Other")
            allocation = float(asset.get("Amount_Invested", 0))
            asset_allocation[asset_class] = asset_allocation.get(asset_class, 0) + allocation

        # Normalize asset allocation to percentages
        total_allocation = sum(asset_allocation.values())
        before_allocation = {k: v / total_allocation * 100 for k, v in asset_allocation.items()}

        # Calculate 'after' allocation using the rebalancing strategy
        after_allocation = rebalancing_strategy(before_allocation)

        # Generate recommendations
        recommendations = [
            f"Reduce allocation in {cat} to below 50%." for cat, perc in before_allocation.items() if perc > 50
        ]
        client_rebalancing_recommendations[client_id] = recommendations

        # Prepare chart data for recommendations impact
        chart_data = {
            "categories": list(before_allocation.keys()),
            "before_values": list(before_allocation.values()),
            "after_values": [after_allocation.get(cat, 0) for cat in before_allocation.keys()],
        }
        dashboard_data["recommendations_impact"][client_id] = chart_data
        print(f"Dashboard data for recommendations impact : {chart_data}")
        
    return {
        "performance_list": performance_list,
        "stress_test_results": stress_test_results,
        "client_goal_progress": client_goal_progress,
        "client_rebalancing_recommendations": client_rebalancing_recommendations,
        "comparative_analysis_data": comparative_analysis_data,
        "portfolio_management_scores": portfolio_management_scores,
        "risk_and_liquidity_data": risk_and_liquidity_data,
        "recommendations_impact": dashboard_data["recommendations_impact"].get(client_id, {}),
    }




# Visualization functions :

# plots correctly with data :

def plot_comparative_analysis(data):
    """
    Create a side-by-side bar graph for comparative analysis of multiple clients.
    :param data: List of dicts with client metrics (e.g., assets, liabilities, net worth, profit/loss).
    """
    categories = ['assets', 'liabilities', 'net_worth', 'profit_loss']
    clients = [item['client_name'] for item in data]

    # Calculate values for each category
    values = {
        category: [client.get(category, 0) for client in data]
        for category in categories
    }

    # x = np.arange(len(clients))  # Client positions
    # width = 0.2  # Bar width

    # fig, ax = plt.subplots(figsize=(12, 8))

    # # Plot bars for each category
    # for i, category in enumerate(categories):
    #     ax.bar(x + i * width, values[category], width, label=category.replace('_', ' ').title())

    # # Adjust tick positions and labels
    # ax.set_xticks(x + (width * (len(categories) - 1)) / 2)  # Center tick labels
    # ax.set_xticklabels(clients, rotation=45, ha='right')

    # ax.set_ylabel('Values')
    # ax.set_title('Comparative Analysis of Clients')
    # ax.legend()

    # plt.tight_layout()
    # plt.show()

    # Prepare chart data for frontend
    chart_data = {
        "labels": clients,
        "datasets": [
            {
                "label": category.replace('_', ' ').title(),
                "data": values[category],
            }
            for category in categories
        ],
    }
    return chart_data


def plot_portfolio_management_scores(data):
    """
    Create a radar chart for portfolio management scores.
    :param data: List of dicts with metrics for each client.
    """
    metrics = ['sharpe_ratio', 'savings_rate', 'volatility', 'diversification_score', 'progress_to_goals_score']
    
    for client in data:
        # Extract progress_to_goals_score safely
        progress_to_goals_score = client.get('progress_to_goals_score', 0)
        
        # If progress_to_goals_score is a dict, calculate a summary metric
        if isinstance(progress_to_goals_score, dict):
            progress_to_goals_score = sum(progress_to_goals_score.values()) / len(progress_to_goals_score)  # Average

        # Validate and prepare metric values
        values = [
            float(client.get(metric, 0)) if not isinstance(client.get(metric), dict) else 0
            for metric in metrics
        ]
        values[metrics.index('progress_to_goals_score')] = progress_to_goals_score  # Insert processed value
        values += values[:1]  # Close the radar chart (repeat the first value)

        labels = metrics  # Labels for the radar chart axes
        angles = np.linspace(0, 2 * np.pi, len(labels), endpoint=False).tolist()
        angles += angles[:1]  # Close the radar chart (repeat the first angle)

        fig, ax = plt.subplots(figsize=(6, 6), subplot_kw=dict(polar=True))
        
        # Plot data and fill area
        ax.plot(angles, values, label=client.get('client_name', 'Unknown Client'))
        ax.fill(angles, values, alpha=0.25)

        # Set the correct number of labels and positions
        ax.set_xticks(angles[:-1])  # Use only the first N angles (one for each metric)
        ax.set_xticklabels(labels)  # Label only the metrics (no duplicate for closing)

        ax.set_title(f"Portfolio Scores: {client.get('client_name', 'Unknown Client')}")
        ax.legend(loc='upper right', bbox_to_anchor=(1.1, 1.1))

        plt.tight_layout()
        plt.show()

# updated :

def plot_goal_achievement_progress(data):
    """
    Create a horizontal bar chart for goal achievement progress.
    :param data: Dict of progress scores for each client and their goals.
                 Example: {'John': {'Home': 50, 'Education': -25}, 'Jane': {'Retirement': 150}}
    :return: Chart data for the frontend.
    """
    if not data:
        print("No data available for plotting goal achievement progress.")
        return {"clients": [], "goals": [], "progress_values": [], "annotations": []}

    fig, ax = plt.subplots(figsize=(12, 8))

    y_positions = []
    labels = []
    progress_values = []
    annotations = []

    # Extract progress scores and prepare data
    for client_name, goals in data.items():
        if not goals:
            print(f"No goals for client {client_name}. Skipping.")
            continue
        for goal_name, progress in goals.items():
            y_positions.append(len(y_positions))  # Incremental index
            labels.append(f"{client_name}: {goal_name or 'Unknown Goal'}")  # Combine client and goal
            progress_values.append(progress)
            
            # Add descriptive annotation for each progress value
            if progress > 50:
                annotations.append("Highly Likely/Already Achieved")
            elif progress > 0:
                annotations.append("Likely to Achieve Goal")
            elif progress > -50:
                annotations.append("At Risk of Missing Goal")
            else:
                annotations.append("Unrealistic Goal")

    if not progress_values:
        print("No valid progress values to plot.")
        return {"clients": [], "goals": [], "progress_values": [], "annotations": []}

    # Bar colors based on progress type
    bar_colors = [
        "darkgreen" if progress > 50 else "lightgreen" if progress > 0 else "orange" if progress > -50 else "red"
        for progress in progress_values
    ]

    # ax.barh(y_positions, progress_values, color=bar_colors, edgecolor="black")

    # # Add labels
    # ax.set_yticks(y_positions)
    # ax.set_yticklabels(labels, fontsize=10)
    # ax.set_xlabel("Progress")
    # ax.set_title("Goal Achievement Progress")
    # ax.axvline(x=0, color="black", linewidth=0.8, linestyle="--")  # Reference line at 0%

    # # Annotate the bars with descriptive text
    # for i, (v, annotation) in enumerate(zip(progress_values, annotations)):
    #     ax.text(
    #         v,
    #         i,
    #         annotation,
    #         va="center",
    #         ha="right" if v < 0 else "left",
    #         fontsize=9,
    #         color="black"
    #     )

    # plt.tight_layout()
    # plt.show()

    # Prepare data for frontend
    chart_data = {
        "clients": [label.split(":")[0] for label in labels],
        "goals": [label.split(":")[1].strip() for label in labels],
        "progress_values": progress_values,
        "annotations": annotations,
    }
    return chart_data



# Prev :

# def plot_goal_achievement_progress(data):
#     """
#     Create a horizontal bar chart for goal achievement progress.
#     :param data: Dict of progress scores for each client and their goals.
#     """
#     if not data:
#         print("No data available for plotting goal achievement progress.")
#         return

#     fig, ax = plt.subplots(figsize=(10, 6))

#     y_positions = []
#     labels = []
#     progress_values = []

#     # Extract progress scores and prepare data
#     for client, goals in data.items():
#         if not goals:
#             print(f"No goals for client {client}. Skipping.")
#             continue
#         for goal_name, progress in goals.items():
#             y_positions.append(len(y_positions))  # Incremental index
#             labels.append(f"{client}: {goal_name or 'Unknown Goal'}")  # Combine client and goal
#             progress_values.append(progress)

#     if not progress_values:
#         print("No valid progress values to plot.")
#         return

#     # Bar colors based on progress type
#     bar_colors = [
#         "green" if progress > 0 else "orange" if -50 <= progress <= 0 else "red"
#         for progress in progress_values
#     ]

#     ax.barh(y_positions, progress_values, color=bar_colors, edgecolor="black")

#     # Add labels
#     ax.set_yticks(y_positions)
#     ax.set_yticklabels(labels, fontsize=10)
#     ax.set_xlabel("Progress (%)")
#     ax.set_title("Goal Achievement Progress")
#     ax.axvline(x=0, color="black", linewidth=0.8, linestyle="--")  # Reference line at 0%

#     # Annotate the bars with progress values
#     for i, v in enumerate(progress_values):
#         ax.text(
#             v,
#             i,
#             f"{v:.1f}%",
#             va="center",
#             ha="right" if v < 0 else "left",
#             fontsize=9,
#         )

#     plt.tight_layout()
#     plt.show()

#updated version :

def plot_risk_and_liquidity(data):
    """
    Create a bubble chart for risk and liquidity analysis and prepare data for frontend.
    :param data: List of dicts with metrics (liquidity_ratio, debt_to_asset_ratio, net_worth).
    """
    x = [item['liquidity_ratio'] for item in data]
    y = [item['debt_to_asset_ratio'] for item in data]
    raw_sizes = [item['net_worth'] / 1e6 for item in data]  # Scale net worth for visualization
    sizes = [max(abs(size), 0.01) * 500 for size in raw_sizes]  # Normalize and scale sizes
    labels = [item['client_name'] for item in data]

    # Plot bubble chart
    # fig, ax = plt.subplots(figsize=(12, 8))
    # scatter = ax.scatter(
    #     x, y, s=sizes, alpha=0.7, c=raw_sizes, cmap='viridis', edgecolors='k', linewidth=0.5
    # )

    # Add labels with dynamic offsets
    # for i, label in enumerate(labels):
    #     ax.text(
    #         x[i], y[i] * 1.05, label, fontsize=9, ha='center', va='bottom',
    #         bbox=dict(facecolor='white', alpha=0.7, edgecolor='none', pad=1)
    #     )

    # Apply logarithmic scale for wide y-axis range
    # ax.set_yscale('log')
    # ax.set_xlabel('Liquidity Ratio (%)')
    # ax.set_ylabel('Debt-to-Asset Ratio (%) (Log Scale)')
    # ax.set_title('Risk and Liquidity Analysis')
    # plt.colorbar(scatter, ax=ax, label='Net Worth (in $M)')
    # plt.grid(True, which="both", linestyle='--', linewidth=0.5, alpha=0.7)
    # plt.tight_layout()
    # plt.show()

    # Prepare chart data for frontend
    chart_data = {
        "x": x,
        "y": y,
        "sizes": sizes,
        "labels": labels,
    }
    return chart_data


# updated :

def plot_recommendations_impact(before, after):
    """
    Create a before-and-after bar chart for recommendations impact and prepare chart data for the frontend.
    :param before: Dict of asset allocation percentages before recommendations.
    :param after: Dict of asset allocation percentages after recommendations.
    :return: Chart data for frontend.
    """
    # Handle empty input gracefully
    if not before or not after:
        print("No data provided for plotting.")
        return {
            "categories": [],
            "before_values": [],
            "after_values": []
        }

    categories = list(before.keys())
    x = np.arange(len(categories))

    before_values = [before.get(cat, 0) for cat in categories]
    after_values = [after.get(cat, 0) for cat in categories]

    # fig, ax = plt.subplots(figsize=(10, 6))
    # ax.bar(x - 0.2, before_values, width=0.4, label='Before', color='skyblue')
    # ax.bar(x + 0.2, after_values, width=0.4, label='After', color='orange')

    # ax.set_xticks(x)
    # ax.set_xticklabels(categories, rotation=45, ha='right')
    # ax.set_ylabel('Allocation (%)')
    # ax.set_title('Impact of Recommendations on Asset Allocation')
    # ax.legend()

    # plt.tight_layout()
    # plt.show()

    # Prepare data for the frontend
    chart_data = {
        "categories": categories,
        "before_values": before_values,
        "after_values": after_values
    }
    return chart_data


# with api :

@app.route('/dashboard_infographics', methods=['POST'])
def dashboard_infographics():
# def dashboard_infographics(clients):
    try:
        clients = request.json.get("data")
        if not clients:
            return jsonify({"message": "No client data provided"}), 400

        client_ids = [client.get("uniqueId") for client in clients if client.get("uniqueId")]
        if not client_ids:
            return jsonify({"message": "No valid client IDs found in the request"}), 400

        portfolios = fetch_portfolios(client_ids)
        

        dashboard_data = calculate_dashboard_metrics(
            clients=clients,
            portfolios=portfolios,
            client_summary_folder=client_summary_folder
        )

        # Call plotting functions here
        
        comparative_analysis_data = plot_comparative_analysis(dashboard_data["comparative_analysis_data"])
        # print(f"comparative_analysis_data: {comparative_analysis_data}")
        
        # plot_portfolio_management_scores(dashboard_data["portfolio_management_scores"])
        
        risk_and_liquidity_data = plot_risk_and_liquidity(dashboard_data["risk_and_liquidity_data"])
        # print(f"risk_and_liquidity_data: {risk_and_liquidity_data}")
        
        print("\nrecommendations_impact:\n")
        print(dashboard_data["recommendations_impact"])
        
        rebalancing_recommendations_data = []
        
        impact = dashboard_data["recommendations_impact"]

        # Transform lists to dictionaries
        before = dict(zip(impact["categories"], impact["before_values"]))
        after = dict(zip(impact["categories"], impact["after_values"]))

        # Call the function with transformed data
        rebalancing_recommendations_data.append(plot_recommendations_impact(before, after))

        # print(f"rebalancing_recommendations_data: {rebalancing_recommendations_data}")
        
        dashboard_data["client_goal_progress"]
        
        client_goal_progress_data = plot_goal_achievement_progress(dashboard_data["client_goal_progress"])
        
        print(f"client_goal_progress :\n{client_goal_progress_data}")

        print("\nSuccess\n")
        
        print(dashboard_data)
        
        return jsonify({"success": True, 
                        "comparative_analysis_data": comparative_analysis_data,
                        "risk_and_liquidity_data": risk_and_liquidity_data,
                        "client_goal_progress":client_goal_progress_data,
                        "rebalancing_recommendations_data" : rebalancing_recommendations_data,
                        }),200
        
        # return jsonify({"success": True, "data": dashboard_data}),200

    except Exception as e:
        print(f"Error calculating dashboard metrics: {e}")
        return jsonify({"success": False, "error": str(e)}), 500

#################################################################################################################################

# Asset Allocation based on INvestor Profile :

@app.route('/asset_allocation_investor_profile',methods=['POST'])
def asset_allocation_investor_profile():

    try:
        # Fetch client data and their portfolios
        clients = request.json.get("data")
        if not clients:
            return jsonify({"message": "No client data provided"}), 400

        # Create lookup for client investment personality
        client_personality_map = {client["uniqueId"]: client.get("investorPersonality", "Unknown") for client in clients if client.get("uniqueId")}
        if not client_personality_map:
            return jsonify({"message": "No valid client data found"}), 400

        # Fetch consolidated portfolio data
        portfolios = fetch_consolidated_portfolio(list(client_personality_map.keys()))
        if not portfolios:
            return jsonify({"message": "No portfolio data found for provided clients"}), 404

        # Initialize personality-based asset allocation data
        personality_data = {"Conservative": 0, "Moderate": 0, "Aggressive": 0}

        # Process portfolios by personality type
        for client_id, portfolio in portfolios.items():
            # Get the investor personality for the client
            personality = client_personality_map.get(client_id, "Unknown")

            # Ensure valid portfolio structure
            assets = portfolio.get("assets", []) if isinstance(portfolio, dict) else []
            if not isinstance(assets, list):
                continue

            # Sum total investment per personality type
            total_investment = sum(float(asset.get("Amount_Invested", 0)) for asset in assets)
            if personality in personality_data:
                personality_data[personality] += total_investment

        # Prepare pie chart data for asset allocation by personality type
        pie_chart_data = {
            "labels": list(personality_data.keys()),
            "datasets": [{
                "data": list(personality_data.values()),
                "label": "Total Asset Allocation by Personality Type",
                "backgroundColor": ["#FF6384", "#36A2EB", "#FFCE56"]  # Default colors for chart
            }]
        }

        # Format response for the frontend
        response = {
            "message": "Personality-Based Asset Infographics generated successfully",
            "pie_chart_data": pie_chart_data,
            "details": personality_data
        }

        return jsonify(response), 200

    except Exception as e:
        print(f"Error in personality_asset_infographics: {e}")
        return jsonify({"message": f"Error in personality_asset_infographics: {e}"}), 500
0
        

#################################################################################################################################

# Run the Flask application
if __name__ == '__main__':
    app.run(host='0.0.0.0',debug=True)

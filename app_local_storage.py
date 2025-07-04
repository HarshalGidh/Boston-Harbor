##########################################################################################################################################
################################################### Start of Using Local Storage ########################################################################################
# import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt

import os
import filetype
import docx
import PyPDF2
import re
from dotenv import load_dotenv
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import Chroma
from langchain_community.docstore.in_memory import InMemoryDocstore
from langchain_community.vectorstores import FAISS
from langchain_community.document_loaders import Docx2txtLoader
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains import create_retrieval_chain
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.memory import ConversationSummaryMemory
import asyncio
import numpy as np
import json

import google.generativeai as genai
import pathlib
import logging
import sys
import io
import matplotlib.pyplot as plt
import seaborn as sns
# Import things that are needed generically
from langchain.pydantic_v1 import BaseModel, Field
from langchain.tools import BaseTool, StructuredTool, tool
# Define functions to generate investment suggestions :\


# # -------------------------------------Start Aws---------------------
# import paramiko

# # Set up the SSH key file, IP, username, and passphrase
# key_path = "keys/aws_key.pem"  # Path to the converted .pem file
# hostname = "172.31.15.173"  # AWS EC2 public IP address
# username = "pragatidhobe"  # EC2 instance username
# passphrase = "12345678"  # Passphrase, if any

# # Create an SSH client instance
# ssh_client = paramiko.SSHClient()
# ssh_client.set_missing_host_key_policy(paramiko.AutoAddPolicy())

# try:
#     # Load SSH key
#     key = paramiko.RSAKey.from_private_key_file(key_path, password=passphrase)

#     # Connect to the instance
#     ssh_client.connect(hostname=hostname, username=username, pkey=key)

#     # Execute a command (example)
#     stdin, stdout, stderr = ssh_client.exec_command("ls")
#     print(stdout.read().decode())  # Print command output

# except Exception as e:
#     print(f"An error occurred: {e}")

# finally:
#     ssh_client.close()

# # -------------------------------------End Aws---------------------



# import boto3
# load_dotenv()

# # AWS keys
# aws_access_key = os.getenv('aws_access_key')
# aws_secret_key = os.getenv('aws_secret_key')
# S3_BUCKET_NAME = os.getenv('S3_BUCKET_NAME')
# client_summary_folder = os.getenv('client_summary_folder') 
# suggestions_folder = os.getenv('suggestions_folder') 
# order_list_folder = os.getenv('order_list_folder')
# portfolio_list_folder = os.getenv('portfolio_list_folder') 
# personality_assessment_folder = os.getenv('personality_assessment_folder') 
# login_folder = os.getenv('login_folder')


# # Connecting to Amazon S3
# s3 = boto3.client(
#     's3',
#     aws_access_key_id=aws_access_key,
#     aws_secret_access_key=aws_secret_key
# )

# def list_s3_keys(bucket_name, prefix=""):
#     try:
#         # List objects in the bucket with the given prefix
#         response = s3.list_objects_v2(Bucket=bucket_name, Prefix=prefix)
#         if 'Contents' in response:
#             print("Keys in the S3 folder:")
#             for obj in response['Contents']:
#                 print(obj['Key'])
#         else:
#             print("No files found in the specified folder.")
#     except Exception as e:
#         print(f"Error listing objects in S3: {e}")

# # Call the function
# list_s3_keys(S3_BUCKET_NAME, order_list_folder)




# =------------------------------------------------------=






GOOGLE_API_KEY = os.getenv('GOOGLE_API_KEY')

from flask import Flask, request, jsonify

# app = Flask(__name__)
from flask import Flask, request, jsonify, send_file
import asyncio
from flask_cors import CORS
app = Flask(__name__)
CORS(app, resources={r"/*": {"origins": "http://localhost:3000"}})

# Configure generativeai with your API key
genai.configure(api_key=GOOGLE_API_KEY)

import markdown

########################### Sign in Sign WithOut using aws ###################################################

from flask import Flask, request, jsonify
from flask_bcrypt import Bcrypt
from flask_mail import Mail, Message
import random
import boto3
import json
from datetime import datetime, timedelta,timezone

bcrypt = Bcrypt(app)

# Email configuration
app.config['MAIL_SERVER'] = 'smtp.gmail.com'
app.config['MAIL_PORT'] = 587
app.config['MAIL_USE_TLS'] = True
app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME') # 'your_email@gmail.com'
app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD') #'your_email_password'

mail = Mail(app)

# In-memory storage for email and OTP (for simplicity)
otp_store = {}

# API Endpoints
from flask import Flask, request, jsonify
import random
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart

# Replace with your email credentials
EMAIL_ADDRESS = os.getenv('MAIL_USERNAME')  #'your-email@gmail.com'
EMAIL_PASSWORD = os.getenv('MAIL_PASSWORD')  #'your-email-password'

import random
from datetime import datetime
import os
import json
from email.mime.text import MIMEText
import smtplib
import jwt
 
# Local storage paths
LOCAL_STORAGE_PATH = "local_storage"
os.makedirs(LOCAL_STORAGE_PATH, exist_ok=True)
# otp_store = {}
 
def load_from_local(filepath):
    try:
        if not os.path.exists(filepath):
            return None
        with open(filepath, 'r') as file:
            return json.load(file)
    except Exception as e:
        print(f"Error loading file: {e}")
        return None
   
def save_to_local(data, filepath):
    try:
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        with open(filepath, 'w') as file:
            json.dump(data, file)
        print(f"Data saved at {filepath}")  # Debug log
    except Exception as e:
        print(f"Error saving file: {e}")  # Debug log
        raise
 
 
def delete_from_local(filename):
    file_path = os.path.join(LOCAL_STORAGE_PATH, filename)
    if os.path.exists(file_path):
        os.remove(file_path)
       
 
def send_email(to_email, otp):
    try:
        # Validate email format
        if not re.match(r"[^@]+@[^@]+\.[^@]+", to_email):
            print(f"Invalid email address: {to_email}")
            return False
 
        # Setup email message
        subject = "Your Reset Password Code"
        message = f"Your Reset Password Code is: {otp}"
        msg = MIMEMultipart()
        msg['From'] = EMAIL_ADDRESS
        msg['To'] = to_email
        msg['Subject'] = subject
        msg.attach(MIMEText(message, 'plain'))
 
        # Send email using SMTP
        with smtplib.SMTP('smtp.gmail.com', 587) as server:
            server.starttls()
            server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
            server.sendmail(EMAIL_ADDRESS, to_email, msg.as_string())
        print("Email sent successfully")
        return True
    except Exception as e:
        print(f"Error sending email: {e}")
        return False
 
 
@app.route('/email-verification', methods=['POST'])
def email_verification():
    try:
        email = request.json.get('email')  # Extract email from the request
        if not email:
            return jsonify({"message": "Email is required"}), 400
 
        print(f"Processing email verification for: {email}")
 
        # Generate the sign-up link
        sign_up_link = f"http://localhost:3000/signUp/{email}"
 
        # Create the email message
        msg = Message(
            "Sign-Up Link - Verify Your Email",
            sender="your_email@gmail.com",
            recipients=[email]
        )
        msg.body = (
            f"Hello,\n\n"
            f"Your email has been successfully verified. Use the following link to complete your sign-up process:\n\n"
            f"{sign_up_link}\n\n"
            f"If you did not request this verification, please ignore this email.\n\n"
            f"Thank you."
        )
        print(f"Sending email to: {email}\nContent: {msg.body}")
       
        # Send the email
        mail.send(msg)
        print("Email sent successfully.")
 
        return jsonify({"message": "Sign-up link sent successfully"}), 200
 
    except Exception as e:
        print(f"Error sending email: {e}")
        return jsonify({"message": f"Error occurred: {str(e)}"}), 500
 
 
 
@app.route('/verify-otp', methods=['POST'])
def verify_otp():
    data = request.get_json()
    email = data.get('email')
    otp = data.get('otp')
 
    if not email or not otp:
        return jsonify({"error": "Email and OTP are required"}), 400
 
    if otp_store.get(email) == int(otp):
        del otp_store[email]
        return jsonify({"message": "Email verified successfully!"}), 200
    else:
        return jsonify({"error": "Invalid OTP"}), 400
 

# 2. Sign Up
@app.route('/sign-up', methods=['POST'])
def sign_up():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"message": "Invalid JSON"}), 400
 
        email = data.get('email')
        password = data.get('password')
        confirm_password = data.get('confirm_password')
        first_name = data.get('first_name')
        last_name = data.get('last_name')
 
        if not all([email, password, confirm_password, first_name, last_name]):
            return jsonify({"message": "All fields are required"}), 400
 
        if password != confirm_password:
            return jsonify({"message": "Passwords do not match"}), 400
 
        if load_from_local(f"users/{email}.json"):
            return jsonify({"message": "User already exists"}), 400
 
        hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
        user_data = {
            "email": email,
            "password": hashed_password,
            "first_name": first_name,
            "last_name": last_name
        }
        save_to_local(user_data, f"users/{email}.json")
 
        return jsonify({"message": "Sign up successful"}), 200
    except Exception as e:
        print(f"Error in sign-up: {e}")
        return jsonify({"message": "Internal server error"}), 500
 
 
@app.route('/sign-in', methods=['POST'])
def sign_in():
    try:
        data = request.get_json()
        email = data.get('email')
        password = data.get('password')
 
        if not all([email, password]):
            return jsonify({"message": "Email and password are required"}), 400
 
        user_data = load_from_local(f"users/{email}.json")
        if not user_data or not bcrypt.check_password_hash(user_data["password"], password):
            return jsonify({"message": "Invalid email or password"}), 401
 
        # Generate JWT Token with corrected datetime usage
        token = jwt.encode(
            {"email": email, "exp": datetime.utcnow() + timedelta(hours=5)},  # Use `datetime.utcnow()` directly
            app.config['JWT_SECRET_KEY'],
            algorithm="HS256"
        )
 
        return jsonify({"message": "Sign in successful", "token": token}), 200
 
    except Exception as e:
        print(f"Error during sign-in: {e}")
        return jsonify({"message": "Internal server error"}), 500
 
 
 
# ------------------------- With Bearer ------------------------------
import requests
 
@app.route('/advisor_profile', methods=['GET'])
def advisor_profile():
    # Get the token from the Authorization header
    token = request.headers.get('Authorization')
 
    # Check if token is present
    if not token:
        return jsonify({"message": "Token is missing"}), 401
 
    try:
        # Extract token from "Bearer <token>"
        token = token.split(" ")[1] if token.startswith("Bearer ") else None
        if not token:
            return jsonify({"message": "Invalid token format"}), 401
 
        # Decode the token and extract the email from the payload
        decoded_token = jwt.decode(token, app.config['JWT_SECRET_KEY'], algorithms=["HS256"])
        email = decoded_token.get('email')
 
        # Check if user data exists for the email
        user_data = load_from_local(f"users/{email}.json")
        if not user_data:
            return jsonify({"message": "User not found"}), 404
 
        # Make a request to the /get-all-client-data API to get the number of clients
        client_data_url = 'http://localhost:5000/get-all-client-data'  # Adjust to your actual endpoint
        response = requests.get(client_data_url)
       
        if response.status_code == 200:
            client_list = response.json().get('data', [])
            client_count = len(client_list)  # Get the length of the client list
        else:
            client_count = 0  # If there's an error fetching client data, assume 0 clients
 
        # Return the profile data along with the client count
        profile_data = {
            "email": user_data["email"],
            "first_name": user_data["first_name"],
            "last_name": user_data["last_name"],
            "client_count": client_count  # Include the client count in the response
        }
 
        return jsonify({"message": "Profile retrieved successfully", "profile": profile_data}), 200
 
    except jwt.ExpiredSignatureError:
        return jsonify({"message": "Token has expired"}), 401
    except jwt.InvalidTokenError:
        return jsonify({"message": "Invalid token"}), 401
    except Exception as e:
        print(f"Error retrieving profile: {e}")
        return jsonify({"message": "Internal server error"}), 500
 
# -------------------------------------------
 
@app.route('/change_password', methods=['POST'])
def change_password():
    try:
        data = request.get_json()
        if not data:
            return jsonify({"message": "Invalid JSON"}), 400
       
        email = data.get('email')
        old_password = data.get('old_password')
        new_password = data.get('new_password')
        confirm_password = data.get('confirm_password')
 
        if not all([email, old_password, new_password, confirm_password]):
            return jsonify({"message": "All fields are required"}), 400
       
        # Check if new password matches the confirm password
        if new_password != confirm_password:
            return jsonify({"message": "Passwords do not match"}), 400
 
        # Load user data
        user_data = load_from_local(f"users/{email}.json")
        if not user_data:
            return jsonify({"message": "User not found"}), 404
       
        # Check if old password is correct
        if not bcrypt.check_password_hash(user_data["password"], old_password):
            return jsonify({"message": "Old password is incorrect"}), 401
       
        # Hash the new password
        hashed_new_password = bcrypt.generate_password_hash(new_password).decode('utf-8')
 
        # Update the user data with the new password
        user_data["password"] = hashed_new_password
        save_to_local(user_data, f"users/{email}.json")
 
        return jsonify({"message": "Password reset successful"}), 200
 
    except Exception as e:
        print(f"Error in reset-password: {e}")
        return jsonify({"message": "Internal server error"}), 500
 
 
 

 
# # 2. Sign Up
# @app.route('/sign-up', methods=['POST'])
# def sign_up():
#     try:
#         data = request.get_json()
#         if not data:
#             return jsonify({"message": "Invalid JSON"}), 400
 
#         email = data.get('email')
#         password = data.get('password')
#         confirm_password = data.get('confirm_password')
 
#         if not all([email, password, confirm_password]):
#             return jsonify({"message": "All fields are required"}), 400
 
#         if password != confirm_password:
#             return jsonify({"message": "Passwords do not match"}), 400
 
#         if load_from_local(f"users/{email}.json"):
#             return jsonify({"message": "User already exists"}), 400
 
#         hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')
#         user_data = {"email": email, "password": hashed_password}
#         save_to_local(user_data, f"users/{email}.json")
 
#         return jsonify({"message": "Sign up successful"}), 200
#     except Exception as e:
#         print(f"Error in sign-up: {e}")
#         return jsonify({"message": "Internal server error"}), 500
 
 
# @app.route('/sign-in', methods=['POST'])
# def sign_in():
#     data = request.get_json()
#     email = data.get('email')
#     password = data.get('password')
 
#     if not all([email, password]):
#         return jsonify({"message": "Email and password are required"}), 400
 
#     user_data = load_from_local(f"users/{email}.json")
#     if not user_data or not bcrypt.check_password_hash(user_data["password"], password):
#         return jsonify({"message": "Invalid email or password"}), 401
 
#     return jsonify({"message": "Sign in successful"}), 200
 
 
 
# @app.route('/send-otp', methods=['POST'])
# def send_otp():
#     data = request.get_json()
#     email = data.get('email')
#     if not email:
#         return jsonify({"error": "Email is required"}), 400
 
#     # Generate and store OTP
#     otp = random.randint(100000, 999999)
#     otp_store[email] = otp
 
#     if send_email(email, "Your Verification Code", f"Your verification code is: {otp}"):
#         return jsonify({"message": "OTP sent successfully!"}), 200
#     else:
#         return jsonify({"error": "Failed to send OTP"}), 500
 
 
#  # 4. Forgot Password
# import traceback
 
# @app.route('/forgot-password', methods=['POST'])
# def forgot_password():
#     try:
#         data = request.get_json()
#         print("Data received:", data)
#         if not data:
#             return jsonify({"message": "Invalid JSON"}), 400
 
#         email = data.get('email')
#         print("Email extracted:", email)
#         if not email:
#             return jsonify({"message": "Email is required"}), 400
 
#         reset_code = random.randint(100000, 999999)
#         print("Reset code generated:", reset_code)
#         reset_data = {"email": email, "reset_code": reset_code, "timestamp": str(datetime.now())}
 
#         save_to_local(reset_data, f"password_resets/{email}.json")
#         print("Reset data saved successfully.")
 
#         # if send_email(email, "Reset Your Password", f"Your reset code is: {reset_code}"):
#         #     print("Email sent successfully.")
#         #     return jsonify({"message": "Password reset code sent successfully"}), 200
        
#         if send_email(email,reset_code):
#             print("Email sent successfully.")
#             return jsonify({"message": "Password reset code sent successfully"}), 200
#         else:
#             print("Failed to send email.")
#             return jsonify({"error": "Failed to send reset code"}), 500
#     except Exception as e:
#         traceback.print_exc()  # Logs the full stack trace
#         return jsonify({"error": "Internal server error"}), 500
 
# @app.route('/reset-password', methods=['POST'])
# def reset_password():
#     data = request.get_json()
#     email = data.get('email')
#     reset_code = data.get('reset_code')
#     new_password = data.get('new_password')
 
#     if not all([email, reset_code, new_password]):
#         return jsonify({"message": "All fields are required"}), 400
 
#     reset_data = load_from_local(f"password_resets/{email}.json")
#     if not reset_data or str(reset_data["reset_code"]) != str(reset_code):
#         return jsonify({"message": "Invalid reset code"}), 400
 
#     user_data = load_from_local(f"users/{email}.json")
#     if not user_data:
#         return jsonify({"message": "User not found"}), 404
 
#     hashed_password = bcrypt.generate_password_hash(new_password).decode('utf-8')
#     user_data["password"] = hashed_password
#     save_to_local(user_data, f"users/{email}.json")
#     delete_from_local(f"password_resets/{email}.json")
 
#     return jsonify({"message": "Password reset successful"}), 200
 




##########################################################################################################


########################### Sign in Sign Out using aws ###################################################


# from flask import Flask, request, jsonify
# from flask_bcrypt import Bcrypt
# from flask_mail import Mail, Message
# import random
# import boto3
# import json
# from datetime import datetime, timedelta,timezone

# bcrypt = Bcrypt(app)

# # Email configuration
# app.config['MAIL_SERVER'] = 'smtp.gmail.com'
# app.config['MAIL_PORT'] = 587
# app.config['MAIL_USE_TLS'] = True
# app.config['MAIL_USERNAME'] = os.getenv('MAIL_USERNAME') # 'your_email@gmail.com'
# app.config['MAIL_PASSWORD'] = os.getenv('MAIL_PASSWORD') #'your_email_password'

# mail = Mail(app)


# # Helper functions
# def upload_to_s3(data, filename):
#     s3.put_object(Bucket=S3_BUCKET_NAME, Key=filename, Body=json.dumps(data))
#     return f"s3://{S3_BUCKET_NAME}/{filename}"

# def download_from_s3(filename):
#     try:
#         response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=filename)
#         return json.loads(response['Body'].read().decode('utf-8'))
#     except Exception as e:
#         return None
    
# def delete_from_s3(key):
#     try:
#         s3.delete_object(Bucket=S3_BUCKET_NAME, Key=key)
#     except Exception as e:
#         print(f"Error deleting {key}: {e}")

# # API Endpoints
# from flask import Flask, request, jsonify
# import random
# import smtplib
# from email.mime.text import MIMEText
# from email.mime.multipart import MIMEMultipart



# # Replace with your email credentials
# EMAIL_ADDRESS = os.getenv('MAIL_USERNAME')  #'your-email@gmail.com'
# EMAIL_PASSWORD = os.getenv('MAIL_PASSWORD')  #'your-email-password'

# # In-memory storage for email and OTP (for simplicity)
# otp_store = {}

# def send_email(to_email, otp):
#     try:
#         # Setup email message
#         subject = "Your Verification Code"
#         message = f"Your verification code is: {otp}"
#         msg = MIMEMultipart()
#         msg['From'] = EMAIL_ADDRESS
#         msg['To'] = to_email
#         msg['Subject'] = subject
#         msg.attach(MIMEText(message, 'plain'))

#         # Send email using SMTP
#         with smtplib.SMTP('smtp.gmail.com', 587) as server:
#             server.starttls()
#             server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
#             server.sendmail(EMAIL_ADDRESS, to_email, msg.as_string())
#         return True
#     except Exception as e:
#         print(f"Error sending email: {e}")
#         return False

# # 1. Email Verification and send otp :

# @app.route('/send-otp', methods=['POST'])
# def send_otp():
#     data = request.get_json()
#     email = data.get('email')

#     if not email:
#         return jsonify({"error": "Email is required"}), 400

#     # Generate a random 6-digit OTP
#     otp = random.randint(100000, 999999)
#     otp_store[email] = otp

#     if send_email(email, otp):
#         return jsonify({"message": "OTP sent successfully!"}), 200
#     else:
#         return jsonify({"error": "Failed to send OTP"}), 500

# @app.route('/verify-otp', methods=['POST'])
# def verify_otp():
#     data = request.get_json()
#     email = data.get('email')
#     otp = data.get('otp')

#     if not email or not otp:
#         return jsonify({"error": "Email and OTP are required"}), 400

#     # Check if the provided OTP matches the stored OTP
#     if otp_store.get(email) == int(otp):
#         del otp_store[email]  # Remove OTP after successful verification
#         return jsonify({"message": "Email verified successfully!"}), 200
#     else:
#         return jsonify({"error": "Invalid OTP"}), 400


# # Previous Version
# # @app.route('/email-verification', methods=['POST'])
# # def email_verification():
# #     try:
# #         email = request.json.get('email')
# #         if not email:
# #             return jsonify({"message": "Email is required"}), 400
        
# #         print(email)
# #         # Generate a 6-digit verification code
# #         verification_code = random.randint(100000, 999999)
# #         sign_up_link = "http://localhost:3000/signUp"
        
# #         # Send the email with the verification code
# #         msg = Message("Sign Up Link",recipients=[email]) # Code", recipients=[email])
# #         msg.body = f"Your Email is Verified.\nUse this Link to Sign Up : {sign_up_link}" #f"Your verification code is: {verification_code}"
# #         print(msg.body)
# #         print(msg)
# #         mail.send(msg)
# #         # msg = Message("Your Verification Code", recipients=[email])
# #         # msg.body = f"Your verification code is: {verification_code}"
# #         # mail.send(msg)

# #         # Save the verification code in S3
# #         # data = {"email": email, "verification_code": verification_code, "timestamp": str(datetime.now())}
# #         # upload_to_s3(data, f"verification_codes/{email}.json")

# #         return jsonify({"message": "Verification code sent successfully"}), 200
# #     except Exception as e:
# #         return jsonify({"message": f"Error occurred: {str(e)}"}), 500


# @app.route('/email-verification', methods=['POST'])
# def email_verification():
#     try:
#         email = request.json.get('email')  # Extract email from the request
#         if not email:
#             return jsonify({"message": "Email is required"}), 400

#         print(f"Processing email verification for: {email}")

#         # Generate the sign-up link
#         sign_up_link = f"http://localhost:3000/signUp/{email}"

#         # Create the email message
#         msg = Message(
#             "Sign-Up Link - Verify Your Email",
#             sender="your_email@gmail.com",
#             recipients=[email]
#         )
#         msg.body = (
#             f"Hello,\n\n"
#             f"Your email has been successfully verified. Use the following link to complete your sign-up process:\n\n"
#             f"{sign_up_link}\n\n"
#             f"If you did not request this verification, please ignore this email.\n\n"
#             f"Thank you."
#         )
#         print(f"Sending email to: {email}\nContent: {msg.body}")
        
#         # Send the email
#         mail.send(msg)
#         print("Email sent successfully.")

#         return jsonify({"message": "Sign-up link sent successfully"}), 200

#     except Exception as e:
#         print(f"Error sending email: {e}")
#         return jsonify({"message": f"Error occurred: {str(e)}"}), 500





# # # 2. Sign Up
# @app.route('/sign-up', methods=['POST'])
# def sign_up():
#     try:
#         email = request.json.get('email')
#         password = request.json.get('password')
#         confirm_password = request.json.get('confirm_password')
#         verification_code = request.json.get('verification_code')

#         if not all([email, password, confirm_password]): #, verification_code]):
#             return jsonify({"message": "All fields are required"}), 400

#         if password != confirm_password:
#             return jsonify({"message": "Passwords do not match"}), 400

#         # Fetch and validate verification code from S3
#         verification_data = download_from_s3(f"verification_codes/{email}.json")
#         if not verification_data or str(verification_data["verification_code"]) != str(verification_code):
#             return jsonify({"message": "Invalid verification code"}), 400

#         # Hash the password
#         hashed_password = bcrypt.generate_password_hash(password).decode('utf-8')

#         # Save user data in S3
#         user_data = {"email": email, "password": hashed_password}
#         upload_to_s3(user_data, f"users/{email}.json")

#         return jsonify({"message": "Sign up successful"}), 200
#     except Exception as e:
#         return jsonify({"message": f"Error occurred: {str(e)}"}), 500

# # 3. Sign In
# import jwt
# # from datetime import datetime, timedelta,timezone

# # Secret key for signing JWT
# JWT_SECRET_KEY =  os.getenv('JWT_SECRET_KEY') 

# @app.route('/sign-in', methods=['POST'])
# def sign_in():
#     try:
#         email = request.json.get('email')
#         password = request.json.get('password')

#         if not all([email, password]):
#             return jsonify({"message": "Email and password are required"}), 400

#         # Fetch user data from S3
#         user_data = download_from_s3(f"users/{email}.json")
#         if not user_data or not bcrypt.check_password_hash(user_data["password"], password):
#             return jsonify({"message": "Invalid email or password"}), 401

#         # Generate a JWT token
#         token_payload = {
#             "email": email,
#             "exp": datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(hours=2),
#             "iat": datetime.datetime.now(datetime.timezone.utc),
#             "sub": "user_authentication"  # Subject of the token
#         }
#         token = jwt.encode(token_payload, JWT_SECRET_KEY, algorithm="HS256")

#         return jsonify({
#             "message": "Sign in successful",
#             "token": token
#         }), 200
#     except Exception as e:
#         return jsonify({"message": f"Error occurred: {str(e)}"}), 500



# # 4. Forgot Password
# @app.route('/forgot-password', methods=['POST'])
# def forgot_password():
#     try:
#         email = request.json.get('email')
#         if not email:
#             return jsonify({"message": "Email is required"}), 400

#         # Generate a 6-digit reset code
#         reset_code = random.randint(100000, 999999)

#         # Send the reset code via email
#         msg = Message(
#             "Reset Your Password",
#             sender="your_email@gmail.com",
#             recipients=[email]
#         )
#         msg.body = (
#             f"Hello,\n\n"
#             f"You are about to Reset Your Password.Use the following Reset Code to Reset Your Password:\n\n"
#             f"{reset_code}\n\n"
#             f"If you did not request this verification, please ignore this email.\n\n"
#             f"Thank you."
#         )
#         print(f"Sending email to: {email}\nContent: {msg.body}")
        
#         mail.send(msg)

#         # Save the reset code and timestamp in S3
#         data = {"email": email, "reset_code": reset_code, "timestamp": str(datetime.datetime.now())}
#         upload_to_s3(data, f"password_resets/{email}.json")

#         return jsonify({"message": "Password reset code sent successfully"}), 200
#     except Exception as e:
#         return jsonify({"message": f"Error occurred while sending reset code: {str(e)}"}), 500

# #5. Reset password
# @app.route('/reset-password', methods=['POST'])
# def reset_password():
#     try:
#         email = request.json.get('email')
#         reset_code = request.json.get('reset_code')
#         new_password = request.json.get('new_password')
#         # confirm_password = request.json.get('confirm_password')
#         if not all([email, reset_code, new_password]):
#             return jsonify({"message": "Email, reset code, and new password are required"}), 400

#         # if new_password != confirm_password:
#         #     return jsonify({"message": "Passwords do not match"}), 400
        
#         # Fetch reset data from S3
#         reset_data = download_from_s3(f"password_resets/{email}.json")
#         if not reset_data:
#             return jsonify({"message": "Invalid email or reset code"}), 400

#         # Validate the reset code
#         if str(reset_data["reset_code"]) != str(reset_code):
#             return jsonify({"message": "Invalid reset code"}), 400

#         # Update the password for the user
#         user_data = download_from_s3(f"users/{email}.json")
#         if not user_data:
#             return jsonify({"message": "User not found"}), 404

#         # Hash the new password
#         hashed_password = bcrypt.generate_password_hash(new_password).decode('utf-8')
#         user_data["password"] = hashed_password

#         # Save the updated user data back to S3
#         upload_to_s3(user_data, f"users/{email}.json")

#         # Remove the reset code entry from S3
#         delete_from_s3(f"password_resets/{email}.json")

#         return jsonify({"message": "Password reset successfully"}), 200
#     except Exception as e:
#         return jsonify({"message": f"Error occurred while resetting password: {str(e)}"}), 500



########################################################################################################################

# def convert_to_markdown(raw_text):
#     # Replace specific text patterns with markdown syntax
#     formatted_text = raw_text.replace('\n', '\n\n')  # Ensure newlines create paragraphs
    
#     # Convert text into markdown format
#     html = markdown.markdown(formatted_text)

#     return html


def markdown_table_to_html(md_table):
    # Split the markdown table by lines
    lines = md_table.strip().split("\n")
    
    # Extract headers and rows
    headers = lines[0].strip('|').split('|')
    rows = [line.strip('|').split('|') for line in lines[2:]]  # Skip the separator line

    # Start creating the HTML table
    html_table = "<table>\n"
    
    # Add headers
    html_table += "  <thead>\n    <tr>\n"
    for header in headers:
        html_table += f"      <th>{header.strip()}</th>\n"
    html_table += "    </tr>\n  </thead>\n"
    
    # Add rows
    html_table += "  <tbody>\n"
    for row in rows:
        html_table += "    <tr>\n"
        for col in row:
            html_table += f"      <td>{col.strip()}</td>\n"
        html_table += "    </tr>\n"
    html_table += "  </tbody>\n</table>"

    return html_table

def process_client_info_and_analysis(content):
    # Identify and extract the client's financial info markdown table part
    client_info_section_start = content.find("| Category | Value |")
    client_info_section_end = content.find("</p>", client_info_section_start) + 4
    
    if client_info_section_start != -1 and client_info_section_end != -1:
        # Extract markdown table
        md_table = content[client_info_section_start:client_info_section_end]
        # Convert only the markdown table portion into an HTML table
        html_table = markdown_table_to_html(md_table)
        # Replace the markdown table part with the generated HTML table
        content = content[:client_info_section_start] + html_table + content[client_info_section_end:]
    
    # Return the rest of the content unchanged
    return content

def generate_final_html(content):
    # Process the content to convert financial info to HTML table while leaving other sections untouched
    html_content = process_client_info_and_analysis(content)
    
    # Any other HTML processing can be done here if needed
    return html_content


# def markdown_table_to_html(md_table):
#     # Split the markdown table by lines
#     lines = md_table.strip().split("\n")
    
#     # Extract headers and rows
#     headers = lines[0].strip('|').split('|')
#     rows = [line.strip('|').split('|') for line in lines[2:]]  # Skip the separator line

#     # Start creating the HTML table
#     html_table = "<table>\n"
    
#     # Add headers
#     html_table += "  <thead>\n    <tr>\n"
#     for header in headers:
#         html_table += f"      <th>{header.strip()}</th>\n"
#     html_table += "    </tr>\n  </thead>\n"
    
#     # Add rows
#     html_table += "  <tbody>\n"
#     for row in rows:
#         html_table += "    <tr>\n"
#         for col in row:
#             html_table += f"      <td>{col.strip()}</td>\n"
#         html_table += "    </tr>\n"
#     html_table += "  </tbody>\n</table>"

#     return html_table



import markdown2
from bs4 import BeautifulSoup

def markdown_to_readable_text(md_text):
    # Convert markdown to HTML
    html = markdown2.markdown(md_text)

    # Parse the HTML
    soup = BeautifulSoup(html, "html.parser")

    # Function to format plain text from tags
    def format_text_from_html(soup):
        formatted_text = ''
        for element in soup:
            if element.name == "h1":
                formatted_text += f"\n\n# {element.text.upper()} #\n\n"
            elif element.name == "h2":
                formatted_text += f"\n\n## {element.text} ##\n\n"
            elif element.name == "h3":
                formatted_text += f"\n\n### {element.text} ###\n\n"
            elif element.name == "strong":
                formatted_text += f"**{element.text}**"
            elif element.name == "em":
                formatted_text += f"_{element.text}_"
            elif element.name == "ul":
                for li in element.find_all("li"):
                    formatted_text += f"\n - {li.text}"
            elif element.name == "ol":
                for idx, li in enumerate(element.find_all("li"), 1):
                    formatted_text += f"\n {idx}. {li.text}"
            elif element.name == "table":
                # Convert markdown table to HTML table
                formatted_text += "<table>\n"
                rows = element.find_all("tr")
                for row in rows:
                    formatted_text += "<tr>\n"
                    cols = row.find_all(["th", "td"])
                    for col in cols:
                        tag = 'th' if col.name == "th" else 'td'
                        formatted_text += f"<{tag}>{col.text.strip()}</{tag}>\n"
                    formatted_text += "</tr>\n"
                formatted_text += "</table>\n"
            else:
                formatted_text += element.text

        return formatted_text.strip()

    return format_text_from_html(soup)



def markdown_to_text(md): # og solution code 
    # Simple conversion for markdown to plain text
    md = md.replace('**', '')
    md = md.replace('*', '')
    md = md.replace('_', '')
    md = md.replace('#', '')
    md = md.replace('`', '')
    return md.strip()



# def extract_responses_from_docx(personality_file):
#     try:
#         doc = docx.Document(personality_file)
#         responses = {}
#         current_question = None

#         # Check paragraphs
#         for para in doc.paragraphs:
#             text = para.text.strip()
#             if text:
#                 # Check if the paragraph contains a question
#                 if "?" in text or text.endswith(":"):
#                     current_question = text
#                 else:
#                     # This is a typed answer
#                     typed_answer = text.strip()
#                     if current_question:
#                         # If the question already has an answer, append to it (handles multiple responses)
#                         if current_question in responses:
#                             responses[current_question] += "; " + typed_answer
#                         else:
#                             responses[current_question] = typed_answer

#         # Check tables for additional responses
#         for table in doc.tables:
#             for row in table.rows:
#                 for cell in row.cells:
#                     text = cell.text.strip()
#                     if text:
#                         if "?" in text or text.endswith(":"):
#                             current_question = text
#                         else:
#                             typed_answer = text.strip()
#                             if current_question:
#                                 if current_question in responses:
#                                     responses[current_question] += "; " + typed_answer
#                                 else:
#                                     responses[current_question] = typed_answer

#         return responses

#     except Exception as e:
#         print(f"Error extracting responses: {e}")
#         return None

import docx



# import asyncio
# # from some_generative_ai_library import GenerativeModel  # Replace with actual import

# async def determine_investment_personality(assessment_data):
#     try:
#         # Prepare input text for the chatbot based on assessment data
#         input_text = "User Profile:\n"
#         for question, answer in assessment_data.items():
#             input_text += f"{question}: {answer}\n"

#         # Introduce the chatbot's task and prompt for classification
#         input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
#                       "- Conservative Investor\n" \
#                       "- Moderate Investor\n" \
#                       "- Aggressive Investor\n\n" \
#                       "Please provide the classification below:\n"

#         # Use your generative AI model to generate a response
#         model = GenerativeModel('gemini-1.5-flash')
#         response = await model.generate_content(input_text)

#         # Determine the investment personality from the chatbot's response
#         response_text = response.text.lower()

#         if "conservative investor" in response_text:
#             personality = "Conservative Investor"
#         elif "moderate investor" in response_text:
#             personality = "Moderate Investor"
#         elif "aggressive investor" in response_text:
#             personality = "Aggressive Investor"
#         else:
#             personality = "Unknown"

#         return personality
#     except Exception as e:
#         print(f"Error generating response: {e}")
#         return "Unknown"


# GET Method
async def determine_investment_personality(assessment_data): # proper code 
    try:
        # Prepare input text for the chatbot based on assessment data
        input_text = "User Profile:\n"
        for question, answer in assessment_data.items():
            input_text += f"{question}: {answer}\n"

        # Introduce the chatbot's task and prompt for classification
        input_text += "\nYou are an investment personality identifier. Based on the user profile, classify the user as:\n" \
                      "- Conservative Investor\n" \
                      "- Moderate Investor\n" \
                      "- Aggressive Investor\n\n" \
                      "Please provide the classification below:\n"

        # Use your generative AI model to generate a response
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(input_text)

        # Determine the investment personality from the chatbot's response
        response_text = response.text.lower()

        if "conservative investor" in response_text:
            personality = "Conservative Investor"
        elif "moderate investor" in response_text:
            personality = "Moderate Investor"
        elif "aggressive investor" in response_text:
            personality = "Aggressive Investor"
        else:
            personality = "Unknown"

        return personality
    except Exception as e:
        print(f"Error generating response: {e}")
        return "Unknown"




#Load the Vector DataBase : # current version :
async def load_vector_db(file_path): # # GET Method 
    try:
        print("Loading vector database...")
        # file_path = os.path.basename(file_path)
        
        # Verify the file path
        if not os.path.isfile(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        print(f"File path: {file_path}")
        
        # Check file permissions
        if not os.access(file_path, os.R_OK):
            raise PermissionError(f"File is not readable: {file_path}")
        
        # print(file_path)
        
        loader = Docx2txtLoader(file_path)
        documents = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        text_chunks = text_splitter.split_documents(documents)
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        # vector_store = Chroma.from_documents(documents=text_chunks, embedding=embeddings)
        
        vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        # index = faiss.IndexFlatL2(len(embeddings.embed_query("hello world")))

        # vector_store = FAISS(
        #     embedding_function=embeddings,
        #     index=index,
        #     docstore=InMemoryDocstore(),
        #     index_to_docstore_id={},
        # )
        
        print("Vector database loaded successfully.") 
        return vector_store.as_retriever(search_kwargs={"k": 1})
    except Exception as e:
        print(f"Error loading vector database: {e}")
        return None

# import os

# async def load_vector_db(file_storage): 
#     try:
#         # Define the destination folder and ensure it exists
#         destination_folder = 'path/to/your/destination/folder'
#         if not os.path.exists(destination_folder):
#             os.makedirs(destination_folder)
        
#         # Construct the destination file path
#         file_path = os.path.join(destination_folder, file_storage.filename)
        
#         # Save the file to the destination folder
#         file_storage.save(file_path)
        
#         print("Loading vector database...")
#         print(f"File path: {file_path}")
        
#         loader = Docx2txtLoader(file_path)
#         documents = loader.load()
        
#         text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
#         text_chunks = text_splitter.split_documents(documents)
        
#         embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
#         vector_store = FAISS.from_documents(documents=text_chunks, embedding=embeddings)
        
#         print("Vector database loaded successfully.") 
#         return vector_store.as_retriever(search_kwargs={"k": 1})
#     except Exception as e:
#         print(f"Error loading vector database: {e}")
#         return None



# investment_personality = "Moderate Investor"
# Retrieval_Chain and Promot Template for Suggestions :
async def make_retrieval_chain(retriever,investmentPersonality,clientName,monthly_investment=10000,investment_period=3): # GET Method
    """
    Create a retrieval chain using the provided retriever.

    Args:
        retriever (RetrievalQA): A retriever object.

    Returns:
        RetrievalQA: A retrieval chain object.
    """
    try:
        # global investment_personality #,summary
        
        print(f"{retriever}\n {investmentPersonality}\n {clientName}\n {monthly_investment}")
        # try:
        #     print(type(investmentPersonality))
        # except Exception as e:
        #     print(f"Error in personality: {e}")
        #     return None
        
        # print(clientName)
        
        llm = ChatGoogleGenerativeAI(
            #model="gemini-pro",
            model = "gemini-1.5-flash",
            temperature = 0.45,
            # temperature=0.7,
            top_p=0.85,
            google_api_key=GOOGLE_API_KEY
        )
        # New Template 
        investment_period = str(investment_period)
        print(investmentPersonality)
        monthly_investment = str(monthly_investment)
        print(monthly_investment)
        print(investment_period)
        
        # New Prompt Template :
        
        prompt_template = """ 
                                You are a Financial Advisor tasked with creating responsible investment suggestions for a client based on their investment personality : """ + investmentPersonality +   "\n" + """ so that the client can reach their Financial Goals, based on their Financial Conditions.
                                Use the following instructions to ensure consistent output:
                                ---

                                ### Required Output Format:
                                
                                #### Client Financial Details:
                                - **Client Name**: """ + clientName + """
                                - **Assets**:
                                - List all asset types, their current values, and annual contributions in a tabular format (columns: "Asset Type", "Current Value", "Annual Contribution").
                                - **Liabilities**:
                                - List all liability types, their balances, interest rates, and monthly payments in a tabular format (columns: "Liability Type", "Balance", "Interest Rate", "Monthly Payment").
                                - **Other Details**:
                                - Retirement plan details, income sources, and goals should be listed in a clear and concise format.
                                - Client's Financial Condition : Analyze the Details and mention the Client's Financial Condition as : Stable/ Currently Stable / Unstable.
                                - **Investment Period** `Z years`
                                
                                #### Investment Allocation:
                                Split investments into **Growth-Oriented Investments** and **Conservative Investments**. Ensure each category includes:
                                - **Investment Type**: Specify the investment type (e.g., "Index Funds", "US Treasury Bonds").
                                - **Allocation Range**: Specify minimum and maximum allocation percentages (e.g., `10% - 20%`).
                                - **Target**: Describe the purpose of the investment.
                                - **How to Invest**: Provide instructions on how to invest in this asset.
                                - **Where to Invest**: Specify platforms or tools for making the investment.

                                **Example**:
                                **Growth-Oriented Investments (Minimum X% - Maximum Y%) **:
                                - **Stocks**: `20% - 30%`
                                - **ETFs**: `10% - 15%`
                                - **Mutual Funds**: `10% - 20%`
                                - **Cryptocurrency**: ` 5% - 10%`
                                - **Real Estates or REITS**: `10% - 20%`
                                - *Target*: Long-term growth potential aligned with the overall market performance tailored to fullfil Client's Financial Goals and manage his Financial Condition.
                                - *How to Invest*: Provide information on how to invest in which market 
                                - *Where to Invest*: Provide Information to buy which assets and how much to invest in terms of amount and percentage(%).Mention 5-6 assets.
                                
                                **Conservative Investments (Minimum X% - Maximum Y%) **:
                                - **High-Yield Savings Account**: `30% - 40%`
                                - **Bonds**: `10% - 20%`
                                - **Commodities**: `5% - 10%`
                                - **Cash**: `5% - 10%`
                                - *Target*: Maintain liquidity for emergencies.
                                - *How to Invest*: Provide information on how to invest.
                                - *Where to Invest*: Mention where to invest and how much to allocate in terms of money and percentage(%). Mention 5-6 assets.

                                #### Returns Overview:
                                - **Minimum Expected Annual Return**: `X% - Y%`
                                - **Maximum Expected Annual Return**: `X% - Y%`
                                - **Minimum Expected Growth in Dollars**: `$X - $Y` (based on the time horizon)
                                - **Maximum Expected Growth in Dollars**: `$X - $Y` (based on the time horizon)
                                - **Time Horizon**: `Z years`

                                ---

                                ### Example Output:
                                
                                #### Client Financial Details:
                                | Asset Type          | Current Value ($) | Annual Contribution ($) |
                                |----------------------|-------------------|--------------------------|
                                | 401(k), 403(b), 457  | 300               | 15                       |
                                | Traditional IRA      | 200               | 15                       |
                                | Roth IRA             | 500               | 28                       |
                                | Cash/Bank Accounts   | 500,000           | 30,000                   |
                                | Real Estate          | 1,000,000         | -                        |
                                | Total Assets Value   | 1,501,000         | -                        |

                                | Liability Type      | Balance ($) | Interest Rate (%) | Monthly Payment ($) |
                                |---------------------|-------------|--------------------|----------------------|
                                | Mortgage            | 1,000       | 10                | 100                  |
                                | Credit Card         | 400         | 8                 | 400                  |
                                | Other Loans         | 500         | 6                 | 100                  |
                                | Total Liabilities   | 1,900       | -                 | -                    |
                                
                                | Investrment Period | 3 years |
                                
                                **Growth-Oriented Investments (Minimum 40% - Maximum 80%)**:
                                - **Stocks**: `20% - 30%`
                                - **ETFs**: `5% - 10%`
                                - **Mutual Funds**: `5% - 20%`
                                - **Cryptocurrency**: ` 5% - 10%`
                                - **Real Estates or REITS**: `5% - 10%`
                                - *Target*: Long-term growth potential aligned with the market.
                                - *How to Invest*: Purchase low-cost index funds.
                                - *Where to Invest*: Stocks such as NVIDIA,AAPL, Vanguard, LiteCoin.

                                **Conservative Investments (Minimum 40% - Maximum 70%)**:
                                - **High-Yield Savings Account**: `20% - 30%`
                                - **Bonds**: `10% - 20%`
                                - **Commodities**: `5% - 10%`
                                - **Cash**: `5% - 10%`
                                - *Target*: Maintain liquidity for emergencies.
                                - *How to Invest*: Deposit funds into an FDIC-insured account.
                                - *Where to Invest*: Ally Bank, Capital One 360.

                                #### Returns Overview:
                                - **Minimum Expected Annual Return**: `4% - 6%`
                                - **Maximum Expected Annual Return**: `8% - 15%`
                                - **Minimum Expected Growth in Dollars**: `$4,000 - $6,000`
                                - **Maximum Expected Growth in Dollars**: `$8,000 - $15,000`
                                - **Time Horizon**: `3 years`

                                ---

                                Ensure the output strictly follows this structure.


                            ### Rationale for Investment Suggestions:
                            Provide a detailed explanation of why these suggestions align with the client’s financial personality and goals.

                            ---
                            <context>
                            {context}
                            </context>
                            Question: {input}

        """



        # #Wasnt consistent for generating the Bar Graph and Pie Chart :
        # prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
                # Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality and Financial Document provided to you.
                # Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
                # Also give the user detailed information about the investment how to invest,where to invest and how much they
                # should invest in terms of percentage of their investment amount based on the clients Financial Conditions and help them to cover up their Mortgage and Debts if any.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
                # Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
                # investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon and how it can help them accumulate wearlth overtime to achive their Financial  goals.
                # Also give the user minimum and maximum expected growth in dollars for the time horizon .
                # Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
                
                # You are a Financial Advisor for question-answering tasks related to the document. Based on the client's investment personality and financial details provided, generate responsible investment suggestions to achieve their financial goals while managing debts.

                # Step-by-Step Guidance:
                # 1. Assets: Calculate total assets by analyzing the provided financial document in the My Assets section. Ensure you include cash, real estate, retirement accounts, brokerage accounts, and any other relevant asset types from the document.
                # 2. Liabilities: Calculate total liabilities by analyzing the provided financial document in the My Liabilities section. Consider mortgages, credit card debts, student loans, car loans, and other liabilities. 
                # 3. Monthly Investment Feasibility: Use the client's assets and liabilities to assess whether their planned monthly investment is feasible. If not feasible, suggest a more realistic monthly investment amount.
                # 4. Analyze Liabilities: Determine if the client's monthly investment plan is feasible after covering liabilities and expected expenses and also considering some amount for savings. If the client's monthly investment plan is not feasible after covering expenses and savings, generate investment suggestions on a smaller monthly investment plan amount if it can help the client else mention amount is too small for the client's requirementys to be made.
                # 5. Investment Strategy: Suggest a strategy where monthly investments can both generate returns and pay off debts effectively and helps client to achieve their financial goals.
                # 6. Allocation: Provide detailed allocations between growth-oriented investments and conservative investments, ensuring the client can meet their monthly debt obligations and save for their future financial goals.
                # 7. Returns: Include minimum and maximum compounded returns over 5-10 years, along with inflation-adjusted returns for clarity.
                # 8. Suggestions: Offer advice on how to use remaining funds to build wealth after clearing liabilities and achive their financial goal.
                
                
        #         Here's an example for the required Output Format(if there are comments indicated by # in the example output format then thats a side note for your reference dont write it in the response that will be generated ) :
                
        #         Client's Financial Information :(# This is a header line have it in bold) 
                
                
        #         Client Name: """ + clientName + """(# have the client name in underline)

        #         Financial Overview: (#the data presented is just an example for your reference do not consider it as factual refere to the document provided to you and generate data based on the provided data and only when nothing is provided assume some data for analysis, This is a header line have it in bold. The data below it should be displayed in a table format so make sure of that data.There must be 2 columns 1 for Category and second for Value.List down all the assets and liabilities along with its values and then Total of assets,liabilities,etc.)
                
        #         - Total Assets: (# Sum of all client assets and Annual Income . Mention all assets and their respected values.if non consider the example assets)
                
        #         - Total Liabilities: (# Sum of all liabilities. Mention all liabilities and their respected values if non consider the example liabilities)
                
                
        #         - Monthly Liabilities: (# Monthly payments derived from liabilities)
                
        #         - Total Annual Income : (# Sum of all client's anual income)
                
        #         - Monthly Investment Amount : """ + monthly_investment + """ (# if no specific amount is specified to you then only assume  10,000 else consider the amount mention to you and just display the amount)
                
        #         - Investment Period : """ + investment_period + """  (# if no specific period is specified to you then only assume 3 years else consider the period mention to you and just display the period)


        #         Financial Analysis :(#Analyse the assets and liabilities and based on that give a suggestion for analysis generate suggestions for one of the following conditions:)
        #         (#1st condition : Everything is Positive)Based on the given Financial Conditions the client is having a good and stable income with great assets and manageable debt and liabilities.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the clients monthly income.
        #         (# if this condition is true then ignore the other conditions and start with the Investment Suggestions)
                
        #         (#2nd condition : Everything is temporarily Negative) Based on the given Financial Conditions the client is facing a low income for now but have great assets and manageable debt and liabilities.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the client's monthly income but the client might not be able to sustain the monthly investment amount that they are planning.)
        #         Instead I would like to recommend this amount to the client for their monthly investment : (#Mention a feasible amount to the client for monthly investment and start suggesting investments based on this amount and not the previous amount being taken into consideration)
                
        #         (#3rd condition : Everything is Negative) Based on the given Financial Conditions the client is facing a low income and doesnt have good assets to manage the debts and liabilities of the client and in such a condition this monthly investment amount is not feasible.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is not manageable for the client's monthly income and so the client might not be able to sustain the monthly investment amount that they are planning to do.)
        #         I would like to recommend this amount to the client for monthly investment : (# Mention a minimum amount to the client for monthly investment if possible else just say the client should first prioritize on savings and generating more income to manage their debts and liabilities first and so dont give any investment suggestions to the client.)
                
        #         (#If the financial is 1 or 2 only then give investment suggestions to the client)
                
                
                
        #         Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

        #         Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

                
        #         Investment Allocation: (#remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

        #         Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
        #         How to Invest: Diversify across various asset classes like:  (#Give allocations % as well)
                
        #         Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
                
        #         ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
                
        #         Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
                
        #         Consider investing in blue-chip companies or growth sectors like technology. 
                
                
        #         Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


        #         Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
        #         Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

        #         Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
        #         How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
        #         Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
        #         Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
        #         Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
        #         Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
        #         Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
        #         Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.


        #         Time Horizon and Expected Returns:

        #         Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
        #         Minimum Expected Annual Return: 4% - 6% 
                
                
        #         Maximum Expected Annual Return: 8% - 10% 
                
                
        #         Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, (# consider the monthly investment amount and give returns based on that only) $10,000 could grow to approximately 17,908 in 10 years.
        #         Minimum Expected Growth in Dollars: 
                
        #         4,000−6,000 (over 10 years) 
                
                
        #         Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
        #         Inflation Adjusted Returns:(#do not write this part inside the bracket just give answer,assume US inflation rate assume 3% if you dont know, and give the investment returns value that was suggested by you for the considered monthly investment amount after 3,5,10years of growth mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
        #         Rationale for Investment Suggestions:

        #         This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
        #         Important Considerations:

        #         Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

        #         Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

        #         Explain how this suggestions can help the client grow their wealth and improve their financial condition and/or cover up thier loans and in turn achive their Financial goals.
        #         <context>
        #         {context}
        #         </context>
        #         Question: {input}"""

        
        # # Without category and value :
        # prompt_template = investmentPersonality +   "\n" + """You are a Financial Advisor for question-answering tasks related to the document.
        #         Give Financial Suggestions to the Wealth Manager so that they could do proper responsible investment based on their client's investment personality and Financial Document provided to you.
        #         Always Mention the Investment for the """ + clientName + """(clientName) provided to you.
        #         Also give the user detailed information about the investment how to invest,where to invest and how much they
        #         should invest in terms of percentage of their investment amount based on the clients Financial Conditions and help them to cover up their Mortgage and Debts if any.Give the user minimum and maximum percentage of growth-oriented investments alloacation.
        #         Give the user detailed information about the returns on their investment by giving them an approximate return based on the time horizon of the investment based on which calculate the compunded returns on their 
        #         investment.Also Give the user minimum and maximum expected annual return percentage for the time horizon and how it can help them accumulate wearlth overtime to achive their Financial  goals.
        #         Also give the user minimum and maximum expected growth in dollars for the time horizon .
        #         Also explain the user why you are giving them that particular investment suggestions for the client with the given investment personality.
                
        #         You are a Financial Advisor for question-answering tasks related to the document. Based on the client's investment personality and financial details provided, generate responsible investment suggestions to achieve their financial goals while managing debts.

        #         Step-by-Step Guidance:
        #         1. Assets: Calculate total assets by analyzing the provided financial document in the My Assets section. Ensure you include cash, real estate, retirement accounts, brokerage accounts, and any other relevant asset types from the document.
        #         2. Liabilities: Calculate total liabilities by analyzing the provided financial document in the My Liabilities section. Consider mortgages, credit card debts, student loans, car loans, and other liabilities. 
        #         3. Monthly Investment Feasibility: Use the client's assets and liabilities to assess whether their planned monthly investment is feasible. If not feasible, suggest a more realistic monthly investment amount.
        #         4. Analyze Liabilities: Determine if the client's monthly investment plan is feasible after covering liabilities and expected expenses and also considering some amount for savings. If the client's monthly investment plan is not feasible after covering expenses and savings, generate investment suggestions on a smaller monthly investment plan amount if it can help the client else mention amount is too small for the client's requirementys to be made.
        #         5. Investment Strategy: Suggest a strategy where monthly investments can both generate returns and pay off debts effectively and helps client to achieve their financial goals.
        #         6. Allocation: Provide detailed allocations between growth-oriented investments and conservative investments, ensuring the client can meet their monthly debt obligations and save for their future financial goals.
        #         7. Returns: Include minimum and maximum compounded returns over 5-10 years, along with inflation-adjusted returns for clarity.
        #         8. Suggestions: Offer advice on how to use remaining funds to build wealth after clearing liabilities and achive their financial goal.
                
                
        #         Here's an example for the required Output Format(if there are comments indicated by # in the example output format then thats a side note for your reference dont write it in the response that will be generated ) :
                
        #         Client's Financial Information :(# This is a header line have it in bold) 
                
                
        #         Client Name: """ + clientName + """(# have the client name in underline)

        #         (#the data presented is just an example for your reference do not consider it as factual refere to the document provided to you and generate data based on the provided data and only when nothing is provided assume some data for analysis.The data below it should be displayed in a table format so make sure of that data.List down all the assets and liabilities along with its values and then Total of assets,liabilities,etc.)
                
        #         - Total Assets: (# Sum of all client assets and Annual Income . Mention all assets and their respected values.if non consider the example assets)
                
        #         - Total Liabilities: (# Sum of all liabilities. Mention all liabilities and their respected values if non consider the example liabilities)
                
                
        #         - Monthly Liabilities: (# Monthly payments derived from liabilities)
                
        #         - Total Annual Income : (# Sum of all client's anual income)
                
        #         - Monthly Investment Amount : """ + monthly_investment + """ (# if no specific amount is specified to you then only assume  10,000 else consider the amount mention to you and just display the amount)
                
        #         - Investment Period : """ + investment_period + """  (# if no specific period is specified to you then only assume 3 years else consider the period mention to you and just display the period)


        #         Financial Analysis :(#Analyse the assets and liabilities and based on that give a suggestion for analysis generate suggestions for one of the following conditions:)
        #         (#1st condition : Everything is Positive)Based on the given Financial Conditions the client is having a good and stable income with great assets and manageable debt and liabilities.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the clients monthly income.
        #         (# if this condition is true then ignore the other conditions and start with the Investment Suggestions)
                
        #         (#2nd condition : Everything is temporarily Negative) Based on the given Financial Conditions the client is facing a low income for now but have great assets and manageable debt and liabilities.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is manageable for the client's monthly income but the client might not be able to sustain the monthly investment amount that they are planning.)
        #         Instead I would like to recommend this amount to the client for their monthly investment : (#Mention a feasible amount to the client for monthly investment and start suggesting investments based on this amount and not the previous amount being taken into consideration)
                
        #         (#3rd condition : Everything is Negative) Based on the given Financial Conditions the client is facing a low income and doesnt have good assets to manage the debts and liabilities of the client and in such a condition this monthly investment amount is not feasible.
        #         Clients monthly expenses on debts is : (#mention the calculated liabilities for a month) , which is not manageable for the client's monthly income and so the client might not be able to sustain the monthly investment amount that they are planning to do.)
        #         I would like to recommend this amount to the client for monthly investment : (# Mention a minimum amount to the client for monthly investment if possible else just say the client should first prioritize on savings and generating more income to manage their debts and liabilities first and so dont give any investment suggestions to the client.)
                
        #         (#If the financial is 1 or 2 only then give investment suggestions to the client)
                
                
        #         Investment Suggestions for """ + clientName + """  with a Moderate Investor Personality(This is just an example for Moderate Investor but you need to generate suggestions for the given investment personality) (This must be like a Header and in Bold)

        #         Based on your provided information, you appear to be a moderate investor with a healthy mix of assets and liabilities. Here's a breakdown of investment suggestions tailored to your profile:

        #         Investment Allocation: (#remember these allocations is just an example you can suggest other investments dpeneding on the details and investor personality provided)

        #         Growth-Oriented Investments (Minimum 40% - Maximum 60%): Target: Focus on investments with the potential for long-term growth while managing risk. 
        #         How to Invest: Diversify across various asset classes like:  (#Give allocations % as well)
                
        #         Mutual Funds(5%-10%): Choose diversified index funds tracking the S&P 500 or broad market indices. 
                
        #         ETFs(10%-20%): Offer similar benefits to mutual funds but with lower fees and more transparency. 
                
        #         Individual Stocks(20%-30%): Carefully select companies with solid financials and growth potential. 
                
        #         Consider investing in blue-chip companies or growth sectors like technology. 
                
                
        #         Where to Invest: Brokerage Accounts: Choose a reputable online broker offering research tools and low fees.


        #         Roth IRA/Roth 401(k): Utilize these tax-advantaged accounts for long-term growth and tax-free withdrawals in retirement. 
                
                
        #         Percentage Allocation for Growth-Oriented Investments: Allocate between 40% and 60% of your investable assets towards these growth-oriented investments. This range allows for flexibility based on your comfort level and market conditions.

        #         Conservative Investments (Minimum 40% - Maximum 60%): Target: Prioritize safety and capital preservation with lower risk. 
        #         How to Invest: Bonds: Invest in government or corporate bonds with varying maturities to match your time horizon. 
                
        #         Cash: Maintain a cash reserve in high-yield savings accounts or short-term CDs for emergencies and upcoming expenses. 
                
        #         Real Estate: Consider investing in rental properties or REITs (Real Estate Investment Trusts) for diversification and potential income generation. 
                
        #         Where to Invest: Brokerage Accounts: Invest in bond mutual funds, ETFs, or individual bonds. 
                
        #         Cash Accounts(20%-30%): Utilize high-yield savings accounts or short-term CDs offered by banks or credit unions. 
                
        #         Real Estate(20%-30%): Invest directly in rental properties or through REITs available through brokerage accounts. 
                
        #         Percentage Allocation for Conservative Investments: Allocate between 40% and 60% of your investable assets towards these conservative investments. This range ensures a balance between growth and security.


        #         Time Horizon and Expected Returns:

        #         Time Horizon: As a moderate investor, your time horizon is likely long-term, aiming for returns over 5-10 years or more. 
                
                
        #         Minimum Expected Annual Return: 4% - 6% 
                
                
        #         Maximum Expected Annual Return: 8% - 10% 
                
                
        #         Compounded Returns: The power of compounding works in your favor over the long term. With a 6% average annual return, (# consider the monthly investment amount and give returns based on that only) $10,000 could grow to approximately 17,908 in 10 years.
        #         Minimum Expected Growth in Dollars: 
                
        #         4,000−6,000 (over 10 years) 
                
                
        #         Maximum Expected Growth in Dollars: 8,000−10,000 (over 10 years)

                
        #         Inflation Adjusted Returns:(#do not write this part inside the bracket just give answer,assume US inflation rate assume 3% if you dont know, and give the investment returns value that was suggested by you for the considered monthly investment amount after 3,5,10years of growth mention the values before adjusting and after adjusting with inflation I want it in a bulleted format)
                   
                    
        #         Rationale for Investment Suggestions:

        #         This investment strategy balances growth potential with risk management. The allocation towards growth-oriented investments allows for potential capital appreciation over time, while the allocation towards conservative investments provides stability and safeguards your principal.

                
        #         Important Considerations:

        #         Regular Review: Periodically review your portfolio and adjust your allocation as needed based on market conditions, your risk tolerance, and your financial goals. Professional Advice: Consider seeking advice from a qualified financial advisor who can provide personalized guidance and help you develop a comprehensive financial plan.

        #         Disclaimer: This information is for educational purposes only and should not be considered financial advice. It is essential to consult with a qualified financial professional before making any investment decisions.

        #         Explain how this suggestions can help the client grow their wealth and improve their financial condition and/or cover up thier loans and in turn achive their Financial goals.
        #         <context>
        #         {context}
        #         </context>
        #         Question: {input}"""
                
        print("Retriever Created ")
        print(f"Investment Personality :{investmentPersonality}")
        
                

        llm_prompt = ChatPromptTemplate.from_template(prompt_template)

        document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
        combine_docs_chain = None  

        if retriever is not None :  
            retriever_chain = create_retrieval_chain(retriever,document_chain) 
            # print(retriever_chain)
            return retriever_chain
        else:
            print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
            return None

    except Exception as e:
        print(f"Error in creating chain: {e}")
        return None


import json
import io

# Process_Documents :
async def process_document(file_path): # GET Method
    try:
        print("Processing the document")
        file_type = filetype.guess(file_path)
        if file_type is not None:
            if file_type.mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                # Await the coroutine to extract text and tables
                return await extract_text_and_tables_from_word(file_path)
            elif file_type.mime == "application/pdf":
                return await extract_text_from_pdf(file_path)
        return None
    except Exception as e:
        print(f"Error processing document: {e}")
        return None

# Async function to extract text from a PDF file
async def extract_text_from_pdf(pdf_file_path): # GET Method
    try:
        print("Processing pdf file")
        with open(pdf_file_path, "rb") as pdf_file:
            pdf_reader = PyPDF2.PdfFileReader(pdf_file)
            text_content = []
            for page_num in range(pdf_reader.numPages):
                page = pdf_reader.getPage(page_num)
                text_content.append(page.extract_text())
            return "\n".join(text_content)
    except Exception as e:
        print(f"Error extracting text from PDF: {e}")
        return None

# Async function to extract text and tables from a Word document
async def extract_text_and_tables_from_word(docx_file_path): # GET Method
    try:
        print("Extracting text and tables from word file")
        doc = docx.Document(docx_file_path)
        text_content = []
        tables_content = []

        for para in doc.paragraphs:
            text_content.append(para.text)

        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables_content.append(table_data)
        print("Extracted text from word file")
        return "\n".join(text_content), tables_content
    except Exception as e:
        print(f"Error extracting text and tables from Word document: {e}")
        return None, None



async def validate_document_content(text, tables):
    """
    Validates the content of the document.

    Args:
        text (str): Extracted text content from the document.
        tables (list): Extracted tables content from the document.

    Returns:
        tuple: Client name and validation errors.
    """
    errors = []
    
    # Extract client name
    client_name_match = re.search(r"Client Name:\s*([^\n]+)", text, re.IGNORECASE)
    client_name = client_name_match.group(1).strip().split(" ")[0] if client_name_match else "Unknown"

    # Define required sections
    required_sections = [
        "YOUR RETIREMENT GOAL",
        "YOUR OTHER MAJOR GOALS",
        "YOUR ASSETS AND LIABILITIES",
        "MY LIABILITIES",
        "YOUR CURRENT ANNUAL INCOME"
    ]

    # Check for the presence of required sections
    for section in required_sections:
        if section not in text:
            errors.append(f"* {section} section missing.")
    
    # Define table field checks
    table_checks = {
        "YOUR RETIREMENT GOAL": [
            r"When do you plan to retire\? \(age or date\)",
            r"Social Security Benefit \(include expected start date\)",
            r"Pension Benefit \(include expected start date\)",
            r"Other Expected Income \(rental, part-time work, etc.\)",
            r"Estimated Annual Retirement Expense"
        ],
        "YOUR OTHER MAJOR GOALS": [
            r"GOAL", r"COST", r"WHEN"
        ],
        "YOUR ASSETS AND LIABILITIES": [
            r"Cash/bank accounts", r"Home", r"Other Real Estate", r"Business",
            r"Current Value", r"Annual Contributions"
        ],
        "MY LIABILITIES": [
            r"Balance", r"Interest Rate", r"Monthly Payment"
        ]
    }

    # Validate table content
    for section, checks in table_checks.items():
        section_found = False
        for table in tables:
            table_text = "\n".join(["\t".join(row) for row in table])
            if section in table_text:
                section_found = True
                for check in checks:
                    if not re.search(check, table_text, re.IGNORECASE):
                        errors.append(f"* Missing or empty field in {section} section: {check}")
                break
        if not section_found:
            errors.append(f"* {section} section missing.")

    return client_name, errors

####################################################################################################################################

################################################## Extract Numerical Data for Pie Chart, Bar Graph and Line Chart #####################################

import re
from collections import defaultdict
import numpy as np
# Updated for Line Chart :
import re
from collections import defaultdict

# def extract_numerical_data(response):
#     # Define patterns to match different sections and their respective allocations
#     patterns = {
#         'Growth-Oriented Investments': re.compile(r'Growth-Oriented Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
#         'Conservative Investments': re.compile(r'Conservative Investments.*?How to Invest:(.*?)Where to Invest:', re.DOTALL),
#         'Time Horizon and Expected Returns': re.compile(r'Time Horizon and Expected Returns:(.*?)$', re.DOTALL)
#     }

#     data = defaultdict(dict)

#     for section, pattern in patterns.items():
#         match = pattern.search(response)
#         if match:
#             investments_text = match.group(1)
#             # Extract individual investment types and their allocations
#             investment_pattern = re.compile(r'(\w[\w\s]+?)\s*\((\d+%)-(\d+%)\)')
#             for investment_match in investment_pattern.findall(investments_text):
#                 investment_type, min_allocation, max_allocation = investment_match
#                 data[section][investment_type.strip()] = {
#                     'min': min_allocation,
#                     'max': max_allocation
#                 }

#     # Extract time horizon and expected returns
#     time_horizon_pattern = re.compile(r'Time Horizon:.*?(\d+)-(\d+) years', re.IGNORECASE)
#     min_return_pattern = re.compile(r'Minimum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
#     max_return_pattern = re.compile(r'Maximum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
#     min_growth_pattern = re.compile(r'Minimum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)
#     max_growth_pattern = re.compile(r'Maximum Expected Growth in Dollars:.*?\$(\d+,\d+)-\$(\d+,\d+)', re.IGNORECASE)

#     time_horizon_match = time_horizon_pattern.search(response)
#     min_return_match = min_return_pattern.search(response)
#     max_return_match = max_return_pattern.search(response)
#     min_growth_match = min_growth_pattern.search(response)
#     max_growth_match = max_growth_pattern.search(response)

#     if time_horizon_match:
#         data['Time Horizon'] = {
#             'min_years': time_horizon_match.group(1),
#             'max_years': time_horizon_match.group(2)
#         }

#     if min_return_match:
#         data['Expected Annual Return'] = {
#             'min': min_return_match.group(1),
#             'max': min_return_match.group(2)
#         }

#     if max_return_match:
#         data['Expected Annual Return'] = {
#             'min': max_return_match.group(1),
#             'max': max_return_match.group(2)
#         }

#     if min_growth_match:
#         data['Expected Growth in Dollars'] = {
#             'min': min_growth_match.group(1),
#             'max': min_growth_match.group(2)
#         }

#     if max_growth_match:
#         data['Expected Growth in Dollars'] = {
#             'min': max_growth_match.group(1),
#             'max': max_growth_match.group(2)
#         }

#     # Extract inflation-adjusted returns
#     inflation_adjusted_pattern = re.compile(r'Inflation Adjusted Returns:.*?Before Inflation:.*?3 Years: \$(\d+,\d+).*?5 Years: \$(\d+,\d+).*?10 Years: \$(\d+,\d+).*?After Inflation.*?3 Years: \$(\d+,\d+).*?5 Years: \$(\d+,\d+).*?10 Years: \$(\d+,\d+)', re.DOTALL)
#     inflation_adjusted_match = inflation_adjusted_pattern.search(response)

#     if inflation_adjusted_match:
#         data['Inflation Adjusted Returns'] = {
#             'Before Inflation': {
#                 '3 Years': inflation_adjusted_match.group(1),
#                 '5 Years': inflation_adjusted_match.group(2),
#                 '10 Years': inflation_adjusted_match.group(3)
#             },
#             'After Inflation': {
#                 '3 Years': inflation_adjusted_match.group(4),
#                 '5 Years': inflation_adjusted_match.group(5),
#                 '10 Years': inflation_adjusted_match.group(6)
#             }
#         }

#     print(f"DATA extracted from Responses : {data}")
#     return data

# new code:
import re
from collections import defaultdict
import re
from collections import defaultdict

# extract numerical data from responses :

def extract_numerical_data(response):
    data = defaultdict(dict)

    # Match Growth-Oriented Investments and Conservative Investments sections
    growth_pattern = re.compile(r"<strong>Growth-Oriented Investments.*?</strong>:\s*(.*?)(<strong>|<h4>)", re.DOTALL)
    conservative_pattern = re.compile(r"<strong>Conservative Investments.*?</strong>:\s*(.*?)(<strong>|<h4>)", re.DOTALL)
    allocation_pattern = re.compile(r"<strong>(.*?)</strong>:\s*<code>(\d+%)\s*-\s*(\d+%)</code>")

    for category, pattern in [("Growth-Oriented Investments", growth_pattern), 
                               ("Conservative Investments", conservative_pattern)]:
        match = pattern.search(response)
        if match:
            investments_text = match.group(1)
            for investment_match in allocation_pattern.findall(investments_text):
                investment_type, min_allocation, max_allocation = investment_match
                data[category][investment_type.strip()] = {
                    'min': min_allocation.strip('%'),
                    'max': max_allocation.strip('%')
                }

    # Match Returns Overview
    returns_pattern = re.compile(r"<h4>Returns Overview:</h4>\s*(.*?)\s*<h4>", re.DOTALL)
    returns_match = returns_pattern.search(response)
    if returns_match:
        returns_text = returns_match.group(1)

        # Extract returns and growth data
        min_return_match = re.search(r"Minimum Expected Annual Return</strong>:\s*<code>(\d+%)\s*-\s*(\d+%)</code>", returns_text)
        max_return_match = re.search(r"Maximum Expected Annual Return</strong>:\s*<code>(\d+%)\s*-\s*(\d+%)</code>", returns_text)
        min_growth_match = re.search(r"Minimum Expected Growth in Dollars</strong>:\s*<code>\$(\d+,\d+)\s*-\s*\$(\d+,\d+)</code>", returns_text)
        max_growth_match = re.search(r"Maximum Expected Growth in Dollars</strong>:\s*<code>\$(\d+,\d+)\s*-\s*\$(\d+,\d+)</code>", returns_text)
        time_horizon_match = re.search(r"Time Horizon</strong>:\s*<code>(\d+ years)</code>", returns_text)

        if min_return_match:
            data['Expected Annual Return'] = {
                'min': min_return_match.group(1),
                'max': min_return_match.group(2)
            }
        if max_return_match:
            data['Expected Annual Return']['max'] = max_return_match.group(2)

        if min_growth_match:
            data['Expected Growth in Dollars'] = {
                'min': min_growth_match.group(1).replace(',', ''),
                'max': min_growth_match.group(2).replace(',', '')
            }
        if max_growth_match:
            data['Expected Growth in Dollars']['max'] = max_growth_match.group(2).replace(',', '')

        if time_horizon_match:
            data['Time Horizon'] = time_horizon_match.group(1)

    return data



# just prev code :
# def extract_numerical_data(response):
#     # Patterns for different sections
#     patterns = {
#         'Growth-Oriented Investments': re.compile(r'Growth-Oriented Investments.*?Target.*?:(.*?)Where to Invest:', re.DOTALL),
#         'Conservative Investments': re.compile(r'Conservative Investments.*?Target.*?:(.*?)Where to Invest:', re.DOTALL),
#         'Time Horizon and Expected Returns': re.compile(r'Time Horizon and Expected Returns.*?:\s*(.*?)$', re.DOTALL)
#     }

#     data = defaultdict(dict)

#     for section, pattern in patterns.items():
#         match = pattern.search(response)
#         if match:
#             investments_text = match.group(1)
#             # Extract investment details
#             investment_pattern = re.compile(r'([\w\s&/-]+?)\s*\((\d+%)-(\d+%)\)')
#             for investment_match in investment_pattern.findall(investments_text):
#                 investment_type, min_allocation, max_allocation = investment_match
#                 data[section][investment_type.strip()] = {
#                     'min': min_allocation.strip(),
#                     'max': max_allocation.strip()
#                 }

#     # Extract additional details
#     time_horizon_pattern = re.compile(r'Time Horizon.*?(\d+)-(\d+)\s*years', re.IGNORECASE)
#     min_return_pattern = re.compile(r'Minimum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
#     max_return_pattern = re.compile(r'Maximum Expected Annual Return:.*?(\d+%)-(\d+%)', re.IGNORECASE)
#     min_growth_pattern = re.compile(r'Minimum Expected Growth in Dollars:.*?\$(\d[\d,]*)-\$(\d[\d,]*)', re.IGNORECASE)
#     max_growth_pattern = re.compile(r'Maximum Expected Growth in Dollars:.*?\$(\d[\d,]*)-\$(\d[\d,]*)', re.IGNORECASE)

#     time_horizon_match = time_horizon_pattern.search(response)
#     if time_horizon_match:
#         data['Time Horizon'] = {
#             'min_years': int(time_horizon_match.group(1)),
#             'max_years': int(time_horizon_match.group(2))
#         }

#     min_return_match = min_return_pattern.search(response)
#     if min_return_match:
#         data['Expected Annual Return'] = {
#             'min': min_return_match.group(1),
#             'max': min_return_match.group(2)
#         }

#     max_growth_match = max_growth_pattern.search(response)
#     if max_growth_match:
#         data['Expected Growth in Dollars'] = {
#             'min': int(max_growth_match.group(1).replace(',', '')),
#             'max': int(max_growth_match.group(2).replace(',', ''))
#         }

#     print("Section Data Extracted:", data)
#     print("Growth-Oriented Investments:", data.get('Growth-Oriented Investments', 'Not Found'))
#     print("Conservative Investments:", data.get('Conservative Investments', 'Not Found'))
#     print("Time Horizon Data:", data.get('Time Horizon', 'Not Found'))

#     return data


def normalize_allocations(allocations):
    total = sum(allocations)
    if total == 100:
        return allocations
    return [round((allocation / total) * 100, 2) for allocation in allocations]

# # Updated Line Chart 
import datetime  # Import the datetime module to get the current year

# line chart data code :

def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4):
    try:
        # Get the current year
        curr_year = datetime.datetime.now().year

        # Print data_extracted to debug the structure
        print("Data extracted:", data_extracted)

        # Check if 'Expected Annual Return' and 'Time Horizon' exist and have the expected keys
        if 'Expected Annual Return' not in data_extracted:
            print("'Expected Annual Return' missing in data_extracted")
            data_extracted['Expected Annual Return'] = {'min': '8%', 'max': '20%'}
            min_return = 8 #6
            max_return = 20 #8
        else:
            min_return = float(data_extracted['Expected Annual Return'].get('min', '0').strip('%'))
            max_return = float(data_extracted['Expected Annual Return'].get('max', '0').strip('%'))

        min_years = int(data_extracted['Time Horizon'].get('min_years', 1))  # Default to 1 year if missing
        max_years = int(data_extracted['Time Horizon'].get('max_years', 10))  # Default to 10 years if missing

        def calculate_compounded_return(principal, rate, years):
            return principal * (1 + rate / 100) ** years

        def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
            return nominal_return / (1 + inflation_rate / 100) ** years

        # Create labels for the next 10 years starting from the current year
        labels = list(range(curr_year, curr_year + max_years))

        min_compounded = []
        max_compounded = []
        min_inflation_adjusted = []
        max_inflation_adjusted = []

        for year in range(1, max_years + 1):
            # Calculate nominal compounded returns
            min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
            max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

            # Calculate inflation-adjusted compounded returns
            min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
            max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

            # Append results
            min_compounded.append(min_compounded_value)
            max_compounded.append(max_compounded_value)
            min_inflation_adjusted.append(min_inflation_value)
            max_inflation_adjusted.append(max_inflation_value)

        # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
        combined_chart_data = {
            'labels': labels,  # Current year and the next 10 years
            'datasets': [
                {
                    'label': 'Minimum Compounded Return',
                    'data': min_compounded,
                    'borderColor': 'rgb(255, 99, 132)',  # Red color
                    'fill': False
                },
                {
                    'label': 'Maximum Compounded Return',
                    'data': max_compounded,
                    'borderColor': 'rgb(54, 162, 235)',  # Blue color
                    'fill': False
                },
                {
                    'label': 'Min Inflation Adjusted Return',
                    'data': min_inflation_adjusted,
                    'borderColor': 'rgb(75, 192, 192)',  # Light blue
                    'borderDash': [5, 5],  # Dashed line for distinction
                    'fill': False
                },
                {
                    'label': 'Max Inflation Adjusted Return',
                    'data': max_inflation_adjusted,
                    'borderColor': 'rgb(153, 102, 255)',  # Light purple
                    'borderDash': [5, 5],  # Dashed line for distinction
                    'fill': False
                }
            ]
        }
    except KeyError as e:
        print(f"KeyError occurred: {e}")
        return jsonify({'message': f'Key Error: {e}'}), 400
    except Exception as e:
        print(f"Error occurred while preparing data for combined line chart: {e}")
        return jsonify({'message': 'Internal Server Error in creating line chart'}), 500

    print(combined_chart_data)
    return combined_chart_data


# import datetime  # Import the datetime module to get the current year
# # uodated to have current year
# def prepare_combined_line_chart_data(data_extracted, initial_investment, inflation_rate=4):
#     try:
#         # Get the current year
#         curr_year = datetime.datetime.now().year

#         # Print data_extracted to debug the structure
#         print("Data extracted:", data_extracted)

#         # Check if 'Expected Annual Return' and 'Time Horizon' exist and have the expected keys
#         if 'Expected Annual Return' not in data_extracted:
#             print("'Expected Annual Return' missing in data_extracted")
#             data_extracted['Expected Annual Return']['min'] = 6
#             data_extracted['Expected Annual Return']['max'] = 8
#             min_return = 6
#             max_return = 8
#         else:
#             min_return = float(data_extracted['Expected Annual Return'].get('min', '0').strip('%'))
#             max_return = float(data_extracted['Expected Annual Return'].get('max', '0').strip('%'))

#         min_years = int(data_extracted['Time Horizon'].get('min_years', 1))  # Default to 1 year if missing
#         max_years = int(data_extracted['Time Horizon'].get('max_years', 10))  # Default to 10 years if missing

#         def calculate_compounded_return(principal, rate, years):
#             return principal * (1 + rate / 100) ** years

#         def calculate_inflation_adjusted_return(nominal_return, inflation_rate, years):
#             return nominal_return / (1 + inflation_rate / 100) ** years

#         # Create labels for the next 10 years starting from the current year
#         labels = list(range(curr_year, curr_year + max_years))

#         min_compounded = []
#         max_compounded = []
#         min_inflation_adjusted = []
#         max_inflation_adjusted = []

#         for year in range(1, max_years + 1):
#             # Calculate nominal compounded returns
#             min_compounded_value = calculate_compounded_return(initial_investment, min_return, year)
#             max_compounded_value = calculate_compounded_return(initial_investment, max_return, year)

#             # Calculate inflation-adjusted compounded returns
#             min_inflation_value = calculate_inflation_adjusted_return(min_compounded_value, inflation_rate, year)
#             max_inflation_value = calculate_inflation_adjusted_return(max_compounded_value, inflation_rate, year)

#             # Append results
#             min_compounded.append(min_compounded_value)
#             max_compounded.append(max_compounded_value)
#             min_inflation_adjusted.append(min_inflation_value)
#             max_inflation_adjusted.append(max_inflation_value)

#         # Combined Line Chart Data for both Nominal and Inflation-Adjusted Compounded Returns
#         combined_chart_data = {
#             'labels': labels,  # Current year and the next 10 years
#             'datasets': [
#                 {
#                     'label': 'Minimum Compounded Return',
#                     'data': min_compounded,
#                     'borderColor': 'rgb(255, 99, 132)',  # Red color
#                     'fill': False
#                 },
#                 {
#                     'label': 'Maximum Compounded Return',
#                     'data': max_compounded,
#                     'borderColor': 'rgb(54, 162, 235)',  # Blue color
#                     'fill': False
#                 },
#                 {
#                     'label': 'Min Inflation Adjusted Return',
#                     'data': min_inflation_adjusted,
#                     'borderColor': 'rgb(75, 192, 192)',  # Light blue
#                     'borderDash': [5, 5],  # Dashed line for distinction
#                     'fill': False
#                 },
#                 {
#                     'label': 'Max Inflation Adjusted Return',
#                     'data': max_inflation_adjusted,
#                     'borderColor': 'rgb(153, 102, 255)',  # Light purple
#                     'borderDash': [5, 5],  # Dashed line for distinction
#                     'fill': False
#                 }
#             ]
#         }
#     except KeyError as e:
#         print(f"KeyError occurred: {e}")
#         return jsonify({'message': f'Key Error: {e}'}), 400
#     except Exception as e:
#         print(f"Error occurred while preparing data for combined line chart: {e}")
#         return jsonify({'message': 'Internal Server Error in creating line chart'}), 500

#     return combined_chart_data





# def plot_investment_allocations(data):
#     # Create subplots with a large figure size
#     fig, axes = plt.subplots(2, 1, figsize= (16,10)) #(28, 15))  # Adjust size as needed

#     # Plot Growth-Oriented Investments
#     growth_data = data['Growth-Oriented Investments']
#     growth_labels = list(growth_data.keys())
#     growth_min = [int(growth_data[label]['min'].strip('%')) for label in growth_labels]
#     growth_max = [int(growth_data[label]['max'].strip('%')) for label in growth_labels]

#     axes[0].bar(growth_labels, growth_min, color='skyblue', label='Min Allocation')
#     axes[0].bar(growth_labels, growth_max, bottom=growth_min, color='lightgreen', label='Max Allocation')
#     axes[0].set_title('Growth-Oriented Investments', fontsize=16)
#     axes[0].set_ylabel('Percentage Allocation', fontsize=14)
#     axes[0].set_xlabel('Investment Types', fontsize=14)
#     axes[0].tick_params(axis='x', rotation=45, labelsize=12)
#     axes[0].tick_params(axis='y', labelsize=12)
#     axes[0].legend()

#     # Plot Conservative Investments
#     conservative_data = data['Conservative Investments']
#     conservative_labels = list(conservative_data.keys())
#     conservative_min = [int(conservative_data[label]['min'].strip('%')) for label in conservative_labels]
#     conservative_max = [int(conservative_data[label]['max'].strip('%')) for label in conservative_labels]

#     axes[1].bar(conservative_labels, conservative_min, color='skyblue', label='Min Allocation')
#     axes[1].bar(conservative_labels, conservative_max, bottom=conservative_min, color='lightgreen', label='Max Allocation')
#     axes[1].set_title('Conservative Investments', fontsize=16)
#     axes[1].set_ylabel('Percentage Allocation', fontsize=14)
#     axes[1].set_xlabel('Investment Types', fontsize=14)
#     axes[1].tick_params(axis='x', rotation=45, labelsize=12)
#     axes[1].tick_params(axis='y', labelsize=12)
#     axes[1].legend()

#     # Tight layout for better spacing
#     plt.tight_layout()
#     plt.show()
#     return fig


# def plot_pie_chart(data):
#     fig, ax = plt.subplots(figsize=(10, 7))  # Increased size

#     # Combine all investment data for pie chart
#     all_data = {**data['Growth-Oriented Investments'], **data['Conservative Investments']}
#     labels = list(all_data.keys())
#     sizes = [int(all_data[label]['max'].strip('%')) for label in labels]
#     colors = plt.cm.Paired(range(len(labels)))

#     wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%', startangle=140)
#     ax.set_title('Investment Allocation')

#     # Add legend
#     ax.legend(wedges, labels, title="Investment Types", loc="center left", bbox_to_anchor=(1, 0, 0.5, 1))

#     return fig



# def bar_chart(data):
#     fig, ax = plt.subplots(figsize=(12, 8))  # Increased size

#     # Data for plotting
#     categories = list(data.keys())
#     values_min = [int(data[cat]['min'].strip('%')) for cat in categories]
#     values_max = [int(data[cat]['max'].strip('%')) for cat in categories]

#     x = range(len(categories))

#     ax.bar(x, values_min, width=0.4, label='Min Allocation', color='skyblue', align='center')
#     ax.bar(x, values_max, width=0.4, label='Max Allocation', color='lightgreen', align='edge')

#     ax.set_xticks(x)
#     ax.set_xticklabels(categories, rotation=45, ha='right')
#     ax.set_xlabel('Investment Categories')
#     ax.set_ylabel('Percentage Allocation')
#     ax.set_title('Investment Allocation')
#     ax.legend()

#     plt.tight_layout()
#     return fig


import random
# generate colors for pie chart :

def generate_colors(n):
    """
    Generate 'n' random RGB colors.

    Args:
        n (int): Number of colors to generate.
    
    Returns:
        list: A list of RGB colors in 'rgb(r, g, b)' format.
    """
    colors = []
    for _ in range(n):
        r = random.randint(0, 255)
        g = random.randint(0, 255)
        b = random.randint(0, 255)
        colors.append(f'rgb({r}, {g}, {b})')
    
    return colors


# import plotly.graph_objects as go
import numpy as np



from datetime import date  # Make sure to import the date class


# Function to parse financial data from the text
import re

def parse_financial_data(text_content):
    assets = []
    liabilities = []

    # Define regex patterns to capture text following headings
    asset_pattern = re.compile(r"MY ASSETS:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)
    liability_pattern = re.compile(r"LIABILITIES:\s*(.+?)(?:YOUR CURRENT ANNUAL INCOME|YOUR PROTECTION PLAN|Securities offered)", re.DOTALL)

    # Extract assets
    asset_matches = asset_pattern.findall(text_content)
    if asset_matches:
        asset_text = asset_matches[0]
        # Further processing to extract individual asset values if they are detailed
        asset_lines = asset_text.split('\n')
        for line in asset_lines:
            match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
            if match:
                asset_value = float(match.group().replace(",", ""))
                assets.append(asset_value)

    # Extract liabilities
    liability_matches = liability_pattern.findall(text_content)
    if liability_matches:
        liability_text = liability_matches[0]
        # Further processing to extract individual liability values if they are detailed
        liability_lines = liability_text.split('\n')
        for line in liability_lines:
            match = re.search(r'\b\d{1,3}(?:,\d{3})*(?:\.\d{1,2})?\b', line)
            if match:
                liability_value = float(match.group().replace(",", ""))
                liabilities.append(liability_value)

    print("Assets Found:", assets)
    print("Liabilities Found:", liabilities)

    return assets, liabilities



# Function to extract numerical values from a text input
def extract_numeric(value):
    try:
        return float(re.sub(r'[^\d.]', '', value))  # Remove non-numeric characters and convert to float
    except ValueError:
        return 0


# plots graph from the details of the form :


def is_float(value):
    try:
        float(value)
        return True
    except ValueError:
        return False



def save_data_to_file(form_data):
    file_path = 'client_data.txt'
    with open(file_path, 'a') as file:
        file.write(str(form_data) + "\n")
    # st.success(f"Form data saved to {file_path}")
    print(f"Form data saved to {file_path}")
    
import math
# calculate compunded amount :
def calculate_compounded_amount(principal, rate, time):
    """
    Calculates the compounded amount using the formula:
    A = P * (1 + r/n)^(nt)
    Assuming n (compounding frequency) is 1 for simplicity (annually).
    """
    if principal == 0 or rate == 0 or time == 0:
        return principal
    else:
        # Using annual compounding
        return principal * (1 + rate / 100) ** time
    
def calculate_totals(assets, liabilities):
    total_assets = sum(extract_numeric(v) for v in assets.values())
    print(f"Total Assets : {total_assets}")
    total_liabilities = 0
    total_liabilities = sum(extract_numeric(v) for v in liabilities.values() )

    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Mortgage']),
    #     liabilities['Annual Mortgage Interest Rate'],
    #     liabilities['Mortagage Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Home Loans']),
    #     liabilities['Home Loans Interest Rate'],
    #     liabilities['Home Loans Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Vehicle Loans']),
    #     liabilities['Vehicle Loans Interest Rate'],
    #     liabilities['Vehicle Loans Time Period']
    # )
    # total_liabilities += calculate_compounded_amount(
    #     extract_numeric(liabilities['Education Loans']),
    #     liabilities['Education Loans Interest Rate'],
    #     liabilities['Education Loans Time Period']
    # )
    
    # For credit card debt, only calculate compounded amount if interest rate > 0

    # credit_card_balance = extract_numeric(liabilities['Credit Card'])
    # credit_card_interest = liabilities['Credit Card Debt Interest Rate']
    # if credit_card_interest > 0:
    #     # Assuming the time period for credit card debt is 1 year for compounding
    #     total_liabilities += calculate_compounded_amount(credit_card_balance, credit_card_interest, 1)
    # else:
    #     total_liabilities += credit_card_balance
    
    # Miscellaneous debts are taken directly as is
    total_liabilities += extract_numeric(liabilities['Miscellaneous'])
    rounded_liabilities = round(total_liabilities,2)

    print(f"Total liabilities :{total_liabilities}")
    print(f"Rounded of Total liabilities :{rounded_liabilities}")

    return total_assets, rounded_liabilities #total_liabilities



from docx import Document
# Define a helper function to read and extract text from a DOCX file
def read_docx(file_path):
    document = Document(file_path)
    extracted_text = "\n".join([para.text for para in document.paragraphs])
    return extracted_text



class TrieNode:
    def __init__(self):
        self.children = {}
        self.client_ids = []
        self.end_of_name = False  # Marks the end of a client's name

class Trie:
    def __init__(self):
        self.root = TrieNode()

    def insert(self, name, client_id):
        node = self.root
        for char in name:
            if char not in node.children:
                node.children[char] = TrieNode()
            node = node.children[char]
        node.client_ids.append(client_id)
        node.end_of_name = True

    def search(self, prefix):
        node = self.root
        for char in prefix:
            if char in node.children:
                node = node.children[char]
            else:
                return []  # Prefix not found
        return self._get_all_names_from_node(prefix, node)

    def _get_all_names_from_node(self, prefix, node):
        suggestions = []
        if node.end_of_name:
            suggestions.append((prefix, node.client_ids))
        for char, child_node in node.children.items():
            suggestions.extend(self._get_all_names_from_node(prefix + char, child_node))
        return suggestions



def preload_trie():
    trie = Trie()
    clients = {
        "John Doe": "C001",
        "Jane Smith": "C002",
        "James Brown": "C003",
        "Jill Johnson": "C004",
        "Jake White": "C005"
    }
    for name, client_id in clients.items():
        trie.insert(name.lower(), client_id)  # Insert in lowercase for case-insensitive search
    return trie

# generate suggestions :
async def generate_investment_suggestions_for_investor(investment_personality,clientName,financial_data,financial_file,monthly_investment=10000,investment_period=3): # # GET Method for py , for front end its Post API
    
    # retriever = asyncio.run(load_vector_db("uploaded_file"))

    # retriever =  await load_vector_db("uploaded_file")
    try:
        retriever =  await load_vector_db(financial_file)
    except Exception as e :
        print(f"Error : {e}")
        return jsonify("Error : Failed to load vector database and to generate suggestions : {e}"),400
    
    if not retriever:
        # await load_vector_db("data\Financial_Investment_new.docx")
        await load_vector_db("data\EW2400.docx")
        # await load_vector_db("data\Financial_Investment_1_new.docx") # doesnt works
        # await load_vector_db("data\Financial_Investment_1.docx")
        if not retriever:
            raise Exception("Failed to load vector database.")
    
    print("VectorDB is created successfully")
    # retriever = await load_vector_db("data\Financial_Investment_1.docx") 
    
    try:
        chain = await make_retrieval_chain(retriever,investment_personality,clientName,monthly_investment,investment_period)
    except Exception as e :
        print(f"Error : {e}")
        return jsonify("Error : Failed to create retrieval chain and generate suggestions : {e}"),400
    
    if not chain:
        raise Exception("Failed to create retrieval chain.")
    print("Chain is created to generate suggestions ")
    
    # chain = asyncio.run(make_retrieval_chain(retriever))
    
    print(f"Financial Data : {financial_data}")
    try :
        print(type(financial_data))
        query = f"The Investment Personality of {clientName} is : {investment_personality}" + f"Consider the Monthly Investment as {monthly_investment} and Investment period as {investment_period}" + f"Financial Data of client is : {financial_data[0]}"
        print(query)
    except Exception as e :
        print(f"Error : {e}")
        return "Error : Failed to load financial data"
    
    if chain is not None:
        # summary = context
        # query = summary + "\n" + investment_personality
        
        # query = str(investment_personality)
        response = chain.invoke({"input": query})
        
        # format_response = markdown_to_text(response['answer'])
        # return format_response
        
        # html_output = markdown.markdown(response['answer'])
        # return html_output
        
        # readable_text = markdown_to_readable_text(response['answer'])
        # print(readable_text)
        # return readable_text

        # format_text = convert_to_markdown(response['answer'])
        # return format_text
        
        return response['answer']
    
        

        # handle_graph(response['answer'])

    else:
        logging.INFO("response is not generated by llm model")
        return jsonify("response is not generated by llm model"),500
        # st.error("Failed to create the retrieval chain. Please upload a valid document.")

####################################################################################################################
# app begining :
# CORS(app,resources={r"/api/*":{"origins":"*"}})
# CORS(app)

# Initialize the Trie with preloaded clients
trie = preload_trie()

@app.route('/')
def home():
    return "Wealth Advisor Chatbot API"




import random

# Generate unique client ID

# def generate_unique_id(name):
#     name_parts = name.split(" ")
#     first_initial = name_parts[0][0] if len(name_parts) > 0 else ""
#     last_initial = name_parts[1][0] if len(name_parts) > 1 else ""
#     random_number = random.randint(1000, 9999)
#     unique_id = f"{first_initial}{last_initial}{random_number}"
#     return unique_id

# # Save details in a Word file
import docx
import os

# #Curr version :

# Financial Form
def save_to_word_file(data, file_name):
    doc = docx.Document()
    doc.add_heading('Client Details', 0)

    # Adding client details
    client_details = data.get('clientDetail', {})
    doc.add_paragraph(f"Client Name: {client_details.get('clientName', '')}")
    doc.add_paragraph(f"Client Mobile: {client_details.get('clientMoNo', '')}")
    doc.add_paragraph(f"Client Age: {client_details.get('clientAge', '')}")
    doc.add_paragraph(f"Co-Client Name: {client_details.get('coClientName', '')}")
    doc.add_paragraph(f"Co-Client Mobile: {client_details.get('coMobileNo', '')}")
    doc.add_paragraph(f"Co-Client Age: {client_details.get('coClientAge', '')}")

    # Retirement Plan
    retirement_goal = data.get('retirementGoal', {})
    retirement_plan = retirement_goal.get('retirementPlan', {})
    doc.add_paragraph(f"Retirement Plan Client Age: {retirement_plan.get('retirementAgeClient', '')}")
    doc.add_paragraph(f"Retirement Plan Co-Client Age: {retirement_plan.get('retirementAgeCoClient', '')}")
    
    social_benefit = retirement_goal.get('socialBenefit', {})
    doc.add_paragraph(f"Social Benefit Client: {social_benefit.get('socialBenefitClient', '')}")
    doc.add_paragraph(f"Social Benefit Co-Client: {social_benefit.get('socialBenefitCoClient', '')}")
    
    pension_benefit = retirement_goal.get('pensionBenefit', {})
    doc.add_paragraph(f"Pension Benefit Client: {pension_benefit.get('pensionBenefitClient', '')}")
    doc.add_paragraph(f"Pension Benefit Co-Client: {pension_benefit.get('pensionBenefitCoClient', '')}")
    
    otherIncome = retirement_goal.get('otherIncome', {})
    doc.add_paragraph(f"Other IncomeClient Client: {otherIncome.get('otherIncomeClient', '')}")
    doc.add_paragraph(f"Other IncomeClient Co-Client: {otherIncome.get('otherIncomeCoClient', '')}")
   
    # Estimated Annual Retirement Expense ($ or % of current salary)
    annualRetirement = retirement_goal.get('annualRetirement', {})
    doc.add_paragraph(f"Estimated Annual Retirement Expense ($ or % of current salary) Client: {annualRetirement.get('annualRetireClient', '')}")
    doc.add_paragraph(f"Estimated Annual Retirement Expense ($ or % of current salary) Co-Client: {annualRetirement.get('annualRetireCoClient', '')}")
    

    # Assets and Liabilities
    assets_liabilities = data.get('assetsLiabilities', {})
    
    # Assets
    
    for asset_key, asset_info in assets_liabilities.items():
        current_value_key = [key for key in asset_info.keys() if key.startswith("current")][0]
        annual_value_key = [key for key in asset_info.keys() if key.startswith("annual")][0]
        assets_name_key = "assetsName"
        doc.add_paragraph(f"Assets - {asset_info[assets_name_key]} : Current Value - {asset_info[current_value_key]} , Annual Contributions - {asset_info[annual_value_key]}")
        
    # Liabilities
    myLiabilities = data.get('myLiabilities', {})
    for liability_key, liability_info in myLiabilities.items():
        balance_key = [key for key in liability_info.keys() if key.endswith("Balance")][0]
        interest_key = [key for key in liability_info.keys() if key.endswith("Interest")][0]
        monthly_key = [key for key in liability_info.keys() if key.endswith("Monthly")][0]
        liability_name_key = "liabilityName"
        doc.add_paragraph(f"Liabilities - {liability_info[liability_name_key]} : Balance - {liability_info[balance_key]} , Interest - {liability_info[interest_key]} , Monthly - {liability_info[monthly_key]}")
        
    # my_liabilities = data.get('myLiabilities', {})
    # for liability_type, liability_info in my_liabilities.items():
    #     doc.add_paragraph(f"Liabilities - {liability_info.get('liabilityName', '')}: Balance - {liability_info.get('mortgageBalance', '')} Interest - {liability_info.get('mortgageInterest', '')} Monthly - {liability_info.get('mortgageMonthly', '')}")

    # Protection Plan
    protection_plan = data.get('protectionPlan', {})
    doc.add_paragraph(f"Check Will: {protection_plan.get('checkWill', False)}")
    doc.add_paragraph(f"Check Healthcare: {protection_plan.get('checkHealthCare', False)}")
    doc.add_paragraph(f"Check Attorney: {protection_plan.get('checkAttorney', False)}")
    doc.add_paragraph(f"Check Trust: {protection_plan.get('checkTrust', False)}")

    # Insurance Coverage
    insurance_coverage = data.get('insuranceCoverage', {})
    life_insurance_client = insurance_coverage.get('lifeInsuranceClient', {})
    doc.add_paragraph(f"Life Insurance Client: Benefit - {life_insurance_client.get('benefitLIClient', '')} Monthly Pay - {life_insurance_client.get('monthlyPayLIClient', '')}")
    
    life_insurance_co_client = insurance_coverage.get('lifeInsuranceCoClient', {})
    doc.add_paragraph(f"Life Insurance Co-Client: Benefit - {life_insurance_co_client.get('benefitLICoClient', '')} Monthly Pay - {life_insurance_co_client.get('monthlyPayLICoClient', '')}")
 
    disableIncome = insurance_coverage.get('disableIncomeClient', {})
    disableIncomeClient = insurance_coverage.get('disableIncomeClient',{})
    doc.add_paragraph(f"Disable Income Client - {disableIncomeClient.get('benefitDisableClient', '')}")
    
    disableIncomeCoClient = insurance_coverage.get('disableIncomeCoClient', {})
    doc.add_paragraph(f"Disable Income Co-Client - {disableIncomeCoClient.get('benefitDisableCoClient', '')}")
    
    longTermCoClient = insurance_coverage.get('longTermCoClient')
    doc.add_paragraph(f"Long Term Client: Benefit - {longTermCoClient.get('benefitLongTermClient', '')} Monthly Pay - {longTermCoClient.get('monthlyPayLongTermClient', '')}")
    
    investmentAmount = insurance_coverage.get('investmentAmount')
    doc.add_paragraph(f"Investment Amount Available : {investmentAmount}")
                      
    # Goal Fields
    goal_fields = data.get('goalFields', [])
    for goal in goal_fields:
        doc.add_paragraph(f"Goal: {goal.get('goal', '')} Cost: {goal.get('cost', '')} When: {goal.get('when', '')}")

    # Income Fields
    income_fields = data.get('incomeFields', [])
    for income in income_fields:
        doc.add_paragraph(f"Income Source: {income.get('sourceIncome', '')} Amount: {income.get('amountIncome', '')}")

    # funds_investment = data.get('Funds',[]) commented for later use
    # Save file
    file_name = os.path.join("data", file_name)
    doc.save(f"{file_name}.docx")

# # store client data in aws :
# @app.route('/submit-client-data', methods=['POST'])
# def submit_client_data():
#     try:
#         # Parse JSON payload
#         data = request.get_json()
#         if not data:
#             return jsonify({'message': 'Invalid or missing request payload'}), 400
        
#         # Extract client details
#         client_name = data.get('clientDetail', {}).get('clientName')
#         unique_id = data.get('uniqueId')
        
#         if not client_name or not unique_id:
#             return jsonify({'message': 'Client name and unique ID are required'}), 400

#         print(f"Processing data for client: {client_name}, ID: {unique_id}")
        
#         # Define the S3 key
#         s3_key = f"{client_summary_folder}client-data/{unique_id}.json"
        
#         # Check if the client data already exists in S3
#         try:
#             response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
#             existing_data = json.loads(response['Body'].read().decode('utf-8'))
#             is_update = True
#             print(f"Existing data found for unique ID: {unique_id}")
#         except s3.exceptions.NoSuchKey:
#             existing_data = {}
#             is_update = False
#             print(f"No existing data found for unique ID: {unique_id}. Creating new record.")
        
#         # Merge or replace the existing data (logic can vary based on requirements)
#         if is_update:
#             existing_data.update(data)
#             data_to_save = existing_data
#         else:
#             data_to_save = data
        
#         # Save the updated or new data back to S3
#         try:
#             s3.put_object(
#                 Bucket=S3_BUCKET_NAME,
#                 Key=s3_key,
#                 Body=json.dumps(data_to_save),
#                 ContentType="application/json"
#             )
#             action = "updated" if is_update else "created"
#             print(f"Client data successfully {action} in S3 for unique ID: {unique_id}")
#         except Exception as s3_error:
#             logging.error(f"Error uploading data to S3: {s3_error}")
#             return jsonify({'message': f"Error uploading data to S3: {s3_error}"}), 500
        
#         # Return a success response
#         return jsonify({
#             'message': f'Client data successfully {action}.',
#             'uniqueId': unique_id
#         }), 200

#     except Exception as e:
#         logging.error(f"An error occurred: {e}")
#         return jsonify({'message': f"An error occurred: {e}"}), 500

# # get all client data :
# @app.route('/get-all-client-data', methods=['GET'])
# def get_all_client_data():
#     try:
#         # List objects in the S3 bucket
#         response = s3.list_objects_v2(Bucket=S3_BUCKET_NAME, Prefix="client_summary_folder")
        
#         # Check if there are any objects in the bucket
#         if 'Contents' in response:
#             all_data = []
#             for obj in response['Contents']:
#                 # Get the object content
#                 try:
#                     file_key = obj['Key']
#                     # Retrieve and decode file content
#                     file_response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=file_key)
#                     file_data = file_response['Body'].read().decode('utf-8')
#                      # Parse the file content as JSON
#                     data_json = json.loads(file_data)
#                     all_data.append(data_json)
#                 except Exception as e:
#                     print(f"Error reading file {obj['Key']}: {e}")
#                     continue
            
#             return jsonify({
#                 # 'message': 'All client data retrieved successfully.',
#                 'data': all_data
#             }), 200
        
#         else:
#             return jsonify({'message': 'No client data found in the bucket.'}), 404

#     except Exception as e:
#         return jsonify({'message': f"Error occurred while retrieving data: {e}"}), 500

# # get client data by client id :
# @app.route('/get-client-data-by-id', methods=['GET'])
# def get_client_data():
#     try:
#         # Retrieve client_id from query parameters
#         client_id = request.args.get('client_id')
        
#         # Validate the client_id
#         if not client_id:
#             return jsonify({'message': 'client_id is required as a query parameter'}), 400

#         # Define the S3 key for the object
#         s3_key = f"{client_summary_folder}client-data/{client_id}.json"

#         # Retrieve the object from S3
#         try:
#             response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
#             # Decode and parse the JSON data
#             client_data = json.loads(response['Body'].read().decode('utf-8'))
            
#             return jsonify({
#                 'message': 'Client data retrieved successfully.',
#                 'data': client_data
#             }), 200
#         except s3.exceptions.NoSuchKey:
#             return jsonify({'message': 'Client data not found for the given client_id.'}), 404
#         except Exception as e:
#             return jsonify({'message': f"Error retrieving data: {e}"}), 500

#     except Exception as e:
#         return jsonify({'message': f"An error occurred: {e}"}), 500

#################################################################################################################################
# storing client data using local storage :
# Local storage directories
LOCAL_STORAGE_DIR = "local_storage"
CLIENT_DATA_DIR = os.path.join(LOCAL_STORAGE_DIR, "client_data")
# Ensure directories exist
os.makedirs(CLIENT_DATA_DIR, exist_ok=True)


@app.route('/submit-client-data', methods=['POST'])
def submit_client_data():
    try:
        # Parse JSON payload
        data = request.get_json()
        if not data:
            return jsonify({'message': 'Invalid or missing request payload'}), 400

        # Extract client details
        client_name = data.get('clientDetail', {}).get('clientName')
        unique_id = data.get('uniqueId')

        if not client_name or not unique_id:
            return jsonify({'message': 'Client name and unique ID are required'}), 400

        print(f"Processing data for client: {client_name}, ID: {unique_id}")

        # Define the file path for local storage
        file_path = os.path.join(CLIENT_DATA_DIR, f"client_data/{unique_id}.json")

        # Check if the client data already exists
        if os.path.exists(file_path):
            with open(file_path, 'r') as f:
                existing_data = json.load(f)
            existing_data.update(data)  # Merge the new data
            is_update = True
        else:
            existing_data = data  # New data
            is_update = False

        # Save the data to local storage
        with open(file_path, 'w') as f:
            json.dump(existing_data, f, indent=4)

        action = "updated" if is_update else "created"
        print(f"Client data successfully {action} for unique ID: {unique_id}")

        return jsonify({
            'message': f'Client data successfully {action}.',
            'uniqueId': unique_id
        }), 200

    except Exception as e:
        return jsonify({'message': f"An error occurred: {e}"}), 500


# Define the directory where client data is stored
CLIENT_DATA_DIR = './client_data/'
CLIENT_SUMMARY_DIR = os.path.join(CLIENT_DATA_DIR, "client_data")

# Ensure the directory exists
if not os.path.exists(CLIENT_DATA_DIR):
    os.makedirs(CLIENT_DATA_DIR)

# Get client data by client ID
@app.route('/get-client-data-by-id', methods=['GET'])
def get_client_data():
    try:
        # Retrieve client_id from query parameters
        client_id = request.args.get('client_id')

        # Validate the client_id
        if not client_id:
            return jsonify({'message': 'client_id is required as a query parameter'}), 400

        # Define the file path for the client data
        file_path = os.path.join(CLIENT_DATA_DIR, f"client_data/{client_id}.json")

        # Check if the file exists and retrieve the data
        if not os.path.exists(file_path):
            return jsonify({'message': 'Client data not found for the given client_id.'}), 404

        with open(file_path, 'r') as f:
            client_data = json.load(f)

        return jsonify({
            'message': 'Client data retrieved successfully.',
            'data': client_data
        }), 200

    except Exception as e:
        return jsonify({'message': f"An error occurred: {e}"}), 500


import os
import json
from flask import Flask, jsonify
import requests  # This is for calling the show_order_list API
 
 
@app.route('/get-all-client-data', methods=['GET'])
def get_all_client_data():
    try:
        all_data = []
        client_data_folder = './client_data/client_data'
 
        # Iterate over all JSON files in the client data folder
        for filename in os.listdir(client_data_folder):
            if filename.endswith(".json"):
                file_path = os.path.join(client_data_folder, filename)
               
                # Load client data
                with open(file_path, 'r') as f:
                    client_data = json.load(f)
                    client_id = client_data.get("uniqueId")
 
                    # Default value for isNewClient
                    client_data["isNewClient"] = True
 
                    # Check if the client has orders by calling /show_order_list API
                    if client_id:
                        # Make a request to /show_order_list API to check if orders exist
                        order_url = f'http://localhost:5000/show_order_list'  # Adjust to your actual endpoint
                        response = requests.post(order_url, json={'client_id': client_id})
 
                        if response.status_code == 200:
                            orders = response.json().get("transaction_data", [])
                            # If orders exist, set isNewClient to False
                            if len(orders) > 0:
                                client_data["isNewClient"] = False
                        else:
                            # If no orders, set isNewClient to True
                            client_data["isNewClient"] = True
 
                    # Append the client data to the result list
                    all_data.append(client_data)
 
        # Handle case when no client data is found
        if not all_data:
            return jsonify({'message': 'No client data found in local storage.'}), 404
 
        # Return all client data
        return jsonify({
            'message': 'All client data retrieved successfully.',
            'data': all_data
        }), 200
 
    except Exception as e:
        return jsonify({'message': f"An error occurred while retrieving data: {e}"}), 500
    
# investment assessment using Local Storage :

import os
import json

# LOCAL_STORAGE_DIR = "local_storage"
CLIENT_DATA_DIR = './client_data/'
CLIENT_SUMMARY_DIR = os.path.join(CLIENT_DATA_DIR, "client_data")
PERSONALITY_ASSESSMENT_DIR = "client_data/personality_assessments"

# Ensure directories exist
os.makedirs(CLIENT_SUMMARY_DIR, exist_ok=True)
os.makedirs(PERSONALITY_ASSESSMENT_DIR, exist_ok=True)



# Personality Assessment using Local Storage :

@app.route('/get-personality-assessment', methods=['POST'])
def get_client_data_by_id():
    try:
        # Parse incoming request data
        payload = request.json

        # Validate the payload
        client_id = payload.get('client_id')
        if not client_id:
            return jsonify({'message': 'client_id is required in the payload.'}), 400

        # Locate the client's assessment data
        file_path = os.path.join(CLIENT_DATA_DIR, f"personality_assessments/{client_id}.json")
        if not os.path.exists(file_path):
            return jsonify({'message': 'No data found for the provided client_id.'}), 404

        with open(file_path, 'r') as f:
            file_content = json.load(f)

        return jsonify({
            'message': 'Data fetched successfully.',
            'data': file_content
        }), 200

    except Exception as e:
        return jsonify({'message': f'Internal Server Error: {str(e)}'}), 500
    


# api for generating suggestions with client id using Local Storage :  
@app.route('/investor-personality-assessment', methods=['POST'])
def investor_personality_assessment():
    try:
        # Parse incoming data
        data = request.json
        client_id = data.get('client_id')
        assessment_data = data.get('assessment_data')
 
        if not client_id or not assessment_data:
            return jsonify({'message': 'Client ID and assessment data are required.'}), 400
 
        logging.info(f"Received assessment data for client ID: {client_id}")
 
        # Determine investment personality
        personality = asyncio.run(determine_investment_personality(assessment_data))
        logging.info(f"Determined personality for client ID {client_id}: {personality}")
 
        # Save assessment data and personality in a dedicated file
        personality_file_path = os.path.join(CLIENT_DATA_DIR, f"personality_assessments/{client_id}.json")
        client_data_dir = os.path.join(CLIENT_DATA_DIR, "client_data")
        client_file_path = os.path.join(client_data_dir, f"{client_id}.json")
 
        # Update or create personality-specific data
        personality_data = {}
        if os.path.exists(personality_file_path):
            with open(personality_file_path, 'r') as f:
                personality_data = json.load(f)
 
        personality_data.update({
            "client_id": client_id,
            "assessment_data": assessment_data,
            "investment_personality": personality
        })
 
        with open(personality_file_path, 'w') as f:
            json.dump(personality_data, f, indent=4)
 
        # Update the main client data file
        if os.path.exists(client_file_path):
            with open(client_file_path, 'r') as f:
                client_data = json.load(f)
            # Update investment personality in the existing client file
            client_data['investment_personality'] = personality
        else:
            # If client file does not exist, create it
            client_data = {
                "client_id": client_id,
                "investment_personality": personality,
            }
 
        with open(client_file_path, 'w') as f:
            json.dump(client_data, f, indent=4)
 
        logging.info(f"Updated client data file for client ID {client_id}")
 
        return jsonify({
            'client_id': client_id,
            'investment_personality': personality
        }), 200
 
    except Exception as e:
        logging.error(f"Error processing investor assessment: {e}")
        return jsonify({'message': 'Internal Server Error'}), 500
 
 
@app.route('/personality-assessment', methods=['POST'])
def personality_selected():
    try:
        # Parse incoming data
        data = request.json
        if not data:
            return jsonify({'message': 'Invalid or missing request payload'}), 400

        investment_personality = data.get('investmentPersonality')
        client_name = data.get('clientName')
        client_id = data.get('clientId')

        print(f"Client Name: {client_name}, Investment Personality: {investment_personality}")

        # Validate required data
        if not client_id or not client_name or not investment_personality:
            return jsonify({'message': 'Missing client_id, clientName, or investmentPersonality.'}), 400

        # Load client data from local storage
        file_path = os.path.join(CLIENT_DATA_DIR, f"{client_id}.json")
        if not os.path.exists(file_path):
            return jsonify({'message': 'Client data not found for the given client_id.'}), 404

        with open(file_path, 'r') as f:
            client_data = json.load(f)

        print(f"Loaded Client Data: {client_data}")

        # Generate suggestions
        try:
            result, pie_chart_data, bar_chart_data, combined_chart_data = asyncio.run(
                make_suggestions_using_clientid(
                    investment_personality,
                    client_name,
                    client_data
                )
            )

            html_suggestions = markdown.markdown(result)
            format_suggestions = markdown_to_text(html_suggestions)

            return jsonify({
                "status": 200,
                "message": "Success",
                "investmentSuggestions": format_suggestions,
                "pieChartData": pie_chart_data,
                "barChartData": bar_chart_data,
                "compoundedChartData": combined_chart_data
            }), 200

        except Exception as e:
            logging.error(f"Error generating suggestions: {e}")
            return jsonify({'message': f"Error generating suggestions: {e}"}), 500

    except Exception as e:
        logging.error(f"Unhandled exception: {e}")
        return jsonify({'message': 'Internal Server Error'}), 500


#################################################END OF Dashboard Local Storage Method #################################################

# investor personality assessment :
# @app.route('/investor-personality-assessment', methods=['POST'])
# async def investor_personality_assessment():
#     try:
#         # Parse incoming request data
#         data = request.json
#         logging.debug(f"Received request data: {data}")
        
#         if not data:
#             logging.error("No data received in the request.")
#             return jsonify({'message': 'Invalid request: No data received.'}), 400

#         client_id = data.get('client_id')
#         assessment_data = data.get('assessment_data')

#         if not client_id or not assessment_data:
#             logging.error("Missing client_id or assessment_data.")
#             return jsonify({'message': 'Client ID and assessment data are required.'}), 400

#         # Determine the investment personality
#         personality = await determine_investment_personality(assessment_data)
#         logging.info(f"Determined personality for client ID {client_id}: {personality}")
        
#           # Define the S3 key for client data
#         s3_key = f"{client_summary_folder}client-data/{client_id}.json"
#         existing_data = None

#         # Check for existing client data in S3 (to store investment_personality in client detail)
#         try:
#             response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
#             existing_data = json.loads(response['Body'].read().decode('utf-8'))
#             logging.info(f"Existing data found for client ID {client_id}: {existing_data}")
#         except s3.exceptions.NoSuchKey:
#             logging.error(f"No existing client data found for client ID {client_id}.")
#             return jsonify({'message': f"No existing client data found for client ID {client_id}."}), 404

#         # Update the existing data with the new investment personality
#         if existing_data:
#             existing_data['investment_personality'] = personality
#             logging.info(f"Updated investment personality for client ID {client_id}: {personality}")
        
#         # Save the updated data back to S3
#         try:
#             s3.put_object(
#                 Bucket=S3_BUCKET_NAME,
#                 Key=s3_key,
#                 Body=json.dumps(existing_data),
#                 ContentType='application/json'
#             )
#             logging.info(f"Client data successfully updated in S3 for client ID: {client_id}")
#         except Exception as e:
#             logging.error(f"Error occurred while saving updated data to S3: {e}")
#             return jsonify({'message': f'Error occurred while saving updated data to S3: {e}'}), 500


#         # Check if the file exists in S3
#         file_key = f"{personality_assessment_folder}{client_id}.json"
#         existing_file_data = None

#         try:
#             file_response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=file_key)
#             file_data = file_response['Body'].read().decode('utf-8')
#             existing_file_data = json.loads(file_data)
#             logging.info(f"Existing file data for client ID {client_id}: {existing_file_data}")
#         except s3.exceptions.NoSuchKey:
#             logging.info(f"No existing file found for client ID {client_id}. Creating a new file.")

#         # Update or create data
#         updated_data = {
#             'client_id': client_id,
#             'assessment_data': assessment_data,
#             'investment_personality': personality
#         }

#         if existing_file_data:
#             # Update the existing file with new data
#             existing_file_data.update(updated_data)
#             updated_data = existing_file_data
#             logging.info(f"Updated data for client ID {client_id}: {updated_data}")

#         # Save the data back to S3
#         try:
#             s3.put_object(
#                 Bucket=S3_BUCKET_NAME,
#                 Key=file_key,
#                 Body=json.dumps(updated_data),
#                 ContentType='application/json'
#             )
#             logging.info(f"Data successfully saved to S3 for clientId: {client_id}")
#         except Exception as e:
#             logging.error(f"Error occurred while saving to S3: {e}")
#             return jsonify({'message': f'Error occurred while saving to S3: {e}'}), 500

#         # Return the result
#         return jsonify(updated_data), 200
#         # return jsonify({
#         #     'message': 'Data saved successfully',
#         #     'client_id': client_id,
#         #     'data': assessment_data,
#         #     'investment_personality': personality
#         # }), 200

#     except Exception as e:
#         logging.error(f"Unhandled exception: {e}")
#         return jsonify({'message': 'Internal Server Error'}), 500
 
# @app.route('/get-personality-assessment', methods=['POST'])
# def get_client_data_by_id():
#     try:
#         # Parse incoming request data
#         payload = request.json
#         logging.info(f"Received request payload: {payload}")

#         # Validate the payload
#         if not payload or 'client_id' not in payload:
#             logging.error("Invalid request: Missing client_id in payload.")
#             return jsonify({'message': 'client_id is required in the payload.'}), 400

#         client_id = payload.get('client_id')

#         # Ensure client_id is a valid non-empty string
#         if not client_id or not isinstance(client_id, str):
#             logging.error("Invalid client_id: Must be a non-empty string.")
#             return jsonify({'message': 'client_id must be a non-empty string.'}), 400

#         # Define folder path for S3
#         folder_path = f"{personality_assessment_folder}"
#         logging.info(f"Looking for files in folder: {folder_path}")

#         # List objects in the folder
#         response = s3.list_objects_v2(Bucket=S3_BUCKET_NAME, Prefix=folder_path)
#         logging.debug(f"S3 list_objects_v2 response: {response}")

#         # Check if the folder contains any objects
#         if 'Contents' not in response:
#             logging.warning(f"No files found in folder: {folder_path}")
#             return jsonify({'message': 'No data found in the specified folder.'}), 404

#         # Iterate through the files to find the matching client_id
#         for obj in response['Contents']:
#             file_key = obj['Key']

#             # Skip the folder itself and non-JSON files
#             if file_key == folder_path or not file_key.endswith('.json'):
#                 continue

#             # Fetch file content if the file matches the client_id
#             if f"{client_id}.json" in file_key:
#                 try:
#                     file_response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=file_key)
#                     file_content = json.loads(file_response['Body'].read().decode('utf-8'))
#                     logging.info(f"Found and retrieved data for client_id {client_id}.")
                    
#                     return jsonify({
#                         'message': 'Data fetched successfully.',
#                         'data': file_content  # Ensure the actual client data is nested in 'data'
#                     }), 200
#                 except Exception as fetch_error:
#                     logging.error(f"Error retrieving file {file_key}: {fetch_error}")
#                     return jsonify({'message': 'Error retrieving client data from S3.'}), 500

#         # If no matching file is found
#         logging.warning(f"No data found for client_id {client_id}.")
#         return jsonify({'message': 'No data found for the provided client_id.'}), 404

#     except Exception as e:
#         logging.error(f"Unhandled exception: {e}")
#         return jsonify({'message': 'Internal Server Error'}), 500
    


# @app.route('/investor-personality-assessment', methods=['POST'])
# async def investor_personality_assessment():
#     try:
#         # Parse incoming request data
#         data = request.json
#         if not data:
#             return jsonify({'message': 'Invalid request: No data received.'}), 400

#         client_id = data.get('client_id')
#         assessment_data = data.get('assessment_data')

#         if not client_id or not assessment_data:
#             return jsonify({'message': 'Client ID and assessment data are required.'}), 400

#         # Determine the investment personality
#         personality = await determine_investment_personality(assessment_data)

#         # Update or create the client data
#         client_file_path = os.path.join(CLIENT_SUMMARY_DIR, f"{client_id}.json")
#         if os.path.exists(client_file_path):
#             with open(client_file_path, 'r') as f:
#                 client_data = json.load(f)
#         else:
#             client_data = {"client_id": client_id}

#         client_data['investment_personality'] = personality

#         with open(client_file_path, 'w') as f:
#             json.dump(client_data, f, indent=4)

#         # Save or update the assessment data
#         assessment_file_path = os.path.join(PERSONALITY_ASSESSMENT_DIR, f"{client_id}.json")
#         updated_data = {
#             'client_id': client_id,
#             'assessment_data': assessment_data,
#             'investment_personality': personality
#         }

#         with open(assessment_file_path, 'w') as f:
#             json.dump(updated_data, f, indent=4)

#         return jsonify(updated_data), 200

#     except Exception as e:
#         return jsonify({'message': f'Internal Server Error: {str(e)}'}), 500



##################################################################################################################################

import logging
# global investmentPersonality  # Global Variable
# investmentPersonality = ""

#prev version
# def generate_chart_data(data):
#     # Pie Chart
#     labels = list(data['Growth-Oriented Investments'].keys()) + list(data['Conservative Investments'].keys())
#     max_allocations = [
#         int(data['Growth-Oriented Investments'][label]['max']) for label in data['Growth-Oriented Investments']
#     ] + [
#         int(data['Conservative Investments'][label]['max']) for label in data['Conservative Investments']
#     ]
#     # Generate colors based on the number of labels
#     all_labels = list({**data['Growth-Oriented Investments'], **data['Conservative Investments']}.keys())
#     num_labels = len(all_labels)
#     dynamic_colors = generate_colors(num_labels)
#     pie_chart_data = {
#         'labels': labels,
#         'datasets': [{
#             'label': 'Investment Allocation',
#             'data': max_allocations,
#             'backgroundColor': dynamic_colors,  # Example colors
#             'hoverOffset': 4
#         }]
#     }
    
#     # pie_chart_data = {
#     #     'labels': labels,
#     #     'datasets': [{
#     #         'label': 'Investment Allocation',
#     #         'data': max_allocations,
#     #         'backgroundColor': ['#FF6384', '#36A2EB', '#FFCE56', '#4BC0C0'],  # Example colors
#     #         'hoverOffset': 4
#     #     }]
#     # }

#     # Bar Chart
#     min_allocations = [
#         int(data['Growth-Oriented Investments'][label]['min']) for label in data['Growth-Oriented Investments']
#     ] + [
#         int(data['Conservative Investments'][label]['min']) for label in data['Conservative Investments']
#     ]
#     bar_chart_data = {
#         'labels': labels,
#         'datasets': [
#             {
#                 'label': 'Allocation for Min returns',
#                 'data': min_allocations,
#                 'backgroundColor': 'skyblue'
#             },
#             {
#                 'label': 'Allocation for Max returns',
#                 'data': max_allocations,
#                 'backgroundColor': 'lightgreen'
#             }
#         ]
#     }
#     print(f"Pie Chart Data : {pie_chart_data}")
#     print(f"Bar Chart Data : {bar_chart_data}")
    
#     return pie_chart_data, bar_chart_data

# generate pie chart data and bar chart data :
def generate_chart_data(data):
    # Pie Chart Data
    labels = list(data['Growth-Oriented Investments'].keys()) + list(data['Conservative Investments'].keys())
    max_allocations = [
        int(data['Growth-Oriented Investments'][label]['max']) for label in data['Growth-Oriented Investments']
    ] + [
        int(data['Conservative Investments'][label]['max']) for label in data['Conservative Investments']
    ]
    num_labels = len(labels)
    dynamic_colors = generate_colors(num_labels)
    pie_chart_data = {
        'labels': labels,
        'datasets': [{
            'label': 'Investment Allocation',
            'data': max_allocations,
            'backgroundColor': dynamic_colors, #['#FF6384', '#36A2EB', '#FFCE56', '#4BC0C0'],  # Example colors
            'hoverOffset': 4
        }]
    }

    # Bar Chart Data
    min_allocations = [
        int(data['Growth-Oriented Investments'][label]['min']) for label in data['Growth-Oriented Investments']
    ] + [
        int(data['Conservative Investments'][label]['min']) for label in data['Conservative Investments']
    ]
    bar_chart_data = {
        'labels': labels,
        'datasets': [
            {
                'label': 'Allocation for Min returns',
                'data': min_allocations,
                'backgroundColor': 'skyblue'
            },
            {
                'label': 'Allocation for Max returns',
                'data': max_allocations,
                'backgroundColor': 'lightgreen'
            }
        ]
    }

    return pie_chart_data, bar_chart_data


#new retrieval_chain code :
# async def generate_prompt_template(retriever,investmentPersonality,clientName,client_data):
#     try:
#         # global investment_personality #,summary
        
#         print(f"{investmentPersonality}\n {clientName}\n {client_data}")
        
        
#         llm = ChatGoogleGenerativeAI(
#             #model="gemini-pro",
#             model = "gemini-1.5-flash",
#             temperature = 0.45,
#             # temperature=0.7,
#             top_p=0.85,
#             google_api_key=GOOGLE_API_KEY
#         )
#         # New Template 
#         investmentPersonality = str(investmentPersonality)
#         print(investmentPersonality)
#         # clientName = str(clientName)
#         print(clientName)
#         context = str(clientName)
        
        
#         # New Prompt Template :
        
#         prompt_template = """ 
#                                 You are a Financial Advisor tasked with creating responsible investment suggestions for a client based on their investment personality : """ + investmentPersonality +   "\n" + """ so that the client can reach their Financial Goals, based on their Financial Conditions.
#                                 Use the following instructions to ensure consistent output:
#                                 ---

#                                 ### Required Output Format:
                                
#                                 #### Client Financial Details:
#                                 - **Client Name**: """ + clientName + f"""
#                                 - **Assets**:
#                                 - List all asset types, their current values, and annual contributions in a tabular format (columns: "Asset Type", "Current Value", "Annual Contribution").
#                                 - **Liabilities**:
#                                 - List all liability types, their balances, interest rates, and monthly payments in a tabular format (columns: "Liability Type", "Balance", "Interest Rate", "Monthly Payment").
#                                 - **Other Details**:
#                                 - Retirement plan details, income sources, and goals should be listed in a clear and concise format.
#                                 - Client's Financial Condition : Analyze the Details and mention the Client's Financial Condition as : Stable/ Currently Stable / Unstable.
#                                 - **Investment Period** `Z years`
                                
#                                 #### Investment Allocation:
#                                 Split investments into **Growth-Oriented Investments** and **Conservative Investments**. Ensure each category includes:
#                                 - **Investment Type**: Specify the investment type (e.g., "Index Funds", "US Treasury Bonds").
#                                 - **Allocation Range**: Specify minimum and maximum allocation percentages (e.g., `10% - 20%`).
#                                 - **Target**: Describe the purpose of the investment.
#                                 - **How to Invest**: Provide instructions on how to invest in this asset.
#                                 - **Where to Invest**: Specify platforms or tools for making the investment.

#                                 **Example**:
#                                 **Growth-Oriented Investments (Minimum X% - Maximum Y%) **:
#                                 - **Stocks**: `20% - 30%`
#                                 - **ETFs**: `10% - 15%`
#                                 - **Mutual Funds**: `10% - 20%`
#                                 - **Cryptocurrency**: ` 5% - 10%`
#                                 - **Real Estates or REITS**: `10% - 20%`
#                                 - *Target*: Long-term growth potential aligned with the overall market performance tailored to fullfil Client's Financial Goals and manage his Financial Condition.
#                                 - *How to Invest*: Provide information on how to invest in which market 
#                                 - *Where to Invest*: Provide Information to buy which assets and how much to invest in terms of amount and percentage(%).Mention 5-6 assets.
                                
#                                 **Conservative Investments (Minimum X% - Maximum Y%) **:
#                                 - **High-Yield Savings Account**: `30% - 40%`
#                                 - **Bonds**: `10% - 20%`
#                                 - **Commodities**: `5% - 10%`
#                                 - **Cash**: `5% - 10%`
#                                 - *Target*: Maintain liquidity for emergencies.
#                                 - *How to Invest*: Provide information on how to invest.
#                                 - *Where to Invest*: Mention where to invest and how much to allocate in terms of money and percentage(%). Mention 5-6 assets.

#                                 #### Returns Overview:
#                                 - **Minimum Expected Annual Return**: `X% - Y%`
#                                 - **Maximum Expected Annual Return**: `X% - Y%`
#                                 - **Minimum Expected Growth in Dollars**: `$X - $Y` (based on the time horizon)
#                                 - **Maximum Expected Growth in Dollars**: `$X - $Y` (based on the time horizon)
#                                 - **Time Horizon**: `Z years`

#                                 ---

#                                 ### Example Output:
                                
#                                 #### Client Financial Details:
#                                 | Asset Type          | Current Value ($) | Annual Contribution ($) |
#                                 |----------------------|-------------------|--------------------------|
#                                 | 401(k), 403(b), 457  | 300               | 15                       |
#                                 | Traditional IRA      | 200               | 15                       |
#                                 | Roth IRA             | 500               | 28                       |
#                                 | Cash/Bank Accounts   | 500,000           | 30,000                   |
#                                 | Real Estate          | 1,000,000         | -                        |
#                                 | Total Assets Value   | 1,501,000         | -                        |

#                                 | Liability Type      | Balance ($) | Interest Rate (%) | Monthly Payment ($) |
#                                 |---------------------|-------------|--------------------|----------------------|
#                                 | Mortgage            | 1,000       | 10                | 100                  |
#                                 | Credit Card         | 400         | 8                 | 400                  |
#                                 | Other Loans         | 500         | 6                 | 100                  |
#                                 | Total Liabilities   | 1,900       | -                 | -                    |
                                
#                                 | Investrment Period | 3 years |
                                
#                                 **Growth-Oriented Investments (Minimum 40% - Maximum 80%)**:
#                                 - **Stocks**: `20% - 30%`
#                                 - **ETFs**: `5% - 10%`
#                                 - **Mutual Funds**: `5% - 20%`
#                                 - **Cryptocurrency**: ` 5% - 10%`
#                                 - **Real Estates or REITS**: `5% - 10%`
#                                 - *Target*: Long-term growth potential aligned with the market.
#                                 - *How to Invest*: Purchase low-cost index funds.
#                                 - *Where to Invest*: Stocks such as NVIDIA,AAPL, Vanguard, LiteCoin.

#                                 **Conservative Investments (Minimum 40% - Maximum 70%)**:
#                                 - **High-Yield Savings Account**: `20% - 30%`
#                                 - **Bonds**: `10% - 20%`
#                                 - **Commodities**: `5% - 10%`
#                                 - **Cash**: `5% - 10%`
#                                 - *Target*: Maintain liquidity for emergencies.
#                                 - *How to Invest*: Deposit funds into an FDIC-insured account.
#                                 - *Where to Invest*: Ally Bank, Capital One 360.

#                                 #### Returns Overview:
#                                 - **Minimum Expected Annual Return**: `4% - 6%`
#                                 - **Maximum Expected Annual Return**: `8% - 15%`
#                                 - **Minimum Expected Growth in Dollars**: `$4,000 - $6,000`
#                                 - **Maximum Expected Growth in Dollars**: `$8,000 - $15,000`
#                                 - **Time Horizon**: `3 years`

#                                 ---

#                                 Ensure the output strictly follows this structure.


#                             ### Rationale for Investment Suggestions:
#                             Provide a detailed explanation of why these suggestions align with the client’s financial personality and goals.

#                             ---
#                             <context>
#                             {context}
#                             </context>
#                             Question: {input}

#         """

#         print(f"Investment Personality :{investmentPersonality}")
        
                

#         llm_prompt = ChatPromptTemplate.from_template(prompt_template)

#         document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
#         combine_docs_chain = None  

#         if retriever is not None :  
#             retriever_chain = create_retrieval_chain(retriever,document_chain) 
#             # print(retriever_chain)
#             return retriever_chain
#         else:
#             print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
#             return None

#     except Exception as e:
#         print(f"Error in creating chain: {e}")
#         return None

########################################################################################################



# Create Vector DB for JSON Data from cloud :
from langchain_community.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document

async def load_vector_db_from_json(json_data):
    try:
        print("Loading vector database from JSON data...")
        
        # Step 1: Convert JSON to a list of Documents
        documents = []
        for key, value in json_data.items():
            if isinstance(value, dict):
                nested_text = "\n".join([f"{nested_key}: {nested_value}" for nested_key, nested_value in value.items()])
                documents.append(Document(page_content=f"{key}:\n{nested_text}"))
            elif isinstance(value, list):
                list_text = "\n".join([str(item) for item in value])
                documents.append(Document(page_content=f"{key}:\n{list_text}"))
            else:
                documents.append(Document(page_content=f"{key}: {value}"))

        print(f"Prepared {len(documents)} documents for FAISS.")

        # Step 2: Split documents into chunks
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        text_chunks = text_splitter.split_documents(documents)

        # Step 3: Embed and load into FAISS
        embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001", google_api_key=GOOGLE_API_KEY)
        vector_store = FAISS.from_documents(text_chunks, embeddings)

        print("Vector database loaded successfully.")
        return vector_store.as_retriever(search_kwargs={"k": 3})  # Top-3 results
    except Exception as e:
        print(f"Error loading vector database: {e}")
        return None




from langchain.prompts.chat import ChatPromptTemplate
from langchain.chains import StuffDocumentsChain, create_retrieval_chain
from langchain.schema.runnable import RunnableConfig, RunnableSequence
# from langchain.chains import LLMChain
# from langchain.schema.runnable import RunnableMap
# from langchain.schema.runnable import RunnableSequence

async def generate_prompt_with_retriever(retriever, investmentPersonality, clientName):
    try:
        llm = ChatGoogleGenerativeAI(
            model="gemini-1.5-flash",
            temperature=0.45,
            top_p=0.85,
            google_api_key=GOOGLE_API_KEY
        )
        clientName = clientName
        # Define the prompt
        prompt_template = """ 
                                You are a Financial Advisor tasked with creating responsible investment suggestions for a client based on their investment personality : """ + investmentPersonality +   "\n" + """ so that the client can reach their Financial Goals, based on their Financial Conditions.
                                Use the following instructions to ensure consistent output:
                                ---

                                ### Required Output Format:
                                
                                #### Client Financial Details:
                                - **Client Name**: """ + clientName + """
                                - **Assets**:
                                - List all asset types, their current values, and annual contributions in a tabular format (columns: "Asset Type", "Current Value", "Annual Contribution").
                                - **Liabilities**:
                                - List all liability types, their balances, interest rates, and monthly payments in a tabular format (columns: "Liability Type", "Balance", "Interest Rate", "Monthly Payment").
                                - **Other Details**:
                                - Retirement plan details, income sources, and goals should be listed in a clear and concise format.
                                - Client's Financial Condition : Analyze the Details and mention the Client's Financial Condition as : Stable/ Currently Stable / Unstable.
                                - **Investment Period** `Z years`
                                
                                #### Investment Allocation:
                                Split investments into **Growth-Oriented Investments** and **Conservative Investments**. Ensure each category includes:
                                - **Investment Type**: Specify the investment type (e.g., "Index Funds", "US Treasury Bonds").
                                - **Allocation Range**: Specify minimum and maximum allocation percentages (e.g., `10% - 20%`).
                                - **Target**: Describe the purpose of the investment.
                                - **How to Invest**: Provide instructions on how to invest in this asset.
                                - **Where to Invest**: Specify platforms or tools for making the investment.

                                **Example**:
                                **Growth-Oriented Investments (Minimum X% - Maximum Y%) **:
                                - **Stocks**: `20% - 30%`
                                - **ETFs**: `10% - 15%`
                                - **Mutual Funds**: `10% - 20%`
                                - **Cryptocurrency**: ` 5% - 10%`
                                - **Real Estates or REITS**: `10% - 20%`
                                - *Target*: Long-term growth potential aligned with the overall market performance tailored to fullfil Client's Financial Goals and manage his Financial Condition.
                                - *How to Invest*: Provide information on how to invest in which market 
                                - *Where to Invest*: Provide Information to buy which assets and how much to invest in terms of amount and percentage(%).Mention 5-6 assets.
                                
                                **Conservative Investments (Minimum X% - Maximum Y%) **:
                                - **High-Yield Savings Account**: `30% - 40%`
                                - **Bonds**: `10% - 20%`
                                - **Commodities**: `5% - 10%`
                                - **Cash**: `5% - 10%`
                                - *Target*: Maintain liquidity for emergencies.
                                - *How to Invest*: Provide information on how to invest.
                                - *Where to Invest*: Mention where to invest and how much to allocate in terms of money and percentage(%). Mention 5-6 assets.

                                #### Returns Overview:
                                - **Minimum Expected Annual Return**: `X% - Y%`
                                - **Maximum Expected Annual Return**: `X% - Y%`
                                - **Minimum Expected Growth in Dollars**: `$X - $Y` (based on the time horizon)
                                - **Maximum Expected Growth in Dollars**: `$X - $Y` (based on the time horizon)
                                - **Time Horizon**: `Z years`

                                ---

                                ### Example Output:
                                
                                #### Client Financial Details:
                                | Asset Type          | Current Value ($) | Annual Contribution ($) |
                                |----------------------|-------------------|--------------------------|
                                | 401(k), 403(b), 457  | 300               | 15                       |
                                | Traditional IRA      | 200               | 15                       |
                                | Roth IRA             | 500               | 28                       |
                                | Cash/Bank Accounts   | 500,000           | 30,000                   |
                                | Real Estate          | 1,000,000         | -                        |
                                | Total Assets Value   | 1,501,000         | -                        |

                                | Liability Type      | Balance ($) | Interest Rate (%) | Monthly Payment ($) |
                                |---------------------|-------------|--------------------|----------------------|
                                | Mortgage            | 1,000       | 10                | 100                  |
                                | Credit Card         | 400         | 8                 | 400                  |
                                | Other Loans         | 500         | 6                 | 100                  |
                                | Total Liabilities   | 1,900       | -                 | -                    |
                                
                                | Investrment Period | 3 years |
                                
                                **Growth-Oriented Investments (Minimum 40% - Maximum 80%)**:
                                - **Stocks**: `20% - 30%`
                                - **ETFs**: `5% - 10%`
                                - **Mutual Funds**: `5% - 20%`
                                - **Cryptocurrency**: ` 5% - 10%`
                                - **Real Estates or REITS**: `5% - 10%`
                                - *Target*: Long-term growth potential aligned with the market.
                                - *How to Invest*: Purchase low-cost index funds.
                                - *Where to Invest*: Stocks such as NVIDIA,AAPL, Vanguard, LiteCoin.

                                **Conservative Investments (Minimum 40% - Maximum 70%)**:
                                - **High-Yield Savings Account**: `20% - 30%`
                                - **Bonds**: `10% - 20%`
                                - **Commodities**: `5% - 10%`
                                - **Cash**: `5% - 10%`
                                - *Target*: Maintain liquidity for emergencies.
                                - *How to Invest*: Deposit funds into an FDIC-insured account.
                                - *Where to Invest*: Ally Bank, Capital One 360.

                                #### Returns Overview:
                                - **Minimum Expected Annual Return**: `4% - 6%`
                                - **Maximum Expected Annual Return**: `8% - 15%`
                                - **Minimum Expected Growth in Dollars**: `$4,000 - $6,000`
                                - **Maximum Expected Growth in Dollars**: `$8,000 - $15,000`
                                - **Time Horizon**: `3 years`

                                ---

                                Ensure the output strictly follows this structure.


                            ### Rationale for Investment Suggestions:
                            Provide a detailed explanation of why these suggestions align with the client’s financial personality and goals.

                            ---
                            <context>
                            {context}
                            </context>
                            Question: {input}

        """
                
        
        print("Prompt Created Successfully")
                

        llm_prompt = ChatPromptTemplate.from_template(prompt_template)

        document_chain = create_stuff_documents_chain(llm, llm_prompt)
        
        combine_docs_chain = None  

        if retriever is not None :  
            retriever_chain = create_retrieval_chain(retriever,document_chain) 
            print("\nRetrieval chain created successfully\n")
            print(retriever_chain)
            return retriever_chain
        else:
            print("Failed to create retrieval chain: Missing retriever or combine_docs_chain")
            return None

    except Exception as e:
        print(f"Error in generating prompt or retrieval chain: {e}")
        return None

# # using aws :
# async def make_suggestions_using_clientid(investmentPersonality, clientName, client_data):
#     try:
#         print(f"Processing client data for {clientName}.")
        
#         # Load vector database
#         retriever = await load_vector_db_from_json(client_data)
#         if not retriever:
#             raise Exception("Failed to load vector database.")

#         print(f"Created Retriever : {retriever}")
#         # Generate retriever-based prompt
#         retrieval_chain = await generate_prompt_with_retriever(retriever, investmentPersonality, clientName)
#         if not retrieval_chain:
#             raise Exception("Failed to create retrieval chain.")

#         # Use the chain to generate a response
#         query = f"""Generate financial suggestions for the client {clientName} based on their investment personality: {investmentPersonality} 
#                 tailored to their Financial Goals and Considering their Financial Situations. Suggest 6-7 assets per category with 6-7 examples per asset."""
        
        
#         # response = retrieval_chain.invoke(query)
#         response = retrieval_chain.invoke({"input": query})
#         answer = response['answer']
#         print("Suggestions generated successfully.")
        
#         # Extract Data from Response

#         data_extracted = extract_numerical_data(answer)
        
#         min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
#                         [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
#         max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
#                         [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

#         # Normalize allocations
#         min_allocations = normalize_allocations(min_allocations)
#         max_allocations = normalize_allocations(max_allocations)

#         bar_chart_data,pie_chart_data = generate_chart_data(data_extracted)
    
#         print(f"Bar chart data: {bar_chart_data}")
#         print(f"Pie chart data: {pie_chart_data}")

        
#         # print(f"Pie Chart Data is : {pie_chart_data}")
#         # Prepare the data for the line chart with inflation adjustment
#         initial_investment = 10000
#         combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
#         print(f"\nThe combined chart data is: {combined_chart_data}")
        
#         print(f"Suggestions : {answer}")
        
#         return answer, pie_chart_data, bar_chart_data, combined_chart_data
            
#     except Exception as e:
#         print(f"Error generating suggestions: {e}")
#         return jsonify({'message': f'Error occurred while generating suggestions: {e}'}), 500



        
# # api for generating suggestions with client id :
# @app.route('/personality-assessment', methods=['POST'])
# def personality_selected():
#     try:
#         data = request.json
#         try :
#             investmentPersonality = data.get('investmentPersonality') # investment_personality
#             clientName = data.get('clientName')
#             print(f"The clients ClientName is : {clientName} ")
#             print(f"InvestmentPersonality received is : {investmentPersonality}")
#             logging.info('Recieved Values')
            
#         except Exception as e:
#             logging.info(f"Error occurred while retrieving client id: {e}")
#             return jsonify({'message': f'Error occurred while retrieving client id: {e}'}), 400

#         # Retrieve Client Financial Form Information :
#         try:
#             # Retrieve client_id from query parameters
#             clientId = data.get('clientId')
#             print(f"Received Client Id : {clientId}")
#             # client_id = request.args.get('clientId')
            
#             # Validate the client_id
#             if not clientId:
#                 return jsonify({'message': 'client_id is required as a query parameter'}), 400

#             # Define the S3 key for the object
#             s3_key = f"{client_summary_folder}client-data/{clientId}.json"

#             # Retrieve the object from S3
#             try:
#                 response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
#                 # Decode and parse the JSON data
#                 client_data = json.loads(response['Body'].read().decode('utf-8'))
#                 print(f"Received Client Data :\n{client_data}")
#                 # return jsonify({
#                 #     'message': 'Client data retrieved successfully.',
#                 #     'data': client_data
#                 # }), 200
                
#                 result,pie_chart_data,bar_chart_data,combined_chart_data = asyncio.run(make_suggestions_using_clientid(investmentPersonality,
#                                                                                                                    clientName,client_data))
                
#                 htmlSuggestions = markdown.markdown(result)
#                 logging.info(f"Suggestions for investor: \n{result}")
                
#                 formatSuggestions = markdown_to_text(htmlSuggestions)
#                 answer = markdown_table_to_html(formatSuggestions)
#                 print(answer)
               
#                 # Return the Results :
                
#                 # return jsonify({
#                 #     "status": 200,
#                 #     "message": "Success",
#                 #     "investmentSuggestions": answer, #formatSuggestions,
#                 #     "pieChartData": pie_chart_data,
#                 #     "barChartData": bar_chart_data,
#                 #     "compoundedChartData":combined_chart_data
#                 # }), 200
                
#                 return jsonify({
#                     "status": 200,
#                     "message": "Success",
#                     "investmentSuggestions": formatSuggestions,
#                     "pieChartData": pie_chart_data,
#                     "barChartData": bar_chart_data,
#                     "compoundedChartData":combined_chart_data
#                 }), 200
                
#             except s3.exceptions.NoSuchKey:
#                 return jsonify({'message': 'Client data not found for the given client_id.'}), 404
#             except Exception as e:
#                 return jsonify({'message': f"Error retrieving data: {e}"}), 500

#         except Exception as e:
#             return jsonify({'message': f"An error occurred: {e}"}), 500
    
#     except Exception as e:
#         print(f"An error occurred while requesting Data: {e}")
#         return jsonify({'message': f"An error occurred while requesting Data :" + str(e)}, 500)
   

#########################################################################################################################

# Route to handle generating investment suggestions
import shutil
import os

def save_file_to_folder(file_storage, destination_folder):
    try:
        # Ensure the destination folder exists
        if not os.path.exists(destination_folder):
            os.makedirs(destination_folder)
        
        # Construct the destination file path
        destination_file_path = os.path.join(destination_folder, file_storage.filename)
        
        # Check if the file already exists
        if not os.path.exists(destination_file_path):
            # Save the file
            file_storage.save(destination_file_path)
            print(f"File saved to {destination_file_path}")
            return destination_file_path
        else:
            print(f"File already exists at {destination_file_path}")
            return destination_file_path
        
    except Exception as e:
        print(f"Error saving file: {e}")


# #Working for both the methods :
# generate_suggestions by taking files as i/p :
@app.route('/generate-investment-suggestions', methods=['POST'])
def generate_investment_suggestions():
    try:
        assessment_file = request.files['assessmentFile']
        financial_file = request.files['financialFile']
        logging.info("Requested files")
        
        responses = extract_responses_from_docx(assessment_file)
        if not responses:
            raise Exception("Failed to extract responses from assessment file.")
        
        destination_folder = 'data'
        file_path = save_file_to_folder(financial_file, destination_folder)
        if not file_path:
            raise Exception("Failed to save financial file.")
        
        financial_data = asyncio.run(process_document(file_path))
        if not financial_data:
            raise Exception("Failed to process financial file.")
        
        logging.info(f"Received Responses from the file {responses}")
        
        personality = asyncio.run(determine_investment_personality(responses))
        if not personality:
            raise Exception("Failed to determine personality.")
        
        logging.info(f"Personality of the user is: {personality}")
        
        clientName = "Rohit Sharma" #"Emilly Watts"
        suggestions = asyncio.run(generate_investment_suggestions_for_investor(personality, clientName, financial_data, file_path))
        if "Error" in suggestions:
            raise Exception(suggestions)
        
        htmlSuggestions = markdown.markdown(suggestions)
        logging.info(f"Suggestions for investor: \n{suggestions}")
        
        formatSuggestions = markdown_to_text(htmlSuggestions)
        answer = markdown_table_to_html(formatSuggestions)
        print(answer)
        
        # need to change the data extraction process : 
        data_extracted = extract_numerical_data(suggestions)
        
        min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                        [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
        max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
                        [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

        # Normalize allocations
        min_allocations = normalize_allocations(min_allocations)
        max_allocations = normalize_allocations(max_allocations)

        bar_chart_data,pie_chart_data = generate_chart_data(data_extracted)
        
        # Sometimes Generating Pie Charts and Bar charts : 
        # data_extracted = extract_numerical_data(suggestions)

        # # Fixing pie and bar chart generation
        # growth_investments = data_extracted.get('Growth-Oriented Investments', {})
        # conservative_investments = data_extracted.get('Conservative Investments', {})

        # # Generate normalized allocations
        # min_allocations = [int(growth_investments[label]['min'].strip('%')) for label in growth_investments] + \
        #                 [int(conservative_investments[label]['min'].strip('%')) for label in conservative_investments]
        # max_allocations = [int(growth_investments[label]['max'].strip('%')) for label in growth_investments] + \
        #                 [int(conservative_investments[label]['max'].strip('%')) for label in conservative_investments]

        # # Normalize
        # min_allocations = normalize_allocations(min_allocations)
        # max_allocations = normalize_allocations(max_allocations)

        # # Bar Chart
        # bar_chart_data = {
        #     'labels': list(growth_investments.keys()) + list(conservative_investments.keys()),
        #     'datasets': [
        #         {'label': 'Allocation for Min returns', 'data': min_allocations, 'backgroundColor': 'skyblue'},
        #         {'label': 'Allocation for Max returns', 'data': max_allocations, 'backgroundColor': 'lightgreen'}
        #     ]
        # }

        # # Pie Chart
        # all_labels = list({**growth_investments, **conservative_investments}.keys())
        # num_labels = len(all_labels)
        # max_allocations_for_pie = normalize_allocations(
        #     [int(growth_investments.get(label, {}).get('max', '0').strip('%')) for label in growth_investments] +
        #     [int(conservative_investments.get(label, {}).get('max', '0').strip('%')) for label in conservative_investments]
        # )

        # # Normalize to 100% for pie chart
        # total = sum(max_allocations_for_pie)
        # max_allocations_for_pie = [(value / total) * 100 for value in max_allocations_for_pie]

        # dynamic_colors = generate_colors(num_labels)
        # pie_chart_data = {
        #     'labels': all_labels,
        #     'datasets': [{'label': 'Investment Allocation', 'data': max_allocations_for_pie, 'backgroundColor': dynamic_colors, 'hoverOffset': 4}]
        # }
    #############################################################################################
    
        print(f"Bar chart data: {bar_chart_data}")
        print(f"Pie chart data: {pie_chart_data}")

        # min_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['min'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
        #                   [int(data_extracted['Conservative Investments'][label]['min'].strip('%')) for label in data_extracted['Conservative Investments']]
        # max_allocations = [int(data_extracted['Growth-Oriented Investments'][label]['max'].strip('%')) for label in data_extracted['Growth-Oriented Investments']] + \
        #                   [int(data_extracted['Conservative Investments'][label]['max'].strip('%')) for label in data_extracted['Conservative Investments']]

        # # Normalize allocations
        # min_allocations = normalize_allocations(min_allocations)
        # max_allocations = normalize_allocations(max_allocations)

        # # Update Bar Chart Data
        
        # bar_chart_data = {
        #     'labels': list(data_extracted['Growth-Oriented Investments'].keys()) + list(data_extracted['Conservative Investments'].keys()),
        #     'datasets': [{
        #         'label': 'Allocation for Min returns',
        #         'data': min_allocations,
        #         'backgroundColor': 'skyblue'
        #     },
        #     {
        #         'label': 'Allocation for Max returns',
        #         'data': max_allocations,
        #         'backgroundColor': 'lightgreen'
        #     }]
        # }

        # # Similar changes can be made for the Pie Chart Data:
        # all_labels = list({**data_extracted['Growth-Oriented Investments'], **data_extracted['Conservative Investments']}.keys())
        # num_labels = len(all_labels)
        # max_allocations_for_pie = normalize_allocations(
        #     [int(data_extracted['Growth-Oriented Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Growth-Oriented Investments']] + 
        #     [int(data_extracted['Conservative Investments'].get(label, {}).get('max', '0').strip('%')) for label in data_extracted['Conservative Investments']]
        # )
        
        # # Generate colors based on the number of labels
        # dynamic_colors = generate_colors(num_labels)

        # # Update Pie Chart Data
        # pie_chart_data = {
        #     'labels': all_labels,
        #     'datasets': [{
        #         'label': 'Investment Allocation',
        #         'data': max_allocations_for_pie,
        #         'backgroundColor': dynamic_colors,
        #         'hoverOffset': 4
        #     }]
        # }
        
        # print(f"Pie Chart Data is : {pie_chart_data}")
        # Prepare the data for the line chart with inflation adjustment
        initial_investment = 10000
        combined_chart_data = prepare_combined_line_chart_data(data_extracted, initial_investment)
        print(f"\nThe combined chart data is: {combined_chart_data}")
        
        return jsonify({
            "status": 200,
            "message": "Success",
            "investmentSuggestions":  answer, #formatSuggestions,
            "pieChartData": pie_chart_data,
            "barChartData": bar_chart_data,
            "compoundedChartData": combined_chart_data
        }), 200
        
        # return jsonify({
        #     "status": 200,
        #     "message": "Success",
        #     "investmentSuggestions": htmlSuggestions,
        #     "pieChartData": pie_chart_data,
        #     "barChartData": bar_chart_data,
        #     "compoundedChartData": combined_chart_data
        # }), 200

    except Exception as e:
        logging.info(f"Error in generating investment suggestions: {e}")
        return jsonify({'message': f'Internal Server Error in Generating responses : {e}'}), 500

################################################################-------------------- Stocks Analysis -------------------------------- #################################
# #Stock analysis code :

from flask import Flask, request, jsonify
import yfinance as yf
import pandas as pd
import requests
import os
import logging


NEWS_API_KEY = os.getenv('NEWS_API_KEY')

# Simulate memory using a file
CHAT_HISTORY_FILE = "chat_history.json"
CHAT_ID_TRACKER_FILE = "chat_id_tracker.json"  # File to track chat_id

# Helper to save chat history to a file
def save_chat_history(chat_id, history):
    if os.path.exists(CHAT_HISTORY_FILE):
        with open(CHAT_HISTORY_FILE, 'r') as f:
            chat_data = json.load(f)
    else:
        chat_data = {}

    chat_data[chat_id] = history

    with open(CHAT_HISTORY_FILE, 'w') as f:
        json.dump(chat_data, f, indent=4)

# Helper to load chat history from a file
def load_chat_history(chat_id):
    if os.path.exists(CHAT_HISTORY_FILE):
        with open(CHAT_HISTORY_FILE, 'r') as f:
            chat_data = json.load(f)
        return chat_data.get(str(chat_id), [])
    return []

# Helper to track chat_id and increment it
def get_next_chat_id():
    if os.path.exists(CHAT_ID_TRACKER_FILE):
        with open(CHAT_ID_TRACKER_FILE, 'r') as f:
            chat_id_data = json.load(f)
        chat_id = chat_id_data.get("chat_id", 1)
    else:
        chat_id = 1

    chat_id_data = {"chat_id": chat_id + 1}
    with open(CHAT_ID_TRACKER_FILE, 'w') as f:
        json.dump(chat_id_data, f, indent=4)

    return chat_id


# # Fetch Stock Data :
def get_stock_data(ticker):
    try:
        # Step 1: Fetch Stock Data :
        stock = yf.Ticker(ticker)
        
        data = {}

        company_details = stock.info.get('longBusinessSummary', 'No details available')
        data['Company_Details'] = company_details
        sector = stock.info.get('sector', 'No sector information available')
        data['Sector'] = sector
        prev_close = stock.info.get('previousClose', 'No previous close price available')
        data['Previous_Closing_Price'] = prev_close
        open_price = stock.info.get('open', 'No opening price available')
        data['Today_Opening_Price'] = open_price
         
        hist = stock.history(period="5d")
        if not hist.empty and 'Close' in hist.columns:
            if hist.index[-1].date() == yf.download(ticker, period="1d").index[-1].date():
                close_price = hist['Close'].iloc[-1]
                data['Todays_Closing_Price'] = close_price
            else:
                data['Todays_Closing_Price'] = "Market is open, no closing price available yet."
        else:
            data['Todays_Closing_Price'] = "No historical data available for closing price."

        day_high = stock.info.get('dayHigh', 'No high price available')
        data['Today_High_Price'] = day_high
        day_low = stock.info.get('dayLow', 'No low price available')
        data['Today_Low_Price'] = day_low
        volume = stock.info.get('volume', 'No volume information available')
        data['Today_Volume'] = volume
        dividends = stock.info.get('dividendRate', 'No dividend information available')
        data['Today_Dividends'] = dividends
        splits = stock.info.get('lastSplitFactor', 'No stock split information available')
        data['Today_Stock_Splits'] = splits
        pe_ratio = stock.info.get('trailingPE', 'No P/E ratio available')
        data['PE_Ratio'] = pe_ratio
        market_cap = stock.info.get('marketCap', 'No market cap available')
        data['Market_Cap'] = market_cap

        # Additional KPIs
        data['EPS'] = stock.info.get('trailingEps', 'No EPS information available')
        data['Book_Value'] = stock.info.get('bookValue', 'No book value available')
        data['ROE'] = stock.info.get('returnOnEquity', 'No ROE information available')
        data['ROCE'] = stock.info.get('returnOnAssets', 'No ROCE information available')  # ROCE is not available directly
        
        # Revenue Growth (CAGR) and Earnings Growth would need to be calculated based on historical data
        earnings_growth = stock.info.get('earningsGrowth', 'No earnings growth available')
        revenue_growth = stock.info.get('revenueGrowth', 'No revenue growth available')

        data['Earnings_Growth'] = earnings_growth
        data['Revenue_Growth'] = revenue_growth
        
        
        income_statement = stock.financials
        balance_sheet = stock.balance_sheet
        cashflow = stock.cashflow

        # Step 2: Get News Related to Stock
        news_url = f'https://newsapi.org/v2/everything?q={ticker}&apiKey={NEWS_API_KEY}&pageSize=3'
        news_response = requests.get(news_url)
        if news_response.status_code == 200:
            news_data = news_response.json()
            articles = news_data.get('articles', [])
            if articles:
                top_news = "\n\n".join([f"{i+1}. {article['title']} - {article['url']}" for i, article in enumerate(articles)])
                data['Top_News'] = top_news
            else:
                data['Top_News'] = "No news articles found."
        else:
            data['Top_News'] = "Failed to fetch news articles."
    except Exception as e:
        logging.info(f"Error occurred while collecting stock data: {e}")
        print(f"Error occurred while collecting stock data: :\n{e}")
        return jsonify({'message': 'Internal Server Error in Stock Data Collection'}), 500
    
    print(data['Top_News'])
    
    try:
            
        # Step 3: Save Financial Data to Excel
        file_path = os.path.join('data', f'{ticker}_financial_data.xlsx')
        with pd.ExcelWriter(file_path, engine='xlsxwriter') as writer:
            income_statement.to_excel(writer, sheet_name='Income Statement')
            balance_sheet.to_excel(writer, sheet_name='Balance Sheet')
            cashflow.to_excel(writer, sheet_name='Cashflow')

        # Step 4: Perform Analysis
        avg_close = hist['Close'].mean()
        formatted_data = extract_excel_data(file_path)
        return data,formatted_data,avg_close,file_path
    except Exception as e:
        logging.info(f"Error occurred while performing analysis: {e}")
        print(f"Error occurred while performing analysis :\n{e}")
        return jsonify({'message': 'Internal Server Error in Stock Analysis'}), 500



# Helper function to extract a ticker from the query

# # Best Code answers the queries properly :)
def extract_ticker(query):
    # Mapping of popular company names to tickers for demonstration (you can expand this)
    companies_to_tickers = {
        "apple": "AAPL",
        "microsoft": "MSFT",
        "amazon": "AMZN",
        "tesla": "TSLA",
        "google": "GOOGL",
        "nvidia": "NVDA"
    }

    # Split the query into words
    words = query.lower().split()
    
    # Check for known company names or tickers
    for word in words:
        if word in companies_to_tickers:
            # word[0] = word[0].upper()
            # word[1:] = word[1:].lower()
            return companies_to_tickers[word] ,word.capitalize() #word.upper() #word 
    
    # Try to find a valid stock ticker by querying Yahoo Finance
    for word in words:
        if word:  # Ensure the word is not None or empty
            try:
                ticker = yf.Ticker(word.upper())
                if ticker.info.get('regularMarketPrice') is not None:
                    return ticker ,word.upper()  # Return the valid ticker
            except Exception as e:
                continue
    
    # Default fallback if no ticker is found
    print("No valid ticker found in the query.")
    return None,None


def format_chat_history_for_llm(chat_history, new_query):
    # Format chat history as a readable conversation for the model
    conversation = ""
    for entry in chat_history:
        user_query = entry.get('user_query', '')
        message = entry.get('message', '')
        
        # Append user query and model's response to the conversation
        conversation += f"User Query: {user_query}\nResponse: {message}\n\n"
    
    # Append the new query
    conversation += f"User Query: {new_query}\n"
    
    return conversation

from flask import jsonify, send_file, make_response

@app.route('/analyze_stock', methods=['POST'])
def analyze_stock():
    try:
        ticker = request.json.get('ticker')
        company = request.json.get('company',None)
        query = request.json.get('query')
        chat_id = request.json.get('chat_id', get_next_chat_id())  # Use auto-incrementing chat ID if not provided
        # chat_id = request.json.get('chat_id', 1)  # Default chat_id to 1 if not provided
        
        # Load chat history
        chat_history = load_chat_history(chat_id)

        # If no ticker provided in the request, try to extract it from the query
        if not ticker and query:
            # ticker = extract_ticker(query)
            
            ticker,company = extract_ticker(query)
        
        # If a valid ticker is found, fetch stock data
        if ticker:
            try:
                data, formatted_data, avg_close,file_path = get_stock_data(ticker)
                user_query = ticker  # Save the ticker as the user query
            except Exception as e:
                print("Error getting the stock data")
                return jsonify({'message': f'Error occurred while fetching stock data: {e}'}), 400
        else:
            # No valid ticker found, generate generic suggestions
            print("No valid ticker found in the query, generating general stock suggestions.")
            data = {}  # No specific stock data need to check for news
            formatted_data = ""  # No financial data
            avg_close = 0
            user_query = query  # Save the original user query if no ticker is found

        # If query is empty, set a default query for stock analysis
        # if not query:
        #     query = "Generate general stock suggestions based on current market trends and give some stock predictions."
        
        

        
         # Save the user's query (ticker or original query) to chat history
        if user_query:
            chat_history.append({"user_query": user_query, "message": query})
        
        # Detect if this is a follow-up query based on previous history
        if chat_history:
            print("This is a follow-up query. Checking previous chat history.")
            # The logic here could vary; you might compare the current query with past responses or check patterns
            query = f"Following up on: {chat_history[-1]['user_query']} \n\n {chat_history[-1]['message']}" + query

        # Save the user's query (ticker or original query) to chat history
        chat_history.append({"user_query": user_query, "message": query})
        
      
            
        # Format the chat history for the LLM
        try :
            formatted_history = format_chat_history_for_llm(chat_history, query)
        except Exception as e:
            logging.error(f"Error while formatting chat history for LLM: {e}")
            return jsonify({'message': 'Internal Server Error in Formatting Chat History'}), 500
        
        
    except Exception as e :
        logging.error(f"Error while fetching stock data: {e}")
        return jsonify({'message': 'Internal Server Error in Stock Data Fetch'}), 500
    
    try:
        if ticker:
            # task = f"""You are a Stock Market Expert. You know everything about stock market trends and patterns.Given a stock related query and if the company's details are provided,
            #             Based on the provided stock data, analyze the stock's performance, including whether it is overvalued or undervalued.
            #             Give the user details and information of all the KPI's related to the compnay such as PE ratio,EPS,Book Value,ROE,ROCE,Ernings Growth and Revenue Growth and give your views on them.
            #             Analyse all the stock information and provide the analysis of the company's performance related to Income Statement,Balance Sheet, and Cashflow.
            #             Predict the stock price range for the next week (if a particular time period is not mentioned) and provide reasons for your prediction.
            #             Advise whether to buy this stock now or not, with reasons for your advice. If no stock data is provided just answer the user's query.
            #             If the user asks for some stock suggestions then provide them a list of stock suggestions based on the query.
            #             If the user has asked a follow up question then provide them a good response by also considering their previous queries
            #             Do not answer any questions unrelated to the stocks."""
                        
            task = f"""You are a Stock Market Expert. You know everything about stock market trends and patterns.Given a stock related query and if the company's details are provided,
                    Based on the provided stock data, analyze the stock's performance, including whether it is overvalued or undervalued.
                    Give the user details and information of all the KPI's related to the compnay such as PE ratio,EPS,Book Value,ROE,ROCE,Ernings Growth and Revenue Growth and give your views on them.
                    Analyse all the stock information and provide the analysis of the company's performance related to Income Statement,Balance Sheet, and Cashflow.
                    Predict the stock price range for the next week (if a particular time period is not mentioned) and provide reasons for your prediction.
                    Advise whether to buy this stock now or not, with reasons for your advice."""
        

            query = task + "\nStock Data: " + str(data) + "\nFinancial Data: " + formatted_data + query
        
        else:
            task = """You are a Stock Market Expert. You know everything about stock market trends and patterns.Given a stock related query.
                        You are the best Stock recommendations AI and you give the best recommendations for stocks.Answer to the questions of the users and help them 
                        with any queries they might have.
                        If the user asks for some stock suggestions or some good stocks then provide them a list of stock suggestions based on the query give them the well known stocks in that sector or whatever the query asks for .
                        If the user has asked a follow up question then provide them a good response by also considering their previous queries
                        Do not answer any questions unrelated to the stocks."""
            
            query = task + query + "\n\nConversation:\n" + formatted_history #+ chat_history
            print(f"The formatted chat history passed to llm is : {formatted_history}")
            print(f"The query passed to llm is : {query}")
         # task = f"""You are a Stock Market Expert. You know everything about stock market trends and patterns.
        #             Based on the provided stock data, analyze the stock's performance, including whether it is overvalued or undervalued.
        #             Predict the stock price range for the next week and provide reasons for your prediction.
        #             Advise whether to buy this stock now or not, with reasons for your advice."""
        
        
        # Use your generative AI model for analysis (example with 'gemini-1.5-flash')
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(query)
        print(response.text)
        print(data)
    
    except Exception as e:
        logging.error(f"Error performing analysis with generative AI: {e}")
        return jsonify({f"error": "Failed to give analysis of stock data : {e}"}), 500
    
    # Extract response from the model
    try:
        html_suggestions = markdown.markdown(response.text)
        
        print(f"Html Suggestions : {html_suggestions}")
        
        logging.info(f"Suggestions for stock: \n{response.text}")
        
        # format_suggestions = markdown_to_text(response)
        print(f"Html Suggestions : {html_suggestions}")
        format_suggestions = markdown_to_text(html_suggestions)
        
    except Exception as e:
        logging.error(f"Error extracting text from response: {e}")
        print(f"Error extracting text from response : {e}")
        return jsonify({"error": "Failed to analyze stock data"}), 500

    # Save the assistant's response to chat history
    chat_history.append({"user_query": user_query, "message": format_suggestions})
    save_chat_history(chat_id, chat_history)

    # Increment chat_id for the next follow-up question
    new_chat_id = get_next_chat_id()
    
    if data == {}:
        data['Top_News'] = None
        
    data['Company'] = company if company else None
      
    # Return all collected and analyzed data
      # Create a response dictionary # gave responses in headers :
    # response_dict = {
    #     "data": data,
    #     "average_closing_price": f"${avg_close:.2f}",
    #     "analysis": format_suggestions,  # Use the response text here
    #     "news": data.get('Top_News'),
    #     "graph_url": f"https://finance.yahoo.com/chart/{ticker}"
    # }
    # # If the Excel file exists, send it as an attachment along with the response
    # if os.path.exists(file_path):
    #         file_response = send_file(file_path, as_attachment=True, download_name=f'{ticker}_financial_data.xlsx')
    #         file_response.headers['Content-Disposition'] = f'attachment; filename={ticker}_financial_data.xlsx'
    #         file_response.headers['X-Stock-Metadata'] = json.dumps(response_dict)  # Add metadata as a custom header
    #         return file_response
    # else:
    #     return jsonify(response_dict)
    
    # if os.path.exists(file_path): # works for either file or response
    #         # Combine the file response and JSON response
    #         file_response = send_file(file_path, as_attachment=True, download_name=f'{ticker}_financial_data.xlsx')
    #         file_response.headers['Content-Disposition'] = f'attachment; filename={ticker}_financial_data.xlsx'
    #         print("File is passed as attachment")
    #         return file_response
    # else:
    #     print("Data is passed")
    #     return jsonify(response_dict)
    
    return jsonify({
        # "Company": company,
        "data": data,
        "average_closing_price": f"${avg_close:.2f}",
        "analysis": format_suggestions,
        "news": data['Top_News'],
        "graph_url": f"https://finance.yahoo.com/chart/{ticker}"
    }) # "chat_history" : chat_history
    # # "new_chat_id" : new_chat_id

def extract_excel_data(file_path):
    financial_data = ""
    xls = pd.ExcelFile(file_path)
    for sheet_name in xls.sheet_names:
        df = pd.read_excel(xls, sheet_name=sheet_name)
        financial_data += f"\n\nSheet: {sheet_name}\n"
        financial_data += df.to_string()
    
    print(f"Financial data of excel file : {financial_data}")
    return financial_data

#########------------------- Portfolio Analysis --------------------------------################

@app.route('/current_stock_price', methods=['POST'])
def current_stock_price():
    try:
        ticker = request.json.get('ticker')
        stock = yf.Ticker(ticker)
        # Fetch the current stock price using the 'regularMarketPrice' field
        current_price = stock.info.get('regularMarketPrice')
        
        if not current_price:
            print(f"Failed to retrieve the current price for {ticker}.\nExtracting closing Price of the Stock")
            current_price = stock.history(period='1d')['Close'].iloc[-1]
            return jsonify({"current_price":current_price})
        
        if current_price is None:
            # If still None, check for mutual fund-specific fields
            print(f"Attempting to retrieve price for Mutual Fund {ticker}...")
            fund_close_price = stock.history(period="1d")['Close']
            if len(fund_close_price) > 0:
                current_price = fund_close_price.iloc[-1]  # Last available closing price
            return jsonify({"current_price":current_price})

        # If everything fails, raise an error
        if current_price is None:
            raise ValueError(f"Unable to retrieve price for {ticker}.")

        return jsonify({"current_price":current_price})
    
    except Exception as e:
        print(f"Failed to retrieve the current price for {ticker} : {e}")
        return jsonify({"error": f"Failed to retrieve the current price for {ticker}"}), 500


@app.route('/dividend_yield', methods=['POST'])
def dividend_yield():
    
    ticker_name = request.json.get('ticker')
    # Create a Ticker object using yfinance
    stock = yf.Ticker(ticker_name)
    
    # Fetch the stock information, including dividend yield
    try:
        dividend_yield = stock.info.get('dividendYield')
        sector = stock.info.get('sector')
        industry = stock.info.get('industry')

        if dividend_yield is not None:
            dividend_yield_percent = dividend_yield * 100  # Convert to percentage
            print(f"The dividend yield for {ticker_name} is: {dividend_yield_percent:.2f}%")
        else:
            print(f"No dividend yield information available for {ticker_name}.")
        
        # Additional information check to verify it's a REIT or commercial real estate company
        if industry and ('reit' in industry.lower() or 'real estate' in industry.lower()):
            print(f"{ticker_name} belongs to the {industry} industry.")
        else:
            print(f"{ticker_name} may not be a REIT or a commercial real estate company.")
        
        return jsonify({'dividend_yield_percent': float(dividend_yield_percent) , "status": 200})
    except Exception as e:
        print(f"Error occurred while fetching data for {ticker_name}: {e}")

# # Works well for real estate as well : 
## Direct Ownership :
def calculate_direct_property_ownership(vacancy_rate, capex, cap_rate, market_value, 
                                        property_management_fees, maintenance_repairs, 
                                        property_taxes, insurance, utilities, hoa_fees):
    # 1. Calculate the Gross Rental Income (assuming 100% occupancy)
    gross_rental_income = market_value * cap_rate
    
    # 2. Adjust for vacancy
    effective_rental_income = gross_rental_income * (1 - vacancy_rate)
    
    # 3. Total Operating Expenses
    operating_expenses = (property_management_fees + maintenance_repairs + property_taxes + 
                          insurance + utilities + hoa_fees)
    
    # 4. Net Operating Income (NOI)
    noi = effective_rental_income - operating_expenses
    
    # 5. Capital Expenditures (CapEx)
    # CapEx are large expenses that increase property value but are not part of NOI
    cash_flow_before_financing = noi - capex
    
    # 6. Return on Investment (ROI) assuming market value as initial investment
    roi = (cash_flow_before_financing / market_value) * 100
    
    # Return a dictionary with all key metrics
    return gross_rental_income,effective_rental_income,operating_expenses,noi,cash_flow_before_financing,roi
    # return {
    #     'Gross Rental Income': gross_rental_income,
    #     'Effective Rental Income': effective_rental_income,
    #     'Operating Expenses': operating_expenses,
    #     'Net Operating Income (NOI)': noi,
    #     'Cash Flow Before Financing': cash_flow_before_financing,
    #     'Return on Investment (ROI)': roi
    # }

#### AWS Method to place order :

# @app.route('/order_placed', methods=['POST'])
# def order_placed():
#     try:
#         # Extract data from the request
#         order_data = request.json.get('order_data')
#         client_name = request.json.get('client_name')
#         client_id = request.json.get('client_id')
#         funds = request.json.get('funds')
#         print(f"Received order for client: {client_name} ({client_id}), Available Funds: {funds}")

#         # File key for the S3 object
#         order_list_key = f"{order_list_folder}{client_id}_orders.json"

#         # Load existing data from S3 if available
#         try:
#             response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=order_list_key)
#             client_transactions = json.loads(response['Body'].read().decode('utf-8'))
#             print(f"Loaded existing transactions for client {client_id}")
#         except s3.exceptions.NoSuchKey:
#             # Initialize a new transaction list if the file doesn't exist
#             client_transactions = []
#             print(f"No existing transactions for client {client_id}. Initializing new list.")

#         # Process Real Estate or other assets based on asset class
#         assetClass = order_data.get('assetClass')
#         print(f"Processing Asset Class: {assetClass}")
        
#         if assetClass == 'Real Estate':
#             ownership = order_data.get('ownership')
#             if ownership in ['REIT/Fund', 'Commercial Real Estate (Triple Net Lease)']:
#                 # Real estate REIT/fund or commercial real estate transaction
#                 new_transaction = {
#                     "AssetClass": assetClass,
#                     "ownership": ownership,
#                     "Date": order_data.get('date'),
#                     "Name": order_data.get('name'),
#                     "TransactionAmount": order_data.get('investmentAmount'),
#                     "DividendYield": order_data.get('dividendYield')
#                 }
#             else:
#                 # Direct real estate transaction
#                 new_transaction = {
#                     "AssetClass": assetClass,
#                     "ownership": ownership,
#                     "Date": order_data.get('date'),
#                     "Name": order_data.get('name'),
#                     "estimated_annual_income": order_data.get('estimated_annual_income'),
#                     "estimated_yield": order_data.get('estimated_yield')
#                 }
#         else:
#             # Standard transaction for Stocks, Bonds, etc.
#             new_transaction = {
#                 "Market": order_data.get('market'),
#                 "AssetClass": assetClass,
#                 "Date": order_data.get('date'),
#                 "Action": order_data.get('buy_or_sell'),
#                 "Name": order_data.get('name'),
#                 "Symbol": order_data.get('symbol'),
#                 "Units": order_data.get('units'),
#                 "UnitPrice": order_data.get('unit_price'),
#                 "TransactionAmount": order_data.get('transactionAmount')
#             }

#         # Append the new transaction to the client's transaction list
#         client_transactions.append(new_transaction)
#         print(f"Appended transaction for client {client_id}: {new_transaction}")

#         # Save the updated data back to S3
#         updated_data = json.dumps(client_transactions, indent=4)
#         s3.put_object(Bucket=S3_BUCKET_NAME, Key=order_list_key, Body=updated_data)
#         print(f"Saved updated transactions for client {client_id} in S3 bucket.")

#         return jsonify({"message": "Order placed successfully", "status": 200})

#     except Exception as e:
#         print(f"Error occurred while placing order: {e}")
#         return jsonify({"message": f"Error occurred while placing order: {str(e)}"}), 500



# # # Updated Local Storage Code :

LOCAL_STORAGE_PATH = "data/orders/"

@app.route('/order_placed', methods=['POST'])
def order_placed():
    try:
        # Extract data from the request
        order_data = request.json.get('order_data')
        client_name = request.json.get('client_name', 'Rohit Sharma')  # Default client name
        client_id = request.json.get('client_id', 'RS4603')  # Default client ID if not provided
        funds = request.json.get('funds')  # Example extra data if needed
        print(f"Received order for client: {client_name} ({client_id}), Available Funds: {funds}")

        # Local file path for storing orders
        order_file_path = os.path.join(LOCAL_STORAGE_PATH, f"{client_id}_orders.json")

        # Load existing data from local storage if available
        if os.path.exists(order_file_path):
            with open(order_file_path, 'r') as file:
                client_transactions = json.load(file)
            print(f"Loaded existing transactions for client {client_id}")
        else:
            # Initialize a new transaction list if the file doesn't exist
            client_transactions = []
            print(f"No existing transactions for client {client_id}. Initializing new list.")

        # Process Real Estate or other assets based on asset class
        assetClass = order_data.get('assetClass')
        print(f"Processing Asset Class: {assetClass}")
        
        if assetClass == 'Real Estate':
            ownership = order_data.get('ownership')
            if ownership in ['REIT/Fund', 'Commercial Real Estate (Triple Net Lease)']:
                # Real estate REIT/fund or commercial real estate transaction
                new_transaction = {
                    "AssetClass": assetClass,
                    "Ownership": ownership,
                    "Date": order_data.get('date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                    "Name": order_data.get('name'),
                    "TransactionAmount": order_data.get('investmentAmount'),
                    "DividendYield": order_data.get('dividendYield')
                }
            else:
                # Direct real estate transaction
                new_transaction = {
                    "AssetClass": assetClass,
                    "Ownership": ownership,
                    "Date": order_data.get('date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                    "Name": order_data.get('name'),
                    "EstimatedAnnualIncome": order_data.get('estimated_annual_income'),
                    "EstimatedYield": order_data.get('estimated_yield')
                }
        else:
            # Standard transaction for Stocks, Bonds, etc.
            new_transaction = {
                "Market": order_data.get('market'),
                "AssetClass": assetClass,
                "Date": order_data.get('date', datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
                "Action": order_data.get('buy_or_sell'),
                "Name": order_data.get('name'),
                "Symbol": order_data.get('symbol'),
                "Units": order_data.get('units'),
                "UnitPrice": order_data.get('unit_price'),
                "TransactionAmount": order_data.get('transactionAmount')
            }

        # Append the new transaction to the client's transaction list
        client_transactions.append(new_transaction)
        print(f"Appended transaction for client {client_id}: {new_transaction}")

        # Save the updated data back to local storage
        with open(order_file_path, 'w') as file:
            json.dump(client_transactions, file, indent=4)
        print(f"Saved updated transactions for client {client_id} in local storage.")

        return jsonify({"message": "Order placed successfully", "status": 200})

    except Exception as e:
        print(f"Error occurred while placing order: {e}")
        return jsonify({"message": f"Error occurred while placing order: {str(e)}"}), 500




# ## Using AWS to Show Order :
# @app.route('/show_order_list', methods=['POST'])
# def show_order_list():
#     try:
#         # Get client_id from the request
#         client_id = request.json.get('client_id')

#         if not client_id:
#             return jsonify({"message": "Client ID is required", "status": 400})

#         # Define the S3 file key for the given client ID
#         order_list_key = f"{order_list_folder}{client_id}_orders.json"
#         print(f"clientIDDDD: {client_id}")

#         try:
#             # Fetch the file from the S3 bucket
#             response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=order_list_key)
#             file_content = response['Body'].read().decode('utf-8')

#             # Parse the file content as JSON
#             client_transactions = json.loads(file_content)
#             print(f"Retrieved transactions for client {client_id}: {client_transactions}")

#             return jsonify({"transaction_data": client_transactions, "status": 200})

#         except s3.exceptions.NoSuchKey:
#             # Handle case where the file does not exist in S3
#             print(f"No transactions found for client ID: {client_id}")
#             return jsonify({"message": "No transactions found for the provided client ID", "status": 404})

#         except Exception as e:
#             print(f"Error occurred while fetching data from S3: {e}")
#             return jsonify({"message": f"Error occurred while fetching data from S3: {str(e)}"}), 500

#     except Exception as e:
#         print(f"Error occurred while retrieving the order list: {e}")
#         return jsonify({"message": f"Error occurred while retrieving order list: {str(e)}"}), 500

# Updated Show Order List for Local Storage :

@app.route('/show_order_list', methods=['POST'])
def show_order_list():
    try:
        # Get client_id from the request
        client_id = request.json.get('client_id')

        if not client_id:
            return jsonify({"message": "Client ID is required", "status": 400})

        # Local file path for storing orders
        order_file_path = os.path.join(LOCAL_STORAGE_PATH, f"{client_id}_orders.json")

        # Check if the order file exists
        if os.path.exists(order_file_path):
            # Load transactions from the local file
            with open(order_file_path, 'r') as file:
                client_transactions = json.load(file)
            print(f"Retrieved transactions for client {client_id}: {client_transactions}")
            return jsonify({"transaction_data": client_transactions, "status": 200})
        else:
            print(f"No transactions found for client ID: {client_id}")
            return jsonify({"message": "No transactions found for the provided client ID", "status": 404})

    except Exception as e:
        print(f"Error occurred while retrieving the order list: {e}")
        return jsonify({"message": f"Error occurred while retrieving order list: {str(e)}"}), 500


### Using AWS to show Portfolio of the user :

# @app.route('/portfolio', methods=['POST'])
# def portfolio():
#     try:
#         # Extract the client_id from the POST request
#         client_id = request.json.get('client_id') #, 'RS4603')
#         curr_date = request.json.get('curr_date', None) # to be used to check market is open or closed
#         # print(f"Portfolio of the client with client id is :{client_id}")
#         order_list_key = f"{order_list_folder}{client_id}_orders.json"
#         # print(f"client_orders {order_list_key}")
            
#         if not client_id:
#             return jsonify({"message": "Client ID is required"}), 400


#         #  Load existing data of order list from S3 if available
#         try:
#             response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=order_list_key)
#             client_orders = json.loads(response['Body'].read().decode('utf-8'))
#             print(f"client_orders {client_orders}")

#         except s3.exceptions.NoSuchKey:
#             # Initialize a new transaction list if the file doesn't exist
#             client_orders = []
#             print(f"No existing transactions for client {client_id}. Initializing new list.")


#            # Read the order_list.json file
#         # with open('order_list.json', 'r') as f:
#         #     order_list = json.load(f)

#         # print(order_list)

#         # # Fetch orders for the client
#         # client_orders = order_list.get(client_id, [])
#         # # print(f"The orders are : {client_orders}")

#         # Check if any orders are found
#         if not client_orders:
            
#             return jsonify({"message": f"No data found for client_id: {client_id}"}), 404


#         # Filter the transactions for the specific client_id
#         # client_orders = [order for order in order_list if order.get('client_id') == client_id]
#         # if not client_orders:
#         #     return jsonify({"message": f"No data found for client_id: {client_id}"}), 404

#         # Initialize an array to store the portfolio data
#         portfolio_data = []
#         print(f"client_ordersclient_orders : {client_orders}")
#         # Iterate over all transactions for the specific client
#         portfolio_current_value,porfolio_daily_change,portfolio_daily_change_perc,portfolio_investment_gain_loss,portfolio_investment_gain_loss_perc,portfolio_daily_value_change = 0,0,0,0,0,0
#         for order in client_orders:
#             assetClass = order.get('AssetClass', 'N/A')
#             name = order.get('Name', 'N/A')  # Stock name
#             # market = order.get('market', 'N/A')
#             symbol = order.get('Symbol', 'N/A')
#             units = order.get('Units', 0)
#             bought_price = order.get('UnitPrice', 0)
#             transaction_type = order.get('Action', 'N/A')
#             transaction_amount = order.get('TransactionAmount', 0)
#             date = order.get('Date', 'N/A')
            
#             print(f"\n{assetClass} \n{name} \n{units} \n{bought_price} \n{transaction_type} \n{transaction_amount} \n{date}")
            
#             if assetClass == 'Real Estate':
#                 ownership = order.get('ownership')
#                 if ownership == 'REIT/Fund' or ownership == 'Commercial Real Estate (Triple Net Lease)':
#                     InvestmentAmount = order.get('TransactionAmount',500)
#                     print(f"Investment amount : {InvestmentAmount}")
#                     DividendYield = order.get('DividendYield',3.2)
#                     print(f"Dividend Yield : {DividendYield}")
#                     estimated_annual_income = InvestmentAmount * DividendYield
#                     print(f"Estimated Annualincome : {estimated_annual_income}")
#                     estimated_yield = round((InvestmentAmount/DividendYield))
#                     print(f"Estimated yield : {estimated_yield}")
                    
#                     current_price = 0 
#                     current_value = 0
#                     daily_price_change = 0
#                     daily_value_change = 0
#                     bought_price = 0
#                     transaction_amount = 0
#                     investment_gain_loss = 0
#                     investment_gain_loss_per = 0
                    
#                 elif ownership == "Direct":
#                     pass
                    
#             else :
#                 # Fetch the current stock price from external source (API, database)
#                 def fetch_current_stock_price(ticker):
#                     stock = yf.Ticker(ticker)
#                     try:
#                         # Fetch the current stock price using the 'regularMarketPrice' field
#                         current_price = stock.info.get('regularMarketPrice')
                        
#                         if current_price is None:
#                             print(f"Failed to retrieve the current price for {ticker}.\nExtracting closing Price of the Stock")
#                             # Fetch the last closing price if the current price is unavailable
#                             current_price = stock.history(period='1d')['Close'].iloc[-1]
                            
#                         # Ensure we have a valid price at this point
#                         if current_price is None:
#                             raise ValueError(f"Unable to fetch current or closing price for {ticker}.")
                        
#                         # print(current_price)
#                         return current_price
                    
#                     except Exception as e:
#                         # Handle exceptions more explicitly
#                         print(f"Error fetching stock price for {ticker}: {str(e)}")
#                         return 0

        
#                 current_price = fetch_current_stock_price(symbol)
#                 print(f"Current Stock Price is :{current_price}")
#                 # Calculate difference in price and percentage
#                 print(f"Bought price is : {bought_price}")
#                 diff_price = current_price - bought_price
#                 percentage_diff = (diff_price / bought_price) * 100 if bought_price > 0 else 0

#                 # Assume daily price change is available (fetch it if possible, or calculate)
#                 daily_price_change =  diff_price #current_price - order.get('previousDayPrice', bought_price)  # Placeholder logic
#                 daily_value_change = daily_price_change * units
#                 current_value = current_price*units

                
#                 # Calculate investment gain/loss and other financial metrics
#                 investment_gain_loss = diff_price * units
#                 investment_gain_loss_per = round(investment_gain_loss/transaction_amount*100,2)
#                 estimated_annual_income = 0 #order.get('estimatedAnnualIncome', 0)
#                 estimated_yield = 0 #(estimated_annual_income / (bought_price * units)) * 100 if bought_price > 0 else 0

#             # Append the transaction details to the portfolio_data array
#             portfolio_data.append({
#                 "assetClass": assetClass,
#                 "name": name,
#                 "symbol": symbol ,
#                 "Quantity": units,
#                 "Delayed_Price": current_price, # Delayed Price
#                 "current_value" : current_value ,
#                 "Daily_Price_Change": daily_price_change,
#                 "Daily_Value_Change" : daily_value_change,
#                 "Amount_Invested_per_Unit" :  bought_price, #transaction_amount/units ,
#                 "Amount_Invested": transaction_amount,
#                 "Investment_Gain_or_Loss_percentage": investment_gain_loss_per ,
#                 "Investment_Gain_or_Loss": investment_gain_loss,
#                 "Estimated_Annual_Income": estimated_annual_income,
#                 "Estimated_Yield": estimated_yield,
#                 "Time_Held": date,
#             })
            
#             print(f"Portfolio Data is : {portfolio_data}")
            
#                 # "Client ID": client_id,
#                 # "Market": market,
#                 # "Transaction Type": transaction_type,
#                 # "Price Per Unit (Bought)": bought_price, 
#                 # "Difference in Price": diff_price,
#                 # "Percentage Difference": f"{percentage_diff:.2f}%",
            
#             portfolio_current_value += current_value
#             porfolio_daily_change += daily_price_change
#             portfolio_daily_value_change += daily_value_change
#             portfolio_investment_gain_loss += investment_gain_loss
        
#         portfolio_daily_change_perc = round(porfolio_daily_change/portfolio_current_value *100 ,2)
#         portfolio_investment_gain_loss_perc = round(portfolio_investment_gain_loss/portfolio_current_value*100,4)
        
#         # Save the portfolio data as a JSON file
#         portfolio_file_path = f'portfolio_{client_id}.json'
#         with open(portfolio_file_path, 'w') as portfolio_file:
#             json.dump(portfolio_data, portfolio_file, indent=4)
            
#             portfolio_response = {
#             "portfolio_current_value":portfolio_current_value,
#             "porfolio_daily_change":porfolio_daily_change,
#             "portfolio_daily_change_perc":portfolio_daily_change_perc,
#             "portfolio_investment_gain_loss":portfolio_investment_gain_loss,
#             "portfolio_investment_gain_loss_perc":portfolio_investment_gain_loss_perc,
#             "portfolio_data": portfolio_data }
            
#         try:
#             s3.put_object(
#                 Bucket=S3_BUCKET_NAME,
#                 # Key=f"responses/{clientId}_response.json",
#                 Key=f"{portfolio_list_folder}/{client_id}.json",
#                 Body=json.dumps(portfolio_response),
#                 ContentType='application/json'
#             )
#             logging.info(f"Response successfully saved to S3 for client_id: {client_id}")
#         except Exception as e:
#             logging.error(f"Error occurred while saving to S3: {e}")
#             return jsonify({'message': f'Error occurred while saving to S3: {e}'}), 500
        

#         return jsonify(portfolio_response), 200

#     except Exception as e:
#         print(f"Error occured in portfolio : {e}")
#         return jsonify({"message": f"Error occurred: {str(e)}"}), 500


# Updated Portfolio List using Local Storage :
@app.route('/download_excel', methods=['GET'])
def download_excel():
    file_path = request.args.get('file_path')
    if os.path.exists(file_path):
        return send_file(file_path, as_attachment=True)
    else:
        return jsonify({"message": "File not found"}), 404


## Collect live news for stocks in portfolio :
# Define a function to fetch news for a given query 
def fetch_news(query):
    news_url = f'https://newsapi.org/v2/everything?q={query}&apiKey={NEWS_API_KEY}&pageSize=3'
    news_response = requests.get(news_url)
    
    if news_response.status_code == 200:
        news_data = news_response.json()
        articles = news_data.get('articles', [])
        if articles:
            top_news = "\n\n".join([f"{i+1}. {article['title']} - {article['url']}" for i, article in enumerate(articles)])
        else:
            top_news = "No news articles found."
    else:
        top_news = "Failed to fetch news articles."
    
    return top_news

# Function to collect news for each asset in the portfolio
def collect_portfolio_news(portfolio_data):
    portfolio_news = {}
    
    for asset in portfolio_data:
        asset_class = asset.get("AssetClass", "Unknown")
        name = asset.get("Name", "")
        symbol = asset.get("Symbol", None)
        
        # Generate a news query based on the asset class and name/symbol
        if asset_class == "Stocks" or asset_class == "Bonds":
            query = symbol if symbol else name
        elif asset_class == "cryptocurrency":
            query = asset.get("Name", "")
        elif asset_class == "Real Estate":
            query = asset.get("Name", "")
        else:
            query = asset.get("Name", "")
        
        # Fetch news for the query
        news = fetch_news(query)
        portfolio_news[name] = news
    
    return portfolio_news

# aws method :
# @app.route('/analyze_portfolio', methods=['POST'])
# def analyze_portfolio():
#     try:
#         # Retrieve the requested asset type
#         assetName = request.json.get('assetName', 'all')
#         client_name = request.json.get('client_name')
#         funds = request.json.get('funds')
#         client_id = request.json.get('client_id')
#         investor_personality = request.json.get('investor_personality', 'Aggressive Investor Personality')

#         # Initialize economic news to pass to LLM
#         topics = ["rising interest rates", "U.S. inflation", "geopolitical tensions", "US Elections", "Global Wars"]
#         economic_news = {topic: fetch_news(topic) for topic in topics}

#         # Load portfolio data for client (if analyzing the whole portfolio)
#         portfolio_data = {}
#         portfolio_news = {}

#         if assetName == 'all':
#             # Load the complete portfolio
#             with open(f'portfolio_{client_id}.json', 'r') as f:
#                 portfolio_data = json.load(f)
#             portfolio_news = collect_portfolio_news(portfolio_data)

#         else:
#             # Extract specific asset data from request if assetName is specific
#             portfolioList = request.json.get('portfolioList', [])
#             portfolio_data = [item for item in portfolioList if item.get('assetClass', '').lower() == assetName.lower()]
            
#             # Fetch news for each asset in the specified list
#             portfolio_news = collect_portfolio_news(portfolio_data)
        
#         # Fetching Client's Financial Data to get Financial 
#         print(f"Received Client Id : {client_id}")
#         # client_id = request.args.get('clientId')
        
#         # Validate the client_id
#         if not client_id:
#             return jsonify({'message': 'client_id is required as a query parameter'}), 400

#         # Define the S3 key for the object
#         s3_key = f"{client_summary_folder}client-data/{client_id}.json"

#         # Retrieve the object from S3
#         try:
#             response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
#             # Decode and parse the JSON data
#             client_data = json.loads(response['Body'].read().decode('utf-8'))
            
#         except Exception as e:
#             logging.error(f"Error occurred while retrieving client data from S3: {e}")
#             return jsonify({'message': f'Error occurred while retrieving client data from S3: {e}'}), 500

#          # Initialize portfolio-level metrics
#         portfolio_current_value = request.json.get('portfolio_current_value') 
#         portfolio_daily_change = request.json.get('porfolio_daily_change')
#         portfolio_daily_change_perc = request.json.get('portfolio_daily_change_perc')
#         portfolio_investment_gain_loss = request.json.get('portfolio_investment_gain_loss')
#         portfolio_investment_gain_loss_perc = request.json.get('portfolio_investment_gain_loss_perc')

#         print(f"{portfolio_current_value} \n{portfolio_daily_change} \n{portfolio_daily_change_perc} \n{portfolio_investment_gain_loss} \n{portfolio_investment_gain_loss_perc}" )


#         # Task prompt for LLM based on the asset name
#         task = f"""
#                 You are the best Stock Market Expert and Portfolio Analyst working for a Wealth Manager on the client: {client_name}.
#                 The portfolio contains several stocks and investments.
#                 Based on the portfolio data provided:

#                 - The available funds for the client are {funds}.
#                 - The current value of the portfolio is {portfolio_current_value}.
#                 - The portfolio's daily change is {portfolio_daily_change}.
#                 - The daily percentage change is {portfolio_daily_change_perc:.2f}%.
#                 - The total gain/loss in the portfolio is {portfolio_investment_gain_loss}.
#                 - The percentage gain/loss in the portfolio is {portfolio_investment_gain_loss_perc:.2f}%.
#                 - The risk tolerance of the client based on their investment personality is {investor_personality}.

#                 Given the Clients Financial Data: {client_data} determine the Financial Situation based on the Assets,Liabilities and Debts of of the Client as : Stable,Currently Stable or Unstable.
#                 Based on the Client's Financial Situation and the Client's Financial Goals,
#                 Provide an in-depth analysis of the portfolio, including an evaluation of performance, suggestions for improvement, 
#                 and detailed stock recommendations to the Wealth Manager for the client based on the Client's Financial Situation and in order to achive their Financial Goal's and the Client's risk tolerance for the given portfolio : {portfolio_data}
#                 and top news of each holdings in the portfolio : {portfolio_news} and the economic news of the US Market : {economic_news}

#                 - If the client has a conservative investment personality, give stocks and low risk assets recommendations that could provide returns with minimal risk.
#                 - If the client has a moderate investment personality, give stocks and medium risk assets recommendations that could provide returns with a moderate level of risk.
#                 - If the client has an aggressive investment personality, give stocks,Real Estate,cryptocurrency,or any High Risk High Reward Assets recommendations that could provide higher returns with higher risk. 
#                 Also, help the Wealth Manager rearrange the funds, including which stocks to sell and when to buy them.

#                 Provide detailed reasons for each stock recommendation based on the funds available to the client and their investor personality in order for the Client to achive their Financial Goals. Include specific suggestions on handling the portfolio, such as when to buy, when to sell, and in what quantities, to maximize the client's profits. Highlight the strengths and weaknesses of the portfolio, and give an overall performance analysis.

#                 Additionally, provide:

#                 1. A risk assessment of the current portfolio composition.
#                 2. Give a proper Analysis and Performance of the current portfolio holdings by considering its current news.
#                 3. Funds Rearrangement of the portfolio if required and give stocks that would give better returns to the client.
#                 4. Recommendations for sector allocation to balance risk and return as per the investor personality and suggest stocks accordingly.
#                 5. Strategies for tax efficiency in the portfolio management.
#                 6. Insights on market trends and current economic news that could impact the portfolio.
#                 7. Explain in brief the Contingency plans for different market scenarios (bullish, bearish, and volatile markets) and suggest some stocks/assets and sectors from which the client can benefit .
#                 8. Explain How the client can achieve their Financial Goals of the client that they have mentioned and whether they can  achieve it/them till the time(if mentioned) they are planning of achieving it/them.

#                 Ensure the analysis is comprehensive and actionable, helping the Wealth Manager make informed decisions to optimize the client's portfolio.
#                 Dont give any Disclaimer as you are providing all the information to a Wealth Manager who is a Financial Advisor and has good amount of knowledge and experience in managing Portfolios.
#                 """

#         # Generate response using LLM
#         try:
#             model = genai.GenerativeModel('gemini-1.5-flash')
#             response = model.generate_content(task)

#             # Process the response
#             html_suggestions = markdown.markdown(response.text)
#             format_suggestions = markdown_to_text(html_suggestions)
            
#             # Return response in JSON format
#             return jsonify({
#                     "portfolio_current_value": portfolio_current_value,
#                     "portfolio_daily_change": portfolio_daily_change,
#                     "portfolio_daily_change_perc": f"{portfolio_daily_change_perc:.2f}%",
#                     "portfolio_investment_gain_loss": portfolio_investment_gain_loss,
#                     "portfolio_investment_gain_loss_perc": f"{portfolio_investment_gain_loss_perc:.2f}%",
#                     "suggestion": format_suggestions,
#                      "assetClass": assetName
#             }), 200

#         except Exception as e:
#             print(f"Error generating suggestions from LLM: {e}")
#             return jsonify({"message": f"Error occurred while analyzing the portfolio: {e}"}), 500

#     except Exception as e:
#         print(f"Error in analyzing portfolio for asset '{assetName}': {e}")
#         return jsonify({"message": f"Error analyzing portfolio for asset '{assetName}'"}), 500


# Analyzing the Portfolio using Local Storage :
# New Version :
# New Version :

# File paths for local storage
LOCAL_STORAGE_PATH = "local_data"
ORDER_LIST_PATH = os.path.join(LOCAL_STORAGE_PATH, "orders")
DAILY_CHANGES_PATH = os.path.join(LOCAL_STORAGE_PATH, "daily_changes")
PORTFOLIO_PATH = os.path.join(LOCAL_STORAGE_PATH, "portfolios")

# Ensure directories exist
os.makedirs(ORDER_LIST_PATH, exist_ok=True)
os.makedirs(DAILY_CHANGES_PATH, exist_ok=True)
os.makedirs(PORTFOLIO_PATH, exist_ok=True)

@app.route('/portfolio', methods=['POST'])
def portfolio():
    try:
        # Extract client ID and current date
        client_id = request.json.get('client_id')
        curr_date = request.json.get('curr_date', datetime.now().strftime('%Y-%m-%d'))
 
        if not client_id:
            return jsonify({"message": "Client ID is required"}), 400
 
        # Load orders from the local file
        order_file_path = os.path.join(LOCAL_STORAGE_PATH, f"{client_id}_orders.json")
 
        if not os.path.exists(order_file_path):
            return jsonify({"message": f"No orders found for client_id: {client_id}"}), 404
 
        with open(order_file_path, 'r') as file:
            client_orders = json.load(file)
 
        # Initialize portfolio data and metrics
        portfolio_data = []
        portfolio_current_value = 0
        porfolio_daily_change = 0
        portfolio_investment_gain_loss = 0
 
        # Load existing daily changes for the quarter
        daily_changes_file = os.path.join(DAILY_CHANGES_PATH, f"{client_id}_daily_changes.json")
        if os.path.exists(daily_changes_file):
            with open(daily_changes_file, 'r') as file:
                daily_changes = json.load(file)
        else:
            daily_changes = {}
 
        # Process client orders
        for order in client_orders:
            asset_class = order.get('AssetClass', 'N/A')
            name = order.get('Name', 'N/A')
            symbol = order.get('Symbol', 'N/A')
            units = order.get('Units', 0)
            bought_price = order.get('UnitPrice', 0)
            transaction_amount = order.get('TransactionAmount', 0)
 
            # Fetch current stock price
            def fetch_current_stock_price(ticker):
                stock = yf.Ticker(ticker)
                try:
                    current_price = stock.history(period='1d')['Close'].iloc[-1]
                    return current_price
                except Exception as e:
                    print(f"Error fetching stock price for {ticker}: {e}")
                    return 0
 
            current_price = fetch_current_stock_price(symbol)
            diff_price = current_price - bought_price
            daily_price_change = diff_price
            daily_value_change = daily_price_change * units
            current_value = current_price * units
 
            # Calculate investment gain/loss and other metrics
            investment_gain_loss = diff_price * units
            investment_gain_loss_per = round((investment_gain_loss / transaction_amount) * 100, 2) if transaction_amount > 0 else 0
 
            # Append data to portfolio
            portfolio_data.append({
                "assetClass": asset_class,
                "name": name,
                "symbol": symbol,
                "Quantity": units,
                "Delayed_Price": current_price,
                "current_value": current_value,
                "Daily_Price_Change": daily_price_change,
                "Daily_Value_Change": daily_value_change,
                "Amount_Invested_per_Unit": bought_price,
                "Amount_Invested": transaction_amount,
                "Investment_Gain_or_Loss_percentage": investment_gain_loss_per,
                "Investment_Gain_or_Loss": investment_gain_loss,
                "Time_Held": order.get('Date', 'N/A'),
            })
 
            # Update portfolio metrics
            portfolio_current_value += current_value
            porfolio_daily_change += daily_price_change
            portfolio_investment_gain_loss += investment_gain_loss
 
        # Calculate daily change percentages
        portfolio_daily_change_perc = round((porfolio_daily_change / portfolio_current_value) * 100, 2) if portfolio_current_value > 0 else 0
        portfolio_investment_gain_loss_perc = round((portfolio_investment_gain_loss / portfolio_current_value) * 100, 4) if portfolio_current_value > 0 else 0
 
        # Update daily changes for the current date
        daily_changes[curr_date] = {
            "portfolio_current_value": portfolio_current_value,
            "porfolio_daily_change": porfolio_daily_change,
            "portfolio_daily_change_perc": portfolio_daily_change_perc,
            "portfolio_investment_gain_loss": portfolio_investment_gain_loss,
            "portfolio_investment_gain_loss_perc": portfolio_investment_gain_loss_perc,
        }
 
        # Save daily changes to a file
        with open(daily_changes_file, 'w') as file:
            json.dump(daily_changes, file, indent=4)
 
        # Save portfolio data as JSON
        portfolio_file_path = os.path.join(PORTFOLIO_PATH, f"portfolio_{client_id}.json")
        with open(portfolio_file_path, 'w') as file:
            json.dump(portfolio_data, file, indent=4)
 
        # Response data
        portfolio_response = {
            "portfolio_current_value": portfolio_current_value,
            "porfolio_daily_change": porfolio_daily_change,
            "portfolio_daily_change_perc": portfolio_daily_change_perc,
            "portfolio_investment_gain_loss": portfolio_investment_gain_loss,
            "portfolio_investment_gain_loss_perc": portfolio_investment_gain_loss_perc,
            "daily_changes": daily_changes,
            "portfolio_data": portfolio_data,
        }
 
        return jsonify(portfolio_response), 200
 
    except Exception as e:
        print(f"Error occurred in portfolio: {e}")
        return jsonify({"message": f"Error occurred: {str(e)}"}), 500


# New Version :
from flask import Flask, request, jsonify
import requests
import os
import json
import markdown

@app.route('/analyze_portfolio', methods=['POST'])
def analyze_portfolio():
    try:
        # Retrieve input data
        assetName = request.json.get('assetName', 'all')
        client_id = request.json.get('client_id')
        client_name = request.json.get('client_name')
        funds = request.json.get('funds')
        investor_personality = request.json.get('investor_personality', 'Aggressive Investor Personality')

        # Validate client_id
        if not client_id:
            return jsonify({"message": "Client ID is required"}), 400

        # Define file path for portfolio data
        portfolio_file_path = f"local_data/portfolios/portfolio_{client_id}.json"

        # Load portfolio data from local file
        if os.path.exists(portfolio_file_path):
            with open(portfolio_file_path, 'r') as f:
                portfolio_data = json.load(f)
        else:
            return jsonify({"message": f"Portfolio file not found for client ID: {client_id}"}), 404

        # Verify portfolio data is a list
        if not isinstance(portfolio_data, list):
            return jsonify({"message": "Portfolio data is not in the expected format"}), 500

        # Initialize variables to calculate portfolio-level metrics
        portfolio_current_value = sum(asset["current_value"] for asset in portfolio_data)
        portfolio_daily_change = sum(asset["Daily_Value_Change"] for asset in portfolio_data)
        portfolio_investment_gain_loss = sum(asset["Investment_Gain_or_Loss"] for asset in portfolio_data)

        if portfolio_current_value != 0:
            portfolio_daily_change_perc = (portfolio_daily_change / portfolio_current_value) * 100
            portfolio_investment_gain_loss_perc = (portfolio_investment_gain_loss / portfolio_current_value) * 100
        else:
            portfolio_daily_change_perc = 0
            portfolio_investment_gain_loss_perc = 0

        # Filter portfolio data if a specific asset type is requested
        if assetName != 'all':
            filtered_portfolio_data = [
                asset for asset in portfolio_data if asset["assetClass"].lower() == assetName.lower()
            ]
        else:
            filtered_portfolio_data = portfolio_data

        # Initialize economic news to pass to LLM
        topics = ["rising interest rates", "U.S. inflation", "geopolitical tensions", "US Elections", "Global Wars"]
        economic_news = {topic: fetch_news(topic) for topic in topics}
        
        # Load client financial data from local storage
        client_data_file_path = f"client_data/client_data/{client_id}.json"
        if os.path.exists(client_data_file_path):
            with open(client_data_file_path, 'r') as f:
                client_data = json.load(f)
        else:
            return jsonify({"message": f"No client data found for client ID: {client_id}"}), 404

        portfolio_news = collect_portfolio_news(filtered_portfolio_data)

        # Task prompt for LLM based on the asset name
        task = f"""
                You are the best Stock Market Expert and Portfolio Analyst working for a Wealth Manager on the client: {client_name}.
                The portfolio contains several stocks and investments.
                Based on the portfolio data provided:

                - The available funds for the client are {funds}.
                - The current value of the portfolio is {portfolio_current_value}.
                - The portfolio's daily change is {portfolio_daily_change}.
                - The daily percentage change is {portfolio_daily_change_perc:.2f}%.
                - The total gain/loss in the portfolio is {portfolio_investment_gain_loss}.
                - The percentage gain/loss in the portfolio is {portfolio_investment_gain_loss_perc:.2f}%.
                - The risk tolerance of the client based on their investment personality is {investor_personality}.
                - These are the relevant news for all the stocks in the portfolio : {portfolio_news}
                - These are the relevant economic news : {economic_news}

                Given the Clients Financial Data: {client_data} determine the Financial Situation based on the Assets, Liabilities, and Debts of the Client as: Stable, Currently Stable, or Unstable.
                Based on the Client's Financial Situation and the Client's Financial Goals, provide an in-depth analysis of the portfolio, 
                including an evaluation of performance, suggestions for improvement, and detailed stock recommendations to the Wealth Manager 
                for the client based on the Client's Financial Situation and risk tolerance for the portfolio: {filtered_portfolio_data}.
                Ensure your analysis includes detailed recommendations, risk assessments, and strategies tailored to the client's financial goals based on the recent news regarding the assets in the portfolio : {portfolio_news} and the economic news : {economic_news} .
                """

        try:
            model = genai.GenerativeModel('gemini-1.5-flash')
            response = model.generate_content(task)

            # Process the response
            html_suggestions = markdown.markdown(response.text)
            format_suggestions = markdown_to_text(html_suggestions)

            # Return the analysis response
            return jsonify({
                "portfolio_current_value": portfolio_current_value,
                "portfolio_daily_change": portfolio_daily_change,
                "portfolio_daily_change_perc": f"{portfolio_daily_change_perc:.2f}%",
                "portfolio_investment_gain_loss": portfolio_investment_gain_loss,
                "portfolio_investment_gain_loss_perc": f"{portfolio_investment_gain_loss_perc:.2f}%",
                "suggestion": format_suggestions,
                "assetClass": assetName
            }), 200

        except Exception as e:
            print(f"Error in generating analysis: {e}")
            return jsonify({"message": f"Error generating analysis: {e}"}), 500

    except Exception as e:
        print(f"Error in analyzing portfolio: {e}")
        return jsonify({"message": f"Error analyzing portfolio: {e}"}), 500


#####################################################################################################################
# Actual vs predicted comparison using local storage :


# Local storage directories
BASE_DIR = "local_data"
DAILY_CHANGES_DIR = os.path.join(BASE_DIR, "daily_changes")
PREDICTIONS_DIR = os.path.join(BASE_DIR, "predictions")
COMPARISONS_DIR = os.path.join(BASE_DIR, "comparisons")
PORTFOLIO_DIR = "local_data/portfolios"
# CLIENT_SUMMARY_DIR = os.path.join(BASE_DIR, "client_summary") 

# Ensure all directories exist
os.makedirs(DAILY_CHANGES_DIR, exist_ok=True)
os.makedirs(PREDICTIONS_DIR, exist_ok=True)
os.makedirs(COMPARISONS_DIR, exist_ok=True)
# os.makedirs(CLIENT_SUMMARY_DIR, exist_ok=True)

# Helper: Fetch current date and determine the start of the quarter
def get_start_of_quarter():
    current_date = datetime.now()
    quarter_start_months = [1, 4, 7, 10]  # January, April, July, October
    start_month = quarter_start_months[(current_date.month - 1) // 3]
    return datetime(current_date.year, start_month, 1)

# Helper: Save data to a file
def save_to_file(filepath, data):
    with open(filepath, 'w') as f:
        json.dump(data, f, indent=4)

# Helper: Load data from a file
def load_from_file(filepath):
    if os.path.exists(filepath):
        with open(filepath, 'r') as f:
            return json.load(f)
    return None

# Function to update daily return data
def update_daily_returns(client_id, portfolio_daily_change, current_date):
    daily_changes_file = os.path.join(DAILY_CHANGES_DIR, f"{client_id}_daily_changes.json")
    
    # Load existing data or initialize if not found
    daily_changes = load_from_file(daily_changes_file) or {
        "start_of_quarter": str(get_start_of_quarter()), 
        "daily_returns": []
    }
    
    # Check if today's return exists and update it if necessary
    last_recorded_date = daily_changes["daily_returns"][-1]["date"] if daily_changes["daily_returns"] else None
    if last_recorded_date == current_date:
        if daily_changes["daily_returns"][-1]["price"] != portfolio_daily_change:
            daily_changes["daily_returns"][-1]["price"] = portfolio_daily_change
    else:
        # Add new entry for today's return
        daily_changes["daily_returns"].append({"date": current_date, "price": portfolio_daily_change})
    
    # Save updated daily changes back to file
    save_to_file(daily_changes_file, daily_changes)

# Function to calculate actual returns
def calculate_actual_returns(client_id):
    start_of_quarter = get_start_of_quarter()
    current_date = datetime.now().strftime("%Y-%m-%d")
    daily_changes_file = os.path.join(DAILY_CHANGES_DIR, f"{client_id}_daily_changes.json")
    
    # Load daily changes
    daily_changes = load_from_file(daily_changes_file)
    if not daily_changes:
        return {"message": "No daily return data found for the client."}
    
    # Filter data from the start of the quarter
    quarter_data = [entry for entry in daily_changes["daily_returns"] if datetime.strptime(entry["date"], "%Y-%m-%d") >= start_of_quarter]
    
    # Calculate total and percentage returns
    if quarter_data:
        initial_price = quarter_data[0]["price"]
        final_price = quarter_data[-1]["price"]
        total_return = sum(entry["price"] for entry in quarter_data)
        percentage_return = (final_price - initial_price) / initial_price * 100
        return {"total_return": total_return, "percentage_return": percentage_return, "daily_returns": quarter_data}
    
    return {"message": "No valid data for the current quarter."}

# Actual vs Predicted Endpoint
@app.route('/actual_vs_predicted', methods=['POST'])
def actual_vs_predicted():
    try:
        # Retrieve client ID and current portfolio daily change
        client_id = request.json.get('client_id')
        portfolio_daily_change = request.json.get('porfolio_daily_change')
        current_date = datetime.now().strftime("%Y-%m-%d")

        current_quarter = "2024_Q4"
        
        # Load previously predicted line chart data
        predicted_file = os.path.join(PREDICTIONS_DIR, f"{client_id}_{current_quarter}_line_chart.json")
        predicted_line_chart_data = load_from_file(predicted_file)
        if not predicted_line_chart_data:
            return jsonify({'message': 'No previous predictions found for this client.'}), 404

        # Fetch and process portfolio data
        PORTFOLIO_DIR_file = os.path.join(PORTFOLIO_DIR, f"portfolio_{client_id}.json")
        portfolio_data = load_from_file(PORTFOLIO_DIR_file)
        if not portfolio_data:
            return jsonify({'message': 'Portfolio data not found for this client.'}), 404

        # Update daily returns if there's a change
        # update_daily_returns(client_id, portfolio_daily_change, current_date)

        # Calculate actual returns
        # actual_line_chart_data = calculate_actual_returns(client_id)
        
        # actual_line_chart_data = [2582.1 - 2209.48 + 2469.66*2 - 4709.36,
        #                           2199.4 - 2209.48 + 2613.03*2 - 4709.36,
        #                           2501.9 - 2209.48 + 2517.45*2 - 4709.36,
        #                           2490.6 - 2209.48 + 2517.45*2 - 4709.36,
        #                           3225.6 - 2209.48 + 3131.91*2 - 4709.36,
        #                           3463.1 - 2209.48 + 3705.41*2 - 4709.36,
        #                           3463.1 - 2209.48 + 3719.06*2 - 4709.36,
        #                           4191.1 - 2209.48 + 3898.17*2 - 4709.36,
        #                           ]
        
        actual_line_chart_data = [602.58,506.62,618.96,606.66,1570.58,3955.08,3982.38,4068.60]
                
                                  
                                
        # Combine actual and predicted data
        comparison_data = {
            "actual": actual_line_chart_data,
            "predicted": predicted_line_chart_data
        }

        # Save comparison data locally
        comparison_file = os.path.join(COMPARISONS_DIR, f"{client_id}_{current_quarter}_comparison_chart.json")
        save_to_file(comparison_file, comparison_data)

        # Return the comparison data
        return jsonify({
            "client_id": client_id,
            "comparison_chart_data": comparison_data
        }), 200

    except Exception as e:
        print(f"Error generating comparison: {e}")
        return jsonify({"message": f"Error generating comparison: {e}"}), 500


######################################################################################################################
# Portfolio Return on Investment Prediction for next quarter Local Storage:
import os
import json
from flask import Flask, request, jsonify
from datetime import datetime, timedelta
import calendar
import markdown


# Generate next quarter's dates
def get_next_quarter_dates():
    current_date = datetime.now()
    current_month = current_date.month

    # Determine the starting month of the next quarter
    if current_month in [1, 2, 3]:  # Q1
        start_month = 4  # Q2
    elif current_month in [4, 5, 6]:  # Q2
        start_month = 7  # Q3
    elif current_month in [7, 8, 9]:  # Q3
        start_month = 10  # Q4
    else:  # Q4
        start_month = 1  # Q1 of the next year

    # Determine the year of the next quarter
    next_quarter_year = current_date.year if start_month != 1 else current_date.year + 1

    # Generate dates for the next quarter
    next_quarter_dates = []
    for month in range(start_month, start_month + 3):
        # Get the first, 15th, and last day of the month
        first_day = datetime(next_quarter_year, month, 1)
        fifteenth_day = datetime(next_quarter_year, month, 15)
        last_day = datetime(next_quarter_year, month, calendar.monthrange(next_quarter_year, month)[1])

        next_quarter_dates.extend([first_day.strftime("%Y-%m-%d"),
                                   fifteenth_day.strftime("%Y-%m-%d"),
                                   last_day.strftime("%Y-%m-%d")])

    return next_quarter_dates

def get_next_quarter():
    current_date = datetime.now()
    current_month = current_date.month

    # Determine the starting month of the next quarter
    if current_month in [1, 2, 3]:  # Q1
        start_month = 4  # Q2
        next_quarter = str(current_date.year) + "_Q2"
    elif current_month in [4, 5, 6]:  # Q2
        start_month = 7  # Q3
        next_quarter = str(current_date.year) + "_Q3"
    elif current_month in [7, 8, 9]:  # Q3
        start_month = 10  # Q4
        next_quarter = str(current_date.year) + "_Q4"
    else:  # Q4
        start_month = 1  # Q1 of the next year
        next_quarter = str(current_date.year + 1) + "_Q1"
        
    # # Determine the year of the next quarter
    # next_quarter_year = current_date.year if start_month != 1 else current_date.year + 1
    
    return next_quarter
    
    

# Function to save data to a file
def save_to_file(filepath, data):
    with open(filepath, 'w') as f:
        json.dump(data, f, indent=4)

# Function to load data from a file
def load_from_file(filepath):
    if os.path.exists(filepath):
        with open(filepath, 'r') as f:
            return json.load(f)
    return None

# Extract line chart data from LLM responseimport re
from datetime import datetime
from bs4 import BeautifulSoup
import re

# Extract line chart data from LLM response
def extract_line_chart_data(llm_response_text):
    try:
        # Parse HTML content
        soup = BeautifulSoup(llm_response_text, "html.parser")
        lines = soup.get_text().split("\n")
        
        line_chart_data = {
            "dates": [],
            "overall_returns": {
                "percentages": [],
                "amounts": []
            }
        }

        # Iterate through lines and match rows with the expected format
        for line in lines:
            match = re.match(r"\|\s*(\d{4}-\d{2}-\d{2})\s*\|\s*([-+]?\d*\.?\d+)%\s*\|\s*\$?(\d+)", line)
            if match:
                date = match.group(1).strip()
                return_percentage = float(match.group(2).strip())
                return_amount = float(match.group(3).strip())
                
                line_chart_data["dates"].append(date)
                line_chart_data["overall_returns"]["percentages"].append(return_percentage)
                line_chart_data["overall_returns"]["amounts"].append(return_amount)
        
        return line_chart_data

    except Exception as e:
        print(f"Error extracting line chart data: {e}")
        return {}

##########################################################################################
# # Prediction Improvements :

# #1. Fetch historical returns for the past 3 months
import yfinance as yf

# Example usage:
# portfolio = [{'ticker': 'AAPL'}, {'ticker': 'MSFT'}, {'ticker': 'GOOGL'}]

# Fetch and store historical returns
# for asset in portfolio:
#     asset['historical_returns'] = fetch_historical_returns(asset['ticker'])
#     print(f"Fetched {len(asset['historical_returns'])} returns for {asset['ticker']}")


def fetch_historical_returns(ticker, period='3mo'):
    """Fetch historical returns using yfinance."""
    stock = yf.Ticker(ticker)
    data = stock.history(period=period)
    data['returns'] = data['Close'].pct_change()
    return data['returns'].dropna()



# 2. Add Quantitative Metrics :
import numpy as np

def compute_volatility(returns):
    """Calculate standard deviation of returns (volatility)."""
    return np.std(returns)

def compute_sharpe_ratio(returns, risk_free_rate=0.0):
    """Calculate Sharpe Ratio (risk-adjusted return)."""
    mean_return = np.mean(returns)
    std_dev = np.std(returns)
    return (mean_return - risk_free_rate) / std_dev if std_dev != 0 else 0

def compute_beta(asset_returns, market_returns):
    """Calculate Beta (sensitivity to market)."""
    # Align lengths of asset_returns and market_returns
    min_length = min(len(asset_returns), len(market_returns))
    asset_returns = asset_returns.iloc[-min_length:]
    market_returns = market_returns.iloc[-min_length:]

    # Calculate covariance and beta
    covariance = np.cov(asset_returns, market_returns)[0][1]
    market_variance = np.var(market_returns)
    return covariance / market_variance if market_variance != 0 else 0


# def compute_beta(asset_returns, market_returns):
#     """Calculate Beta (sensitivity to market)."""
#     covariance = np.cov(asset_returns, market_returns)[0][1]
#     market_variance = np.var(market_returns)
#     return covariance / market_variance

# Fetch market index returns (S&P 500)
market_returns = fetch_historical_returns('^GSPC')

# # Compute metrics for each asset
# for asset in portfolio:
#     returns = asset['historical_returns']
#     asset['volatility'] = compute_volatility(returns)
#     asset['sharpe_ratio'] = compute_sharpe_ratio(returns)
#     asset['beta'] = compute_beta(returns, market_returns)
#     print(f"{asset['ticker']} -> Volatility: {asset['volatility']:.4f}, Sharpe: {asset['sharpe_ratio']:.4f}, Beta: {asset['beta']:.4f}")

# 3. Add ARIMA Forecasting for Returns
from statsmodels.tsa.arima.model import ARIMA

# def arima_forecast(returns, forecast_days=30):
#     """Use ARIMA to forecast future returns."""
#     model = ARIMA(returns, order=(1, 1, 1))  # ARIMA(1,1,1)
#     model_fit = model.fit()
#     return model_fit.forecast(steps=forecast_days)
from statsmodels.tsa.arima.model import ARIMA
import pandas as pd

def arima_forecast(returns, forecast_days=30):
    """Use ARIMA to forecast future returns."""
    # Ensure `returns` is a Pandas Series
    if not isinstance(returns, pd.Series):
        returns = pd.Series(returns)

    # Assign a dummy date index with frequency if missing
    if not isinstance(returns.index, pd.DatetimeIndex) or returns.index.freq is None:
        returns.index = pd.date_range(start="2023-01-01", periods=len(returns), freq="D")

    # ARIMA Model
    model = ARIMA(returns, order=(1, 1, 1))
    model_fit = model.fit()

    # Forecast
    forecast = model_fit.forecast(steps=forecast_days)
    return forecast


# Forecast for each asset
# forecast_days = 30
# for asset in portfolio:
#     returns = asset['historical_returns']
#     asset['forecasted_returns'] = arima_forecast(returns, forecast_days)
#     print(f"Forecasted {forecast_days} returns for {asset['ticker']}")

# 4. Simulate Realistic Fluctuations
import random

def simulate_fluctuations(base_value, volatility, days=30):
    """Simulate fluctuations based on volatility."""
    simulated = [base_value]
    for _ in range(1, days):
        noise = random.uniform(-volatility, volatility)
        simulated.append(simulated[-1] * (1 + noise))
    return simulated

# Simulate for each asset
# for asset in portfolio:
#     volatility = asset['volatility']
#     base_return = asset['forecasted_returns'].iloc[0]
#     asset['simulated_returns'] = simulate_fluctuations(base_return, volatility)
#     print(f"Simulated returns for {asset['ticker']}: {asset['simulated_returns'][:5]}")

# 5. Validate Predictions
def validate_predictions(predictions, historical_returns, threshold=0.1):
    """Smooth out predictions exceeding a threshold deviation."""
    validated = []
    # last_known = historical_returns[-1]
    last_known = historical_returns.iloc[-1]
    for pred in predictions:
        if abs(pred - last_known) > threshold:
            pred = (pred + last_known) / 2  # Smooth deviations
        validated.append(pred)
        last_known = pred
    return validated

# Validate for each asset
# for asset in portfolio:
#     asset['validated_returns'] = validate_predictions(
#         asset['simulated_returns'], asset['historical_returns'])
#     print(f"Validated returns for {asset['ticker']}")


# 6. Visualize Predictions with Confidence Bands
import matplotlib.pyplot as plt

def plot_with_confidence(asset_name, best_case, worst_case):
    """Plot predictions with confidence bands."""
    days = len(best_case)
    dates = [datetime.datetime.today() + datetime.timedelta(days=i) for i in range(days)]
    
    lower_bound = [min(b, w) for b, w in zip(best_case, worst_case)]
    upper_bound = [max(b, w) for b, w in zip(best_case, worst_case)]
    
    plt.figure(figsize=(10, 6))
    plt.fill_between(dates, lower_bound, upper_bound, color='pink', alpha=0.2, label="Confidence Band")
    plt.plot(dates, best_case, color='red', label='Best Case')
    plt.plot(dates, worst_case, color='blue', linestyle='dashed', label='Worst Case')
    plt.title(f"{asset_name} Return Predictions")
    plt.xlabel("Date")
    plt.ylabel("Returns")
    plt.legend()
    plt.show()

# # # # Plot for each asset
# # for asset in portfolio:
# #     best_case = [r * 1.05 for r in asset['validated_returns']]
# #     worst_case = [r * 0.95 for r in asset['validated_returns']]
# #     plot_with_confidence(asset['ticker'], best_case, worst_case)






############################################################################################
# Endpoint to predict returns

# Define directories
PORTFOLIO_DIR = "local_data/portfolios"
CLIENT_SUMMARY_DIR = "client_data/client_data"
PREDICTIONS_DIR = "local_data/predictions"
os.makedirs(PREDICTIONS_DIR, exist_ok=True)

# # Updated New Implementation Version :

#  getting error

# @app.route('/predict_returns', methods=['POST'])
# def predict_returns():
#     try:
#         # Retrieve client and portfolio details
#         client_id = request.json.get('client_id')
#         client_name = request.json.get('client_name')
#         funds = request.json.get('funds')
#         investor_personality = request.json.get('investor_personality', 'Aggressive Investor Personality')

#         # Load portfolio data
#         portfolio_file = os.path.join(PORTFOLIO_DIR, f"portfolio_{client_id}.json")
#         portfolio_data = load_from_file(portfolio_file)
#         if not portfolio_data:
#             return jsonify({"message": f"No portfolio data found for client ID: {client_id}"}), 404

#         # Calculate portfolio-level metrics
#         total_current_value = sum(asset["current_value"] for asset in portfolio_data)
#         total_daily_change = sum(asset["Daily_Value_Change"] for asset in portfolio_data)
#         total_investment_gain_loss = sum(asset["Investment_Gain_or_Loss"] for asset in portfolio_data)

#         # Ensure calculations are meaningful
#         total_daily_change_perc = (total_daily_change / total_current_value * 100) if total_current_value else 0
#         total_investment_gain_loss_perc = (total_investment_gain_loss / total_current_value * 100) if total_current_value else 0

#         # Load client financial data
#         client_summary_file = os.path.join(CLIENT_SUMMARY_DIR, f"{client_id}.json")
#         client_financial_data = load_from_file(client_summary_file)
#         if not client_financial_data:
#             return jsonify({"message": f"No client financial data found for client ID: {client_id}"}), 404
        
        
#         # Initialize economic news to pass to LLM
#         topics = ["rising interest rates", "U.S. inflation", "geopolitical tensions", "US Elections", "Global Wars"]
#         economic_news = {topic: fetch_news(topic) for topic in topics}
#         portfolio_news = collect_portfolio_news(portfolio_data)

#         # Generate date intervals for next quarter
#         date_intervals = get_next_quarter_dates()
    
#         # next_quarter = "2024_Q4" 
#         next_quarter = get_next_quarter()
#         print(f"Next Quarter : {next_quarter}")
        
#         # Fetch market index returns for Beta calculation
#         market_returns = fetch_historical_returns('^GSPC')

#         forecast_days = 90
#         threshold = 0.1  # Validation threshold

#         confidence_data = []
        
#         for asset in portfolio_data:
#             ticker = asset.get('symbol')  # Extract the symbol from the asset
#             if not ticker:
#                 return jsonify({"error": "Missing ticker in portfolio asset."}), 400

#             # Fetch historical returns
#             historical_returns = fetch_historical_returns(ticker)
#             asset['historical_returns'] = historical_returns.tolist()

#             # Calculate metrics
#             volatility = compute_volatility(historical_returns)
#             sharpe_ratio = compute_sharpe_ratio(historical_returns)

#             # Align market returns for beta calculation
#             min_length = min(len(historical_returns), len(market_returns))
#             aligned_historical_returns = historical_returns.iloc[-min_length:]
#             aligned_market_returns = market_returns.iloc[-min_length:]

#             beta = compute_beta(aligned_historical_returns, aligned_market_returns)

#             asset['volatility'] = volatility
#             asset['sharpe_ratio'] = sharpe_ratio
#             asset['beta'] = beta

#             # ARIMA Forecast
#             try:
#                 forecasted_returns = arima_forecast(historical_returns, forecast_days)
#                 asset['forecasted_returns'] = forecasted_returns.tolist()
#             except Exception as e:
#                 print(f"ARIMA failed for {ticker}: {e}")
#                 asset['forecasted_returns'] = [0] * forecast_days

#             # Simulate fluctuations
#             base_return = forecasted_returns.iloc[0] if len(forecasted_returns) > 0 else 0
#             simulated_returns = simulate_fluctuations(base_return, volatility, forecast_days)
#             asset['simulated_returns'] = simulated_returns

#             # Validate predictions
#             validated_returns = validate_predictions(simulated_returns, historical_returns, threshold)
#             asset['validated_returns'] = validated_returns

#         # Prepare the task prompt for LLM
#         task = f"""
#             You are the world's best financial advisor tasked with predicting the next quarter's (3-month) returns for a client's portfolio containing various assets.
#             You are the best at Predicting client's portfolio returns and nothing is impossible for you to predict.
#             The client, {client_name}, has the following portfolio:

#             Portfolio Details: {portfolio_data}
#             Financial Situation: {client_financial_data}
#             Available Funds: ${funds}
#             Investor Personality: {investor_personality}
#             Portfolio News: {portfolio_news}
#             Economic News: {economic_news}

#             Confidence Data for Assets: {confidence_data}

#             Predict the expected returns (in percentages and dollar amounts) for the overall portfolio at the following dates:
#             {date_intervals}
            
#             Example of simulated_response = 
#             | Date       | Total Return (%) | Total Return ($) |
#             |------------|------------------|------------------|
#             | 2024-04-01 | 4.5%             | $10,500          |
#             | 2024-04-15 | 5.0%             | $10,800          |
#             | 2024-04-30 | 5.2%             | $11,000          |
#             |------------|------------------|------------------|
            
#             Your Response must be in the above table format no messages is required just table format data.
#         """

#         # Simulate LLM prediction
#         model = genai.GenerativeModel('gemini-1.5-flash')
#         response = model.generate_content(task)

#         # Process the response
#         simulated_response = markdown_to_text(response.text)

#         # Extract line chart data from the simulated response
#         line_chart_data = extract_line_chart_data(simulated_response)

#         # Save line chart data locally
#         prediction_file = os.path.join(PREDICTIONS_DIR, f"{client_id}_{next_quarter}_line_chart.json")
#         save_to_file(prediction_file, line_chart_data)

#         # Return the response
#         return jsonify({
#             "client_id": client_id,
#             "client_name": client_name,
#             "portfolio_value": total_current_value,
#             "daily_change_percentage": total_daily_change_perc,
#             "gain_loss_percentage": total_investment_gain_loss_perc,
#             "predicted_returns": simulated_response,
#             "line_chart_data": line_chart_data,
#             "confidence_data": confidence_data
#         }), 200

#     except Exception as e:
#         print(f"Error in predicting returns: {e}")
#         return jsonify({"message": f"Error predicting returns: {e}"}), 500



# original version :

@app.route('/predict_returns', methods=['POST'])
def predict_returns():
    try:
        # Retrieve client and portfolio details
        client_id = request.json.get('client_id')
        client_name = request.json.get('client_name')
        funds = request.json.get('funds')
        investor_personality = request.json.get('investor_personality', 'Aggressive Investor Personality')

        # Load portfolio data
        portfolio_file = os.path.join(PORTFOLIO_DIR, f"portfolio_{client_id}.json")
        portfolio_data = load_from_file(portfolio_file)
        if not portfolio_data:
            return jsonify({"message": f"No portfolio data found for client ID: {client_id}"}), 404

        # Calculate portfolio-level metrics
        total_current_value = sum(asset["current_value"] for asset in portfolio_data)
        total_daily_change = sum(asset["Daily_Value_Change"] for asset in portfolio_data)
        total_investment_gain_loss = sum(asset["Investment_Gain_or_Loss"] for asset in portfolio_data)

        # Ensure calculations are meaningful
        total_daily_change_perc = (total_daily_change / total_current_value * 100) if total_current_value else 0
        total_investment_gain_loss_perc = (total_investment_gain_loss / total_current_value * 100) if total_current_value else 0

        # Load client financial data
        client_summary_file = os.path.join(CLIENT_SUMMARY_DIR, f"{client_id}.json")
        client_financial_data = load_from_file(client_summary_file)
        if not client_financial_data:
            return jsonify({"message": f"No client financial data found for client ID: {client_id}"}), 404
        
         # Initialize economic news to pass to LLM
        topics = ["rising interest rates", "U.S. inflation", "geopolitical tensions", "US Elections", "Global Wars"]
        economic_news = {topic: fetch_news(topic) for topic in topics}
        portfolio_news = collect_portfolio_news(portfolio_data)

        # Generate date intervals for next quarter
        # date_intervals = [
        #     "2024-10-01", "2024-10-15", "2024-10-31",
        #     "2024-11-01", "2024-11-15", "2024-11-30",
        #     "2024-12-01", "2024-12-15", "2024-12-31"
        # ] 
        date_intervals = get_next_quarter_dates()
    
        # next_quarter = "2024_Q4" 
        next_quarter = get_next_quarter()
        print(f"Next Quarter : {next_quarter}")
        
        # Prepare the task prompt for LLM
        task = f"""
            You are the world's best financial advisor tasked with predicting the next quarter's (3-month) returns for a client's portfolio containing various assets.
            You are the best at Predicting client's portfolio returns and nothing is impossible for you to predict.
            The client, {client_name}, has the following portfolio:

            Portfolio Details: {portfolio_data}
            Financial Situation: {client_financial_data}
            Available Funds: ${funds}
            Investor Personality: {investor_personality}
            Portfolio News: {portfolio_news}
            Economic News: {economic_news}

            Analyze the portfolio and each assets in the portfolio properly and also refer to the Portfolio news and Economic News for your reference and Performance of the assets.
            Predict the expected returns (in percentages and dollar amounts) for the overall portfolio at the following dates:
            {date_intervals}
            
            Example of simulated_response = 
            | Date       | Total Return (%) | Total Return ($) |
            |------------|------------------|------------------|
            | 2024-04-01 | 4.5%             | $10,500          |
            | 2024-04-15 | 5.0%             | $10,800          |
            | 2024-04-30 | 5.2%             | $11,000          |
            |------------|------------------|------------------|
            
            Your Response must be in the above table format no messages is required just table format data.
        """

        # Simulate LLM prediction
        model = genai.GenerativeModel('gemini-1.5-flash')
        response = model.generate_content(task)

        # Process the response
        html_suggestions = markdown.markdown(response.text)
        
        print(f"\nHTML Suggestions : {html_suggestions}")
        
        simulated_response = markdown_to_text(html_suggestions)
        
        print(f"\nSimulated Response : {simulated_response}")
        # Extract line chart data from the simulated response
        line_chart_data = extract_line_chart_data(simulated_response)
        
        print(f"\nLine Chart Data : {line_chart_data}")
        
        # Save line chart data locally
        prediction_file = os.path.join(PREDICTIONS_DIR, f"{client_id}_{next_quarter}_line_chart.json")
        save_to_file(prediction_file, line_chart_data)

        # Return the response
        return jsonify({
            "client_id": client_id,
            "client_name": client_name,
            "predicted_returns": simulated_response,
            "line_chart_data": line_chart_data
        }), 200

    except Exception as e:
        print(f"Error in predicting returns: {e}")
        return jsonify({"message": f"Error predicting returns: {e}"}), 500



########################################################################################################################
# Portfolio Return on Investment Prediction for next quarter using aws :

# Determine Next Quarter Date Intervals :
# from datetime import datetime, timedelta
# import calendar

# def get_next_quarter_dates():
#     current_date = datetime.now()
#     current_month = current_date.month

#     # Determine the starting month of the next quarter
#     if current_month in [1, 2, 3]:  # Q1
#         start_month = 4  # Q2
#     elif current_month in [4, 5, 6]:  # Q2
#         start_month = 7  # Q3
#     elif current_month in [7, 8, 9]:  # Q3
#         start_month = 10  # Q4
#     else:  # Q4
#         start_month = 1  # Q1 of the next year

#     # Determine the year of the next quarter
#     next_quarter_year = current_date.year if start_month != 1 else current_date.year + 1

#     # Generate dates for the next quarter
#     next_quarter_dates = []
#     for month in range(start_month, start_month + 3):
#         # Get the first, 15th, and last day of the month
#         first_day = datetime(next_quarter_year, month, 1)
#         fifteenth_day = datetime(next_quarter_year, month, 15)
#         last_day = datetime(next_quarter_year, month, calendar.monthrange(next_quarter_year, month)[1])

#         next_quarter_dates.extend([first_day.strftime("%Y-%m-%d"), 
#                                    fifteenth_day.strftime("%Y-%m-%d"), 
#                                    last_day.strftime("%Y-%m-%d")])

#     return next_quarter_dates



# # Line Chart for Predcting Next quarter Returns :
# def extract_line_chart_data(llm_response_text):
#     """
#     Extracts line chart data from the LLM's response for plotting.
#     """
#     try:
#         # Example parsing logic for the response (modify as needed)
#         lines = llm_response_text.split("\n")
#         line_chart_data = {
#             "dates": [],
#             "overall_returns": {"percentages": [], "amounts": []}
#         }
#         current_date = datetime.now()
#         current_year = current_date.year
#         for line in lines:
#             if line.startswith(f"| {current_year}-"):  # Ex: "| 2024-01-01 |"
#                 parts = line.split("|")
#                 date = parts[1].strip()
#                 return_percentage = float(parts[2].replace("%", "").strip())
#                 return_amount = float(parts[3].replace("$", "").strip())
#                 line_chart_data["dates"].append(date)
#                 line_chart_data["overall_returns"]["percentages"].append(return_percentage)
#                 line_chart_data["overall_returns"]["amounts"].append(return_amount)
#         return line_chart_data
#     except Exception as e:
#         print(f"Error extracting line chart data: {e}")
#         return {}


# # V-3 : Actual Line Chart 
# import hashlib
# import json
# from datetime import datetime

# Global variable to store the hash of the last portfolio
# last_portfolio_hash = None

# V-3 : Check for Portdolio Changes implemented :
# @app.route('/predict_returns', methods=['POST'])
# def predict_returns():
#     global last_portfolio_hash
#     try:
#         # Retrieve portfolio and client data from the request
#         client_id = request.json.get('client_id')
#         client_name = request.json.get('client_name')
#         funds = request.json.get('funds')
#         investor_personality = request.json.get('investor_personality', 'Aggressive Investor Personality')

#         # Load portfolio data
#         with open(f'portfolio_{client_id}.json', 'r') as f:
#             portfolio_data = json.load(f)

#         # Calculate the hash of the current portfolio data
#         current_portfolio_hash = hashlib.sha256(json.dumps(portfolio_data, sort_keys=True).encode()).hexdigest()

#         # Check if the portfolio is the same as before
#         if last_portfolio_hash == current_portfolio_hash:
#             print("Portfolio data is unchanged. Using previously generated predictions.")
#             # Load previously stored predictions and line chart data
#             try:
#                 with open(f'predictions_{client_id}.json', 'r') as f:
#                     previous_predictions = json.load(f)
#                 return jsonify(previous_predictions), 200
#             except FileNotFoundError:
#                 return jsonify({"message": "No previous predictions found. Please update the portfolio."}), 404

#         # Update the global portfolio hash
#         last_portfolio_hash = current_portfolio_hash

#         # Load financial data from S3
#         s3_key = f"{client_summary_folder}client-data/{client_id}.json"
#         try:
#             response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
#             client_financial_data = json.loads(response['Body'].read().decode('utf-8'))
#         except Exception as e:
#             logging.error(f"Error retrieving client financial data: {e}")
#             return jsonify({'message': f'Error retrieving client financial data: {e}'}), 500

#         # Fetch news for the portfolio assets
#         portfolio_news = collect_portfolio_news(portfolio_data)

#         # Prepare date intervals for predictions
#         date_intervals = get_next_quarter_dates()

#         # Prepare prompt for the LLM
#         task = f"""
#             You are a financial advisor tasked with predicting the next quarter's (3-month) returns for a client's portfolio.
#             The client, {client_name}, has the following portfolio:
            
#             Portfolio Details: {portfolio_data}
#             Financial Situation: {client_financial_data}
#             Available Funds: ${funds}
#             Investor Personality: {investor_personality}
            
#             Consider these factors:
#             1. Economic trends such as inflation, interest rates, and geopolitical events.
#             2. Past performance of assets in the portfolio.
#             3. Risk tolerance based on investor personality.
#             4. The current news and economic news for the assets in the portfolio: {portfolio_news}

#             Predict the expected returns (in percentages and dollar amounts) for each asset and the overall portfolio at the following dates:
#             {date_intervals}

#             Provide the output in the following format:
            
#             #### Predicted Returns:
#             - **Asset-wise Predictions (Per Date)**:
#               | Date       | Asset Name | Predicted Return (%) | Predicted Return ($) |
#               |------------|------------|----------------------|-----------------------|
#               | 2024-01-01 | Asset 1    | 5.5%                | $500                 |
#               | 2024-01-15 | Asset 1    | 5.8%                | $520                 |
#               | ...        | ...        | ...                 | ...                  |

#             - **Overall Portfolio Return**:
#               | Date       | Total Return (%) | Total Return ($) |
#               |------------|------------------|------------------|
#               | 2024-01-01 | 4.5%            | $10,500          |
#               | ...        | ...             | ...              |

#             Ensure the output is comprehensive and formatted for easy parsing into a line chart.
#         """

#         # Call the LLM model to generate predictions
#         try:
#             model = genai.GenerativeModel('gemini-1.5-flash')
#             response = model.generate_content(task)

#             # Process the LLM response
#             html_predictions = markdown.markdown(response.text)
#             formatted_predictions = markdown_to_text(html_predictions)

#             # Extract line chart data from the response
#             predicted_line_chart_data = extract_line_chart_data(response.text)

#             # Fetch actual returns for the portfolio
#             actual_line_chart_data = get_actual_returns(client_id)

#             # Combine actual and predicted data for line chart
#             line_chart_data = {
#                 "actual": actual_line_chart_data,
#                 "predicted": predicted_line_chart_data
#             }

#             # Save predictions and line chart data for reuse
#             result = {
#                 "client_id": client_id,
#                 "client_name": client_name,
#                 "predicted_returns": formatted_predictions,
#                 "line_chart_data": line_chart_data
#             }
#             with open(f'predictions_{client_id}.json', 'w') as f:
#                 json.dump(result, f, indent=4)

#             # Return response with line chart data
#             return jsonify(result), 200

#         except Exception as e:
#             print(f"Error generating predictions from LLM: {e}")
#             return jsonify({"message": f"Error generating predictions: {e}"}), 500

#     except Exception as e:
#         print(f"Error in predicting returns: {e}")
#         return jsonify({"message": f"Error predicting returns: {e}"}), 500


# # V-2 : with Line Chart Data
# # # Predict Next Quarter Returns
# @app.route('/predict_returns', methods=['POST'])
# def predict_returns():
#     try:
#         # Retrieve portfolio and client data from the request
#         client_id = request.json.get('client_id')
#         client_name = request.json.get('client_name')
#         funds = request.json.get('funds')
#         investor_personality = request.json.get('investor_personality', 'Aggressive Investor Personality')

#         # Load portfolio data
#         with open(f'portfolio_{client_id}.json', 'r') as f:
#             portfolio_data = json.load(f)

#         # Load financial data from S3
#         s3_key = f"{client_summary_folder}client-data/{client_id}.json"
#         try:
#             response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
#             client_financial_data = json.loads(response['Body'].read().decode('utf-8'))
#         except Exception as e:
#             logging.error(f"Error retrieving client financial data: {e}")
#             return jsonify({'message': f'Error retrieving client financial data: {e}'}), 500

#         # Fetch news for the portfolio assets
#         portfolio_news = collect_portfolio_news(portfolio_data)

#         # Prepare date intervals for predictions
        
#         # date_intervals = [
#         #     "2024-01-01", "2024-01-15", "2024-01-31",
#         #     "2024-02-15", "2024-02-29",
#         #     "2024-03-15", "2024-03-31"
#         # ]
        
#         date_intervals = get_next_quarter_dates()

#         # Prepare prompt for the LLM
#         task = f"""
#             You are a financial advisor tasked with predicting the next quarter's (3-month) returns for a client's portfolio.
#             The client, {client_name}, has the following portfolio:
            
#             Portfolio Details: {portfolio_data}
#             Financial Situation: {client_financial_data}
#             Available Funds: ${funds}
#             Investor Personality: {investor_personality}
            
#             Consider these factors:
#             1. Economic trends such as inflation, interest rates, and geopolitical events.
#             2. Past performance of assets in the portfolio.
#             3. Risk tolerance based on investor personality.
#             4. The current news and economic news for the assets in the portfolio: {portfolio_news}

#             Predict the expected returns (in percentages and dollar amounts) for each asset and the overall portfolio at the following dates:
#             {date_intervals}

#             Provide the output in the following format:
            
#             #### Predicted Returns:
#             - **Asset-wise Predictions (Per Date)**:
#               | Date       | Asset Name | Predicted Return (%) | Predicted Return ($) |
#               |------------|------------|----------------------|-----------------------|
#               | 2024-01-01 | Asset 1    | 5.5%                | $500                 |
#               | 2024-01-15 | Asset 1    | 5.8%                | $520                 |
#               | ...        | ...        | ...                 | ...                  |

#             - **Overall Portfolio Return**:
#               | Date       | Total Return (%) | Total Return ($) |
#               |------------|------------------|------------------|
#               | 2024-01-01 | 4.5%            | $10,500          |
#               | ...        | ...             | ...              |

#             Ensure the output is comprehensive and formatted for easy parsing into a line chart.
#         """

#         # Call the LLM model to generate predictions
#         try:
#             model = genai.GenerativeModel('gemini-1.5-flash')
#             response = model.generate_content(task)

#             # Process the LLM response
#             html_predictions = markdown.markdown(response.text)
#             formatted_predictions = markdown_to_text(html_predictions)

#             # Extract line chart data from the response
#             line_chart_data = extract_line_chart_data(response.text)

#             # Return response with line chart data
#             return jsonify({
#                 "client_id": client_id,
#                 "client_name": client_name,
#                 "predicted_returns": formatted_predictions,
#                 "line_chart_data": line_chart_data
#             }), 200

#         except Exception as e:
#             print(f"Error generating predictions from LLM: {e}")
#             return jsonify({"message": f"Error generating predictions: {e}"}), 500

#     except Exception as e:
#         print(f"Error in predicting returns: {e}")
#         return jsonify({"message": f"Error predicting returns: {e}"}), 500


# # V-1 : without line chart
# @app.route('/predict_returns', methods=['POST'])
# def predict_returns():
#     try:
#         # Retrieve portfolio data from the request
#         client_id = request.json.get('client_id')
#         client_name = request.json.get('client_name')
#         funds = request.json.get('funds')
#         investor_personality = request.json.get('investor_personality', 'Aggressive Investor Personality')
        
#         # Load portfolio data
#         with open(f'portfolio_{client_id}.json', 'r') as f:
#             portfolio_data = json.load(f)

#         # Load financial data from S3
#         s3_key = f"{client_summary_folder}client-data/{client_id}.json"
#         try:
#             response = s3.get_object(Bucket=S3_BUCKET_NAME, Key=s3_key)
#             client_financial_data = json.loads(response['Body'].read().decode('utf-8'))
#         except Exception as e:
#             logging.error(f"Error retrieving client financial data: {e}")
#             return jsonify({'message': f'Error retrieving client financial data: {e}'}), 500
        
#         # Fetch news for each asset in the specified list
#         portfolio_news = collect_portfolio_news(portfolio_data)
        
#         # Prepare prompt for LLM
#         task = f"""
#             You are a financial advisor tasked with predicting the next quarter's (3-month) returns for a client's portfolio.
#             The client, {client_name}, has the following portfolio:
            
#             Portfolio Details: {portfolio_data}
#             Financial Situation: {client_financial_data}
#             Available Funds: ${funds}
#             Investor Personality: {investor_personality}
            
#             Consider these factors:
#             1. Economic trends such as inflation, interest rates, and geopolitical events.
#             2. Past performance of assets in the portfolio.
#             3. Risk tolerance based on investor personality.
#             4. The current news and economic news for the assets in the portfolio : {portfolio_news}

#             Predict the expected returns for each asset (in both percentages and dollar amounts) and the overall portfolio. 
#             Include insights on how market conditions and client financial goals may affect these predictions.

#             Provide the output in the following format:
            
#             #### Predicted Returns:
#             - **Asset-wise Predictions**:
#               | Asset Name | Predicted Return (%) | Predicted Return ($) |
#               |------------|----------------------|-----------------------|
#               | Asset 1    | 5.5%                | $500                 |
#               | ...        | ...                 | ...                  |

#             - **Overall Portfolio Return**:
#               | Metric              | Value   |
#               |---------------------|---------|
#               | Total Return (%)    | 8.5%    |
#               | Total Return ($)    | $10,500 |
#         """

#         # Call the LLM model to generate predictions
#         try:
#             model = genai.GenerativeModel('gemini-1.5-flash')
#             response = model.generate_content(task)

#             # Process the LLM response
#             html_predictions = markdown.markdown(response.text)
#             formatted_predictions = markdown_to_text(html_predictions)

#             # Return response
#             return jsonify({
#                 "client_id": client_id,
#                 "client_name": client_name,
#                 "predicted_returns": formatted_predictions
#             }), 200

#         except Exception as e:
#             print(f"Error generating predictions from LLM: {e}")
#             return jsonify({"message": f"Error generating predictions: {e}"}), 500

#     except Exception as e:
#         print(f"Error in predicting returns: {e}")
#         return jsonify({"message": f"Error predicting returns: {e}"}), 500



# Run the Flask application
if __name__ == '__main__':
    app.run(host='0.0.0.0',debug=True)
